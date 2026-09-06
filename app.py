import streamlit as st
import streamlit.components.v1 as components
import os
import requests
import re
import json
import csv
import html
import time
import datetime
import copy
import base64
import hashlib
import math
import tempfile
import uuid
from fpdf import FPDF
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, ns
from io import BytesIO, StringIO
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
import PyPDF2  # Libreria necessaria per leggere i PDF caricati
from PIL import Image
from editorial_rules import (
    PROFILI_LUNGHEZZA_STESURA,
    chiave_sezione,
    chiave_sezione_precedente,
    classifica_sezione,
    controllo_completezza_testi_gratuito as controllo_completezza_core,
    minimo_parole_per_sezione_editoriale as minimo_parole_core,
    motivo_chiusura_tecnica,
    stati_sezioni_editoriali as stati_sezioni_core,
    vincolo_parole_con_tolleranza,
)
from project_csv import (
    LIMITE_CAMPO_CSV_PROGETTO,
    esporta_fotografia_csv,
    importa_fotografia_csv,
    imposta_limite_lettura_csv_progetto,
)
import project_memory as memoria_core
from commercial_layer import (
    AI_REQUEST_CREDITS,
    CREDIT_COSTS,
    COMMERCIAL_TEST_VERSION,
    CommercialCreditError,
    bootstrap_commercial_test,
    charge_credits,
    completa_logout_sicuro,
    annulla_logout_sicuro,
    logout_sicuro_richiesto,
    mostra_crediti_esauriti,
    registra_utilizzo_ai,
    refund_credits,
)

# Compatibilità temporanea: l'app resta avviabile anche se il file commerciale
# su GitHub viene aggiornato qualche istante dopo app.py.
try:
    from commercial_layer import (
        carica_progetto_automatico,
        elimina_progetto_automatico,
        salva_progetto_automatico,
    )
except ImportError:
    def carica_progetto_automatico():
        return {}

    def salva_progetto_automatico(_snapshot):
        return False

    def elimina_progetto_automatico():
        return None

# ======================================================================================================================
# 0. GESTIONE MEMORIA DI STATO E PREVENZIONE AUTO-RESET
# ======================================================================================================================
# Questo blocco garantisce che l'applicazione mantenga i dati in memoria durante le elaborazioni lunghe
# e i cambi di tab. I dati verranno azzerati SOLO tramite l'esplicito pulsante di RESET.
if "memoria_blindata" not in st.session_state:
    st.session_state["memoria_blindata"] = True
    # Ogni browser riceve una cartella temporanea privata: nessun file o stato viene condiviso tra utenti.
    st.session_state["id_sessione_utente"] = uuid.uuid4().hex
    st.session_state["tmp_dir"] = os.path.join(
        tempfile.gettempdir(), "ebook_creator_sessions", st.session_state["id_sessione_utente"]
    )
    st.session_state["indice_raw"] = ""
    st.session_state["lista_capitoli"] = []
    st.session_state["conoscenza_extra"] = ""
    st.session_state["immagini_capitoli"] = {}

# Compatibilità con sessioni aperte prima dell'introduzione delle immagini.
if "immagini_capitoli" not in st.session_state:
    st.session_state["immagini_capitoli"] = {}

# ======================================================================================================================
# FUNZIONI DI SUPPORTO PER ANALISI DOCUMENTI (NUOVO MODULO)
# ======================================================================================================================
def estrai_testo_da_files(caricati):
    testo_totale = ""
    for file in caricati:
        try:
            if file.name.lower().endswith('.pdf'):
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    testo_totale += (page.extract_text() or "") + "\n"
            elif file.name.lower().endswith('.docx'):
                doc = Document(file)
                for para in doc.paragraphs:
                    testo_totale += para.text + "\n"
        except Exception as e:
            st.error(f"Errore nella lettura di {file.name}: {e}")
    return testo_totale


def firma_fonti_esterne(caricati):
    """Riconosce i file caricati senza rileggerli a ogni aggiornamento della pagina."""
    digest = hashlib.sha256()
    for file in caricati or []:
        contenuto = file.getvalue()
        digest.update(file.name.encode("utf-8", "ignore"))
        digest.update(str(len(contenuto)).encode("ascii"))
        digest.update(hashlib.sha256(contenuto).digest())
    return digest.hexdigest()


def crea_scheda_fonti(testo, limite=2600):
    """Sintesi locale e rapida delle fonti, senza chiamate API supplementari."""
    paragrafi = [re.sub(r"\s+", " ", p).strip() for p in re.split(r"\n\s*\n|(?<=\.)\s{2,}", testo or "")]
    paragrafi = [p for p in paragrafi if len(p) > 80]
    scelti, usati = [], 0
    for paragrafo in paragrafi:
        if usati + len(paragrafo) > limite:
            break
        scelti.append(paragrafo)
        usati += len(paragrafo)
    return "\n".join(scelti) or (testo or "")[:limite]


MODELLO_STESURA = os.getenv("WRITING_MODEL", "gpt-5.4-mini")
# Il modello completo viene usato solo dove il ragionamento editoriale pesa davvero.
MODELLO_EDITORIALE = os.getenv("EDITORIAL_MODEL", "gpt-5.4")
MODELLO_ANALISI_FONTI = os.getenv("SOURCE_ANALYSIS_MODEL", MODELLO_EDITORIALE)
# Il copyright completo usa il mini su ogni blocco e riserva il modello completo
# ai soli lotti che presentano un rischio: qualità dove serve, costi contenuti.
MODELLO_CONTROLLO_COPYRIGHT_COMPLETO = os.getenv("COPYRIGHT_SCREENING_MODEL", MODELLO_STESURA)
MODELLO_CONTROLLO_COPYRIGHT_APPROFONDITO = os.getenv("COPYRIGHT_REVIEW_MODEL", MODELLO_EDITORIALE)
MODELLO_DEEPSEEK_PRO = os.getenv("DEEPSEEK_MODEL", "deepseek-v4-pro")
# Le API client possono attendere fino a dieci minuti per impostazione
# predefinita. Per indice e fonti è un'attesa eccessiva: una risposta che si
# blocca deve tornare alla UI con un esito chiaro, non lasciare lo spinner
# attivo senza fine. I retry sono gestiti dall'app con logica editoriale,
# quindi vengono disattivati nelle chiamate con un limite esplicito.
TIMEOUT_RICERCA_WEB_SECONDI = 90.0
TIMEOUT_INDICE_SECONDI = 150.0


def provider_ia_selezionato():
    """La scelta è dell'utente e viene conservata nel progetto esportabile."""
    return str(st.session_state.get("provider_ia", "GPT-5.4")).strip()


def usa_deepseek_pro():
    return provider_ia_selezionato().casefold().startswith("deepseek")


def studia_fonti_con_ai(testo, limite_input=30000):
    """Trasforma le fonti in una mappa concettuale interna e indipendente.

    La mappa non viene mai mostrata nel libro e non conserva formulazioni delle fonti:
    serve soltanto come base concettuale per indice e stesura.
    """
    testo = (testo or "").strip()
    if not testo:
        return ""
    estratto = testo[:limite_input]
    try:
        dossier = chiedi_gpt(
            (
                "Sei un ricercatore editoriale e un progettista didattico. Trasforma le fonti "
                "qui sotto in una MAPPA CONCETTUALE INTERNA per un libro davvero autonomo. "
                "Non fare un riassunto lineare, non mantenere ordine, titoli, esempi distintivi "
                "o formulazioni delle fonti e non copiare mai sequenze di sei o più parole. "
                "Scomponi invece la conoscenza in principi, relazioni causali, domande del lettore, "
                "procedure generiche, limiti e verifiche. Ricombina i concetti in una progressione "
                "editoriale nuova, adatta a un testo originale. Non inventare fatti.\n\n"
                "Restituisci esclusivamente queste voci interne: 1) concetti da spiegare con parole "
                "nuove; 2) fatti e dati da verificare prima della pubblicazione; 3) nessi logici; "
                "4) esempi nuovi che l'autore può costruire; 5) confini e cautele; 6) competenze da "
                "assegnare a capitoli e sottocapitoli. Questo materiale non sarà pubblicato.\n\nFONTI:\n" + estratto
            ),
            "Produci solo una mappa concettuale interna, originale e non pubblicabile.",
            addebita=False,
            model=MODELLO_ANALISI_FONTI,
        )
        return (dossier or "").strip() if not str(dossier).startswith("ERRORE:") else ""
    except Exception:
        # In caso di indisponibilità dell'AI non passiamo brani originali alla
        # stesura: è più prudente sospendere l'uso editoriale delle fonti.
        return ""


def firma_ricerca_preliminare(titolo, genere, trama, obiettivo, lingua, approfondimenti):
    """La ricerca viene riutilizzata finché il brief non cambia."""
    base = "\n".join([titolo or "", genere or "", trama or "", obiettivo or "", lingua or "", approfondimenti or ""])
    return hashlib.sha256(base.encode("utf-8", "ignore")).hexdigest()


def separa_mappa_e_registro_fonti_web(testo):
    """Separa la mappa usata dall'AI dal registro leggibile delle fonti web."""
    testo = (testo or "").strip()
    # I due cervelli possono rendere il titolo con Markdown o con una
    # formulazione leggermente diversa. Il parser resta quindi tollerante,
    # altrimenti il registro veniva confuso con il dossier interno e non era
    # visibile né salvabile come elenco delle fonti web.
    marcatore = re.search(
        r"(?im)^\s*(?:#{1,6}\s*)?REGISTRO\s+(?:DELLE\s+)?FONTI(?:\s+WEB)?\s*:?[ \t]*$",
        testo,
    )
    if marcatore:
        return testo[:marcatore.start()].strip(), testo[marcatore.end():].strip()

    # Estrema tutela: se il modello ha dimenticato il titolo del registro ma
    # ha comunque restituito URL, li conserviamo in una sezione interna invece
    # di perderli. La mappa rimane separata dai collegamenti.
    righe = testo.splitlines()
    righe_fonti = [riga for riga in righe if re.search(r"https?://\S+", riga)]
    if righe_fonti:
        mappa = "\n".join(riga for riga in righe if riga not in righe_fonti).strip()
        return mappa or testo, "\n".join(righe_fonti).strip()
    return testo, ""


def conserva_solo_fonti_web_selezionate(registro, massimo_fonti=5):
    """Conserva solo le fonti scelte dal ricercatore per questo brief.

    Il modello riceve la consegna di ordinare prima le fonti per pertinenza e
    autorevolezza. Qui manteniamo soltanto le prime voci selezionate: le fonti
    consultate ma non scelte non sono mostrate nell'interfaccia, non entrano
    nella memoria del progetto e non finiscono nel CSV.
    """
    righe = [
        riga.strip() for riga in str(registro or "").splitlines()
        if re.search(r"https?://\S+", riga)
    ]
    return "\n".join(righe[:max(1, int(massimo_fonti))])


def ricerca_preliminare_per_indice(titolo, genere, trama, obiettivo, lingua, approfondimenti, forza=False):
    """Cerca fonti con un tempo massimo, senza bloccare la generazione dell'indice.

    La ricerca migliora l'indice, ma non può diventare un punto morto: se il
    provider non chiude la richiesta entro il limite, il credito web viene
    restituito e l'indice prosegue usando il brief della sidebar.
    """
    firma = firma_ricerca_preliminare(titolo, genere, trama, obiettivo, lingua, approfondimenti)
    if not forza and st.session_state.get("firma_ricerca_preliminare") == firma:
        return st.session_state.get("dossier_ricerca_preliminare", "")

    riferimento = addebita_azione_diretta("ricerca_preliminare_indice", amount=CREDIT_COSTS["indice_ricerca_web"])
    try:
        istruzioni_ricerca = (
            "Sei un ricercatore editoriale. Cerca sul web fonti autorevoli e aggiornate utili "
            "per progettare un libro. Non scrivere il libro e non produrre citazioni per il lettore. "
            "Restituisci prima una MAPPA CONCETTUALE INTERNA concisa, non un riassunto delle pagine lette: "
            "concetti e definizioni affidabili; fatti, norme, date o dati da verificare; controversie o limiti; "
            "progressione didattica consigliata; aspetti concreti da assegnare all'indice. Riformula tutto con "
            "parole indipendenti, senza citazioni, titoli o formulazioni riconoscibili delle fonti. Non inserire URL, "
            "link Markdown o bibliografie nella MAPPA.\n\n"
            "Prima esplora liberamente le fonti necessarie, poi seleziona in autonomia SOLO le 3-5 più autorevoli, "
            "aggiornate e direttamente pertinenti a titolo, argomento, genere e obiettivo del brief. Scarta le fonti "
            "generiche, duplicate, deboli o marginali: non elencarle e non citarle. Alla fine inserisci obbligatoriamente "
            "una riga sola 'REGISTRO FONTI WEB' e sotto esclusivamente le fonti selezionate, in ordine di utilità, nel "
            "formato: - Titolo della fonte | URL completo | motivo di utilità. Il registro è solo per la schermata interna "
            "dell'utente: non inserire estratti testuali, citazioni o frasi tratte dalle pagine.\n\n"
            f"Titolo: {titolo}\nGenere: {genere}\nLingua del libro: {lingua}\n"
            f"Argomento: {trama}\nObiettivo: {obiettivo}\n"
            f"Approfondimenti: {approfondimenti or 'Nessuno'}"
        )
        if usa_deepseek_pro():
            # Ricerca nativa DeepSeek: nessuna chiamata a GPT. Il registro
            # viene conservato e mostrato nella stessa area delle fonti GPT.
            client_ricerca = client_deepseek.with_options(
                timeout=TIMEOUT_RICERCA_WEB_SECONDI, max_retries=0
            )
            risposta = client_ricerca.responses.create(
                model=MODELLO_DEEPSEEK_PRO,
                tools=[{"type": "web_search"}],
                tool_choice={"type": "web_search"},
                instructions="Usa la ricerca web integrata prima di rispondere. "
                "Le fonti servono solo alla progettazione interna e devono essere elencate nel registro richiesto.",
                input=istruzioni_ricerca,
            )
        else:
            client_ricerca = client.with_options(
                timeout=TIMEOUT_RICERCA_WEB_SECONDI, max_retries=0
            )
            risposta = client_ricerca.responses.create(
                model=MODELLO_ANALISI_FONTI,
                tools=[{"type": "web_search_preview"}],
                input=istruzioni_ricerca,
            )
        registra_esito_chiamata_ai(
            risposta,
            riferimento=riferimento,
            reason="ricerca_preliminare_indice",
            amount=CREDIT_COSTS["indice_ricerca_web"],
            model=MODELLO_DEEPSEEK_PRO if usa_deepseek_pro() else MODELLO_ANALISI_FONTI,
        )
        risposta_testo = (getattr(risposta, "output_text", "") or "").strip()
        mappa, registro = separa_mappa_e_registro_fonti_web(risposta_testo)
        registro = conserva_solo_fonti_web_selezionate(registro)
        if not mappa:
            refund_credits(riferimento, reason="ricerca_preliminare_vuota", amount=CREDIT_COSTS["indice_ricerca_web"])
            return ""
        st.session_state["firma_ricerca_preliminare"] = firma
        st.session_state["dossier_ricerca_preliminare"] = mappa
        st.session_state["registro_fonti_web"] = registro
        # La ricerca è un'operazione a consumo: proteggiamo subito dossier e
        # registro nel cloud, anche se l'utente aggiorna la pagina prima di
        # generare l'indice. I contenuti manuali restano invece soggetti al
        # pulsante SALVA SESSIONE.
        try:
            sezioni_correnti = list(dict.fromkeys([
                *st.session_state.get("lista_capitoli", []),
                *st.session_state.get(CHIAVE_MEMORIA_SEZIONI, {}).keys(),
            ]))
            if salva_progetto_corrente(sidebar_memorizzata_corrente(), sezioni_correnti):
                st.session_state["fonti_web_salvate"] = True
        except Exception:
            st.session_state["fonti_web_salvate"] = False
        return mappa
    except Exception as exc:
        refund_credits(riferimento, reason="ricerca_preliminare_fallita", amount=CREDIT_COSTS["indice_ricerca_web"])
        registra_esito_chiamata_ai(
            None,
            riferimento=riferimento,
            reason="ricerca_preliminare_indice",
            amount=CREDIT_COSTS["indice_ricerca_web"],
            model=MODELLO_DEEPSEEK_PRO if usa_deepseek_pro() else MODELLO_ANALISI_FONTI,
            riuscita=False,
            rimborsata=True,
            errore=type(exc).__name__,
        )
        st.session_state["ultimo_esito_ricerca_preliminare"] = (
            "Ricerca preliminare non disponibile o scaduta: l'indice verrà creato "
            "dal brief della sidebar. Il credito della ricerca è stato riaccreditato."
        )
        st.session_state["ultimo_errore_ricerca_preliminare"] = str(exc)
        return ""


def estratti_fonti_pertinenti(sezione, argomento, limite=3500):
    """Restituisce solo la mappa concettuale, mai i brani originali caricati."""
    mappa = st.session_state.get("brief_fonti_originale") or st.session_state.get("dossier_fonti_ai", "")
    if not mappa:
        return ""
    parole = set(re.findall(r"[a-zA-ZÀ-ÖØ-öø-ÿ0-9]{4,}", f"{sezione} {argomento}".lower()))
    paragrafi = [re.sub(r"\s+", " ", p).strip() for p in re.split(r"\n\s*\n|(?<=\.)\s{2,}", mappa)]
    valutati = []
    for posizione, paragrafo in enumerate(paragrafi):
        if len(paragrafo) < 80:
            continue
        punteggio = len(parole & set(re.findall(r"[a-zA-ZÀ-ÖØ-öø-ÿ0-9]{4,}", paragrafo.lower())))
        valutati.append((punteggio, posizione, paragrafo))
    valutati.sort(key=lambda elemento: (-elemento[0], elemento[1]))
    scelti, usati = [], 0
    for punteggio, _, paragrafo in valutati:
        if punteggio == 0 and scelti:
            break
        if usati + len(paragrafo) > limite:
            continue
        scelti.append(paragrafo)
        usati += len(paragrafo)
    materiali = "\n\n".join(scelti) or mappa[:limite]
    return f"MAPPA CONCETTUALE INTERNA (NON È TESTO DA RIPRENDERE):\n{materiali[:limite]}"


def controllo_originalita_fonti(contenuti, fonti, parole_per_sequenza=9, sezioni=None):
    """Controllo locale di sequenze identiche rispetto ai soli documenti caricati.

    Non sostituisce una verifica legale o una banca dati editoriale globale, ma
    individua con precisione le somiglianze letterali che non devono arrivare
    nell'esportazione.
    """
    def parole(testo):
        return re.findall(r"[a-zA-ZÀ-ÖØ-öø-ÿ0-9]{2,}", (testo or "").lower())

    parole_fonti = parole(fonti)[:80000]
    parole_libro = parole(contenuti)[:120000]
    if len(parole_fonti) < parole_per_sequenza or len(parole_libro) < parole_per_sequenza:
        return {"eseguito": False, "messaggio": "Servono fonti caricate e almeno una sezione scritta per eseguire il confronto."}

    sequenze_fonti = {
        " ".join(parole_fonti[indice:indice + parole_per_sequenza])
        for indice in range(len(parole_fonti) - parole_per_sequenza + 1)
    }
    trovate, viste = [], set()
    for indice in range(len(parole_libro) - parole_per_sequenza + 1):
        sequenza = " ".join(parole_libro[indice:indice + parole_per_sequenza])
        if sequenza in sequenze_fonti and sequenza not in viste:
            viste.add(sequenza)
            trovate.append(sequenza)
            if len(trovate) >= 12:
                break
    segnalazioni_sezioni = {}
    for titolo, testo_sezione in sezioni or []:
        parole_sezione = " ".join(parole(testo_sezione))
        corrispondenze = [sequenza for sequenza in trovate if sequenza in parole_sezione]
        if corrispondenze:
            segnalazioni_sezioni[titolo] = corrispondenze
    return {
        "eseguito": True,
        "trovate": trovate,
        "segnalazioni_sezioni": segnalazioni_sezioni,
        "messaggio": (
            "Nessuna sequenza letterale di almeno 9 parole in comune con le fonti caricate: controllo locale superato."
            if not trovate else
            f"Rilevate {len(trovate)} sequenze letterali di almeno 9 parole in comune con le fonti caricate. Rigenerale o riscrivile prima della pubblicazione."
        ),
    }


def sezioni_segnalate_per_originalita(sezioni, report_locale=None, *report_web):
    """Collega i risultati del controllo alle singole sezioni da rielaborare.

    Il controllo locale restituisce sequenze precise; i controlli web sono
    istruiti a riportare il titolo della sezione. Non vengono mai incluse
    sezioni non realmente indicate dal report.
    """
    segnalate = set()
    sequenze = (report_locale or {}).get("trovate", []) if isinstance(report_locale, dict) else []
    segnalate.update((report_locale or {}).get("segnalazioni_sezioni", {}).keys() if isinstance(report_locale, dict) else [])
    rapporti = "\n".join(str(report or "") for report in report_web).casefold()
    for titolo, testo in sezioni:
        testo_norm = re.sub(r"\s+", " ", (testo or "").casefold())
        if any(re.sub(r"\s+", " ", sequenza.casefold()) in testo_norm for sequenza in sequenze):
            segnalate.add(titolo)
        # Il titolo viene aggiunto solo se il report lo cita esplicitamente.
        if titolo and titolo.casefold() in rapporti:
            segnalate.add(titolo)
    return [titolo for titolo, _ in sezioni if titolo in segnalate]


def contesto_report_web_per_sezione(sezione, *report_web):
    """Restituisce solo i passaggi del report web associati al titolo indicato."""
    estratti = []
    for report in report_web:
        report = str(report or "")
        for corrispondenza in re.finditer(re.escape(sezione), report, flags=re.IGNORECASE):
            inizio = max(0, corrispondenza.start() - 180)
            fine = min(len(report), corrispondenza.end() + 1100)
            estratto = report[inizio:fine].strip()
            if estratto and estratto not in estratti:
                estratti.append(estratto)
    return "\n\n---\n\n".join(estratti)[:3500]


def rielabora_sezione_per_originalita(sezione, indice, trama, genere, stile, narrativa, pov,
                                      obiettivo, lingua, approfondimenti, profilo_lunghezza,
                                      report_locale=None, *report_web):
    """Riscrive da zero una sola sezione segnalata, senza alterare le altre."""
    testo_precedente = leggi_sezione_memorizzata(sezione)
    prompt_base = crea_prompt_stesura_sezione(
        sezione, indice, trama, genere, stile, narrativa, pov, obiettivo,
        lingua, approfondimenti, profilo_lunghezza
    )
    frasi_vietate = []
    if isinstance(report_locale, dict):
        frasi_vietate = report_locale.get("segnalazioni_sezioni", {}).get(sezione, [])
    contesto_web = contesto_report_web_per_sezione(sezione, *report_web)
    blocco_frasi_vietate = (
        "\nSEQUENZE ESATTE SEGNALATE DAL CONTROLLO LOCALE — NON RIPRODURLE, "
        "NÉ RIPRENDERNE L'ORDINE O LA FORMA:\n- " + "\n- ".join(frasi_vietate[:12])
        if frasi_vietate else ""
    )
    blocco_web = (
        "\nCONTESTO DEL CONTROLLO WEB PER QUESTA SOLA SEZIONE — usa queste indicazioni "
        "per cambiare radicalmente formulazione, esempio e sviluppo, senza citarle nel libro:\n"
        f"{contesto_web}"
        if contesto_web else ""
    )
    richiesta = prompt_base + f"""

RIELABORAZIONE PER ORIGINALITÀ E COPYRIGHT
Questa sezione è stata segnalata da un controllo prudenziale. Riscrivila integralmente da zero:
- conserva esclusivamente tema, obiettivo, dati necessari e livello di approfondimento della sezione;
- crea una nuova progressione di idee, nuovi esempi, nuove frasi e nuovi collegamenti;
- non conservare struttura dei paragrafi, elenchi, metafore o formulazioni del testo precedente;
- non riprodurre alcuna sequenza segnalata qui sotto, né una sua parafrasi ravvicinata; usa un percorso esplicativo alternativo;
- non citare fonti, URL, autori, controlli o copyright nel libro;
- rispetta rigorosamente la lunghezza prevista e chiudi il ragionamento con una frase completa.

{blocco_frasi_vietate}
{blocco_web}

TESTO PRECEDENTE DA SUPERARE (non copiarne la forma):
{testo_precedente}
"""
    return genera_contenuto_editoriale(
        richiesta, S_PROMPT, sezione, indice, trama, genere, obiettivo, lingua, profilo_lunghezza
    )


def verifica_originalita_web_con_ai(testo_libro, registro_fonti):
    """Schermo supplementare sul web: segnala rischi, non certifica diritti."""
    if usa_deepseek_pro():
        return "La verifica copyright sul web è disponibile solo con il cervello GPT: DeepSeek Pro resta separato e non usa strumenti GPT. Puoi comunque usare gratuitamente il controllo locale sulle fonti caricate."
    testo_libro = (testo_libro or "").strip()
    if len(testo_libro.split()) < 80:
        return "Servono almeno una sezione sostanziale già scritta per la verifica web."
    # Campione distribuito: evita di inviare un manoscritto intero e permette
    # al modello di cercare le formulazioni più distintive nella ricerca web.
    passo = max(1, len(testo_libro) // 10)
    campioni = [testo_libro[posizione:posizione + 650] for posizione in range(0, len(testo_libro), passo)][:10]
    riferimento = addebita_azione_diretta("verifica_originalita_copyright_web", amount=CREDIT_COSTS["copyright_web_rapido"])
    try:
        risposta = client.responses.create(
            model=MODELLO_ANALISI_FONTI,
            tools=[{"type": "web_search_preview"}],
            input=(
                "Agisci come revisore editoriale prudente. Cerca sul web possibili corrispondenze letterali, "
                "parafrasi molto vicine o struttura troppo derivativa per i campioni di un manoscritto, dando priorità "
                "alle fonti consultate nel registro. Confronta anche titoli, esempi distintivi, elenchi e sequenze "
                "argomentative. Non dichiarare mai che il libro è libero da copyright e non dare pareri legali. "
                "Restituisci solo: ESITO (nessuna corrispondenza evidente / attenzione / rischio elevato); "
                "PASSAGGI DA RIVEDERE (massimo 5, senza riportare oltre 12 parole per passaggio); "
                "MOTIVO; AZIONE CONSIGLIATA. Se non trovi riscontri, spiega che è uno screening limitato e "
                "che non equivale a una verifica antiplagio completa.\n\n"
                f"REGISTRO FONTI CONSULTATE:\n{registro_fonti or 'Nessun registro disponibile'}\n\n"
                "CAMPIONI DEL MANOSCRITTO DA VERIFICARE:\n" + "\n\n---\n\n".join(campioni)
            ),
        )
        registra_esito_chiamata_ai(
            risposta,
            riferimento=riferimento,
            reason="verifica_originalita_copyright_web",
            amount=CREDIT_COSTS["copyright_web_rapido"],
            model=MODELLO_ANALISI_FONTI,
        )
        esito = (getattr(risposta, "output_text", "") or "").strip()
        if not esito:
            refund_credits(riferimento, reason="verifica_originalita_web_vuota", amount=CREDIT_COSTS["copyright_web_rapido"])
            return "La verifica web non ha prodotto un esito. Nessun credito è stato addebitato."
        return esito
    except Exception as errore:
        refund_credits(riferimento, reason="verifica_originalita_web_fallita", amount=CREDIT_COSTS["copyright_web_rapido"])
        registra_esito_chiamata_ai(
            None,
            riferimento=riferimento,
            reason="verifica_originalita_copyright_web",
            amount=CREDIT_COSTS["copyright_web_rapido"],
            model=MODELLO_ANALISI_FONTI,
            riuscita=False,
            rimborsata=True,
            errore=type(errore).__name__,
        )
        return "La verifica web non è riuscita. Nessun credito è stato addebitato; puoi riprovare più tardi."


def prepara_blocchi_verifica_web_completa(sezioni, massimo_caratteri=2400):
    """Suddivide tutte le sezioni, conservando il riferimento editoriale di origine."""
    blocchi = []
    for titolo, testo in sezioni:
        testo = (testo or "").strip()
        if not testo:
            continue
        for inizio in range(0, len(testo), massimo_caratteri):
            estratto = testo[inizio:inizio + massimo_caratteri]
            blocchi.append((titolo, estratto))
    return blocchi


def richiede_revisione_copyright(esito):
    """Individua gli esiti del mini che meritano una seconda lettura più rigorosa."""
    testo = (esito or "").lower()
    if "rischio elevato" in testo or "attenzione" in testo:
        return True
    indicatori = (
        "passaggi da rivedere", "somiglianza",
        "corrispondenza", "parafrasi", "troppo vicin", "derivativ",
    )
    return any(indicatore in testo for indicatore in indicatori) and "nessuna corrispondenza evidente" not in testo


def verifica_originalita_web_completa(sezioni, registro_fonti, aggiorna=None):
    """Analizza tutto il manoscritto: mini su ogni lotto, completo solo sui dubbi."""
    if usa_deepseek_pro():
        return "La verifica copyright completa sul web è disponibile solo con il cervello GPT. DeepSeek Pro non utilizza il motore web GPT.", 0
    blocchi = prepara_blocchi_verifica_web_completa(sezioni)
    if not blocchi:
        return "Servono sezioni già scritte per eseguire la verifica completa.", 0
    dimensione_lotto = 8
    lotti = [blocchi[indice:indice + dimensione_lotto] for indice in range(0, len(blocchi), dimensione_lotto)]
    esiti, completati, revisioni_approfondite, crediti_effettivi = [], 0, 0, 0
    for numero_lotto, lotto in enumerate(lotti, start=1):
        if aggiorna:
            aggiorna(numero_lotto - 1, len(lotti))
        riferimento = addebita_azione_diretta(
            "verifica_copyright_web_completa",
            amount=CREDIT_COSTS["copyright_lotto_screening_mini"],
        )
        testo_lotto = "\n\n---\n\n".join(
            f"SEZIONE: {titolo}\nBLOCCO: {testo}" for titolo, testo in lotto
        )
        try:
            risposta = client.responses.create(
                model=MODELLO_CONTROLLO_COPYRIGHT_COMPLETO,
                tools=[{"type": "web_search_preview"}],
                input=(
                    "Agisci come revisore editoriale prudente. Analizza TUTTI i blocchi ricevuti in questo lotto "
                    "e cerca sul web corrispondenze letterali, parafrasi troppo vicine, titoli, esempi, elenchi o "
                    "sviluppi argomentativi distintivi, con priorità alle fonti del registro. Non emettere pareri legali "
                    "né certificazioni di assenza copyright. Non citare più di 12 parole di un testo potenzialmente protetto. "
                    "Per ogni rischio indica la sezione esatta, il livello (attenzione/rischio elevato), il motivo e una "
                    "azione concreta di riscrittura. Se non trovi elementi, scrivi: 'Nessuna corrispondenza evidente nel lotto'.\n\n"
                    f"REGISTRO FONTI CONSULTATE:\n{registro_fonti}\n\n"
                    f"LOTTO {numero_lotto}/{len(lotti)}\n{testo_lotto}"
                ),
            )
            esito = (getattr(risposta, "output_text", "") or "").strip()
            if esito:
                registra_esito_chiamata_ai(
                    risposta,
                    riferimento=riferimento,
                    reason="verifica_copyright_web_completa",
                    amount=CREDIT_COSTS["copyright_lotto_screening_mini"],
                    model=MODELLO_CONTROLLO_COPYRIGHT_COMPLETO,
                )
                crediti_effettivi += CREDIT_COSTS["copyright_lotto_screening_mini"]
                esito_finale = f"Screening GPT-5.4 mini\n{esito}"
                if richiede_revisione_copyright(esito):
                    riferimento_revisione = addebita_azione_diretta(
                        "verifica_copyright_web_revisione_gpt54",
                        amount=CREDIT_COSTS["copyright_lotto_revisione_gpt54"],
                    )
                    try:
                        revisione = client.responses.create(
                            model=MODELLO_CONTROLLO_COPYRIGHT_APPROFONDITO,
                            tools=[{"type": "web_search_preview"}],
                            input=(
                                "Sei un revisore editoriale senior specializzato in rischi di originalità. Riesamina "
                                "il lotto segnalato da un primo screening. Cerca sul web e nelle fonti registrate solo "
                                "riscontri concreti; riduci i falsi positivi. Non dare pareri legali né certificazioni. "
                                "Restituisci: ESITO FINALE (nessuna corrispondenza evidente / attenzione / rischio elevato); "
                                "SEZIONE O BLOCCO; MOTIVO; AZIONE DI RISCRITTURA. Non riportare più di 12 parole di "
                                "materiale potenzialmente protetto.\n\n"
                                f"REGISTRO FONTI CONSULTATE:\n{registro_fonti}\n\n"
                                f"PRIMO SCREENING:\n{esito}\n\n"
                                f"LOTTO DA RIESAMINARE:\n{testo_lotto}"
                            ),
                        )
                        esito_revisione = (getattr(revisione, "output_text", "") or "").strip()
                        if esito_revisione:
                            registra_esito_chiamata_ai(
                                revisione,
                                riferimento=riferimento_revisione,
                                reason="verifica_copyright_web_revisione_gpt54",
                                amount=CREDIT_COSTS["copyright_lotto_revisione_gpt54"],
                                model=MODELLO_CONTROLLO_COPYRIGHT_APPROFONDITO,
                            )
                            esito_finale += f"\n\nRevisione mirata GPT-5.4\n{esito_revisione}"
                            revisioni_approfondite += 1
                            crediti_effettivi += CREDIT_COSTS["copyright_lotto_revisione_gpt54"]
                        else:
                            refund_credits(
                                riferimento_revisione,
                                reason="verifica_copyright_revisione_vuota",
                                amount=CREDIT_COSTS["copyright_lotto_revisione_gpt54"],
                            )
                            registra_esito_chiamata_ai(
                                revisione,
                                riferimento=riferimento_revisione,
                                reason="verifica_copyright_web_revisione_gpt54",
                                amount=CREDIT_COSTS["copyright_lotto_revisione_gpt54"],
                                model=MODELLO_CONTROLLO_COPYRIGHT_APPROFONDITO,
                                riuscita=False,
                                rimborsata=True,
                                errore="risposta_vuota",
                            )
                    except Exception as errore_revisione:
                        refund_credits(
                            riferimento_revisione,
                            reason="verifica_copyright_revisione_fallita",
                            amount=CREDIT_COSTS["copyright_lotto_revisione_gpt54"],
                        )
                        registra_esito_chiamata_ai(
                            None,
                            riferimento=riferimento_revisione,
                            reason="verifica_copyright_web_revisione_gpt54",
                            amount=CREDIT_COSTS["copyright_lotto_revisione_gpt54"],
                            model=MODELLO_CONTROLLO_COPYRIGHT_APPROFONDITO,
                            riuscita=False,
                            rimborsata=True,
                            errore=type(errore_revisione).__name__,
                        )
                        esito_finale += "\n\nRevisione mirata GPT-5.4 non disponibile: conserva lo screening mini e valuta manualmente il lotto."
                esiti.append(f"LOTTO {numero_lotto}/{len(lotti)}\n{esito_finale}")
                completati += 1
            else:
                refund_credits(
                    riferimento,
                    reason="verifica_copyright_completa_vuota",
                    amount=CREDIT_COSTS["copyright_lotto_screening_mini"],
                )
                registra_esito_chiamata_ai(
                    risposta,
                    riferimento=riferimento,
                    reason="verifica_copyright_web_completa",
                    amount=CREDIT_COSTS["copyright_lotto_screening_mini"],
                    model=MODELLO_CONTROLLO_COPYRIGHT_COMPLETO,
                    riuscita=False,
                    rimborsata=True,
                    errore="risposta_vuota",
                )
        except Exception as errore_lotto:
            refund_credits(
                riferimento,
                reason="verifica_copyright_completa_fallita",
                amount=CREDIT_COSTS["copyright_lotto_screening_mini"],
            )
            registra_esito_chiamata_ai(
                None,
                riferimento=riferimento,
                reason="verifica_copyright_web_completa",
                amount=CREDIT_COSTS["copyright_lotto_screening_mini"],
                model=MODELLO_CONTROLLO_COPYRIGHT_COMPLETO,
                riuscita=False,
                rimborsata=True,
                errore=type(errore_lotto).__name__,
            )
            esiti.append(f"LOTTO {numero_lotto}/{len(lotti)}\nVerifica non completata: nessun credito addebitato per questo lotto.")
    if aggiorna:
        aggiorna(len(lotti), len(lotti))
    intestazione = (
        f"VERIFICA WEB COMPLETA — {len(blocchi)} blocchi controllati in {len(lotti)} lotti; "
        f"{completati} lotti completati; {revisioni_approfondite} revisioni mirate GPT-5.4; "
        f"{crediti_effettivi} crediti effettivamente addebitati.\n\n"
        "Metodo e costi: GPT-5.4 mini analizza tutti i lotti (1 credito ciascuno); "
        "GPT-5.4 viene usato solo per i lotti segnalati (2 crediti aggiuntivi per revisione completata).\n\n"
        "Questo è uno screening di rischio sul web, non una certificazione legale né una verifica antiplagio universale.\n\n"
    )
    return intestazione + "\n\n".join(esiti), len(lotti)


def notifica_sonora(evento, lingua="Italiano", ripeti=False):
    """Emette un segnale sonoro e un avviso testuale localizzato in alto a destra."""
    chiave = f"notifica_emessa_{evento}"
    if not ripeti and st.session_state.get(chiave):
        return
    st.session_state[chiave] = True
    messaggi = {
        "Italiano": {
            "sidebar_pronta": "Brief completo: puoi generare l’indice.", "voto_indice_completato": "Voto dell’indice completato.",
            "avvio_scrittura_completa": "Avviata la scrittura completa del libro.", "errore_scrittura": "La scrittura si è interrotta: controlla lo stato.",
            "libro_completato": "Scrittura completa del libro terminata.", "coerenza_completata": "Controllo coerenza completo terminato.",
            "word_pronto": "File Word pronto per il download.", "pdf_pronto": "File PDF pronto per il download.", "formattazione_completata": "Formattazione del documento completata."
        },
        "English": {
            "sidebar_pronta": "Brief complete: you can generate the index.", "voto_indice_completato": "Index score completed.",
            "avvio_scrittura_completa": "Full-book writing has started.", "errore_scrittura": "Writing stopped: check the status.",
            "libro_completato": "Full-book writing completed.", "coerenza_completata": "Full consistency check completed.",
            "word_pronto": "Word file ready to download.", "pdf_pronto": "PDF file ready to download.", "formattazione_completata": "Document formatting completed."
        },
        "Español": {"sidebar_pronta": "Brief completo: puedes generar el índice.", "voto_indice_completato": "Evaluación del índice completada.", "avvio_scrittura_completa": "Ha comenzado la escritura completa.", "errore_scrittura": "La escritura se interrumpió: revisa el estado.", "libro_completato": "Escritura completa terminada.", "coerenza_completata": "Control de coherencia completado.", "word_pronto": "Archivo Word listo para descargar.", "pdf_pronto": "Archivo PDF listo para descargar.", "formattazione_completata": "Formato del documento completado."},
        "Français": {"sidebar_pronta": "Brief complet : vous pouvez générer l’index.", "voto_indice_completato": "Évaluation de l’index terminée.", "avvio_scrittura_completa": "Rédaction complète commencée.", "errore_scrittura": "La rédaction est interrompue : vérifiez l’état.", "libro_completato": "Rédaction complète terminée.", "coerenza_completata": "Contrôle de cohérence terminé.", "word_pronto": "Fichier Word prêt à télécharger.", "pdf_pronto": "Fichier PDF prêt à télécharger.", "formattazione_completata": "Mise en forme terminée."},
        "Deutsch": {"sidebar_pronta": "Brief vollständig: Der Index kann erstellt werden.", "voto_indice_completato": "Indexbewertung abgeschlossen.", "avvio_scrittura_completa": "Vollständige Bucherstellung gestartet.", "errore_scrittura": "Schreiben wurde unterbrochen: Status prüfen.", "libro_completato": "Vollständige Bucherstellung abgeschlossen.", "coerenza_completata": "Vollständige Kohärenzprüfung abgeschlossen.", "word_pronto": "Word-Datei zum Download bereit.", "pdf_pronto": "PDF-Datei zum Download bereit.", "formattazione_completata": "Dokumentformatierung abgeschlossen."},
        "Română": {"sidebar_pronta": "Brief complet: poți genera cuprinsul.", "voto_indice_completato": "Evaluarea cuprinsului s-a încheiat.", "avvio_scrittura_completa": "A început scrierea completă a cărții.", "errore_scrittura": "Scrierea s-a oprit: verifică starea.", "libro_completato": "Scrierea completă s-a încheiat.", "coerenza_completata": "Controlul complet de coerență s-a încheiat.", "word_pronto": "Fișierul Word este gata de descărcare.", "pdf_pronto": "Fișierul PDF este gata de descărcare.", "formattazione_completata": "Formatarea documentului s-a încheiat."},
        "Русский": {"sidebar_pronta": "Бриф готов: можно создать оглавление.", "voto_indice_completato": "Оценка оглавления завершена.", "avvio_scrittura_completa": "Начато написание всей книги.", "errore_scrittura": "Написание остановлено: проверьте состояние.", "libro_completato": "Написание всей книги завершено.", "coerenza_completata": "Полная проверка согласованности завершена.", "word_pronto": "Файл Word готов к скачиванию.", "pdf_pronto": "Файл PDF готов к скачиванию.", "formattazione_completata": "Форматирование документа завершено."},
        "العربية": {"sidebar_pronta": "اكتمل الملخص: يمكنك إنشاء الفهرس.", "voto_indice_completato": "اكتمل تقييم الفهرس.", "avvio_scrittura_completa": "بدأت كتابة الكتاب كاملاً.", "errore_scrittura": "توقفت الكتابة: راجع الحالة.", "libro_completato": "اكتملت كتابة الكتاب.", "coerenza_completata": "اكتمل فحص الاتساق الكامل.", "word_pronto": "ملف Word جاهز للتنزيل.", "pdf_pronto": "ملف PDF جاهز للتنزيل.", "formattazione_completata": "اكتمل تنسيق المستند."},
        "中文": {"sidebar_pronta": "简介已完成：可以生成目录。", "voto_indice_completato": "目录评分已完成。", "avvio_scrittura_completa": "整本书写作已开始。", "errore_scrittura": "写作已中断：请检查状态。", "libro_completato": "整本书写作已完成。", "coerenza_completata": "完整一致性检查已完成。", "word_pronto": "Word 文件可供下载。", "pdf_pronto": "PDF 文件可供下载。", "formattazione_completata": "文档格式化已完成。"}
    }
    messaggi_prova = {
        "Italiano": "Notifiche attive: messaggio e suono di prova.",
        "English": "Notifications enabled: test message and sound.",
        "Español": "Notificaciones activas: mensaje y sonido de prueba.",
        "Français": "Notifications actives : message et son de test.",
        "Deutsch": "Benachrichtigungen aktiv: Testmeldung und Ton.",
        "Română": "Notificări active: mesaj și sunet de test.",
        "Русский": "Уведомления активны: тестовое сообщение и звук.",
        "العربية": "الإشعارات مفعلة: رسالة وصوت تجريبي.",
        "中文": "通知已启用：测试消息和声音。",
    }
    messaggio = messaggi_prova.get(lingua, messaggi_prova["Italiano"]) if evento == "test_notifiche" else messaggi.get(lingua, messaggi["Italiano"]).get(evento, "Operazione completata.")
    icona = "⚠️" if evento == "errore_scrittura" else "✅"
    # Durata lunga: il messaggio resta visibile anche quando l'operazione termina velocemente.
    st.toast(messaggio, icon=icona, duration="long")
    # Il token cambia a ogni evento: forza il browser a rieseguire l'audio anche per notifiche consecutive.
    token_audio = uuid.uuid4().hex
    html_audio = """
    <script>
    (() => {
      try {
        const eventId = '__TOKEN_AUDIO__';
        const AudioCtx = window.AudioContext || window.webkitAudioContext;
        const ctx = new AudioCtx();
        [[988, 0.00, 0.15, 0.88], [1319, 0.16, 0.15, 0.92], [1760, 0.32, 0.42, 0.96]].forEach(([f, start, duration, volume]) => {
          const oscillator = ctx.createOscillator();
          const gain = ctx.createGain();
          oscillator.frequency.value = f;
          oscillator.type = 'sine';
          gain.gain.setValueAtTime(0.001, ctx.currentTime + start);
          gain.gain.exponentialRampToValueAtTime(volume, ctx.currentTime + start + 0.02);
          gain.gain.exponentialRampToValueAtTime(0.001, ctx.currentTime + start + duration);
          oscillator.connect(gain); gain.connect(ctx.destination);
          oscillator.start(ctx.currentTime + start);
          oscillator.stop(ctx.currentTime + start + duration + 0.03);
        });
      } catch (error) { console.log('Audio notification unavailable', error); }
    })();
    </script>
    """.replace("__TOKEN_AUDIO__", token_audio)
    components.html(html_audio, height=0)

# ======================================================================================================================
# 1. ARCHITETTURA DI SISTEMA E SICUREZZA API
# ======================================================================================================================
# Nome Applicazione: AI di Antonino: Ebook Mondiale Creator PRO
# Developer: Antonino & Gemini Collaboration
# Core Update: Integrazione Neuromarketing (Triune Brain Methodology) con Motore Decisionale Dinamico.
# Identificativo visibile: permette di verificare che Streamlit stia eseguendo l'ultimo deploy.
VERSIONE_DEPLOY = f"Scrittore Site commerciale — {COMMERCIAL_TEST_VERSION}"
VERSIONE_AUDIT_COHERENZA = "3"

# --- CONNESSIONI AI: due cervelli separati, selezionabili dall'utente ---
client_openai = None
client_deepseek = None
try:
    client_openai = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except Exception:
    pass
try:
    client_deepseek = OpenAI(
        api_key=st.secrets["DEEPSEEK_API_KEY"],
        base_url="https://api.deepseek.com",
    )
except Exception:
    pass

# Compatibilità con gli strumenti esclusivi OpenAI già presenti (immagini e
# ricerca web). Non viene mai usato per la scrittura se l'utente sceglie
# DeepSeek Pro.
client = client_openai

st.set_page_config(
    page_title="Scrittore Site",
    layout="wide",
    initial_sidebar_state="expanded",
    page_icon="✒️"
)

# AREA COMMERCIALE: accesso, crediti e pagamenti restano separati dalla release personale.
bootstrap_commercial_test()

# ======================================================================================================================
# COLLAUDO RISERVATO ALL'AMMINISTRATORE
# ======================================================================================================================
# Il comando viene creato esclusivamente dal pannello amministratore di
# commercial_layer.py. Qui consumiamo la richiesta prima che Streamlit disegni
# i widget della sidebar, così il progetto di prova può riempirli in sicurezza.
_PREFISSI_PROGETTO_COLLAUDO = (
    "book_", "txt_", "mod_", "indice_", "lista_", "memoria_", "conoscenza_",
    "scheda_", "dossier_", "brief_", "registro_", "firma_", "job_", "immagini_",
    "audit_", "report_", "cache_", "autosave_", "ultimo_", "analisi_",
)
_CHIAVI_PROGETTO_COLLAUDO = {
    "editor_language", "provider_ia", "profilo_lunghezza_stesura",
    "sezione_editor_attiva", "sezione_editor_selezionata", "fonti_web_salvate",
    "commercial_project_reset_requested",
}


def _chiave_progetto_collaudo(chiave):
    return str(chiave).startswith(_PREFISSI_PROGETTO_COLLAUDO) or str(chiave) in _CHIAVI_PROGETTO_COLLAUDO


def _pulisci_progetto_per_collaudo():
    for chiave in list(st.session_state.keys()):
        if _chiave_progetto_collaudo(chiave) and not str(chiave).startswith("commercial_"):
            del st.session_state[chiave]


def _profilo_collaudo_breve(nome_profilo):
    profili = {
        "Manuale e controlli": {
            "titolo": "Riunioni di progetto efficaci",
            "genere": "Manuale Pratico",
            "tipologia": "Scientifico Divulgativo",
            "stile": "Diretto e Pratico (Action-oriented)",
            "pov": "Tu (Diretto, confidenziale e personale)",
            "obiettivo": "Aiutare titolari di piccole imprese, coordinatori, freelance e responsabili di team da 3 a 12 persone a progettare riunioni che producano decisioni chiare, responsabilità assegnate e scadenze verificabili.",
            "risultato": "Al termine il lettore saprà valutare se una riunione è necessaria, definire uno scopo misurabile, selezionare i partecipanti, preparare un ordine del giorno con tempi, condurre un confronto, registrare una decisione e controllare il follow-up.",
            "argomento": "Guida pratica alle riunioni di progetto da 30 a 60 minuti in piccole aziende e gruppi di lavoro. Il libro affronta diagnosi delle riunioni inconcludenti, preparazione, agenda, ruoli, facilitazione, gestione delle deviazioni, decisioni, verbale essenziale, responsabilità e controllo delle azioni concordate. Non è un manuale di leadership generico e non si rivolge a dirigenti di grandi multinazionali.",
            "approfondimenti": "Il pubblico ha esperienza operativa ma poca formazione manageriale. Usa esempi realistici di un progetto digitale, di un piccolo cantiere e di un'attività commerciale; includi un modello di agenda, una checklist di chiusura, un esempio di verbale breve e criteri concreti per distinguere una decisione da una semplice discussione. Se menzioni strumenti o metodi aggiornabili, segnala che vanno verificati prima della pubblicazione. Evita promesse garantite, frasi riempitive e contenuti ripetuti.",
            "personalizzazione": {
                "voce": "Tono pratico di chi coordina progetti reali, con attenzione alle persone coinvolte.",
                "materiale": "Un caso verosimile di riunione con tempi limitati, decisioni poco chiare e responsabilità distribuite.",
                "priorita": "Il lettore deve uscire con una procedura applicabile già alla riunione successiva.",
                "confini": "Evita promesse assolute, gergo superfluo e soluzioni valide solo per grandi aziende.",
            },
        },
        "Test Prep e simulazioni": {
            "titolo": "HACCP: preparazione al test per addetti alimentari",
            "genere": "Test Prep (Preparazione Esami)",
            "tipologia": "Professionale Accademico",
            "stile": "Tecnico e Analitico",
            "pov": "Voi (Plurale, autorevole e rispettoso)",
            "obiettivo": "Preparare addetti di bar, ristoranti, gastronomie e piccoli laboratori alimentari a comprendere i principi base dell'igiene alimentare e ad allenarsi per una prova formativa con domande, casi pratici e simulazioni.",
            "risultato": "Il candidato saprà distinguere pericoli biologici, chimici e fisici, riconoscere comportamenti igienici corretti, ragionare su contaminazione crociata, pulizia, conservazione e controllo operativo, quindi completare una simulazione prima di leggere soluzioni commentate.",
            "argomento": "Test Prep introduttivo per addetti alla manipolazione degli alimenti: principi HACCP, igiene personale, allergeni, contaminazione crociata, ricevimento e conservazione, pulizia, sanificazione e controlli quotidiani. Il testo deve separare teoria, esercizi, quiz, simulazione e soluzioni, senza sostituire corsi obbligatori o procedure aziendali.",
            "approfondimenti": "Usa un linguaggio tecnico comprensibile a chi lavora in cucina o al banco. Inserisci quiz a risposta multipla separati dalle soluzioni, casi realistici su frigo, utensili, allergeni e pulizia, una simulazione finale realmente svolgibile e correzioni commentate. Ogni informazione normativa, temperatura, requisito locale o indicazione sanitaria deve essere verificata e aggiornata prima della pubblicazione. Non promettere certificazione, conformità legale assoluta o superamento garantito dell'esame.",
            "personalizzazione": {
                "voce": "Tono tecnico chiaro di un formatore che spiega le ragioni delle regole senza allarmismo.",
                "materiale": "Situazioni verosimili di banco gastronomia, piccola cucina e ricevimento merci.",
                "priorita": "Il candidato deve saper ragionare sui casi, non soltanto ricordare definizioni.",
                "confini": "Non sostituire procedure aziendali o formazione obbligatoria; evita prescrizioni normative non verificate.",
            },
        },
        "Narrativa e stile": {
            "titolo": "Il taccuino della casa sul porto",
            "genere": "Narrativo",
            "tipologia": "Storytelling Immersivo",
            "stile": "Storytelling Emozionale",
            "pov": "Tu (Diretto, confidenziale e personale)",
            "obiettivo": "Coinvolgere il lettore in un romanzo contemporaneo di scoperta e riparazione, mettendo alla prova coerenza del punto di vista, progressione emotiva, dialoghi, scene concrete e conclusione non tronca.",
            "risultato": "Il lettore segue una vicenda completa: ritorno alla città d'origine, scoperta di un taccuino nascosto, conflitto con una persona del passato, scelta rischiosa e risoluzione che cambia il protagonista senza cancellare le conseguenze delle sue decisioni.",
            "argomento": "Dopo la morte della nonna, Mara, restauratrice trentacinquenne, torna nella casa sul porto di una piccola città adriatica per venderla. Dietro una parete trova un taccuino che documenta una promessa non mantenuta tra due famiglie e mette in dubbio la storia che le è stata raccontata. Per ricostruire la verità deve affrontare Elia, figlio dell'uomo che ritiene responsabile, e scegliere se rendere pubblico il contenuto, rischiando di ferire persone ancora vive. Atmosfera mediterranea contemporanea, conflitto umano credibile, nessun elemento soprannaturale.",
            "approfondimenti": "Rivolgiti a lettori adulti che apprezzano narrativa introspettiva ma scorrevole. Ogni scena deve cambiare informazioni, relazione o posta in gioco; alterna azioni osservabili, dialoghi essenziali e brevi momenti interiori. Evita riassunti degli eventi, coincidenze risolutive e melodramma. La struttura deve accompagnare situazione iniziale, ostacoli, scelta e conseguenza finale, chiudendo davvero ogni sezione.",
            "personalizzazione": {
                "voce": "Voce intima e concreta, attenta ai dettagli quotidiani della costa adriatica senza lirismo eccessivo.",
                "materiale": "Un ricordo sensoriale della casa sul porto: sale sulle persiane, carta ingiallita e rumore delle barche.",
                "priorita": "Il lettore deve comprendere la posta emotiva di Mara prima della scelta conclusiva.",
                "confini": "Evita cliché romantici, spiegazioni onniscienti e svolte risolte da coincidenze.",
            },
        },
    }
    return profili.get(nome_profilo, profili["Manuale e controlli"])


def _fotografia_locale_progetto_collaudo():
    """Copia il lavoro locale senza fallire se un widget conserva un oggetto speciale."""
    fotografia = {}
    for chiave, valore in st.session_state.items():
        if _chiave_progetto_collaudo(chiave) and not str(chiave).startswith("commercial_"):
            try:
                fotografia[chiave] = copy.deepcopy(valore)
            except Exception:
                fotografia[chiave] = valore
    return fotografia


def avvia_collaudo_amministratore_se_richiesto():
    richiesta = st.session_state.pop("commercial_admin_test_request", None)
    if not isinstance(richiesta, dict):
        return
    if not st.session_state.get("admin_test_mode"):
        st.session_state["admin_test_backup"] = _fotografia_locale_progetto_collaudo()
    _pulisci_progetto_per_collaudo()
    st.session_state.pop("admin_test_verifica_visiva", None)
    st.session_state.pop("admin_test_run_requested", None)
    st.session_state.pop("admin_test_run_report", None)
    provider = str(richiesta.get("provider") or "GPT-5.4 (OpenAI)")
    profilo_nome = str(richiesta.get("profilo") or "Manuale e controlli")
    profilo = _profilo_collaudo_breve(profilo_nome)
    st.session_state.update({
        "editor_language": "Italiano",
        "provider_ia": provider,
        "book_title": profilo["titolo"],
        "book_author": "Redazione interna Scrittore Site — non pubblicare",
        "book_genre": profilo["genere"],
        "book_writing_style": profilo["tipologia"],
        "book_narrative_style": profilo["stile"],
        "book_point_of_view": profilo["pov"],
        "book_goal": profilo["obiettivo"],
        "book_desired_result": profilo["risultato"],
        "book_plot": profilo["argomento"],
        "book_further_details": profilo["approfondimenti"],
        # Valori realistici e coerenti con la modalità scelta: il collaudo
        # esercita la personalizzazione senza cambiare parametri o regole del test.
        "book_personal_voice": profilo.get("personalizzazione", {}).get("voce", ""),
        "book_personal_material": profilo.get("personalizzazione", {}).get("materiale", ""),
        "book_personal_priorities": profilo.get("personalizzazione", {}).get("priorita", ""),
        "book_personal_boundaries": profilo.get("personalizzazione", {}).get("confini", ""),
        "book_personal_checkpoint_mode": "",
        "book_personal_checkpoint_notes": "",
        # Il laboratorio usa un brief nuovo e isolato, ma le stesse opzioni
        # produttive disponibili agli utenti: nessun limite editoriale speciale.
        "profilo_lunghezza_stesura": "Standard KDP",
        "indice_raw": "",
        "lista_capitoli": [],
        "memoria_sezioni_editor": {},
        "immagini_capitoli": {},
        "admin_test_mode": True,
        "admin_test_provider": provider,
        "admin_test_profile": profilo_nome,
        "admin_test_started_at": datetime.datetime.now().isoformat(timespec="seconds"),
    })
    storico = st.session_state.setdefault("commercial_admin_test_history", [])
    storico.append({
        "avviato": st.session_state["admin_test_started_at"],
        "cervello": provider,
        "modalita": profilo_nome,
        "indice": False,
        "sezioni": 0,
    })
    st.session_state["commercial_admin_test_history"] = storico[-20:]


def termina_collaudo_amministratore():
    """Chiude il laboratorio senza eliminare né sovrascrivere la bozza reale."""
    storico = st.session_state.get("commercial_admin_test_history", [])
    if storico:
        storico[-1]["concluso"] = datetime.datetime.now().isoformat(timespec="seconds")
        storico[-1]["indice"] = bool(st.session_state.get("indice_raw", "").strip())
        storico[-1]["sezioni"] = len(st.session_state.get("memoria_sezioni_editor", {}) or {})
    backup = st.session_state.get("admin_test_backup", {})
    _pulisci_progetto_per_collaudo()
    for chiave in ("admin_test_mode", "admin_test_provider", "admin_test_profile", "admin_test_started_at", "admin_test_backup", "admin_test_verifica_visiva", "admin_test_run_requested", "admin_test_run_report", "admin_test_pausa_ripresa_verificata", "admin_test_prefazione_pausa_verificata", "admin_test_personalizzazione_verificata"):
        st.session_state.pop(chiave, None)
    st.session_state.update(backup)


avvia_collaudo_amministratore_se_richiesto()

# ======================================================================================================================
# 2. DIZIONARIO MULTILINGUA INTEGRALE (9 LINGUE GLOBALI - ESPANSO)
# ======================================================================================================================
TRADUZIONI = {
    "Italiano": {
        "side_tit": "⚙️ Configurazione Editor",
        "lbl_tit": "Titolo del Libro", "lbl_auth": "Nome Autore", "lbl_lang": "Lingua", 
        "lbl_gen": "Genere Letterario", "lbl_style": "Tipologia Scrittura", "lbl_plot": "Trama o Argomento",
        "lbl_narrative": "Stile di Racconto", "lbl_goal": "Obiettivo del Libro", "lbl_pov": "Punto di Vista (Pronome)",
        "btn_res": "🔄 RESET PROGETTO", "tabs": ["📊 1. Indice", "✍️ 2. Scrittura & Quiz", "📖 3. Anteprima", "📑 4. Importa / Esporta / Copyright"],
        "btn_idx": "🚀 Genera Indice Professionale", "btn_sync": "✅ Salva e Sincronizza Capitoli",
        "lbl_sec": "Seleziona sezione:", "btn_write": "✨ SCRIVI CONTENUTO (Dettagliato)",
        "btn_quiz": "🧠 AGGIUNGI QUIZ AL LIBRO", "btn_edit": "🚀 RIELABORA CON IA",
        "msg_run": "Il neuro-linguista sta analizzando gerarchia, stile e target emotivo...", "preface": "Prefazione", "ack": "Ringraziamenti",
        "preview_tit": "📖 Vista Lettura Professionale", "btn_word": "📥 Scarica Word (.docx)", "btn_pdf": "📥 Scarica PDF (.pdf)",
        "msg_err_idx": "Genera l'indice nella Tab 1 prima di procedere.", "msg_success_sync": "Capitoli sincronizzati!",
        "label_editor": "Editor di Testo Professionale", "welcome": "👋 Benvenuto nell'Ebook Creator di Antonino.",
        "guide": "Usa la sidebar a sinistra per impostare i parametri del tuo libro."
    },
    "English": {
        "side_tit": "⚙️ Editor Setup", "lbl_tit": "Book Title", "lbl_auth": "Author Name", "lbl_lang": "Language", 
        "lbl_gen": "Genre", "lbl_style": "Writing Style", "lbl_plot": "Plot", "lbl_narrative": "Narrative Style", "lbl_goal": "Book Goal", "lbl_pov": "Point of View (Pronoun)",
        "btn_res": "🔄 RESET PROJECT", "tabs": ["📊 1. Index", "✍️ 2. Write & Quiz", "📖 3. Preview", "📑 4. Import / Export / Copyright"],
        "btn_idx": "🚀 Generate Index", "btn_sync": "✅ Sync Chapters", "lbl_sec": "Select section:",
        "btn_write": "✨ WRITE CONTENT", "btn_quiz": "🧠 ADD QUIZ", "btn_edit": "🚀 REWRITE",
        "msg_run": "Native expert analyzing hierarchy, style and goal...", "preface": "Preface", "ack": "Acknowledgements",
        "preview_tit": "📖 Reading View", "btn_word": "📥 Word", "btn_pdf": "📥 PDF",
        "msg_err_idx": "Generate index first.", "msg_success_sync": "Synced!",
        "label_editor": "Editor", "welcome": "👋 Welcome.", "guide": "Use sidebar."
    },
    "Español": {
        "side_tit": "⚙️ Configuración del Editor", "lbl_tit": "Título del Libro", "lbl_auth": "Nombre del Autor", "lbl_lang": "Idioma", 
        "lbl_gen": "Género Literario", "lbl_style": "Estilo de Escritura", "lbl_plot": "Trama o Argumento", "lbl_narrative": "Estilo Narrativo", "lbl_goal": "Objetivo del Libro", "lbl_pov": "Punto de Vista (Pronombre)",
        "btn_res": "🔄 RESETEAR PROYECTO", "tabs": ["📊 1. Índice", "✍️ 2. Escritura y Quiz", "📖 3. Vista Previa", "📑 4. Importar / Exportar / Copyright"],
        "btn_idx": "🚀 Generar Índice Profesional", "btn_sync": "✅ Guardar y Sincronizar", "lbl_sec": "Seleccionar sección:",
        "btn_write": "✨ ESCRIBIR CONTENIDO", "btn_quiz": "🧠 AÑADIR QUIZ", "btn_edit": "🚀 REESCRIBIR",
        "msg_run": "Analizando jerarquía y estilo...", "preface": "Prefacio", "ack": "Agradecimientos",
        "preview_tit": "📖 Vista de Lectura", "btn_word": "📥 Descargar Word", "btn_pdf": "📥 Descargar PDF",
        "msg_err_idx": "Genera el índice primero.", "msg_success_sync": "¡Sincronizado!", "label_editor": "Editor Profesional", "welcome": "👋 Bienvenido.", "guide": "Usa la barra lateral."
    },
    "Français": {
        "side_tit": "⚙️ Configuration de l'Éditeur", "lbl_tit": "Titre du Livre", "lbl_auth": "Nom de l'Auteur", "lbl_lang": "Langue", 
        "lbl_gen": "Genre Littéraire", "lbl_style": "Style d'Écriture", "lbl_plot": "Intrigue ou Sujet", "lbl_narrative": "Style Narratif", "lbl_goal": "Objectif du Livre", "lbl_pov": "Point de Vue (Pronom)",
        "btn_res": "🔄 RÉINITIALISER", "tabs": ["📊 1. Index", "✍️ 2. Écriture & Quiz", "📖 3. Aperçu", "📑 4. Importer / Exporter / Copyright"],
        "btn_idx": "🚀 Générer l'Index", "btn_sync": "✅ Synchroniser", "lbl_sec": "Sélectionner la section:",
        "btn_write": "✨ ÉCRIRE LE CONTENU", "btn_quiz": "🧠 AJOUTER UN QUIZ", "btn_edit": "🚀 RÉÉCRIRE",
        "msg_run": "Analyse de la hiérarchie et du style...", "preface": "Préface", "ack": "Remerciements",
        "preview_tit": "📖 Aperçu de Lecture", "btn_word": "📥 Télécharger Word", "btn_pdf": "📥 Télécharger PDF",
        "msg_err_idx": "Générez l'index d'abord.", "msg_success_sync": "Synchronisé!", "label_editor": "Éditeur Professionnel", "welcome": "👋 Bienvenue.", "guide": "Utilisez la barre latérale."
    },
    "Deutsch": {
        "side_tit": "⚙️ Editor-Setup", "lbl_tit": "Buchtitel", "lbl_auth": "Autorenname", "lbl_lang": "Sprache", 
        "lbl_gen": "Genre", "lbl_style": "Schreibstil", "lbl_plot": "Handlung", "lbl_narrative": "Erzählstil", "lbl_goal": "Buchziel", "lbl_pov": "Erzählperspektive (Pronomen)",
        "btn_res": "🔄 PROJEKT ZURÜCKSETZEN", "tabs": ["📊 1. Index", "✍️ 2. Schreiben & Quiz", "📖 3. Vorschau", "📑 4. Importieren / Exportieren / Copyright"],
        "btn_idx": "🚀 Index Generieren", "btn_sync": "✅ Synchronisieren", "lbl_sec": "Abschnitt wählen:",
        "btn_write": "✨ INHALT SCHREIBEN", "btn_quiz": "🧠 QUIZ HINZUFÜGEN", "btn_edit": "🚀 UMSCHREIBEN",
        "msg_run": "Analysiere Hierarchie und Stil...", "preface": "Vorwort", "ack": "Danksagungen",
        "preview_tit": "📖 Leseansicht", "btn_word": "📥 Word Herunterladen", "btn_pdf": "📥 PDF Herunterladen",
        "msg_err_idx": "Generiere zuerst den Index.", "msg_success_sync": "Synchronisiert!", "label_editor": "Professioneller Editor", "welcome": "👋 Willkommen.", "guide": "Nutze die Seitenleiste."
    },
    "Română": {
        "side_tit": "⚙️ Configurare Editor", "lbl_tit": "Titlul Cărții", "lbl_auth": "Nume Autor", "lbl_lang": "Limbă", 
        "lbl_gen": "Gen Literar", "lbl_style": "Stil de Scriere", "lbl_plot": "Subiect", "lbl_narrative": "Stil Narativ", "lbl_goal": "Obiectivul Cărții", "lbl_pov": "Punct de Vedere (Pronume)",
        "btn_res": "🔄 RESETARE PROIECT", "tabs": ["📊 1. Cuprins", "✍️ 2. Scriere & Quiz", "📖 3. Previzualizare", "📑 4. Import / Export / Copyright"],
        "btn_idx": "🚀 Generare Cuprins", "btn_sync": "✅ Sincronizare", "lbl_sec": "Selectează secțiunea:",
        "btn_write": "✨ SCRIE CONȚINUT", "btn_quiz": "🧠 ADAUGĂ QUIZ", "btn_edit": "🚀 RESCRIE",
        "msg_run": "Se analizează ierarhia și stilul...", "preface": "Prefață", "ack": "Mulțumiri",
        "preview_tit": "📖 Mod Citire", "btn_word": "📥 Descarcă Word", "btn_pdf": "📥 Descarcă PDF",
        "msg_err_idx": "Generează cuprinsul mai întâi.", "msg_success_sync": "Sincronizat!", "label_editor": "Editor Profesional", "welcome": "👋 Bun venit.", "guide": "Folosește bara lateral."
    },
    "Русский": {
        "side_tit": "⚙️ Настройки Редактора", "lbl_tit": "Название Книги", "lbl_auth": "Имя Автора", "lbl_lang": "Язык", 
        "lbl_gen": "Жанр", "lbl_style": "Стиль Написания", "lbl_plot": "Сюжет", "lbl_narrative": "Стиль Повествования", "lbl_goal": "Цель Книги", "lbl_pov": "Точка зрения (Местоимение)",
        "btn_res": "🔄 СБРОСИТЬ ПРОЕКТ", "tabs": ["📊 1. Оглавление", "✍️ 2. Текст и Тест", "📖 3. Просмотр", "📑 4. Импорт / Экспорт / Copyright"],
        "btn_idx": "🚀 Создать Оглавление", "btn_sync": "✅ Синхронизировать", "lbl_sec": "Выберите раздел:",
        "btn_write": "✨ НАПИСАТЬ ТЕКСТ", "btn_quiz": "🧠 ДОБАВИТЬ ТЕСТ", "btn_edit": "🚀 ПЕРЕПИСАТЬ",
        "msg_run": "Анализ иерархии и стиля...", "preface": "Предисловие", "ack": "Благодарности",
        "preview_tit": "📖 Режим Чтения", "btn_word": "📥 Скачать Word", "btn_pdf": "📥 Скачать PDF",
        "msg_err_idx": "Сначала создайте оглавление.", "msg_success_sync": "Синхронизировано!", "label_editor": "Профессиональный Редактор", "welcome": "👋 Добро пожаловать.", "guide": "Используйте боковую панель."
    },
    "العربية": {
        "side_tit": "⚙️ إعدادات المحرر", "lbl_tit": "عنوان الكتاب", "lbl_auth": "اسم المؤلف", "lbl_lang": "اللغة", 
        "lbl_gen": "النوع الأدبي", "lbl_style": "أسلوب الكتابة", "lbl_plot": "الحبكة أو الموضوع", "lbl_narrative": "الأسلوب السردي", "lbl_goal": "هدف الكتاب", "lbl_pov": "وجهة النظر (الضمير)",
        "btn_res": "🔄 إعادة ضبط المشروع", "tabs": ["📊 1. الفهرس", "✍️ 2. الكتابة والاختبار", "📖 3. معاينة", "📑 4. استيراد / تصدير / Copyright"],
        "btn_idx": "🚀 إنشاء فهرس احترافي", "btn_sync": "✅ حفظ ومزامنة الفصول", "lbl_sec": "اختر القسم:",
        "btn_write": "✨ كتابة المحتوى", "btn_quiz": "🧠 إضافة اختبار", "btn_edit": "🚀 إعادة صياغة",
        "msg_run": "جاري تحليل التسلسل الهرمي والأسلوب...", "preface": "مقدمة", "ack": "شكر وتقدير",
        "preview_tit": "📖 عرض القراءة الاحترافي", "btn_word": "📥 تحميل Word", "btn_pdf": "📥 تحميل PDF",
        "msg_err_idx": "قم بإنشاء الفهرس أولاً.", "msg_success_sync": "تمت المزامنة!", "label_editor": "محرر نصوص احترافي", "welcome": "👋 مرحباً بك.", "guide": "استخدم الشريط الجانبي."
    },
    "中文": {
        "side_tit": "⚙️ 编辑器设置", "lbl_tit": "书名", "lbl_auth": "作者姓名", "lbl_lang": "语言", 
        "lbl_gen": "文学体裁", "lbl_style": "写作类型", "lbl_plot": "情节或主题", "lbl_narrative": "叙事风格", "lbl_goal": "书籍目标", "lbl_pov": "叙事视角 (代词)",
        "btn_res": "🔄 重置项目", "tabs": ["📊 1. 目录", "✍️ 2. 写作与测试", "📖 3. 预览", "📑 4. 导入 / 导出 / Copyright"],
        "btn_idx": "🚀 生成专业目录", "btn_sync": "✅ 保存并同步章节", "lbl_sec": "选择章节:",
        "btn_write": "✨ 编写内容", "btn_quiz": "🧠 添加测试", "btn_edit": "🚀 用AI重写",
        "msg_run": "正在分析层级、风格和情感目标...", "preface": "前言", "ack": "致谢",
        "preview_tit": "📖 专业阅读视图", "btn_word": "📥 下载 Word", "btn_pdf": "📥 下载 PDF",
        "msg_err_idx": "请先生成目录。", "msg_success_sync": "已同步！", "label_editor": "专业文本编辑器", "welcome": "👋 欢迎。", "guide": "请使用左侧边栏设置书籍参数。"
    }
}

# Testi dell'interfaccia separati dai dati editoriali: cambiano soltanto ciò
# che l'utente vede, mai i valori tecnici salvati nei progetti e nei CSV.
TESTI_UI = {
    "Italiano": {
        "seleziona": "— Seleziona —", "lingua": "🌐 Lingua", "cervello": "🧠 Cervello AI",
        "fonti_carica": "Carica fonti esterne", "risultato": "Risultato finale desiderato",
        "risultato_placeholder": "Es: Alla fine il lettore deve saper applicare il metodo in autonomia e verificare il risultato.",
        "approfondimenti": "Approfondimenti (facoltativo)", "approfondimenti_placeholder": "Inserisci istruzioni, aspetti da trattare con maggiore attenzione, vincoli, esempi o temi obbligatori.",
        "lunghezza": "Lunghezza delle sezioni", "lunghezza_help": "Definisce la lunghezza del testo generato per ogni sezione, senza modificare il costo in crediti.",
        "salva": "💾 SALVA SESSIONE", "ripristina": "🔄 RIAGGIORNA ALL'ULTIMA STESURA", "formattazione": "🛠️ 5. Formattazione",
    },
    "English": {
        "seleziona": "— Select —", "lingua": "🌐 Language", "cervello": "🧠 AI engine",
        "fonti_carica": "Upload external sources", "risultato": "Desired final result",
        "risultato_placeholder": "Example: By the end, the reader can apply the method independently and verify the result.",
        "approfondimenti": "Further details (optional)", "approfondimenti_placeholder": "Add instructions, priority topics, constraints, examples or required themes.",
        "lunghezza": "Section length", "lunghezza_help": "Sets the generated-text length for each section without changing the credit cost.",
        "salva": "💾 SAVE SESSION", "ripristina": "🔄 RESTORE LATEST DRAFT", "formattazione": "🛠️ 5. Formatting",
    },
    "Español": {
        "seleziona": "— Selecciona —", "lingua": "🌐 Idioma", "cervello": "🧠 Motor de IA",
        "fonti_carica": "Cargar fuentes externas", "risultato": "Resultado final deseado",
        "risultato_placeholder": "Ejemplo: Al final, el lector podrá aplicar el método de forma autónoma y verificar el resultado.",
        "approfondimenti": "Profundización (opcional)", "approfondimenti_placeholder": "Incluye instrucciones, aspectos prioritarios, límites, ejemplos o temas obligatorios.",
        "lunghezza": "Longitud de las secciones", "lunghezza_help": "Define la longitud del texto generado por sección sin cambiar el coste en créditos.",
        "salva": "💾 GUARDAR SESIÓN", "ripristina": "🔄 RESTAURAR ÚLTIMO BORRADOR", "formattazione": "🛠️ 5. Formato",
    },
    "Français": {
        "seleziona": "— Sélectionnez —", "lingua": "🌐 Langue", "cervello": "🧠 Moteur IA",
        "fonti_carica": "Importer des sources externes", "risultato": "Résultat final souhaité",
        "risultato_placeholder": "Exemple : à la fin, le lecteur pourra appliquer la méthode de façon autonome et vérifier le résultat.",
        "approfondimenti": "Approfondissements (facultatif)", "approfondimenti_placeholder": "Ajoutez des consignes, priorités, contraintes, exemples ou thèmes obligatoires.",
        "lunghezza": "Longueur des sections", "lunghezza_help": "Définit la longueur du texte généré pour chaque section sans modifier le coût en crédits.",
        "salva": "💾 ENREGISTRER LA SESSION", "ripristina": "🔄 RESTAURER LA DERNIÈRE VERSION", "formattazione": "🛠️ 5. Mise en forme",
    },
    "Deutsch": {
        "seleziona": "— Auswählen —", "lingua": "🌐 Sprache", "cervello": "🧠 KI-Modell",
        "fonti_carica": "Externe Quellen hochladen", "risultato": "Gewünschtes Endergebnis",
        "risultato_placeholder": "Beispiel: Am Ende kann der Leser die Methode selbstständig anwenden und das Ergebnis prüfen.",
        "approfondimenti": "Vertiefungen (optional)", "approfondimenti_placeholder": "Fügen Sie Anweisungen, Schwerpunkte, Vorgaben, Beispiele oder Pflichthemen hinzu.",
        "lunghezza": "Abschnittslänge", "lunghezza_help": "Legt die Textlänge je Abschnitt fest, ohne die Kreditkosten zu ändern.",
        "salva": "💾 SITZUNG SPEICHERN", "ripristina": "🔄 LETZTEN ENTWURF WIEDERHERSTELLEN", "formattazione": "🛠️ 5. Formatierung",
    },
    "Română": {
        "seleziona": "— Selectează —", "lingua": "🌐 Limbă", "cervello": "🧠 Motor IA",
        "fonti_carica": "Încarcă surse externe", "risultato": "Rezultatul final dorit",
        "risultato_placeholder": "Exemplu: La final, cititorul poate aplica metoda autonom și poate verifica rezultatul.",
        "approfondimenti": "Aprofundări (opțional)", "approfondimenti_placeholder": "Adaugă instrucțiuni, aspecte prioritare, limite, exemple sau teme obligatorii.",
        "lunghezza": "Lungimea secțiunilor", "lunghezza_help": "Stabilește lungimea textului generat pentru fiecare secțiune fără a schimba costul în credite.",
        "salva": "💾 SALVEAZĂ SESIUNEA", "ripristina": "🔄 RESTABILEȘTE ULTIMA VERSIUNE", "formattazione": "🛠️ 5. Formatare",
    },
    "Русский": {
        "seleziona": "— Выберите —", "lingua": "🌐 Язык", "cervello": "🧠 Модель ИИ",
        "fonti_carica": "Загрузить внешние источники", "risultato": "Желаемый итоговый результат",
        "risultato_placeholder": "Пример: в конце читатель сможет самостоятельно применить метод и проверить результат.",
        "approfondimenti": "Дополнительные сведения (необязательно)", "approfondimenti_placeholder": "Добавьте инструкции, приоритетные аспекты, ограничения, примеры или обязательные темы.",
        "lunghezza": "Длина разделов", "lunghezza_help": "Задаёт длину текста для каждого раздела без изменения стоимости в кредитах.",
        "salva": "💾 СОХРАНИТЬ СЕССИЮ", "ripristina": "🔄 ВОССТАНОВИТЬ ПОСЛЕДНЮЮ ВЕРСИЮ", "formattazione": "🛠️ 5. Форматирование",
    },
    "العربية": {
        "seleziona": "— اختر —", "lingua": "🌐 اللغة", "cervello": "🧠 محرك الذكاء الاصطناعي",
        "fonti_carica": "تحميل مصادر خارجية", "risultato": "النتيجة النهائية المطلوبة",
        "risultato_placeholder": "مثال: في النهاية يستطيع القارئ تطبيق المنهج بشكل مستقل والتحقق من النتيجة.",
        "approfondimenti": "تفاصيل إضافية (اختياري)", "approfondimenti_placeholder": "أضف تعليمات أو جوانب ذات أولوية أو قيوداً أو أمثلة أو موضوعات إلزامية.",
        "lunghezza": "طول الأقسام", "lunghezza_help": "يحدد طول النص المُنشأ لكل قسم من دون تغيير تكلفة الأرصدة.",
        "salva": "💾 حفظ الجلسة", "ripristina": "🔄 استعادة آخر مسودة", "formattazione": "🛠️ 5. التنسيق",
    },
    "中文": {
        "seleziona": "— 请选择 —", "lingua": "🌐 语言", "cervello": "🧠 AI 模型",
        "fonti_carica": "上传外部资料", "risultato": "期望的最终结果",
        "risultato_placeholder": "示例：完成后，读者能够独立运用该方法并验证结果。",
        "approfondimenti": "补充说明（可选）", "approfondimenti_placeholder": "添加说明、重点、限制、示例或必需主题。",
        "lunghezza": "章节长度", "lunghezza_help": "设定每个章节生成文本的长度，不会改变积分费用。",
        "salva": "💾 保存会话", "ripristina": "🔄 恢复最新草稿", "formattazione": "🛠️ 5. 格式化",
    },
}


def testo_ui(chiave, lingua=None):
    """Restituisce testo UI localizzato con ripiego sicuro in italiano."""
    lingua_effettiva = lingua or st.session_state.get("editor_language", "Italiano")
    return TESTI_UI.get(lingua_effettiva, TESTI_UI["Italiano"]).get(
        chiave, TESTI_UI["Italiano"].get(chiave, chiave)
    )


CAMPI_UI_AVANZATI = {
    "Italiano": {"indice_salvato": "Indice conservato", "fonti_salvate": "Fonti web conservate", "sezione_salvata": "Sezione conservata", "testo_salvato": "Testo conservato", "indice": "Indice gerarchico:", "messaggio": "Il tuo messaggio", "cerca": "🔎 Cerca nel manoscritto", "sostituisci": "Sostituisci con", "maiuscole": "Distingui maiuscole e minuscole", "anteprima_sezione": "Vai a una sezione dell'anteprima", "immagine": "Carica un'immagine PNG, JPG o WEBP", "import_csv": "Importa progetto CSV", "manoscritto": "Carica manoscritto", "lingua_metadati": "Lingua dei metadati"},
    "English": {"indice_salvato": "Saved index", "fonti_salvate": "Saved web sources", "sezione_salvata": "Saved section", "testo_salvato": "Saved text", "indice": "Hierarchical index:", "messaggio": "Your message", "cerca": "🔎 Search in the manuscript", "sostituisci": "Replace with", "maiuscole": "Match case", "anteprima_sezione": "Go to a preview section", "immagine": "Upload a PNG, JPG or WEBP image", "import_csv": "Import CSV project", "manoscritto": "Upload manuscript", "lingua_metadati": "Metadata language"},
    "Español": {"indice_salvato": "Índice guardado", "fonti_salvate": "Fuentes web guardadas", "sezione_salvata": "Sección guardada", "testo_salvato": "Texto guardado", "indice": "Índice jerárquico:", "messaggio": "Tu mensaje", "cerca": "🔎 Buscar en el manuscrito", "sostituisci": "Sustituir por", "maiuscole": "Distinguir mayúsculas y minúsculas", "anteprima_sezione": "Ir a una sección de la vista previa", "immagine": "Cargar una imagen PNG, JPG o WEBP", "import_csv": "Importar proyecto CSV", "manoscritto": "Cargar manuscrito", "lingua_metadati": "Idioma de metadatos"},
    "Français": {"indice_salvato": "Index enregistré", "fonti_salvate": "Sources web enregistrées", "sezione_salvata": "Section enregistrée", "testo_salvato": "Texte enregistré", "indice": "Index hiérarchique :", "messaggio": "Votre message", "cerca": "🔎 Rechercher dans le manuscrit", "sostituisci": "Remplacer par", "maiuscole": "Respecter la casse", "anteprima_sezione": "Aller à une section de l'aperçu", "immagine": "Importer une image PNG, JPG ou WEBP", "import_csv": "Importer un projet CSV", "manoscritto": "Importer un manuscrit", "lingua_metadati": "Langue des métadonnées"},
    "Deutsch": {"indice_salvato": "Gespeicherter Index", "fonti_salvate": "Gespeicherte Webquellen", "sezione_salvata": "Gespeicherter Abschnitt", "testo_salvato": "Gespeicherter Text", "indice": "Hierarchischer Index:", "messaggio": "Ihre Nachricht", "cerca": "🔎 Im Manuskript suchen", "sostituisci": "Ersetzen durch", "maiuscole": "Groß- und Kleinschreibung beachten", "anteprima_sezione": "Zu einem Vorschauabschnitt gehen", "immagine": "PNG-, JPG- oder WEBP-Bild hochladen", "import_csv": "CSV-Projekt importieren", "manoscritto": "Manuskript hochladen", "lingua_metadati": "Metadatensprache"},
    "Română": {"indice_salvato": "Cuprins salvat", "fonti_salvate": "Surse web salvate", "sezione_salvata": "Secțiune salvată", "testo_salvato": "Text salvat", "indice": "Cuprins ierarhic:", "messaggio": "Mesajul tău", "cerca": "🔎 Caută în manuscris", "sostituisci": "Înlocuiește cu", "maiuscole": "Distinge majuscule/minuscule", "anteprima_sezione": "Mergi la o secțiune din previzualizare", "immagine": "Încarcă o imagine PNG, JPG sau WEBP", "import_csv": "Importă proiect CSV", "manoscritto": "Încarcă manuscrisul", "lingua_metadati": "Limba metadatelor"},
    "Русский": {"indice_salvato": "Сохранённое оглавление", "fonti_salvate": "Сохранённые веб-источники", "sezione_salvata": "Сохранённый раздел", "testo_salvato": "Сохранённый текст", "indice": "Иерархическое оглавление:", "messaggio": "Ваше сообщение", "cerca": "🔎 Поиск в рукописи", "sostituisci": "Заменить на", "maiuscole": "Учитывать регистр", "anteprima_sezione": "Перейти к разделу предпросмотра", "immagine": "Загрузить изображение PNG, JPG или WEBP", "import_csv": "Импортировать проект CSV", "manoscritto": "Загрузить рукопись", "lingua_metadati": "Язык метаданных"},
    "العربية": {"indice_salvato": "الفهرس المحفوظ", "fonti_salvate": "مصادر الويب المحفوظة", "sezione_salvata": "القسم المحفوظ", "testo_salvato": "النص المحفوظ", "indice": "الفهرس الهرمي:", "messaggio": "رسالتك", "cerca": "🔎 البحث في المخطوط", "sostituisci": "استبدال بـ", "maiuscole": "مطابقة حالة الأحرف", "anteprima_sezione": "الانتقال إلى قسم من المعاينة", "immagine": "تحميل صورة PNG أو JPG أو WEBP", "import_csv": "استيراد مشروع CSV", "manoscritto": "تحميل المخطوط", "lingua_metadati": "لغة البيانات الوصفية"},
    "中文": {"indice_salvato": "已保存的目录", "fonti_salvate": "已保存的网络资料", "sezione_salvata": "已保存的章节", "testo_salvato": "已保存的文本", "indice": "层级目录：", "messaggio": "你的消息", "cerca": "🔎 在稿件中搜索", "sostituisci": "替换为", "maiuscole": "区分大小写", "anteprima_sezione": "前往预览章节", "immagine": "上传 PNG、JPG 或 WEBP 图片", "import_csv": "导入 CSV 项目", "manoscritto": "上传稿件", "lingua_metadati": "元数据语言"},
}


def campo_ui(chiave, lingua=None):
    lingua_effettiva = lingua or st.session_state.get("editor_language", "Italiano")
    return CAMPI_UI_AVANZATI.get(lingua_effettiva, CAMPI_UI_AVANZATI["Italiano"]).get(
        chiave, CAMPI_UI_AVANZATI["Italiano"].get(chiave, chiave)
    )


ETICHETTE_CAMPI_CONTROLLO = {
    "Italiano": ("Ho verificato anteprima, lettore vocale, CSV e PDF", "Analisi e voto dell'indice", "Analisi completa del libro", "Proposta di indice migliorata (l'indice attuale resta al sicuro)", "Confermo la sostituzione in tutte le sezioni trovate.", "Metadati generati", "Prompt pronti per Rigenera con AI (il software non modifica nulla automaticamente)"),
    "English": ("I have checked preview, voice reader, CSV and PDF", "Index analysis and score", "Complete book analysis", "Improved index proposal (the current index remains safe)", "I confirm replacement in every matched section.", "Generated metadata", "Prompts ready for Rewrite with AI (the software changes nothing automatically)"),
    "Español": ("He comprobado la vista previa, lector de voz, CSV y PDF", "Análisis y valoración del índice", "Análisis completo del libro", "Propuesta de índice mejorada (el índice actual permanece seguro)", "Confirmo la sustitución en todas las secciones encontradas.", "Metadatos generados", "Prompts listos para Regenerar con IA (el software no modifica nada automáticamente)"),
    "Français": ("J'ai vérifié l'aperçu, le lecteur vocal, le CSV et le PDF", "Analyse et évaluation de l'index", "Analyse complète du livre", "Proposition d'index améliorée (l'index actuel reste protégé)", "Je confirme le remplacement dans toutes les sections trouvées.", "Métadonnées générées", "Prompts prêts pour Réécrire avec l'IA (le logiciel ne modifie rien automatiquement)"),
    "Deutsch": ("Ich habe Vorschau, Sprachleser, CSV und PDF geprüft", "Indexanalyse und Bewertung", "Vollständige Buchanalyse", "Verbesserter Indexvorschlag (der aktuelle Index bleibt sicher)", "Ich bestätige den Ersatz in allen gefundenen Abschnitten.", "Generierte Metadaten", "Prompts bereit für KI-Umschreibung (die Software ändert nichts automatisch)"),
    "Română": ("Am verificat previzualizarea, cititorul vocal, CSV-ul și PDF-ul", "Analiza și evaluarea cuprinsului", "Analiza completă a cărții", "Propunere de cuprins îmbunătățită (cuprinsul actual rămâne în siguranță)", "Confirm înlocuirea în toate secțiunile găsite.", "Metadate generate", "Prompturi pregătite pentru Rescrie cu IA (software-ul nu modifică nimic automat)"),
    "Русский": ("Я проверил предпросмотр, голосовой чтец, CSV и PDF", "Анализ и оценка оглавления", "Полный анализ книги", "Улучшенное предложение оглавления (текущее оглавление сохранено)", "Я подтверждаю замену во всех найденных разделах.", "Сгенерированные метаданные", "Промпты готовы для переписывания с ИИ (программа ничего не меняет автоматически)"),
    "العربية": ("لقد تحققت من المعاينة والقارئ الصوتي وCSV وPDF", "تحليل وتقييم الفهرس", "التحليل الكامل للكتاب", "اقتراح فهرس محسّن (يبقى الفهرس الحالي آمناً)", "أؤكد الاستبدال في جميع الأقسام التي تم العثور عليها.", "بيانات وصفية تم إنشاؤها", "موجّهات جاهزة لإعادة الصياغة بالذكاء الاصطناعي (لا يعدّل البرنامج شيئاً تلقائياً)"),
    "中文": ("我已检查预览、语音朗读、CSV 和 PDF", "目录分析与评分", "完整图书分析", "改进的目录建议（当前目录保持安全）", "我确认替换所有找到的章节。", "已生成的元数据", "可用于 AI 重写的提示词（软件不会自动修改任何内容）"),
}


def campo_controllo_ui(indice, lingua=None):
    lingua_effettiva = lingua or st.session_state.get("editor_language", "Italiano")
    valori = ETICHETTE_CAMPI_CONTROLLO.get(lingua_effettiva, ETICHETTE_CAMPI_CONTROLLO["Italiano"])
    return valori[indice]

# La Prefazione fa parte del libro e viene generata prima della Parte I.
# I Ringraziamenti restano invece dismessi: non vengono più aggiunti né
# richiesti alla stesura completa, ma non cancelliamo eventuali vecchi testi.
SEZIONI_DISMESSE = frozenset(
    str(traduzione.get("ack", "")).strip()
    for traduzione in TRADUZIONI.values()
    if str(traduzione.get("ack", "")).strip()
)


def titolo_prefazione(lingua=None):
    """Restituisce il titolo localizzato della Prefazione."""
    lingua_effettiva = lingua or st.session_state.get("editor_language", "Italiano")
    return str(TRADUZIONI.get(lingua_effettiva, TRADUZIONI["Italiano"]).get("preface", "Prefazione")).strip()


def sezione_prefazione(sezione):
    valore = re.sub(r"\s+", " ", str(sezione or "").strip()).casefold()
    titoli = {
        re.sub(r"\s+", " ", str(traduzione.get("preface", "")).strip()).casefold()
        for traduzione in TRADUZIONI.values()
    }
    return bool(valore and valore in titoli)


def sezione_dismessa(sezione):
    return str(sezione or "").strip() in SEZIONI_DISMESSE

# ======================================================================================================================
# 3. BLOCCO CSS: SIDEBAR SCURA E PULSANTI SCURI (FORZATURA !IMPORTANT)
# ======================================================================================================================
st.markdown("""
<style>
#MainMenu, footer, header, [data-testid="stHeader"] {visibility: hidden;}
[data-testid="collapsedControl"] { display: none !important; }

section[data-testid="stSidebar"] { 
    min-width: 370px !important; max-width: 370px !important;
    width: 370px !important; display: block !important; visibility: visible !important;
    transform: none !important; background: linear-gradient(180deg, #0b172a 0%, #101d31 100%) !important;
    border-right: 1px solid #29405f;
}
/* Sidebar fissa: impedisce il collasso accidentale della barra laterale. */
section[data-testid="stSidebar"][aria-expanded="false"],
[data-testid="stSidebar"][aria-expanded="false"] {
    min-width: 370px !important; max-width: 370px !important; width: 370px !important;
    display: block !important; visibility: visible !important; transform: none !important;
}
section[data-testid="stSidebar"] > div:first-child {
    width: 370px !important; min-width: 370px !important; display: block !important;
    padding: 1rem .9rem 2rem !important;
}
section[data-testid="stSidebar"] .stMarkdown, section[data-testid="stSidebar"] label, 
section[data-testid="stSidebar"] p, section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2, section[data-testid="stSidebar"] h3 {
    color: #ffffff !important;
}
section[data-testid="stSidebar"] [data-testid="stWidgetLabel"] p {
    font-size: .83rem !important; font-weight: 700 !important; letter-spacing: .01em;
}
section[data-testid="stSidebar"] .stTextInput input,
section[data-testid="stSidebar"] .stTextArea textarea,
section[data-testid="stSidebar"] div[data-baseweb="select"] > div {
    background: #16263d !important; border-color: #385574 !important;
    border-radius: 9px !important;
}
.stButton>button {
    width: 100% !important; border-radius: 10px !important; height: 3.15em !important;
    font-weight: 700 !important; background: #14243a !important; color: #f8fbff !important;
    font-size: .98rem !important; border: 1px solid #355273 !important;
    box-shadow: none !important; transition: all .18s ease !important;
}
.stButton>button:hover { 
    background: #1d3655 !important; border-color: #65a8ff !important;
    color: #ffffff !important; transform: translateY(-1px) !important;
}
.stButton>button[kind="primary"], .stButton>button[data-testid="baseButton-primary"] {
    background: linear-gradient(135deg, #1976e9, #2997ef) !important;
    border-color: #4aa8ff !important;
}
.stButton>button[kind="primary"]:hover, .stButton>button[data-testid="baseButton-primary"]:hover {
    background: linear-gradient(135deg, #1669cf, #1e88e5) !important;
}
.st-key-avvia_chat_sidebar_guidata .stButton>button,
.st-key-avvia_chat_sidebar_guidata button {
    background: linear-gradient(135deg, #15803d, #22a75a) !important;
    border-color: #52d486 !important;
    color: #ffffff !important;
}
.st-key-avvia_chat_sidebar_guidata .stButton>button:hover,
.st-key-avvia_chat_sidebar_guidata button:hover {
    background: linear-gradient(135deg, #126b33, #178747) !important;
    border-color: #86efac !important;
}
.st-key-avvia_chat_sidebar_subito .stButton>button,
.st-key-avvia_chat_sidebar_subito button {
    background: linear-gradient(135deg, #15803d, #22a75a) !important;
    border-color: #52d486 !important;
    color: #ffffff !important;
}
.st-key-avvia_chat_sidebar_subito .stButton>button:hover,
.st-key-avvia_chat_sidebar_subito button:hover {
    background: linear-gradient(135deg, #126b33, #178747) !important;
    border-color: #86efac !important;
}
/* Il reset è l'unico comando distruttivo del progetto: deve distinguersi
   chiaramente da salvataggio, aggiornamento e normali azioni editoriali. */
.st-key-reset_progetto_distruttivo .stButton>button,
.st-key-reset_progetto_distruttivo button {
    background: linear-gradient(135deg, #b4232b, #dc3545) !important;
    border-color: #ff8792 !important;
    color: #ffffff !important;
}
.st-key-reset_progetto_distruttivo .stButton>button:hover,
.st-key-reset_progetto_distruttivo button:hover {
    background: linear-gradient(135deg, #901b23, #b91c2b) !important;
    border-color: #fecdd3 !important;
}
[data-testid="stAppViewContainer"] { background: radial-gradient(circle at 60% -20%, #1b3454 0%, #0b1423 42%, #08111e 100%) !important; }
[data-testid="stMainBlockContainer"] { max-width: 1540px !important; padding-top: 1.35rem !important; }
[data-testid="stTabs"] [data-baseweb="tab-list"] { gap: .4rem; border-bottom: 1px solid #2c405b; }
[data-testid="stTabs"] button[role="tab"] { border-radius: 9px 9px 0 0; padding: .65rem .9rem; color: #b7c8df; }
[data-testid="stTabs"] button[aria-selected="true"] { color: #ffffff; background: #162a44; }
.ss-workspace-header {
    display:flex; align-items:center; justify-content:space-between; gap:1rem;
    padding:1.05rem 1.25rem; margin:0 0 1.15rem; border-radius:16px;
    border:1px solid #314a68; background:linear-gradient(135deg, rgba(22,42,68,.96), rgba(12,27,47,.96));
    box-shadow:0 14px 36px rgba(0,0,0,.20);
}
.ss-workspace-brand { display:flex; align-items:center; gap:.9rem; min-width:0; }
.ss-workspace-mark { width:38px; height:38px; display:grid; place-items:center; border-radius:11px; background:#e7bd69; color:#17253a; font-size:1.3rem; }
.ss-workspace-title { color:#fff; font-size:1.25rem; font-weight:800; white-space:nowrap; overflow:hidden; text-overflow:ellipsis; }
.ss-workspace-subtitle { color:#aac0da; font-size:.78rem; margin-top:.18rem; }
.ss-workspace-meta { display:flex; gap:.55rem; flex-wrap:wrap; justify-content:flex-end; }
.ss-workspace-chip { color:#e8f3ff; background:#12243b; border:1px solid #3a5878; padding:.42rem .65rem; border-radius:999px; font-size:.78rem; font-weight:700; }
.ss-workspace-chip.ok { color:#a6efb7; border-color:#397c53; background:#102d23; }
.ss-workspace-chip.ai { color:#bcdcff; }
.ss-workspace-chip.credit { color:#ffe3a1; }
.ss-project-panel { margin:.2rem 0 1.15rem; padding:1.15rem; border:1px solid #2c4563; border-radius:16px; background:linear-gradient(145deg,rgba(15,32,53,.92),rgba(10,23,40,.90)); box-shadow:0 14px 34px rgba(0,0,0,.16); }
.ss-project-title { margin:0 0 .95rem; color:#fff; font-size:1.55rem; font-weight:850; letter-spacing:-.02em; }
.ss-project-journey { display:grid; grid-template-columns:repeat(5,minmax(0,1fr)); gap:.48rem; margin-bottom:1rem; }
.ss-project-step { min-width:0; display:flex; align-items:center; gap:.48rem; padding:.65rem .6rem; border:1px solid #2b4968; border-radius:11px; background:#0d2037; color:#a8bed5; font-size:.78rem; font-weight:750; }
.ss-project-step span { flex:0 0 1.5rem; width:1.5rem; height:1.5rem; display:grid; place-items:center; border-radius:50%; background:#29445f; color:#dcecff; font-size:.72rem; }
.ss-project-step.done { color:#b8f6c7; border-color:#357754; background:#102c24; }.ss-project-step.done span { background:#1f9d55; color:#fff; }
.ss-project-step.current { color:#fff; border-color:#3ea7fa; background:linear-gradient(135deg,#165a91,#187fc8); box-shadow:0 6px 15px rgba(33,150,243,.18); }.ss-project-step.current span { background:#fff; color:#1477c1; }
.ss-project-metrics { display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); gap:.62rem; margin:.15rem 0 .95rem; }
.ss-project-metric { min-width:0; padding:.78rem .82rem; border-radius:12px; border:1px solid #294764; background:rgba(9,25,43,.72); }.ss-project-metric span { display:block; color:#9eb8d3; font-size:.75rem; font-weight:700; }.ss-project-metric b { display:block; margin-top:.25rem; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; color:#fff; font-size:1.55rem; letter-spacing:-.04em; }
.ss-next-action { display:flex; gap:.75rem; align-items:flex-start; padding:.88rem 1rem; border-radius:12px; background:#153a60; border:1px solid #28699b; color:#e8f5ff; }.ss-next-action .ss-next-icon { flex:0 0 2rem; width:2rem; height:2rem; display:grid; place-items:center; border-radius:9px; background:#2a9ded; font-size:1.08rem; }.ss-next-action small { display:block; color:#a7d4f6; font-weight:800; margin-bottom:.16rem; }.ss-next-action strong { display:block; font-size:1rem; line-height:1.38; }.ss-next-action p { margin:.38rem 0 0; color:#d7ecfc; font-size:.88rem; line-height:1.45; }.ss-next-action p b { color:#fff; }
.ss-next-action.pause { background:#4e4217; border-color:#a98629; }.ss-next-action.pause .ss-next-icon { background:#d8a621; color:#1c1708; }.ss-next-action.complete { background:#103d2e; border-color:#2c8d61; }.ss-next-action.complete .ss-next-icon { background:#28ad70; }
.ss-section-card { background:rgba(14,29,48,.72); border:1px solid #2c4563; border-radius:14px; padding:.45rem 1rem 1rem; }
.ss-section-card h2, .ss-section-card h3 { color:#fff !important; }
.preview-box {
    background-color: #fffdf7 !important; padding: min(6vw,80px); border: 1px solid #d8c9a8;
    border-radius: 12px; height: 900px; overflow-y: scroll;
    font-family: 'Times New Roman', serif; line-height: 2.0; 
    color: #182236 !important; box-shadow: 0px 25px 60px rgba(0,0,0,.28); margin: 0 auto;
}
.custom-title {
    display:none;
}
div[data-baseweb="select"] > div { background-color: #16263d !important; color: white !important; }
@media (max-width: 900px) {
    section[data-testid="stSidebar"], section[data-testid="stSidebar"] > div:first-child { min-width: 320px !important; width:320px !important; }
    .ss-workspace-header { align-items:flex-start; flex-direction:column; }
    .ss-workspace-meta { justify-content:flex-start; }
}
/* Layout mobile: modifiche esclusivamente grafiche e responsive. */
@media (max-width: 768px) {
    [data-testid="stMainBlockContainer"] {
        padding: 3.55rem .7rem 5.6rem !important;
        max-width: 100% !important;
    }
    /* Su telefono la sidebar torna richiudibile: non sottrae spazio alla
       scrittura e si apre dal normale comando laterale di Streamlit. */
    header[data-testid="stHeader"], [data-testid="stHeader"] {
        visibility: visible !important; background: transparent !important;
        pointer-events: none !important;
    }
    [data-testid="collapsedControl"] {
        display: flex !important; visibility: visible !important; pointer-events: auto !important;
        position: fixed !important; top: .55rem !important; left: .55rem !important;
        z-index: 1001 !important; border-radius: 10px !important;
        background: #142b48 !important; border: 1px solid #4aa8ff !important;
        box-shadow: 0 5px 18px rgba(0,0,0,.28) !important;
    }
    section[data-testid="stSidebar"][aria-expanded="false"] {
        min-width: 0 !important; max-width: 0 !important; width: 0 !important;
        display: none !important; visibility: hidden !important;
    }
    section[data-testid="stSidebar"][aria-expanded="true"],
    section[data-testid="stSidebar"][aria-expanded="true"] > div:first-child {
        min-width: min(92vw, 360px) !important; max-width: min(92vw, 360px) !important;
        width: min(92vw, 360px) !important; padding: .7rem .65rem 1.2rem !important;
    }
    section[data-testid="stSidebar"] .stElementContainer { margin-bottom: .38rem !important; }
    section[data-testid="stSidebar"] [data-testid="stWidgetLabel"] p { font-size: .78rem !important; }
    section[data-testid="stSidebar"] .stTextInput input,
    section[data-testid="stSidebar"] .stTextArea textarea,
    section[data-testid="stSidebar"] div[data-baseweb="select"] > div {
        min-height: 2.35rem !important; font-size: .92rem !important;
    }
    section[data-testid="stSidebar"] .stTextArea textarea { min-height: 5.6rem !important; }

    /* Le colonne diventano blocchi verticali: nessun pulsante o campo resta
       troppo stretto per essere usato con il pollice. */
    [data-testid="stHorizontalBlock"] {
        flex-direction: column !important; gap: .65rem !important;
    }
    [data-testid="column"], [data-testid="stColumn"], [data-testid="stHorizontalBlock"] > div {
        width: 100% !important; min-width: 0 !important; flex: 1 1 100% !important;
    }
    .stButton, .stButton > button { width: 100% !important; }
    .stButton > button {
        min-height: 3.45rem !important; height: auto !important;
        padding: .65rem .85rem !important; font-size: 1rem !important;
    }
    /* Salva sessione resta sempre raggiungibile dentro la sidebar quando
       l'utente scorre un progetto lungo da telefono. */
    section[data-testid="stSidebar"] [data-testid="stButton"]:has(button[kind="primary"]) {
        position: sticky !important; bottom: .55rem !important; z-index: 50 !important;
        padding: .4rem 0 !important; background: #0f1d31 !important;
    }

    /* Le tab non vanno a capo né vengono tagliate: scorrono orizzontalmente. */
    [data-testid="stTabs"] [data-baseweb="tab-list"] {
        overflow-x: auto !important; overflow-y: hidden !important; flex-wrap: nowrap !important;
        gap: .3rem !important; -webkit-overflow-scrolling: touch !important;
        scrollbar-width: thin !important;
    }
    [data-testid="stTabs"] button[role="tab"] {
        flex: 0 0 auto !important; white-space: nowrap !important;
        padding: .6rem .72rem !important; font-size: .84rem !important;
    }

    /* Editor, controlli e anteprima mantengono una dimensione leggibile e
       l'anteprima scorre senza creare una pagina eccessivamente lunga. */
    .preview-box {
        padding: 1.15rem .95rem !important; height: auto !important;
        max-height: 68vh !important; min-height: 50vh !important;
        overflow-y: auto !important; font-size: 1rem !important; line-height: 1.78 !important;
    }
    [data-testid="stTextArea"] textarea { font-size: 1rem !important; line-height: 1.55 !important; }
    .ss-workspace-header { padding: .85rem .9rem !important; gap: .7rem !important; }
    .ss-workspace-title { font-size: 1.05rem !important; white-space: normal !important; }
    .ss-workspace-meta { gap: .35rem !important; }
    .ss-workspace-chip { font-size: .7rem !important; padding: .36rem .5rem !important; }
    .ss-project-panel { margin:.1rem 0 .9rem !important; padding:.86rem .75rem !important; border-radius:13px !important; }
    .ss-project-title { font-size:1.22rem !important; margin-bottom:.72rem !important; }
    .ss-project-journey { grid-template-columns:1fr !important; gap:.38rem !important; margin-bottom:.75rem !important; }
    .ss-project-step { padding:.58rem .65rem !important; font-size:.83rem !important; }
    .ss-project-metrics { grid-template-columns:1fr 1fr !important; gap:.45rem !important; margin-bottom:.72rem !important; }
    .ss-project-metric { padding:.66rem .7rem !important; }.ss-project-metric b { font-size:1.28rem !important; }
    .ss-next-action { padding:.75rem .78rem !important; gap:.58rem !important; }.ss-next-action strong { font-size:.93rem !important; }
    .ss-section-card { padding: .35rem .7rem .75rem !important; border-radius: 11px !important; }
}
</style>
""", unsafe_allow_html=True)

# Tema chiaro opzionale: il tema scuro predefinito resta esattamente quello
# storico. Questa sovrascrittura è soltanto estetica e non modifica widget,
# crediti, dati della sidebar o funzioni dell'editor.
if st.session_state.get("commercial_ui_theme", "Scuro") == "Chiaro":
    st.markdown("""
    <style>
    /* Tema chiaro azzurrino: sostituisce integralmente solo i colori della UI. */
    .stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"], section.main {
      background:radial-gradient(circle at 7% 4%, #d8f0ff 0%, transparent 31%),radial-gradient(circle at 96% 17%, #c8e8ff 0%, transparent 27%),linear-gradient(135deg,#eaf7ff 0%,#f6fbff 48%,#dfefff 100%) !important;
      color:#142b43 !important;
    }
    [data-testid="stAppViewBlockContainer"], section.main > div.block-container { background:transparent !important; }
    section[data-testid="stSidebar"] { background:linear-gradient(180deg,#e9f6ff 0%,#d6ecff 100%) !important; border-right:1px solid #9dcaeb !important; }
    section[data-testid="stSidebar"] .stMarkdown, section[data-testid="stSidebar"] label, section[data-testid="stSidebar"] p, section[data-testid="stSidebar"] h1, section[data-testid="stSidebar"] h2, section[data-testid="stSidebar"] h3,
    [data-testid="stMain"] .stMarkdown, [data-testid="stMain"] label, [data-testid="stMain"] p, [data-testid="stMain"] h1, [data-testid="stMain"] h2, [data-testid="stMain"] h3 { color:#142b43 !important; }
    section[data-testid="stSidebar"] .stTextInput input, section[data-testid="stSidebar"] .stTextArea textarea,
    [data-testid="stMain"] .stTextInput input, [data-testid="stMain"] .stTextArea textarea,
    section[data-testid="stSidebar"] div[data-baseweb="select"] > div, div[data-baseweb="select"] > div { background:#ffffff !important; color:#142b43 !important; border-color:#9fc6e4 !important; }
    [data-testid="stMain"] [data-testid="stAlert"], section[data-testid="stSidebar"] [data-testid="stAlert"],
    [data-testid="stExpander"], [data-testid="stDataFrame"] { background:rgba(255,255,255,.84) !important; border-color:#b8d7ed !important; color:#142b43 !important; }
    .stButton>button { background:#ffffff !important; color:#16466f !important; border-color:#8dbde1 !important; box-shadow:0 3px 9px rgba(38,114,166,.10) !important; }
    .stButton>button:hover { background:#dcefff !important; color:#0b365d !important; border-color:#1976e9 !important; }
    .stButton>button[kind="primary"], .stButton>button[data-testid="baseButton-primary"] { background:linear-gradient(135deg,#087cc7,#28a8ed) !important; color:#ffffff !important; border-color:#087cc7 !important; }
    .st-key-avvia_chat_sidebar_guidata .stButton>button, .st-key-avvia_chat_sidebar_guidata button { background:linear-gradient(135deg,#15803d,#22a75a) !important; color:#ffffff !important; border-color:#52d486 !important; }
    .st-key-avvia_chat_sidebar_guidata .stButton>button:hover, .st-key-avvia_chat_sidebar_guidata button:hover { background:linear-gradient(135deg,#126b33,#178747) !important; border-color:#86efac !important; }
    [data-testid="stTabs"] [data-baseweb="tab-list"] { border-bottom-color:#a8cce7 !important; }
    [data-testid="stTabs"] button[role="tab"] { color:#2a587e !important; }
    [data-testid="stTabs"] button[aria-selected="true"] { color:#0b3f6b !important; background:#cfeaff !important; }
    .ss-workspace-header { border-color:#9fcbe8 !important; background:linear-gradient(135deg,#ffffff,#dff2ff) !important; box-shadow:0 12px 30px rgba(38,114,166,.16) !important; }
    .ss-workspace-title, .ss-section-card h2, .ss-section-card h3 { color:#142b43 !important; }
    .ss-workspace-subtitle { color:#476b88 !important; }
    .ss-workspace-chip { color:#174f7a !important; background:#eff9ff !important; border-color:#a6cce7 !important; }
    .ss-section-card { background:rgba(255,255,255,.88) !important; border-color:#b5d6ed !important; box-shadow:0 8px 20px rgba(38,114,166,.08) !important; }
    </style>
    """, unsafe_allow_html=True)

# ======================================================================================================================
# 4. GESTIONE EXPORT PDF (CHIRURGIA: FIX TITOLI LUNGHI E MARGINI)
# ======================================================================================================================
class EbookPDF(FPDF):
    def __init__(self, titolo, autore):
        super().__init__()
        self.titolo = self._clean(titolo)
        self.autore = self._clean(autore)
        
        # --- FIX MARGINI: Imposta margini espliciti e interruzione pagina automatica ---
        # Imposta margine sinistro, superiore e destro a 15 mm
        self.set_margins(15, 15, 15)
        # Forza il salto pagina automatico quando si arriva a 15 mm dal fondo
        self.set_auto_page_break(auto=True, margin=15)
        
    def _clean(self, txt):
        """Sanitizzazione forzata per FPDF latin-1. Evita crash da smart quotes e unicode."""
        if not txt: return ""
        replacements = {'“': '"', '”': '"', '‘': "'", '’': "'", '—': '-', '–': '-', '…': '...'}
        for k, v in replacements.items(): 
            txt = txt.replace(k, v)
        return txt.encode('latin-1', 'replace').decode('latin-1')

    def header(self):
        if self.page_no() > 1:
            self.set_font('Arial', 'I', 9); self.set_text_color(150)
            self.cell(0, 10, f"{self.titolo} - {self.autore}", 0, 0, 'R'); self.ln(15)
            
    def footer(self):
        if self.page_no() <= 1:
            return
        self.set_y(-20); self.set_font('Arial', 'I', 9)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
        
    def cover_page(self):
        self.add_page(); self.set_font('Arial', 'B', 32); self.ln(100)
        self.multi_cell(0, 15, self.titolo.upper(), 0, 'C'); self.ln(20)
        self.set_font('Arial', 'I', 20); self.cell(0, 10, f"di {self.autore}", 0, 1, 'C')
        
    def add_content(self, title, content, image_bytes=None, image_caption=None):
        self.add_page(); self.ln(15); self.set_font('Arial', 'B', 22)
        # FIX: Sostituito cell() con multi_cell() per il titolo, per mandare a capo i titoli lunghi!
        self.multi_cell(0, 15, self._clean(title).upper(), 0, 'L'); self.ln(10); self.set_font('Arial', '', 12)
        # multi_cell con w=0 ora calcola la larghezza rispettando il margine destro (15mm)
        if image_bytes:
            session_dir = st.session_state.get("tmp_dir", os.path.abspath("tmp"))
            image_path = os.path.join(session_dir, f"ebook_creator_image_{uuid.uuid4().hex}.png")
            os.makedirs(os.path.dirname(image_path), exist_ok=True)
            with open(image_path, "wb") as f:
                f.write(image_bytes)
            try:
                # Immagine compatta: lascia spazio alla prosa e mantiene proporzioni corrette.
                self.image(image_path, x=48, w=115)
                self.ln(4)
                if image_caption:
                    self.set_font('Arial', 'I', 9)
                    self.multi_cell(0, 6, self._clean(image_caption), 0, 'C')
                    self.ln(5)
            finally:
                try: os.remove(image_path)
                except OSError: pass
        self.set_font('Arial', '', 12)
        self.multi_cell(0, 10, self._clean(content))

# ======================================================================================================================
# 5. CORE LOGIC DI STESURA E ANALISI QUALITÀ (POTENZIATA) E DECISIONE NEURALE
# ======================================================================================================================
def _client_e_modello_testuale(modello_richiesto=None):
    """Restituisce uno solo dei due cervelli, senza fallback nascosti."""
    if usa_deepseek_pro():
        if not client_deepseek:
            raise RuntimeError("DeepSeek Pro non è configurato. Aggiungi DEEPSEEK_API_KEY nei Secrets di Streamlit oppure seleziona GPT-5.4.")
        return client_deepseek, MODELLO_DEEPSEEK_PRO
    if not client_openai:
        raise RuntimeError("GPT non è configurato. Aggiungi OPENAI_API_KEY nei Secrets di Streamlit oppure seleziona DeepSeek Pro.")
    return client_openai, (modello_richiesto or MODELLO_STESURA)


def registra_esito_chiamata_ai(risposta, *, riferimento=None, reason="generazione_testo", amount=0,
                               model=None, riuscita=True, rimborsata=False, errore=""):
    """Invia al registro amministrativo solo i dati tecnici della chiamata.

    Prompt e testo generato restano esclusivamente nella sessione/progetto
    dell'utente: nel database dei costi finiscono soltanto token, modello,
    operazione, crediti e risultato tecnico.
    """
    try:
        registra_utilizzo_ai(
            reference=riferimento,
            operation=reason,
            model=model or (MODELLO_DEEPSEEK_PRO if usa_deepseek_pro() else MODELLO_STESURA),
            usage=getattr(risposta, "usage", None) if risposta is not None else None,
            credits_requested=amount,
            success=riuscita,
            refunded=rimborsata,
            error_code=str(errore or "")[:160],
        )
    except Exception:
        # Il monitoraggio non deve interferire con l'esito editoriale.
        pass


def chiedi_gpt(prompt, system_prompt, *, addebita=True, amount=AI_REQUEST_CREDITS, max_completion_tokens=None,
               model=None, reason="generazione_testo", timeout_seconds=None):
    """Invia una richiesta testuale con timeout esplicito per i flussi editoriali.

    La progettazione di indice usa il modello completo e può includere audit e
    correzione. Senza un limite il client SDK può attendere dieci minuti e il
    pulsante sembra bloccato. Il timeout ferma la sola richiesta in corso e
    lascia intatti progetto, crediti e contenuti già salvati.
    """
    riferimento = None
    try:
        if addebita:
            riferimento = charge_credits(reason, amount=amount)
        client_testuale, modello_effettivo = _client_e_modello_testuale(model or MODELLO_STESURA)
        richiesta = {
            "model": modello_effettivo,
            "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],
        }
        if max_completion_tokens:
            richiesta["max_tokens" if usa_deepseek_pro() else "max_completion_tokens"] = max_completion_tokens
        # Per la stesura ordinaria DeepSeek Pro lavora senza ragionamento
        # esteso: testo più rapido, costo più prevedibile. I prompt editoriali
        # complessi attivano invece la modalità high.
        if usa_deepseek_pro() and model in {MODELLO_EDITORIALE, MODELLO_ANALISI_FONTI, MODELLO_CONTROLLO_COPYRIGHT_APPROFONDITO}:
            richiesta["extra_body"] = {"thinking": {"type": "enabled", "reasoning_effort": "high"}}
        elif usa_deepseek_pro():
            richiesta["extra_body"] = {"thinking": {"type": "disabled"}}
        if timeout_seconds is None and model in {
            MODELLO_EDITORIALE, MODELLO_ANALISI_FONTI, MODELLO_CONTROLLO_COPYRIGHT_APPROFONDITO,
        }:
            timeout_seconds = TIMEOUT_INDICE_SECONDI
        client_richiesta = client_testuale
        if timeout_seconds:
            client_richiesta = client_testuale.with_options(
                timeout=float(timeout_seconds), max_retries=0
            )
        response = client_richiesta.chat.completions.create(**richiesta)
        testo = response.choices[0].message.content.strip()
        registra_esito_chiamata_ai(
            response,
            riferimento=riferimento,
            reason=reason,
            amount=amount if addebita else 0,
            model=modello_effettivo,
        )
        prefissi = ["ecco", "certamente", "sicuramente", "ok", "here is", "sure"]
        righe = [l for l in testo.split("\n") if not any(l.lower().startswith(p) for p in prefissi)]
        return "\n".join(righe).strip()
    except CommercialCreditError as e:
        # Non inserisce mai un messaggio d'errore nel manoscritto e ferma
        # l'azione prima che una sezione esistente possa essere sovrascritta.
        mostra_crediti_esauriti()
        st.stop()
    except Exception as e:
        if riferimento:
            refund_credits(riferimento, amount=amount)
        registra_esito_chiamata_ai(
            None,
            riferimento=riferimento,
            reason=reason,
            amount=amount if addebita else 0,
            model=model or (MODELLO_DEEPSEEK_PRO if usa_deepseek_pro() else MODELLO_STESURA),
            riuscita=False,
            rimborsata=bool(riferimento),
            errore=type(e).__name__,
        )
        return f"ERRORE: {str(e)}"


def stima_crediti_per_cervello(azione_id, stima_gpt):
    """Mostra un preventivo coerente con il cervello selezionato."""
    if not usa_deepseek_pro():
        return stima_gpt
    azione = str(azione_id).casefold()
    # Per i comandi che redigono più sezioni il valore ricevuto è del tipo
    # "fino a 71". Il preventivo DeepSeek deve ridurre il totale, non limitarsi
    # a descrivere il costo della singola sezione.
    numeri_stima = re.findall(r"\d+", str(stima_gpt))
    totale_gpt = int(numeri_stima[0]) if numeri_stima else 0
    if "rigenera_indice" in azione:
        return "circa 2"
    if "genera_indice" in azione:
        return "circa 3⅓"
    if "fonti" in azione:
        return "circa 1⅓"
    if any(parola in azione for parola in ("scrivi_tutto", "scrivi_sottocapitoli", "rielabora_sezioni_originalita")):
        return f"fino a {max(1, math.ceil(totale_gpt / 3))}"
    if any(parola in azione for parola in ("scrivi_sezione", "rigenera_sezione", "rielabora", "quiz", "esempi")):
        return "1 ogni 3 operazioni"
    if "ricette" in azione:
        return "circa 3⅓"
    if "coerenza" in azione:
        return "circa 3⅓ (primo controllo); poi 1 credito ogni 3 blocchi"
    if "voto_indice" in azione:
        return "1 credito"
    if any(parola in azione for parola in ("report_sintattico", "metadati", "controlla_fatti")):
        return "1 ogni 3 controlli equivalenti"
    if "copyright" in azione or "immagine" in azione:
        return "non disponibile con DeepSeek Pro"
    return stima_gpt


def applica_sostituzioni_tariffario(testo, sostituzioni):
    """Aggiorna un valore numerico già tradotto senza alterare il resto della UI."""
    risultato = str(testo or "")
    for precedente, aggiornato in sostituzioni:
        risultato = risultato.replace(precedente, aggiornato)
    return risultato


def tutela_azione_preventivo(azione_id):
    """Spiega con precisione cosa un'azione a credito non può alterare."""
    azione = str(azione_id).casefold()
    if "scrivi_tutto" in azione or "scrivi_sottocapitoli" in azione:
        return "Le sezioni già presenti restano invariate; vengono elaborate solo quelle ancora vuote."
    if "rigenera_sezione" in azione or "scrivi_sezione" in azione:
        return "Interviene solo sulla sezione selezionata; le altre sezioni restano invariate."
    if "rielabora_sezioni_originalita" in azione:
        return "Interviene solo sulle sezioni segnalate dal controllo; il resto del manoscritto resta invariato."
    if "indice" in azione:
        return "Non modifica le sezioni già scritte, le immagini o le fonti caricate dall'utente."
    if "fonti" in azione:
        return "Aggiorna solo il registro delle fonti; indice, testi e immagini restano invariati."
    if "copyright" in azione or "controlla_fatti" in azione or "coerenza" in azione:
        return "È un controllo: non modifica automaticamente il manoscritto."
    if "immagine" in azione:
        return "Genera o aggiorna solo l'immagine richiesta; il testo del manoscritto resta invariato."
    return "Non pubblica nulla e non modifica altre parti del progetto senza una tua conferma successiva."


def mostra_avviso_operativo(livello, sezione, dettaglio, azione_consigliata, *, contenuti_sicuri=True):
    """Avviso leggibile che distingue problema, dati al sicuro e recupero.

    È una presentazione UI: non modifica code, crediti, salvataggi o testi.
    """
    icona = {"errore": "🔴", "attenzione": "🟡", "ok": "🟢"}.get(livello, "ℹ️")
    messaggio = f"{icona} **{sezione}**\n\n**Cosa è successo:** {dettaglio}"
    if contenuti_sicuri:
        messaggio += "\n\n**Cosa resta al sicuro:** le sezioni già completate non vengono cancellate."
    messaggio += f"\n\n**Cosa fare ora:** {azione_consigliata}"
    if livello == "errore":
        st.error(messaggio)
    elif livello == "attenzione":
        st.warning(messaggio)
    else:
        st.success(messaggio)


def pulsante_con_preventivo(azione_id, etichetta, stima_crediti, descrizione, *,
                            use_container_width=False, disabled=False, tipo=None):
    """Mostra una conferma preventiva prima delle sole azioni che consumano crediti."""
    chiave_confermata = st.session_state.get("azione_crediti_confermata")
    if chiave_confermata == azione_id:
        st.session_state.pop("azione_crediti_confermata", None)
        return True

    parametri = {
        "label": etichetta,
        "key": f"preventivo_avvio_{azione_id}",
        "use_container_width": use_container_width,
        "disabled": disabled,
    }
    if tipo:
        parametri["type"] = tipo
    if st.button(**parametri):
        st.session_state["preventivo_crediti_attesa"] = {
            "azione_id": azione_id,
            "stima": str(stima_crediti_per_cervello(azione_id, stima_crediti)),
            "descrizione": descrizione,
        }

    preventivo = st.session_state.get("preventivo_crediti_attesa", {})
    if preventivo.get("azione_id") != azione_id:
        return False

    motore_preventivo = "DeepSeek V4 Pro" if usa_deepseek_pro() else "GPT-5.4 (OpenAI)"
    st.warning(f"**Preventivo prima dell'avvio — {preventivo['stima']} crediti al massimo**")
    col_cosa_fa, col_tutela = st.columns(2)
    with col_cosa_fa:
        st.caption("**Cosa verrà fatto**")
        st.write(preventivo["descrizione"])
    with col_tutela:
        st.caption("**Cosa resterà invariato**")
        st.write(tutela_azione_preventivo(azione_id))
    st.caption(
        f"**Cervello selezionato:** {motore_preventivo}. "
        "Se una richiesta non restituisce un risultato utilizzabile, il costo della fase non completata non viene trattenuto oppure viene riaccreditato."
    )
    col_conferma, col_annulla = st.columns(2)
    with col_conferma:
        if st.button("Conferma e avvia", key=f"preventivo_conferma_{azione_id}",
                     type="primary", use_container_width=True):
            st.session_state["azione_crediti_confermata"] = azione_id
            st.session_state.pop("preventivo_crediti_attesa", None)
            st.rerun()
    with col_annulla:
        if st.button("Annulla", key=f"preventivo_annulla_{azione_id}", use_container_width=True):
            st.session_state.pop("preventivo_crediti_attesa", None)
            st.rerun()
    return False


def addebita_azione_diretta(reason, amount):
    """Addebito per azioni locali o raggruppate che non passano da chiedi_gpt."""
    try:
        return charge_credits(reason, amount=amount)
    except CommercialCreditError:
        mostra_crediti_esauriti()
        st.stop()

def verifica_e_correggi_fatti_online(testo, sezione, lingua):
    """Verifica soltanto i fatti aggiornabili che meritano una ricerca online."""
    if usa_deepseek_pro():
        st.info("Il controllo fatti online richiede GPT perché utilizza ricerca web. Con DeepSeek Pro il testo resta invariato.")
        return pulisci_testo_editoriale(testo)
    riferimento = None
    try:
        riferimento = charge_credits("verifica_fatti", amount=CREDIT_COSTS["verifica_fatti_web"])
        risposta = client.responses.create(
            model=MODELLO_STESURA,
            tools=[{"type": "web_search_preview"}],
            input=(
                f"Verifica il testo seguente in lingua {lingua} relativo alla sezione '{sezione}'. "
                "Cerca online fonti autorevoli e aggiornate per ogni fatto verificabile, soprattutto "
                "leggi, normative, prezzi, licenze, specifiche, date, software e dati numerici. "
                "Correggi soltanto le affermazioni non aggiornate o non supportate; non inventare dati. "
                "Mantieni struttura e stile, ma NON inserire nel testo URL, link Markdown, citazioni, "
                "note bibliografiche, nomi di fonti o una sezione 'Fonti verificate'. Le fonti servono "
                "esclusivamente per il controllo interno e non devono comparire nell'opera destinata al lettore. "
                "Distingui i fatti verificati dagli esempi ipotetici senza apporre etichette tecniche o note di fonte. "
                "Restituisci solo il testo editoriale revisionato e pulito.\n\n"
                f"TESTO:\n{testo}"
            )
        )
        registra_esito_chiamata_ai(
            risposta,
            riferimento=riferimento,
            reason="verifica_fatti",
            amount=CREDIT_COSTS["verifica_fatti_web"],
            model=MODELLO_STESURA,
        )
        return pulisci_testo_editoriale(getattr(risposta, "output_text", None) or testo)
    except Exception as e:
        if riferimento:
            refund_credits(riferimento, amount=CREDIT_COSTS["verifica_fatti_web"])
        registra_esito_chiamata_ai(
            None,
            riferimento=riferimento,
            reason="verifica_fatti",
            amount=CREDIT_COSTS["verifica_fatti_web"],
            model=MODELLO_STESURA,
            riuscita=False,
            rimborsata=bool(riferimento),
            errore=type(e).__name__,
        )
        st.warning(f"Verifica online non disponibile: {e}")
        return pulisci_testo_editoriale(testo)


def richiede_verifica_fatti(testo, sezione=""):
    """Evita ricerche inutili per scene, esercizi e spiegazioni stabili."""
    campione = f"{sezione}\n{testo}".lower()
    indicatori = (
        "legge", "normativa", "regolamento", "decreto", "licenza", "prezzo", "tariffa",
        "syllabus", "soglia di superamento", "punteggio minimo", "durata dell'esame",
        "versione", "requisiti di sistema", "compatibilità", "aggiornamento software",
        "aggiornata al", "in vigore", "20/", "€", "$"
    )
    return any(indicatore in campione for indicatore in indicatori)


def audit_fatti_capitolo(capitolo, contenuti, lingua):
    """Esegue un solo controllo online sul capitolo completo, senza riscrivere le singole sezioni."""
    if usa_deepseek_pro():
        return "Controllo fatti online disponibile con il cervello GPT. DeepSeek Pro rimane separato e non usa la ricerca web GPT."
    testo = "\n\n".join(f"SEZIONE: {nome}\n{contenuto}" for nome, contenuto in contenuti if contenuto.strip())
    if not testo or not richiede_verifica_fatti(testo, capitolo):
        return "Controllo fatti del capitolo non necessario: nessun dato variabile rilevato."
    riferimento = None
    try:
        riferimento = charge_credits("audit_fatti", amount=CREDIT_COSTS["audit_fatti_capitolo"])
        risposta = client.responses.create(
            model=MODELLO_STESURA,
            tools=[{"type": "web_search_preview"}],
            input=(
                f"Controlla i soli fatti aggiornabili nel capitolo '{capitolo}', in lingua {lingua}. "
                "Verifica esclusivamente regole, norme, date, soglie, prezzi, versioni software, licenze e specifiche. "
                "Non riscrivere il capitolo e non citare fonti o URL. Restituisci soltanto: "
                "ESITO: nessuna correzione necessaria oppure una lista di correzioni puntuali con sezione e formulazione da aggiornare.\n\n"
                f"CAPITOLO:\n{testo}"
            )
        )
        registra_esito_chiamata_ai(
            risposta,
            riferimento=riferimento,
            reason="audit_fatti",
            amount=CREDIT_COSTS["audit_fatti_capitolo"],
            model=MODELLO_STESURA,
        )
        return pulisci_testo_editoriale(getattr(risposta, "output_text", "") or "Controllo non disponibile.")
    except Exception as e:
        if riferimento:
            refund_credits(riferimento, amount=CREDIT_COSTS["audit_fatti_capitolo"])
        registra_esito_chiamata_ai(
            None,
            riferimento=riferimento,
            reason="audit_fatti",
            amount=CREDIT_COSTS["audit_fatti_capitolo"],
            model=MODELLO_STESURA,
            riuscita=False,
            rimborsata=bool(riferimento),
            errore=type(e).__name__,
        )
        return f"Controllo fatti del capitolo non disponibile: {e}"

def pulisci_testo_editoriale(testo):
    """Rimuove fonti tecniche dal testo destinato ad anteprima ed esportazione."""
    if not testo:
        return ""
    testo = str(testo)
    # Rimuove Markdown e segni di formattazione tecnica: l'editor impagina il testo in modo nativo.
    testo = re.sub(r"(?m)^\s{0,3}#{1,6}\s+", "", testo)
    testo = testo.replace("**", "").replace("__", "")
    testo = re.sub(r"(?m)^\s*>\s?", "", testo)
    testo = re.sub(r"(?is)(?:^|\n)\s{0,3}(?:#+\s*)?(?:fonti verificate|fonti consultate|riferimenti bibliografici|sources|references)\s*:?.*$", "", testo)
    testo = re.sub(r"\[([^\]]+)\]\(https?://[^)]+\)", r"\1", testo)
    testo = re.sub(r"https?://[^\s)\]>]+", "", testo)
    # Elimina le attribuzioni residue prodotte dai modelli, incluse fonti senza URL
    # completo come "(esempio.net)" o domini nazionali. Le parentesi tecniche normali
    # restano invece intatte.
    testo = re.sub(
        r"\s*\([^\n()]{0,180}(?:\b[a-z0-9-]+\.)+(?:com|org|net|gov|edu|io|co\.uk|it|fr|de|es|ai|info|biz|co)[^\n()]*\)",
        "", testo, flags=re.I
    )
    testo = re.sub(r"(?im)^\s*\[?(?:informazione|fatto|esempio|fonte)[^\n]{0,120}(?:da verificare|verificato|ipotetico|di carattere generale)[^\n]*\]?\s*$", "", testo)
    testo = re.sub(r"(?m)^\s*[-_*]{3,}\s*$", "", testo)
    testo = re.sub(r"\n{3,}", "\n\n", testo)
    return testo.strip()


def dividi_blocchi_lettura(testo, limite=480):
    """Divide il testo con la stessa logica del lettore browser.

    I blocchi ricevono poi un riferimento nell'anteprima, così l'evidenziazione
    resta sincronizzata con la frase effettivamente pronunciata.
    """
    normalizzato = re.sub(r"\s+", " ", str(testo or "")).strip()
    if not normalizzato:
        return []
    frasi = re.findall(r"[^.!?…]+[.!?…]+|[^.!?…]+$", normalizzato) or [normalizzato]
    blocchi, corrente = [], ""
    for frase in frasi:
        candidata = (corrente + " " + frase).strip()
        if len(candidata) > limite and corrente:
            blocchi.append(corrente)
            corrente = frase.strip()
        else:
            corrente = candidata
    if corrente:
        blocchi.append(corrente)
    return blocchi


def mostra_lettore_vocale_gratuito(testo_libro, lingua, sezioni=None):
    """Legge nel browser il manoscritto completo senza API, crediti o file audio."""
    testo = pulisci_testo_editoriale(testo_libro or "")
    if not testo:
        st.info("Il lettore vocale sarà disponibile dopo la generazione di almeno una sezione.")
        return

    parti = []
    for parte in sezioni or []:
        contenuto = pulisci_testo_editoriale(parte.get("testo", ""))
        if contenuto:
            parti.append({
                "titolo": str(parte.get("titolo", "")).strip() or "Libro",
                "testo": contenuto,
                "anchor_prefix": str(parte.get("anchor_prefix", "")).strip(),
            })
    if not parti:
        parti = [{"titolo": "Libro", "testo": testo}]

    etichette = {
        "Italiano": ("Lettore vocale gratuito", "Ascolta il libro", "Pausa", "Riprendi", "Ferma", "Velocità", "Voce automatica", "Pronto a leggere il manoscritto.", "Lettura in corso", "Lettura terminata.", "Nessun credito o costo API.", "Sezione in lettura"),
        "English": ("Free voice reader", "Listen to the book", "Pause", "Resume", "Stop", "Speed", "Automatic voice", "Ready to read the manuscript.", "Reading in progress", "Reading finished.", "No credits or API cost.", "Reading section"),
        "Español": ("Lector de voz gratuito", "Escuchar el libro", "Pausa", "Reanudar", "Detener", "Velocidad", "Voz automática", "Listo para leer el manuscrito.", "Lectura en curso", "Lectura terminada.", "Sin créditos ni coste de API.", "Sección en lectura"),
        "Français": ("Lecteur vocal gratuit", "Écouter le livre", "Pause", "Reprendre", "Arrêter", "Vitesse", "Voix automatique", "Prêt à lire le manuscrit.", "Lecture en cours", "Lecture terminée.", "Sans crédit ni coût API.", "Section en cours de lecture"),
        "Deutsch": ("Kostenloser Vorleser", "Buch anhören", "Pause", "Fortsetzen", "Stoppen", "Geschwindigkeit", "Automatische Stimme", "Bereit, das Manuskript vorzulesen.", "Wiedergabe läuft", "Wiedergabe beendet.", "Keine Credits oder API-Kosten.", "Gelesener Abschnitt"),
        "Română": ("Cititor vocal gratuit", "Ascultă cartea", "Pauză", "Reia", "Oprește", "Viteză", "Voce automată", "Gata să citească manuscrisul.", "Citire în curs", "Citire încheiată.", "Fără credite sau cost API.", "Secțiunea citită"),
        "Русский": ("Бесплатный голосовой читатель", "Слушать книгу", "Пауза", "Продолжить", "Остановить", "Скорость", "Автоматический голос", "Готов к чтению рукописи.", "Чтение выполняется", "Чтение завершено.", "Без кредитов и затрат API.", "Текущий раздел"),
        "العربية": ("قارئ صوتي مجاني", "استمع إلى الكتاب", "إيقاف مؤقت", "متابعة", "إيقاف", "السرعة", "صوت تلقائي", "جاهز لقراءة المخطوطة.", "القراءة جارية", "انتهت القراءة.", "بلا أرصدة أو تكلفة API.", "القسم الجاري قراءته"),
        "中文": ("免费语音朗读器", "朗读全书", "暂停", "继续", "停止", "语速", "自动选择声音", "准备朗读书稿。", "正在朗读", "朗读完成。", "不消耗积分或 API 费用。", "正在朗读的章节"),
    }
    codici_lingua = {
        "Italiano": "it-IT", "English": "en-US", "Español": "es-ES", "Français": "fr-FR",
        "Deutsch": "de-DE", "Română": "ro-RO", "Русский": "ru-RU", "العربية": "ar-SA", "中文": "zh-CN",
    }
    labels = etichette.get(lingua, etichette["Italiano"])
    testo_json = json.dumps(testo, ensure_ascii=False).replace("</", "<\\/")
    parti_json = json.dumps(parti, ensure_ascii=False).replace("</", "<\\/")
    labels_json = json.dumps(labels, ensure_ascii=False).replace("</", "<\\/")
    lingua_json = json.dumps(codici_lingua.get(lingua, "it-IT"))
    components.html(
        f"""
        <style>
          body{{margin:0;font-family:Arial,sans-serif;background:#fff;color:#102a43}}
          .box{{border:1px solid #cbd5e1;border-radius:12px;padding:16px;background:#f8fafc}}
          h3{{margin:0 0 5px}} p{{margin:0 0 12px;color:#486581;font-size:13px}}
          .row{{display:flex;gap:8px;flex-wrap:wrap;align-items:center}}
          button,select{{border-radius:8px;padding:9px 11px;border:1px solid #cbd5e1;font-weight:700}}
          button{{background:#1689e8;color:#fff;border:0;cursor:pointer}} .stop{{background:#cf3345}}
          #status{{margin-top:10px;font-size:13px;font-weight:700;color:#1269ae}}
          #currentSection{{margin-top:10px;padding:8px 10px;border-radius:7px;background:#dbeafe;color:#0f3f68;font-size:13px;font-weight:700}}
          #currentExcerpt{{margin-top:8px;padding:11px 12px;border-left:5px solid #f59e0b;border-radius:7px;background:#fff7d6;color:#4a3200;font-size:14px;line-height:1.5;font-style:italic}}
          .progress{{height:7px;border-radius:999px;background:#d9e2ec;margin-top:8px;overflow:hidden}}
          #progressBar{{height:100%;width:0;background:#1689e8;transition:width .25s ease}}
        </style>
        <div class="box">
          <h3>🔊 <span id="title"></span></h3><p id="note"></p>
          <div class="row">
            <button id="start"></button><button id="pause"></button><button id="resume"></button><button class="stop" id="stop"></button>
            <label><span id="speedText"></span> <select id="speed"><option value="0.8">0,8×</option><option value="1" selected>1×</option><option value="1.2">1,2×</option><option value="1.4">1,4×</option></select></label>
            <select id="voice"></select>
          </div>
          <div class="row" style="margin-top:10px"><label><span id="startFromText"></span> <select id="startFrom"></select></label></div>
          <div id="status"></div><div id="currentSection"></div><div id="currentExcerpt"></div><div class="progress"><div id="progressBar"></div></div>
        </div>
        <script>
          const bookText = {testo_json}, bookParts = {parti_json}, L = {labels_json}, bookLanguage = {lingua_json};
          // In alcune versioni di Chrome il componente Streamlit vive in un
          // iframe: proviamo prima la sua API e, se disponibile, quella della
          // pagina ospitante. Non usa alcun servizio esterno o credito.
          function speechEngine() {{
            try {{ if (window.speechSynthesis) return window.speechSynthesis; }} catch (_) {{}}
            try {{ if (window.parent && window.parent.speechSynthesis) return window.parent.speechSynthesis; }} catch (_) {{}}
            return null;
          }}
          function utteranceConstructor() {{
            try {{ if (window.SpeechSynthesisUtterance) return window.SpeechSynthesisUtterance; }} catch (_) {{}}
            try {{ if (window.parent && window.parent.SpeechSynthesisUtterance) return window.parent.SpeechSynthesisUtterance; }} catch (_) {{}}
            return null;
          }}
          const synth = speechEngine(), Utterance = utteranceConstructor();
          let chunks = [], position = 0, active = false, paused = false, voices = [], utteranceId = 0;
          let currentUtterance = null, keepAliveTimer = null;
          const el = (id) => document.getElementById(id);
          el("title").textContent=L[0]; el("note").textContent=L[10]; el("start").textContent="▶ "+L[1];
          el("pause").textContent="⏸ "+L[2]; el("resume").textContent="▶ "+L[3]; el("stop").textContent="■ "+L[4];
          el("speedText").textContent=L[5]; el("status").textContent=L[7];
          const startLabels={{"it-IT":"Inizia da", "en-US":"Start from", "es-ES":"Empezar desde", "fr-FR":"Commencer à partir de", "de-DE":"Start ab", "ro-RO":"Începe de la", "ru-RU":"Начать с", "ar-SA":"ابدأ من", "zh-CN":"从这里开始"}};
          el("startFromText").textContent=startLabels[bookLanguage]||"Start from";
          function clearKeepAlive() {{
            if(keepAliveTimer) {{ clearInterval(keepAliveTimer); keepAliveTimer=null; }}
          }}
          function startKeepAlive() {{
            clearKeepAlive();
            // Chrome può interrompere in silenzio le letture molto lunghe.
            // Un resume periodico mantiene viva la lettura senza cambiare testo.
            keepAliveTimer=setInterval(()=>{{ if(active&&!paused&&synth) synth.resume(); }}, 10000);
          }}
          function clearPreviewHighlight() {{
            try {{
              const root=window.parent.document;
              root.querySelectorAll(".voice-preview-active").forEach((node)=>{{
                node.classList.remove("voice-preview-active");
                node.style.background=""; node.style.boxShadow=""; node.style.borderRadius=""; node.style.padding="";
              }});
            }} catch (_) {{}}
          }}
          function highlightPreview(anchor) {{
            clearPreviewHighlight();
            if(!anchor) return;
            try {{
              const node=window.parent.document.getElementById(anchor);
              if(node) {{
                node.classList.add("voice-preview-active");
                node.style.background="#fff3b0"; node.style.boxShadow="0 0 0 3px #f59e0b";
                node.style.borderRadius="4px"; node.style.padding="2px 3px";
              }}
            }} catch (_) {{}}
          }}
          function stopReading() {{
            active=false; paused=false; position=0; utteranceId++; currentUtterance=null;
            clearKeepAlive(); if(synth) synth.cancel();
            clearPreviewHighlight();
            el("status").textContent=L[7]; el("currentSection").textContent=""; el("currentExcerpt").textContent=""; el("progressBar").style.width="0";
          }}
          function split(value, sectionTitle, anchorPrefix="") {{
            const sentences=value.replace(/\\s+/g," ").match(/[^.!?…]+[.!?…]+|[^.!?…]+$/g)||[value], result=[]; let current="";
            // Blocchi brevi: Chrome può interrompere gli utterance troppo lunghi.
            sentences.forEach((sentence)=>{{if((current+" "+sentence).length>480&&current){{result.push({{text:current,section:sectionTitle,anchor:anchorPrefix?anchorPrefix+"_"+result.length:""}});current=sentence.trim();}}else{{current=(current+" "+sentence).trim();}}}});
            if(current)result.push({{text:current,section:sectionTitle,anchor:anchorPrefix?anchorPrefix+"_"+result.length:""}}); return result;
          }}
          function loadVoices() {{
            if(!synth) return;
            voices=synth.getVoices(); const select=el("voice"), old=select.value; select.innerHTML="";
            const automatic=document.createElement("option");automatic.value="";automatic.textContent=L[6];select.appendChild(automatic);
            voices.forEach((voice,index)=>{{const option=document.createElement("option");option.value=index;option.textContent=voice.name+" ("+voice.lang+")";select.appendChild(option);}});
            select.value=old;
          }}
          function loadStartPoints() {{
            const select=el("startFrom"); select.innerHTML="";
            bookParts.forEach((part,index)=>{{
              const option=document.createElement("option"); option.value=String(index);
              option.textContent=part.titolo||("Sezione "+(index+1)); select.appendChild(option);
            }});
          }}
          function next() {{
            if(!active||paused)return;
            if(position>=chunks.length){{active=false;currentUtterance=null;clearKeepAlive();clearPreviewHighlight();el("status").textContent=L[9];return;}}
            const id=++utteranceId;
            const chunk=chunks[position];
            if(!Utterance) {{ active=false; el("status").textContent="Lettore vocale non disponibile in questo browser."; return; }}
            const utterance=new Utterance(chunk.text); currentUtterance=utterance;
            utterance.lang=bookLanguage; utterance.rate=Number(el("speed").value);
            const chosen=el("voice").value; const voice=chosen!==""?voices[Number(chosen)]:voices.find(v=>v.lang.toLowerCase().startsWith(bookLanguage.slice(0,2).toLowerCase()));
            if(voice)utterance.voice=voice;
            utterance.onend=()=>{{
              if(id!==utteranceId||paused||!active)return;
              currentUtterance=null; position+=1; next();
            }};
            utterance.onerror=(event)=>{{
              // L'evento "interrupted" viene emesso da alcuni browser durante
              // una pausa o dopo un cancel volontario: non deve azzerare il libro.
              if(id!==utteranceId||paused||!active||event.error==="interrupted"||event.error==="canceled") return;
              active=false; currentUtterance=null; clearKeepAlive();
              el("status").textContent="Il browser ha interrotto la lettura ("+(event.error||"errore sconosciuto")+"). Premi Ascolta il libro per riprovare.";
            }};
            el("status").textContent=L[8]+" ("+(position+1)+"/"+chunks.length+")";
            el("currentSection").textContent=L[11]+": "+chunk.section;
            el("currentExcerpt").textContent="▶ "+chunk.text;
            highlightPreview(chunk.anchor);
            el("progressBar").style.width=Math.round(((position+1)/chunks.length)*100)+"%";
            try {{
              synth.speak(utterance);
              startKeepAlive();
            }} catch(error) {{
              active=false; currentUtterance=null; clearKeepAlive();
              el("status").textContent="Impossibile avviare il lettore vocale: "+(error.message||error);
            }}
          }}
          el("start").onclick=()=>{{
            if(!synth||!Utterance){{el("status").textContent="Lettore vocale non disponibile: apri l'app con Chrome, Edge o Safari aggiornato.";return;}}
            try {{
              utteranceId++; currentUtterance=null; clearKeepAlive(); synth.cancel(); synth.resume();
              // Le sezioni inviate da Python usano la chiave italiana
              // "testo"; supportiamo anche "text" per vecchi progetti.
              const startAt=Math.max(0, Number(el("startFrom").value||0));
              chunks=bookParts.slice(startAt).flatMap((part)=>split(part.testo||part.text||"",part.titolo||"Libro",part.anchor_prefix||""));
              if(!chunks.length)chunks=split(bookText,"Libro","");
              position=0;active=true;paused=false;
              // speak deve avvenire nello stesso click dell'utente: alcuni
              // browser bloccano il lettore se lo rimandiamo con setTimeout.
              next();
            }} catch(error) {{
              active=false;
              el("status").textContent="Impossibile preparare il lettore vocale: "+(error.message||error);
            }}
          }};
          el("pause").onclick=()=>{{if(active){{paused=true;synth.pause();}}}};
          el("resume").onclick=()=>{{
            if(!active||!paused)return;
            paused=false;
            synth.resume();
            // Alcuni browser perdono la coda dopo una pausa: riparte dallo stesso blocco,
            // senza saltare testo, solo se la ripresa non è effettiva.
            setTimeout(()=>{{
              if(active&&!paused&&!synth.speaking){{utteranceId++;currentUtterance=null;next();}}
            }}, 650);
          }};
          el("stop").onclick=stopReading;
          if(!synth||!Utterance){{
            el("status").textContent="Lettore vocale non disponibile: usa Chrome, Edge o Safari aggiornato.";
            el("start").disabled=el("pause").disabled=el("resume").disabled=true;
          }} else {{
            loadVoices(); loadStartPoints(); if("onvoiceschanged" in synth)synth.onvoiceschanged=loadVoices;
          }}
        </script>
        """,
        # Altezza sufficiente per mostrare sempre il selettore di partenza,
        # anche quando l'estratto in lettura occupa più righe.
        height=330,
        scrolling=False,
    )


def genera_immagine_capitolo(sezione, titolo, genere, trama, contenuto, lingua):
    """GPT-4o-mini prepara il brief; GPT-Image-1 Mini genera il visual economico."""
    if usa_deepseek_pro():
        st.warning("La generazione immagini richiede il cervello GPT. Seleziona GPT-5.4 per creare immagini del capitolo.")
        return None, ""
    contesto_basso = f"{titolo} {trama} {sezione}".lower()
    if "fusion 360" in contesto_basso:
        vincoli_dominio = (
            "Per Fusion 360, se il testo descrive l'interfaccia, rappresenta barra degli strumenti in alto, "
            "browser a sinistra, area di lavoro 3D centrale, timeline in basso e pannello contestuale a destra. "
            "Se descrive una procedura CAD, mostra invece soltanto le fasi e gli oggetti realmente nominati."
        )
    elif any(x in contesto_basso for x in ("ricetta", "cucina", "ricettario")):
        vincoli_dominio = "Per ricette e cucina, mostra ingredienti, utensili e passaggi culinari realmente descritti, senza testo nell'immagine."
    elif any(x in contesto_basso for x in ("romanzo", "thriller", "fantasy", "rosa", "narrativo")):
        vincoli_dominio = "Per narrativa, mostra una scena coerente con luogo, personaggi, atmosfera e azione del brano, senza inserire eventi non presenti."
    elif any(x in contesto_basso for x in ("business", "marketing", "finanza", "economia")):
        vincoli_dominio = "Per business ed economia, mostra relazioni, flussi, strumenti o situazioni operative citate, senza numeri o dati inventati."
    else:
        vincoli_dominio = "Adatta la rappresentazione al dominio del sottocapitolo: mostra esclusivamente oggetti, persone, processi o relazioni realmente descritti nel testo."
    descrizione = chiedi_gpt(
        f"Analizza esclusivamente il sottocapitolo '{sezione}' del libro '{titolo}'. "
        f"Argomento generale: {trama}. Genere: {genere}. Lingua: {lingua}.\n\n"
        "Crea un brief visivo strutturato e concreto con queste voci: "
        "CONCETTO CENTRALE; ELEMENTI OBBLIGATORI (solo quelli realmente descritti); "
        "POSIZIONE E RELAZIONI SPAZIALI; AZIONE O PROCEDURA DA MOSTRARE; "
        "DETTAGLI TECNICI DA RENDERE VISIBILI; ELEMENTI DA ESCLUDERE. "
        "Ogni elemento dell'immagine deve corrispondere a un'informazione del testo. "
        "Non creare una schermata CAD generica e non inventare pannelli, icone o funzioni. "
        "Se il testo descrive un'interfaccia, rappresenta chiaramente le zone nominate "
        "(browser, area di modellazione, barra strumenti, pannello proprietà) nella posizione coerente. "
        "Se descrive una procedura, mostra le fasi in sequenza con forme e frecce non testuali. "
        "VIETATO inserire parole, titoli, paragrafi, numeri, etichette, didascalie, loghi "
        "o schermate con testo nell'immagine. Restituisci solo il brief.\n\n"
        f"{vincoli_dominio}\nContenuto già scritto: {contenuto[-2500:]}",
        "Sei un instructional designer tecnico: produci brief visivi accurati e verificabili."
    )
    riferimento = None
    try:
        riferimento = charge_credits("generazione_immagine", amount=CREDIT_COSTS["immagine_capitolo"])
        risposta = client.images.generate(model="gpt-image-1-mini", prompt=f"Crea un'immagine didattica di alta qualità per il genere '{genere}' e il sottocapitolo '{sezione}'. Segui alla lettera questo brief visivo, senza aggiungere elementi non richiesti:\n{descrizione}\n\n{vincoli_dominio}\nLa scena deve avere corrispondenza uno-a-uno con il testo. Scegli composizione, livello di dettaglio e linguaggio visivo appropriati al dominio e al pubblico: diagramma o tavola tecnica per manuali, scena concreta per procedure, composizione narrativa per narrativa, visualizzazione concettuale per saggistica. Non creare immagini generiche o astratte e non inventare funzioni, dati, persone o oggetti. Nessun testo, lettera, numero, titolo, didascalia o logo nell'immagine. Mantieni sfondo bianco, tratto nero, scala di grigi e stile monocromatico pulito.", size="1024x1024", quality="medium")
        dato = risposta.data[0]
        raw = None
        if getattr(dato, "b64_json", None): raw = base64.b64decode(dato.b64_json)
        elif getattr(dato, "url", None): raw = requests.get(dato.url, timeout=60).content
        if raw:
            # Riduce risoluzione/peso e converte sempre in bianco e nero prima di salvare.
            img = Image.open(BytesIO(raw)).convert("L")
            img.thumbnail((600, 600), Image.Resampling.LANCZOS)
            out = BytesIO(); img.save(out, format="PNG", optimize=True)
            return out.getvalue(), descrizione
        raise ValueError("Risposta immagini priva di dati utilizzabili")
    except Exception as e:
        if riferimento:
            refund_credits(riferimento, reason="generazione_immagine_fallita", amount=CREDIT_COSTS["immagine_capitolo"])
        st.error(f"Errore nella generazione dell'immagine: {e}")
        return None, None

def normalizza_immagine_caricata(file_caricato):
    """Prepara un'immagine caricata dall'utente per anteprima, Word e PDF."""
    try:
        sorgente = Image.open(BytesIO(file_caricato.getvalue()))
        if sorgente.mode in ("RGBA", "LA"):
            sfondo = Image.new("RGB", sorgente.size, "white")
            sfondo.paste(sorgente, mask=sorgente.getchannel("A"))
            sorgente = sfondo
        else:
            sorgente = sorgente.convert("RGB")
        sorgente.thumbnail((1400, 1400), Image.Resampling.LANCZOS)
        output = BytesIO()
        sorgente.save(output, format="PNG", optimize=True)
        return output.getvalue()
    except Exception as e:
        st.error(f"Il file caricato non è un'immagine valida: {e}")
        return None

def elimina_paragrafo_docx(paragrafo):
    elemento = paragrafo._element
    elemento.getparent().remove(elemento)
    paragrafo._p = paragrafo._element = None

def aggiungi_numeri_pagina_docx(documento):
    """Inserisce il campo numero pagina nel piè di pagina di ogni sezione Word."""
    for sezione in documento.sections:
        paragrafo = sezione.footer.paragraphs[0]
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        campo_inizio = OxmlElement('w:fldChar')
        campo_inizio.set(ns.qn('w:fldCharType'), 'begin')
        istruzione = OxmlElement('w:instrText')
        istruzione.text = 'PAGE'
        campo_fine = OxmlElement('w:fldChar')
        campo_fine.set(ns.qn('w:fldCharType'), 'end')
        run = paragrafo.add_run()
        run._r.append(campo_inizio)
        run._r.append(istruzione)
        run._r.append(campo_fine)

def formatta_manoscritto_kdp(file_docx):
    """Applica un formato Word pulito 6x9 per il manoscritto KDP caricato dall'utente."""
    documento = Document(BytesIO(file_docx.getvalue()))
    for nome_stile in ('Heading 1', 'Heading 2'):
        try:
            documento.styles[nome_stile]
        except KeyError:
            documento.styles.add_style(nome_stile, WD_STYLE_TYPE.PARAGRAPH)

    for sezione in documento.sections:
        sezione.page_width = Inches(6)
        sezione.page_height = Inches(9)
        sezione.top_margin = Inches(0.75)
        sezione.bottom_margin = Inches(0.75)
        sezione.left_margin = Inches(0.75)
        sezione.right_margin = Inches(0.75)

    for paragrafo in list(documento.paragraphs):
        testo = pulisci_testo_editoriale(paragrafo.text).strip()
        if not testo:
            elimina_paragrafo_docx(paragrafo)
            continue
        paragrafo.text = ' '.join(testo.split())
        if len(paragrafo.text) < 80 and re.search(r'(?i)\b(capitolo|chapter|parte|part)\b', paragrafo.text):
            paragrafo.style = 'Heading 1'
            paragrafo.paragraph_format.page_break_before = True
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragrafo.paragraph_format.space_before = Pt(0)
            paragrafo.paragraph_format.space_after = Pt(30)
        elif len(paragrafo.text) < 100 and re.match(r'^\d+(?:\.\d+)?\s+', paragrafo.text):
            paragrafo.style = 'Heading 2'
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragrafo.paragraph_format.first_line_indent = Inches(0)
            paragrafo.paragraph_format.space_before = Pt(18)
            paragrafo.paragraph_format.space_after = Pt(10)
        else:
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragrafo.paragraph_format.first_line_indent = Inches(0.25)
            paragrafo.paragraph_format.space_after = Pt(6)

    stile_normale = documento.styles['Normal']
    stile_normale.font.name = 'Georgia'
    stile_normale.font.size = Pt(11)
    aggiungi_numeri_pagina_docx(documento)
    output = BytesIO()
    documento.save(output)
    output.seek(0)
    return output

def estrai_anteprima_manoscritto(file_caricato):
    """Estrae una porzione di testo da DOCX o PDF per la generazione dei metadati."""
    dati = BytesIO(file_caricato.getvalue())
    if file_caricato.name.lower().endswith('.docx'):
        documento = Document(dati)
        return '\n'.join(p.text for p in documento.paragraphs[:100])
    lettore = PyPDF2.PdfReader(dati)
    return '\n'.join((pagina.extract_text() or '') for pagina in lettore.pages[:15])

def analizza_qualita_prosa(testo):
    """
    Motore Linter NLP Potenziato: analizza densità, lunghezza frasi e vocabolario.
    """
    if not testo or len(testo) < 50: 
        return "⚠️ Testo troppo breve per un'analisi sintattica significativa."
    
    risultati = ["📊 **REPORT LINTER AVANZATO E ANALISI SINTATTICA**\n"]
    
    # 1. Parsing base
    parole = re.findall(r'\b\w+\b', testo.lower())
    frasi = [f.strip() for f in re.split(r'[.!?]+', testo) if len(f.strip()) > 5]
    
    tot_parole = len(parole)
    tot_frasi = len(frasi) if len(frasi) > 0 else 1
    
    # 2. Diversità Lessicale (Ricchezza del vocabolario)
    vocabolo_unico = len(set(parole))
    indice_diversita = (vocabolo_unico / tot_parole) * 100 if tot_parole > 0 else 0
    if indice_diversita < 35:
        risultati.append(f"⚠️ **Vocabolario Ripetitivo**: Indice di diversità lessicale basso ({indice_diversita:.1f}%). Valuta di usare più sinonimi.")
    else:
        risultati.append(f"✅ **Ricchezza Lessicale**: Ottima diversità ({indice_diversita:.1f}%). Il testo risulta stimolante.")

    # 3. Lunghezza Media delle Frasi (Pacing e Affaticamento Neocorteccia)
    parole_per_frase = tot_parole / tot_frasi
    if parole_per_frase > 30:
        risultati.append(f"⚠️ **Sintassi Pesante**: Le frasi sono troppo lunghe (media {parole_per_frase:.1f} parole/frase). Rischio di affaticamento cognitivo: spezza i periodi.")
    elif parole_per_frase < 8:
        risultati.append(f"⚠️ **Ritmo Frammentato**: Frasi molto brevi (media {parole_per_frase:.1f} parole/frase). Il testo potrebbe risultare troppo robotico o telegrafico.")
    else:
        risultati.append(f"✅ **Ritmo e Leggibilità**: Lunghezza frasi perfettamente bilanciata (media {parole_per_frase:.1f} parole/frase).")

    # 4. Ripetizioni Ravvicinate Fastidiose (Finestra Mobile)
    ripetizioni = []
    for i in range(len(parole) - 15):
        target = parole[i]
        # Escludiamo congiunzioni e preposizioni comuni basandoci sulla lunghezza della parola
        if len(target) > 4 and target in parole[i+1 : i+15]: 
            ripetizioni.append(target)
            
    if ripetizioni:
        comuni = [p[0] for p in Counter(ripetizioni).most_common(5)]
        risultati.append(f"🔍 **Allerta Ripetizioni Ravvicinate**: Le seguenti parole si ripetono troppo vicine tra loro: *{', '.join(comuni)}*")
    else:
        risultati.append("✅ **Fluidità Testuale**: Nessuna ripetizione fastidiosa o eco ravvicinata rilevata.")

    return "\n\n".join(risultati)

def sync_capitoli():
    """Costruisce la lista scrivibile dall'indice, inclusa la Prefazione."""
    testo_indice = str(st.session_state.get("indice_raw", "") or "")
    if not testo_indice.strip():
        st.session_state["lista_capitoli"] = []
        return []
    regex = r"(?i)^(?:capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|раздел|章节|secţiune|parte|part|partie|teil|partea|часть|الجزء|部分|\d+\.)"
    lista = []
    for riga in testo_indice.splitlines():
        voce = riga.strip()
        if voce and (sezione_prefazione(voce) or re.search(regex, voce)) and voce not in lista:
            lista.append(voce)
    st.session_state["lista_capitoli"] = lista
    return lista


def imposta_indice_progetto(testo_indice):
    """Salva un indice completo, con Prefazione sempre come prima voce."""
    righe = [riga.strip() for riga in str(testo_indice or "").splitlines() if riga.strip()]
    # Rimuove eventuali Prefazioni duplicate o in posizione errata e ne mette
    # una sola all'inizio, nella lingua attiva del progetto.
    righe = [riga for riga in righe if not sezione_prefazione(riga)]
    indice = "\n".join([titolo_prefazione(), *righe]).strip() if righe else ""
    try:
        memoria_progetto_unica()["indice"] = indice
    except NameError:
        pass
    st.session_state["indice_raw"] = indice
    st.session_state["indice_editoriale"] = indice
    st.session_state["indice_widget_version"] = int(st.session_state.get("indice_widget_version", 0)) + 1
    sync_capitoli()
    return indice


# ======================================================================================================================
# PROFILI EDITORIALI: REGOLE SPECIFICHE PER GENERE, TIPOLOGIA E STRUTTURA
# ======================================================================================================================
def profilo_tipologia_stesura(stile):
    """Restituisce istruzioni di stesura realmente diverse per ogni tipologia selezionabile."""
    profili = {
        "Standard": "Esponi con chiarezza e ordine. Alterna spiegazione, esempio e applicazione senza estremi retorici.",
        "Professionale Accademico": "Definisci termini, separa fatti, metodo, interpretazioni e limiti. Usa un registro preciso e prudente; non trasformare il testo in un elenco di istruzioni quando il contenuto richiede argomentazione.",
        "Persuasivo (Neuromarketing Applicato)": "Parti da un problema concreto, chiarisci valore e prove, affronta obiezioni e guida verso una scelta o un'azione. Non usare pressione, manipolazione o promesse garantite.",
        "Conversazionale ed Empatico": "Accompagna il lettore con un linguaggio umano e rispettoso. Anticipa dubbi reali, normalizza gli ostacoli e offri indicazioni applicabili senza toni paternalistici.",
        "Scientifico Divulgativo": "Rendi comprensibili concetti complessi attraverso definizioni semplici, meccanismi, esempi e limiti. Distingui sempre dati, ipotesi, analogie e aspetti da verificare.",
        "Storytelling Immersivo": "Costruisci scene, azioni, conseguenze e dettagli sensoriali coerenti. Ogni sezione deve far evolvere conflitto, personaggio, relazione o posta in gioco; non riassumere ciò che può essere mostrato.",
        "Giornalistico d'Inchiesta": "Mantieni una linea di verifica: fatti documentabili, fonti da controllare, contraddizioni, contesto e conseguenze. Non presentare ipotesi come prove e non inventare testimonianze.",
        "Socratico (Dialogico / Riflessivo)": "Organizza la sezione attorno a una domanda reale. Esplora presupposti, dubbi e obiezioni, quindi porta il lettore a una conclusione argomentata o a una riflessione verificabile.",
        "Epico ed Evocativo": "Usa immagini e ritmo evocativi senza perdere chiarezza. La trasformazione, le prove e il significato devono essere concreti e adeguati al genere, non formule decorative.",
        "Minimalista ed Essenziale": "Elimina tutto ciò che non serve. Usa frasi sobrie, titoli funzionali, esempi strettamente necessari e una sola idea centrale per blocco di testo."
    }
    return profili.get(stile, profili["Standard"])


def profilo_genere_stesura(genere):
    """Regole di forma e contenuto per tutti i generi offerti dall'interfaccia."""
    profili = {
        "Saggio Scientifico": "Sostieni una tesi con definizioni, metodo, evidenze, controargomentazioni, limiti e implicazioni. Non inventare dati o studi.",
        "Quiz Scientifico": "Alterna spiegazione essenziale, domande verificabili, soluzioni motivate e chiarimento degli errori più probabili.",
        "Manuale Tecnico": "Fornisci prerequisiti, strumenti, parametri, sequenze operative, controlli, errori e criteri di riuscita. Se software o norme possono cambiare, segnala cosa verificare.",
        "Religioso / Teologico": "Distingui testi, interpretazioni, tradizioni e opinioni. Mantieni rispetto, precisione storica e nessuna affermazione dogmatica non attribuita.",
        "Spirituale / Esoterico": "Usa un tono rispettoso e non prescrittivo. Presenta pratiche come esperienze personali o tradizionali, non come cure o certezze scientifiche.",
        "Meditazione / Mindfulness": "Offri pratiche graduali, istruzioni sicure, durata indicativa, osservazioni e alternative. Evita promesse terapeutiche o risultati garantiti.",
        "Business & Marketing": "Usa obiettivi, pubblico, casi, metriche, scelte operative e criteri di verifica. Se i dati non sono forniti, usa esempi dichiaratamente ipotetici.",
        "Economia e Finanza": "Separa educazione generale da consulenza personalizzata. Spiega rischio, limiti, dati e ipotesi; non dare raccomandazioni finanziarie individuali.",
        "Romanzo Rosa": "Sviluppa desiderio, relazione, vulnerabilità, ostacoli e scelta emotiva attraverso scene, dialoghi e trasformazione dei personaggi.",
        "Thriller / Noir": "Costruisci tensione con indizi, conseguenze, conflitti e rivelazioni coerenti. Ogni capitolo deve cambiare le informazioni disponibili o aumentare la posta in gioco.",
        "Fantasy": "Mantieni coerenti mondo, regole, conflitti e conseguenze. Mostra il worldbuilding dentro azioni e scene, senza blocchi enciclopedici.",
        "Fantascienza": "Rendi coerente la premessa speculativa e mostra come modifica società, tecnologia, personaggi e conflitto. Non sostituire la storia con spiegazioni astratte.",
        "Manuale Psicologico": "Spiega modelli e pratiche in modo accessibile, con limiti chiari. Non fare diagnosi, non promettere cura e invita a rivolgersi a professionisti quando necessario.",
        "Biografia": "Segui una cronologia significativa, usando fonti verificabili e distinguendo fatti, testimonianze e interpretazioni. Privilegia svolte e contesto rispetto a elenchi di date.",
        "Ricettario": "Ogni capitolo-ricetta deve contenere porzioni, tempi, ingredienti con dosi, procedimento numerato, segnali di riuscita, errore e correzione, variante e conservazione solo se verificata. Non duplicare la stessa ricetta in forma breve ed estesa.",
        "Test Prep (Preparazione Esami)": "Spiega soltanto le competenze pertinenti alla prova, poi fornisci esercizi reali, soluzioni ragionate, errori tipici e criteri di autovalutazione. Quando una sezione promette quiz, test o simulazioni, deve contenere le domande effettive e non istruzioni generiche su come studiare. Mantieni separati quesiti e soluzioni, verifica il numero richiesto, evita duplicati e non inventare regole d'esame non verificate.",
        "Narrativo": "Sviluppa personaggi, conflitto, cause e conseguenze in scene concrete. Ogni capitolo deve avere una funzione narrativa distinta.",
        "Romanzo Classico": "Usa una costruzione narrativa solida, personaggi coerenti, ambientazione e temi sviluppati attraverso azioni e dialoghi; evita imitazioni di autori viventi.",
        "Contemporaneo": "Racconta conflitti e relazioni con voce naturale, dettagli specifici e temi attuali trattati attraverso la storia, non con prediche.",
        "Self-Help": "Definisci problemi realistici, pratiche graduali, esempi e criteri di verifica. Evita promesse di trasformazione garantita o consigli clinici.",
        "Manuale Pratico": "Fornisci un percorso eseguibile: materiali o prerequisiti, passaggi, controlli, errori, alternative e risultato finale verificabile.",
        "Storico": "Ordina il racconto per nessi causali e cronologia, distinguendo fonti, fatti, interpretazioni e controversie. Non inventare citazioni o date."
    }
    return profili.get(genere, "Mantieni una struttura coerente con pubblico, obiettivo, genere e limiti dichiarati.")


def estrai_numero_ricette(titolo, trama, obiettivo):
    testo = f"{titolo} {trama} {obiettivo}".lower()
    # Il confine di parola deve essere una vera espressione regolare, non il testo letterale "\\b".
    # Supporta le principali lingue offerte dalla sidebar.
    match = re.search(r"\b(\d{1,3})\s+(?:ricette|recipes|recetas|recettes|rezepte|rețete|рецептов|وصفات|个食谱)\b", testo)
    return int(match.group(1)) if match else None


def profilo_struttura_indice(genere, titolo, trama, obiettivo):
    """Evita che una stessa gabbia 15-18 capitoli venga applicata a libri incompatibili."""
    if genere == "Ricettario":
        numero = estrai_numero_ricette(titolo, trama, obiettivo)
        quantità = f"esattamente {numero}" if numero else "un numero coerente con la richiesta"
        return f"""RICETTARIO: crea {quantità} ricette effettive, distribuite in parti tematiche coerenti. Ogni ricetta è un Capitolo autonomo e completo. Se è richiesto un numero preciso di ricette, crea esattamente quel numero di Capitoli e ciascun Capitolo deve avere il nome di una ricetta: non usare Capitoli per introduzione, ingredienti, attrezzatura, tecniche o consigli. Le Parti possono orientare il lettore senza aggiungere Capitoli introduttivi. Non creare sottocapitoli 1.1, 1.2 o 1.3 per espandere la stessa ricetta. Il numero delle ricette nell'indice deve coincidere con il numero richiesto."""
    if genere in {"Romanzo Rosa", "Thriller / Noir", "Fantasy", "Fantascienza", "Narrativo", "Romanzo Classico", "Contemporaneo", "Biografia"}:
        return "NARRATIVA E BIOGRAFIA: organizza 3-6 Parti e un numero di capitoli proporzionato all'arco narrativo. Non imporre sottocapitoli a ogni capitolo: usali solo se sono necessari e non spezzano artificialmente scene o svolte. Ogni titolo deve nominare una scena, una scelta, un luogo, un personaggio, un oggetto o una conseguenza specifici del brief. Almeno un terzo dei titoli deve contenere parole concrete tratte dal titolo o dalla trama. Evita titoli generici come 'Il ritorno', 'La scoperta', 'L'incontro inaspettato', 'Il richiamo del passato', 'Riflessioni' o 'La fine'."
    if genere in {"Quiz Scientifico", "Test Prep (Preparazione Esami)"}:
        return "QUIZ E TEST PREP: organizza fondamenti, esercitazione graduata, quiz/domande commentate, almeno una simulazione esplicitamente nominata e correzioni. Nella lingua scelta usa le parole equivalenti a ‘quiz/questions’ e ‘simulation’, così che lo scopo delle sezioni sia leggibile. Ogni unità deve indicare una competenza verificabile; non creare capitoli riempitivi."
    return "SAGGISTICA E MANUALI: distribuisci fondamenti, metodo, applicazione, verifica e sintesi in una struttura proporzionata al brief. Crea sottocapitoli solo per concetti o passaggi realmente distinti; il budget di sezioni indicato nel prompt prevale su ogni schema numerico generale."


def normalizza_indice_generato(indice):
    """Rimuove solo rumore di formattazione, senza alterare l'architettura proposta."""
    righe = []
    for riga in (indice or "").splitlines():
        pulita = re.sub(r"^\s*[-*#]+\s*", "", riga).strip()
        if pulita.lower() in {"indice", "table of contents", "sommaire", "inhaltsverzeichnis"}:
            continue
        if pulita:
            righe.append(pulita)
    return "\n".join(righe).strip()


def criticita_indice_generato(indice, genere, titolo, trama, obiettivo, minimo_parti=4, minimo_capitoli=None):
    """Controllo deterministico leggero: intercetta gli errori che il modello tende a ripetere."""
    testo = normalizza_indice_generato(indice)
    righe = testo.splitlines()
    capitoli = [riga for riga in righe if re.match(r"(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+", riga)]
    parti = [riga for riga in righe if re.match(r"(?i)^(parte|part|partie|teil|partea|часть|الجزء|部分)\s+", riga)]
    if not capitoli:
        return ["non sono stati riconosciuti capitoli nel formato richiesto"]

    problemi = []
    narrativi = {"Romanzo Rosa", "Thriller / Noir", "Fantasy", "Fantascienza", "Narrativo", "Romanzo Classico", "Contemporaneo", "Biografia"}
    if genere != "Ricettario" and len(parti) < minimo_parti:
        problemi.append(f"struttura troppo breve: sono presenti solo {len(parti)} Parti, ne servono almeno {minimo_parti}")
    minimo_capitoli_effettivo = (12 if genere not in {"Ricettario"} else 0) if minimo_capitoli is None else minimo_capitoli
    if len(capitoli) < minimo_capitoli_effettivo:
        problemi.append(f"struttura troppo breve: sono presenti solo {len(capitoli)} Capitoli, ne servono almeno {minimo_capitoli_effettivo}")
    if genere not in narrativi and genere != "Ricettario":
        capitoli_senza_sviluppo = []
        for posizione, capitolo in enumerate(capitoli):
            inizio = righe.index(capitolo)
            fine = next((i for i in range(inizio + 1, len(righe)) if re.match(r"(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+", righe[i]) or re.match(r"(?i)^(parte|part|partie|teil|partea|часть|الجزء|部分)\s+", righe[i])), len(righe))
            sottosezioni = sum(1 for riga in righe[inizio + 1:fine] if re.match(r"^\d+\.\d+\s+", riga))
            if sottosezioni < 2:
                capitoli_senza_sviluppo.append(capitolo)
        if capitoli_senza_sviluppo:
            problemi.append("capitoli senza almeno due sottocapitoli distinti: " + "; ".join(capitoli_senza_sviluppo[:3]))

    if genere == "Ricettario":
        richieste = estrai_numero_ricette(titolo, trama, obiettivo)
        if richieste and len(capitoli) != richieste:
            problemi.append(f"sono richieste {richieste} ricette, ma l'indice contiene {len(capitoli)} capitoli")
        titoli_capitoli = " ".join(capitoli).lower()
        non_ricette = (
            "introduzione", "ingredient", "attrezz", "tecniche", "consigli", "dispensa",
            "sostituz", "substitut", "preparazione di ingredient", "nutrient", "planific", "consejos",
            "conservación", "alternativas", "erreurs", "conseils", "substitutions", "grundlagen"
        )
        if any(parola in titoli_capitoli for parola in non_ricette):
            problemi.append("un capitolo del ricettario è introduttivo o tecnico invece di essere una ricetta")
        if any(re.match(r"^\d+\.\d+\s+", riga) for riga in righe):
            problemi.append("il ricettario contiene sottocapitoli: ogni capitolo deve essere una ricetta completa e autonoma")

    if genere in narrativi:
        titoli_generici = {
            "il ritorno", "la scoperta", "l'inizio", "la fine", "il conflitto", "la scelta", "la crisi",
            "riflessioni", "sogni e memorie", "nuovi inizi", "l'incontro inaspettato", "il richiamo del passato",
            "la dolcezza del ricordo", "il richiamo della tradizione", "riscoprire se stessi", "la verità", "il segreto"
        }
        trovati = []
        for capitolo in capitoli:
            nome = re.sub(r"(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+\s*:\s*", "", capitolo).strip().lower()
            if nome in titoli_generici:
                trovati.append(capitolo)
        if len(trovati) >= 2:
            problemi.append("titoli narrativi troppo generici: " + "; ".join(trovati[:3]))
        parole_da_escludere = {
            "della", "delle", "dello", "degli", "dalla", "nelle", "nello", "come", "con", "una", "uno", "per",
            "che", "del", "dei", "gli", "le", "il", "la", "un", "e", "di", "da", "in", "su", "tra", "fra",
            "storia", "romanzo", "guida", "raccontare", "lettore", "lettori", "obiettivo", "titolo", "libro"
        }
        parole_brief = {
            parola for parola in re.findall(r"[a-zàèéìòóù]{4,}", f"{titolo} {trama}".lower())
            if parola not in parole_da_escludere
        }
        titoli_con_ancora = 0
        for capitolo in capitoli:
            nome = re.sub(r"(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+\s*:\s*", "", capitolo).lower()
            if any(parola in nome for parola in parole_brief):
                titoli_con_ancora += 1
        soglia = max(3, (len(capitoli) + 2) // 3)
        if titoli_con_ancora < soglia:
            problemi.append(
                f"titoli narrativi poco ancorati agli elementi concreti del brief ({titoli_con_ancora}/{len(capitoli)} titoli specifici)"
            )
    if genere in {"Quiz Scientifico", "Test Prep (Preparazione Esami)"}:
        testo_minuscolo = testo.lower()
        parole_quiz = ("quiz", "domand", "question", "pregunta", "frage", "вопрос", "سؤال", "问题")
        parole_simulazione = ("simulaz", "simulation", "simulación", "simulare", "симуля", "محاك", "模拟")
        if not any(parola in testo_minuscolo for parola in parole_quiz):
            problemi.append("manca una sezione con quiz o domande effettive")
        if not any(parola in testo_minuscolo for parola in parole_simulazione):
            problemi.append("manca una sezione di simulazione")
    return problemi


def audit_editoriale_indice_generato(indice, genere, titolo, trama, obiettivo, lingua, stile, narrativa, pov, *, addebita=True):
    """Usa esattamente lo stesso metro del pulsante 'Voto indice', evitando approvazioni incoerenti."""
    risposta = valuta_indice_editoriale(
        indice, titolo, trama, genere, stile, narrativa, pov, obiettivo, lingua, "", addebita=addebita
    ).strip()
    match = re.search(r"(?im)^\s*(?:voto\s+complessivo|voto)\s*:\s*(10|[0-9])\s*(?:/\s*10)?\b", risposta)
    voto = int(match.group(1)) if match else 0
    difetti = re.search(r"(?ims)^\s*(?:miglioramenti consigliati|difetti)\s*:\s*(.+?)(?:\n\s*[A-ZÀ-Ú][A-ZÀ-Ú ]+\s*:|$)", risposta)
    return voto, (difetti.group(1).strip() if difetti else risposta)


def firma_indice(indice):
    """Confronto robusto: ignora maiuscole e spazi, non le differenze editoriali reali."""
    return re.sub(r"\s+", " ", (indice or "").strip().lower())


def conta_sezioni_indice(indice):
    """Conta le voci che saranno realmente disponibili nell'editor, in tutte le lingue supportate."""
    regex = r'(?i)(Capitolo|Chapter|Kapitel|Capítulo|Chapitre|Capitolul|Глава|الفصل|Раздел|章节|Secţiune|Parte|Part|Partie|Teil|Partea|Часть|الجزء|部分|\d+\.)'
    return sum(1 for riga in (indice or "").splitlines() if re.search(regex, riga.strip()))


def superamenti_budget_editoriale(indice, massimo_parti, massimo_capitoli, massimo_sottocapitoli):
    """Rileva un indice troppo esteso o spezzato rispetto al profilo scelto."""
    righe = [riga.strip() for riga in normalizza_indice_generato(indice).splitlines() if riga.strip()]
    regex_capitolo = r"(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+"
    regex_limite = r"(?i)^(parte|part|partie|teil|partea|часть|الجزء|部分)\s+"
    superamenti = []
    parti = [riga for riga in righe if re.match(regex_limite, riga)]
    capitoli = [riga for riga in righe if re.match(regex_capitolo, riga)]
    if massimo_parti and len(parti) > massimo_parti:
        superamenti.append(f"{len(parti)} Parti invece di massimo {massimo_parti}")
    if massimo_capitoli and len(capitoli) > massimo_capitoli:
        superamenti.append(f"{len(capitoli)} Capitoli invece di massimo {massimo_capitoli}")
    if not massimo_sottocapitoli:
        return superamenti
    for posizione, capitolo in enumerate(righe):
        if not re.match(regex_capitolo, capitolo):
            continue
        fine = next(
            (
                indice_successivo for indice_successivo in range(posizione + 1, len(righe))
                if re.match(regex_capitolo, righe[indice_successivo])
                or re.match(regex_limite, righe[indice_successivo])
            ),
            len(righe),
        )
        sottocapitoli = [
            riga for riga in righe[posizione + 1:fine]
            if re.match(r"^\d+\.\d+\s+", riga)
        ]
        if len(sottocapitoli) > massimo_sottocapitoli:
            superamenti.append(
                f"{capitolo} con {len(sottocapitoli)} sottocapitoli invece di massimo {massimo_sottocapitoli}"
            )
    return superamenti


def genera_indice_controllato(prompt, system_prompt, genere, titolo, trama, obiettivo, lingua, stile, narrativa, pov,
                              indice_da_superare="", massimo_sezioni=None, minimo_parti=4, minimo_capitoli=None,
                              budget_strutturale=""):
    """Genera un indice robusto: nessuna proposta valida sparisce dopo l'attesa.

    Il controllo 8/10 continua a certificare un indice "approvato", ma non è
    più un interruttore che cancella una struttura già completa e nei limiti.
    Un indice strutturalmente valido ma ancora migliorabile viene mostrato,
    salvato e segnalato con trasparenza: l'utente può valutarlo o rigenerarlo
    dal voto, senza ripetere la ricerca né pagare una seconda volta.
    """
    costo = CREDIT_COSTS["indice_generazione_editoriale"]
    riferimento = addebita_azione_diretta("genera_indice_controllato", amount=costo)
    indice_di_partenza = firma_indice(indice_da_superare)
    difetti_bloccanti = (
        "non sono stati riconosciuti capitoli", "capitoli senza almeno due sottocapitoli",
        "sono richieste", "un capitolo del ricettario", "il ricettario contiene sottocapitoli",
        "manca una sezione con quiz", "manca una sezione di simulazione", "struttura troppo breve",
        "oltre il massimo consentito",
    )

    def valuta_candidata(indice_candidato):
        """Restituisce validità strutturale e rilievi senza far cadere la UI."""
        candidata = normalizza_indice_generato(indice_candidato)
        problemi = criticita_indice_generato(
            candidata, genere, titolo, trama, obiettivo,
            minimo_parti=minimo_parti, minimo_capitoli=minimo_capitoli,
        )
        if massimo_sezioni and conta_sezioni_indice(candidata) > massimo_sezioni:
            problemi.append(
                f"l'indice contiene {conta_sezioni_indice(candidata)} sezioni, "
                f"oltre il massimo consentito di {massimo_sezioni}"
            )
        identica = bool(indice_di_partenza and firma_indice(candidata) == indice_di_partenza)
        if identica:
            problemi.append(
                "la proposta è identica all'indice valutato: applica concretamente i miglioramenti richiesti"
            )
        blocchi = any(
            any(blocco in problema for blocco in difetti_bloccanti) for problema in problemi
        )
        return candidata, problemi, blocchi, identica

    try:
        corrente = normalizza_indice_generato(
            chiedi_gpt(prompt, system_prompt, addebita=False, model=MODELLO_EDITORIALE)
        )
    except Exception as exc:
        refund_credits(riferimento, reason="genera_indice_fallito", amount=costo)
        st.session_state["ultimo_controllo_indice"] = (
            "Indice non generato: la richiesta non ha restituito una risposta utilizzabile. "
            "Il credito di progettazione è stato riaccreditato."
        )
        st.session_state["ultimo_errore_indice"] = str(exc)
        return ""

    migliore_strutturale = ""
    ultimo_problema = ""
    # Bozza iniziale + una correzione mirata: non moltiplica le attese API.
    for tentativo in range(2):
        corrente, problemi, ha_blocchi, proposta_identica = valuta_candidata(corrente)
        if corrente and not ha_blocchi and not proposta_identica:
            migliore_strutturale = corrente
            try:
                voto_editoriale, difetti_editoriali = audit_editoriale_indice_generato(
                    corrente, genere, titolo, trama, obiettivo, lingua, stile, narrativa, pov,
                    addebita=False,
                )
            except Exception as exc:
                voto_editoriale, difetti_editoriali = 0, f"audit non disponibile ({exc})"
            if voto_editoriale >= 8:
                messaggio = f"Indice approvato: {voto_editoriale}/10 nel controllo strutturale ed editoriale automatico."
                if tentativo:
                    messaggio = f"Indice corretto automaticamente e approvato {voto_editoriale}/10."
                st.session_state["ultimo_controllo_indice"] = messaggio
                return corrente
            ultimo_problema = f"audit editoriale {voto_editoriale}/10: {difetti_editoriali}"
        else:
            voto_editoriale, difetti_editoriali = 0, "vincoli strutturali da correggere prima della valutazione editoriale"
            ultimo_problema = "; ".join(problemi) or "risposta senza una struttura riconoscibile"

        if tentativo == 1:
            break

        supera_limite = bool(massimo_sezioni and conta_sezioni_indice(corrente) > massimo_sezioni)
        if supera_limite:
            revisione = f"""Riduci e riorganizza l'indice qui sotto. Restituisci SOLO l'indice gerarchico pulito.

VINCOLO INDEROGABILE: al massimo {massimo_sezioni} voci totali tra Parti, Capitoli e sottocapitoli.
BUDGET DA RISPETTARE: {budget_strutturale or 'struttura compatta senza voci ridondanti'}.
Unisci argomenti contigui e cancella i sottocapitoli ripetitivi: non limitarti ad abbreviare i titoli.
Prima di rispondere conta le voci; se sono oltre il massimo, continua ad accorpare finché rientrano.

INDICE DA COMPRIMERE
{corrente}
"""
        else:
            revisione = prompt + f"""

REVISIONE OBBLIGATORIA DELL'INDICE
Correggi questa proposta senza commenti, saluti, Markdown o testo esterno all'indice.
Difetti da risolvere: {ultimo_problema}.
Mantieni soltanto argomenti attinenti al brief, applica i limiti di struttura e restituisci
solo l'indice gerarchico pulito. Non ripetere l'indice precedente senza modifiche visibili.

INDICE DA CORREGGERE
{corrente}
"""
        try:
            corrente = normalizza_indice_generato(
                chiedi_gpt(revisione, system_prompt, addebita=False, model=MODELLO_EDITORIALE)
            )
        except Exception as exc:
            ultimo_problema = f"correzione non disponibile ({exc})"
            break

    if migliore_strutturale:
        # Risposta affidabile anche quando il voto automatico è prudente: il
        # testo non sparisce, il bottone VOTO INDICE resta disponibile per il
        # perfezionamento e l'utente non deve ricliccare la generazione.
        st.session_state["ultimo_controllo_indice"] = (
            "Attenzione: indice strutturalmente valido e pubblicato, ma il controllo editoriale "
            f"richiede un miglioramento prima della stesura: {ultimo_problema}. "
            "Usa VOTO INDICE e, se necessario, RIGENERA INDICE SEGUENDO IL VOTO."
        )
        return migliore_strutturale

    refund_credits(riferimento, reason="genera_indice_non_pubblicabile", amount=costo)
    st.session_state["ultimo_controllo_indice"] = (
        "Indice non pubblicato: non è stata prodotta una struttura utilizzabile nei limiti richiesti. "
        "Il credito di progettazione è stato riaccreditato. Riprova senza pagare due volte."
    )
    return ""


def genera_indice_controllato(prompt, system_prompt, genere, titolo, trama, obiettivo, lingua, stile, narrativa, pov,
                              indice_da_superare="", massimo_sezioni=None, minimo_parti=4, minimo_capitoli=None,
                              budget_strutturale="", profilo_lunghezza="Standard KDP",
                              massimo_parti_editoriali=None, massimo_capitoli_editoriali=None,
                              massimo_sottocapitoli=None, aggiorna_stato=None):
    """Genera, verifica e corregge l'indice con avanzamento leggibile."""
    def avanza(percentuale, testo):
        if callable(aggiorna_stato):
            try:
                aggiorna_stato(percentuale, testo)
            except Exception:
                # La barra e' solo informativa: non deve mai interrompere la
                # generazione se il browser viene aggiornato nel frattempo.
                pass

    riferimento = addebita_azione_diretta("genera_indice_controllato", amount=CREDIT_COSTS["indice_generazione_editoriale"])
    # L'indice resta affidato al modello editoriale completo, ma un timeout o
    # una risposta non utilizzabile non deve lasciare lo spinner senza esito.
    avanza(42, "Generazione della struttura editoriale in corso...")
    risposta_iniziale = chiedi_gpt(
        prompt, system_prompt, addebita=False, model=MODELLO_EDITORIALE,
        timeout_seconds=TIMEOUT_INDICE_SECONDI, reason="genera_indice",
    )
    if not str(risposta_iniziale or "").strip() or str(risposta_iniziale).startswith("ERRORE:"):
        refund_credits(riferimento, reason="genera_indice_fallito", amount=CREDIT_COSTS["indice_generazione_editoriale"])
        st.session_state["ultimo_controllo_indice"] = (
            "Indice non generato: il cervello non ha concluso la risposta entro il tempo previsto. "
            "Nessun credito di progettazione è stato trattenuto: riprova tra poco."
        )
        st.session_state["ultimo_errore_indice"] = str(risposta_iniziale or "risposta vuota")
        return ""
    corrente = normalizza_indice_generato(risposta_iniziale)
    avanza(58, "Struttura ricevuta: controllo gerarchia, limiti e coerenza...")
    indice_di_partenza = firma_indice(indice_da_superare)
    massimo_tentativi = 2  # bozza + una sola correzione mirata: evita attese inutili
    for tentativo in range(massimo_tentativi):
        avanza(
            64 + tentativo * 17,
            "Valutazione editoriale dell'indice..." if not tentativo
            else "Verifica della struttura corretta...",
        )
        problemi = criticita_indice_generato(
            corrente, genere, titolo, trama, obiettivo, minimo_parti=minimo_parti, minimo_capitoli=minimo_capitoli
        )
        stima_capacita = stima_budget_parole_indice(corrente, profilo_lunghezza)
        superamenti_budget = superamenti_budget_editoriale(
            corrente, massimo_parti_editoriali, massimo_capitoli_editoriali,
            massimo_sottocapitoli,
        )
        budget_editoriale_superato = bool(superamenti_budget)
        if budget_editoriale_superato:
            problemi.append(
                "indice troppo esteso o dettagliato per il profilo: "
                + "; ".join(superamenti_budget[:4])
            )
        sezioni_generate = conta_sezioni_indice(corrente)
        if massimo_sezioni and sezioni_generate > massimo_sezioni:
            problemi.append(
                f"l'indice contiene {sezioni_generate} sezioni, oltre il massimo consentito di {massimo_sezioni}"
            )
        proposta_identica = bool(indice_di_partenza and firma_indice(corrente) == indice_di_partenza)
        if proposta_identica:
            problemi.append(
                "la proposta è identica all'indice valutato: applica concretamente i miglioramenti richiesti "
                "modificando struttura e titoli pertinenti"
            )
        supera_limite = bool(massimo_sezioni and sezioni_generate > massimo_sezioni)
        problemi_senza_estensione = [
            problema for problema in problemi
            if "oltre il massimo consentito" not in problema
        ]
        if supera_limite and not problemi_senza_estensione and not proposta_identica:
            # Una struttura già leggibile e completa non deve sparire perché
            # il modello ha superato soltanto una fascia numerica. Pubblicarla
            # evita una seconda richiesta AI, non spreca crediti e lascia
            # all'utente la scelta di rigenerarla più compatta in seguito.
            st.session_state["ultimo_controllo_indice"] = (
                f"Indice pubblicato con avviso: {sezioni_generate + 1} sezioni totali inclusa la Prefazione, "
                f"oltre la fascia massima del profilo ({massimo_sezioni + 1}). "
                "La struttura è stata mantenuta perché non presenta altre criticità oggettive; "
                "puoi conservarla o rigenerarla più compatta senza perdere il progetto."
            )
            avanza(100, "Indice pubblicato con avviso sulla lunghezza.")
            return corrente
        difetti_bloccanti = (
            "non sono stati riconosciuti capitoli", "capitoli senza almeno due sottocapitoli",
            "sono richieste", "un capitolo del ricettario", "il ricettario contiene sottocapitoli", "manca una sezione con quiz", "manca una sezione di simulazione",
            "struttura troppo breve", "oltre il massimo consentito"
        )
        ha_blocchi = any(any(blocco in problema for blocco in difetti_bloccanti) for problema in problemi)
        # Non avviamo il lento audit editoriale su un indice già fuori misura:
        # prima lo rendiamo valido sul piano oggettivo.
        if ha_blocchi or proposta_identica:
            voto_editoriale, difetti_editoriali = 0, "vincoli strutturali da correggere prima della valutazione editoriale"
        else:
            voto_editoriale, difetti_editoriali = audit_editoriale_indice_generato(
                corrente, genere, titolo, trama, obiettivo, lingua, stile, narrativa, pov, addebita=False
            )
        if voto_editoriale >= 8 and not ha_blocchi and not proposta_identica and not (
            budget_editoriale_superato and tentativo < massimo_tentativi - 1
        ):
            esito = f"Indice approvato: {voto_editoriale}/10 nel controllo strutturale ed editoriale automatico."
            if problemi:
                esito += " Note qualitative considerate: " + "; ".join(problemi)
            if not stima_capacita["raggiunge_soglia"]:
                esito += (
                    " Avviso di lunghezza: con i limiti delle sezioni questo indice stima "
                    f"{stima_capacita['pagine_minime']}–{stima_capacita['pagine_massime']} pagine 6×9, "
                    f"sotto la soglia con tolleranza ({stima_capacita['soglia_pagine']}/"
                    f"{stima_capacita['obiettivo_pagine']}). È stato pubblicato perché completo e di qualità: "
                    "il software non aggiunge riempitivi né blocca la pubblicazione."
                )
            if tentativo:
                esito = (
                    f"Indice corretto automaticamente al controllo {tentativo} e approvato "
                    f"{voto_editoriale}/10. " + esito
                )
            st.session_state["ultimo_controllo_indice"] = esito
            avanza(100, "Indice approvato e pronto per il salvataggio.")
            return corrente
        if voto_editoriale < 8:
            problemi.append(f"audit editoriale {voto_editoriale}/10: {difetti_editoriali}")
        if tentativo == massimo_tentativi - 1:
            st.session_state["ultimo_controllo_indice"] = "Attenzione: l'indice non ha raggiunto la soglia minima di 8/10 e richiede una verifica manuale: " + "; ".join(problemi)
            return ""
        if supera_limite:
            # Una richiesta dedicata alla compressione è più affidabile del prompt editoriale completo,
            # che potrebbe contenere molte istruzioni e spingere il modello a espandere l'indice.
            revisione = f"""Riduci e riorganizza l'indice qui sotto. Restituisci SOLO l'indice gerarchico pulito.

VINCOLO INDEROGABILE: al massimo {massimo_sezioni} voci totali tra Parti, Capitoli e sottocapitoli.
BUDGET DA RISPETTARE: {budget_strutturale or 'struttura compatta senza voci ridondanti'}.
Conserva soltanto i passaggi indispensabili al titolo, al genere, al brief, ai quiz e alla simulazione se richiesti.
Unisci argomenti contigui e cancella i sottocapitoli ripetitivi: non limitarti ad abbreviare i titoli.
Prima di rispondere conta le voci; se sono oltre il massimo, continua ad accorpare finché rientrano.

INDICE DA COMPRIMERE
{corrente}
"""
        elif budget_editoriale_superato:
            # Il budget protegge il lettore da indici chilometrici: Parti,
            # Capitoli e sottocapitoli devono restare proporzionati insieme.
            revisione = prompt + f"""

COMPRESSIONE DEL BUDGET EDITORIALE — TENTATIVO {tentativo + 1}
L'indice supera il budget del profilo. Riscrivi l'intera struttura rispettando contemporaneamente:
- massimo {massimo_parti_editoriali} Parti;
- massimo {massimo_capitoli_editoriali} Capitoli;
- massimo {massimo_sottocapitoli} sottocapitoli per Capitolo.

Accorpa nello stesso sottocapitolo esempi, varianti, errori comuni, strumenti, verifiche e casi che sviluppano il medesimo
tema. Mantieni separate soltanto competenze, scene, decisioni, problemi o risultati veramente diversi. Preferisci titoli
ampi e concreti a molte micro-voci.

Non aggiungere nuove voci per inseguire la stima delle pagine; non eliminare invece passaggi indispensabili al brief.
Non superare il massimo complessivo già indicato e non trasformare titoli diversi in duplicati generici.

Restituisci SOLO l'indice gerarchico pulito.

INDICE DA COMPATTARE
{corrente}
"""
        else:
            revisione = prompt + f"""

REVISIONE OBBLIGATORIA DELL'INDICE — TENTATIVO {tentativo + 1}
La bozza precedente non rispetta questi vincoli oggettivi/editoriali: {'; '.join(problemi)}.
Usa i difetti elencati come requisiti di correzione concreti. Riscrivi l'intero indice, senza commenti e senza la parola 'Indice' in apertura.
Correggi tutti i punti segnalati, inclusi grammatica, gerarchia, ripetizioni e aderenza al brief; non limitarti a rinominare i titoli.
Mantieni soltanto argomenti attinenti al brief.

Se il difetto riguarda il numero di sezioni, non accorciare soltanto i titoli: unisci gli argomenti contigui
e rimuovi le voci ridondanti fino a rispettare il massimo indicato. Conta nuovamente le voci prima di rispondere.

DIVIETO ASSOLUTO: non restituire l'indice di partenza né una sua copia cosmetica. Ogni miglioramento
indicato dall'editor deve produrre una modifica visibile nella struttura, nella sequenza, nei titoli o
nella copertura dei contenuti. Prima di rispondere confronta internamente la nuova proposta con l'indice
da correggere e verifica di aver applicato almeno le correzioni necessarie.

INDICE RIFIUTATO DA CORREGGERE
{corrente}
"""
        risposta_revisione = chiedi_gpt(
            revisione, system_prompt, addebita=False, model=MODELLO_EDITORIALE,
            timeout_seconds=TIMEOUT_INDICE_SECONDI, reason="correzione_indice",
        )
        avanza(92, "Correzione completata: validazione finale in corso...")
        if not str(risposta_revisione or "").strip() or str(risposta_revisione).startswith("ERRORE:"):
            st.session_state["ultimo_controllo_indice"] = (
                "La bozza dell'indice è stata creata, ma la correzione automatica non ha concluso "
                "la risposta entro il tempo previsto. Riprova la generazione tra poco."
            )
            st.session_state["ultimo_errore_indice"] = str(risposta_revisione or "risposta vuota")
            return ""
        corrente = normalizza_indice_generato(risposta_revisione)
    return ""

def tipo_sezione_editoriale(sezione):
    """Classifica tutte le voci scrivibili, compresa la Prefazione."""
    return classifica_sezione(sezione, sezione_prefazione)


PAROLE_PER_PAGINA_6X9 = 275


def stima_budget_parole_indice(indice, profilo_lunghezza):
    """Stima la capacità reale di un indice prima della stesura.

    Il calcolo non usa un numero teorico di capitoli: somma i limiti reali
    della Prefazione, delle Parti, dei capitoli-cornice e delle sezioni
    sostanziali. Serve a guidare una sola eventuale revisione dell'indice,
    mai a imporre voci vuote o a bloccare un libro già valido.
    """
    profilo = PROFILI_LUNGHEZZA_STESURA.get(
        profilo_lunghezza, PROFILI_LUNGHEZZA_STESURA["Standard KDP"]
    )
    regex = (
        r"(?i)^(?:capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|"
        r"الفصل|раздел|章节|secţiune|parte|part|partie|teil|partea|часть|الجزء|部分|\d+\.)"
    )
    sezioni = []
    for riga in str(indice or "").splitlines():
        voce = riga.strip()
        if voce and (sezione_prefazione(voce) or re.search(regex, voce)) and voce not in sezioni:
            sezioni.append(voce)
    if not any(sezione_prefazione(sezione) for sezione in sezioni):
        # Il generatore restituisce il corpo dell'indice e il software
        # aggiunge la Prefazione dopo: la stima deve già considerarla.
        sezioni.insert(0, titolo_prefazione())

    parole_minime = 0
    parole_massime = 0
    dettaglio = {"prefazione": 0, "parti": 0, "capitoli_cornice": 0, "sezioni": 0}
    for sezione in sezioni:
        tipo = tipo_sezione_editoriale(sezione)
        if tipo == "prefazione":
            minimo, massimo, categoria = 140, 220, "prefazione"
        elif tipo == "parte":
            minimo, massimo, categoria = 110, 180, "parti"
        elif tipo == "capitolo" and individua_sottocapitoli_del_capitolo(sezione, sezioni):
            minimo, massimo, categoria = 160, 260, "capitoli_cornice"
        else:
            minimo, massimo, categoria = profilo["min_parole"], profilo["max_parole"], "sezioni"
        parole_minime += minimo
        parole_massime += massimo
        dettaglio[categoria] += 1

    obiettivo_pagine = int(profilo["pagine_minime"])
    tolleranza = float(profilo.get("tolleranza_pagine", 0.15))
    obiettivo_parole = obiettivo_pagine * PAROLE_PER_PAGINA_6X9
    soglia_parole = math.ceil(obiettivo_parole * (1 - tolleranza))
    return {
        "sezioni_totali": len(sezioni),
        "parole_minime": parole_minime,
        "parole_massime": parole_massime,
        "pagine_minime": math.ceil(parole_minime / PAROLE_PER_PAGINA_6X9) if parole_minime else 0,
        "pagine_massime": math.ceil(parole_massime / PAROLE_PER_PAGINA_6X9) if parole_massime else 0,
        "obiettivo_pagine": obiettivo_pagine,
        "soglia_pagine": math.ceil(obiettivo_pagine * (1 - tolleranza)),
        "soglia_parole": soglia_parole,
        "raggiunge_soglia": parole_massime >= soglia_parole,
        "dettaglio": dettaglio,
    }


def riepilogo_stima_pagine_manoscritto(sezioni, contenuti, indice, profilo_lunghezza):
    """Calcola pagine attuali e previsione prudente senza gonfiare il testo.

    Le sezioni già scritte usano sempre il loro conteggio reale. Per le voci
    ancora vuote viene mostrato un intervallo editoriale: Prefazione, Parti e
    capitoli-cornice restano volutamente più brevi, mentre le sezioni
    sostanziali seguono il profilo selezionato. È una stima trasparente, non
    un invito ad aggiungere riempitivi per raggiungere un numero di pagine.
    """
    profilo = PROFILI_LUNGHEZZA_STESURA.get(
        profilo_lunghezza, PROFILI_LUNGHEZZA_STESURA["Standard KDP"]
    )
    parole_reali = 0
    parole_minime_rimanenti = 0
    parole_massime_rimanenti = 0
    sezioni_mancanti = 0
    righe_indice = str(indice or "").splitlines()
    for sezione in sezioni or []:
        testo = pulisci_testo_editoriale(
            str((contenuti or {}).get(sezione, "") or "")
        ).strip()
        if testo:
            parole_reali += len(testo.split())
            continue
        sezioni_mancanti += 1
        tipo = tipo_sezione_editoriale(sezione)
        if tipo == "prefazione":
            minimo, massimo = 140, 220
        elif tipo == "parte":
            minimo, massimo = 110, 180
        elif tipo == "capitolo" and individua_sottocapitoli_del_capitolo(sezione, righe_indice):
            minimo, massimo = 160, 260
        else:
            minimo, massimo = profilo["min_parole"], profilo["max_parole"]
        parole_minime_rimanenti += minimo
        parole_massime_rimanenti += massimo

    obiettivo_pagine = int(profilo["pagine_minime"])
    obiettivo_parole = obiettivo_pagine * PAROLE_PER_PAGINA_6X9
    tolleranza_pagine = float(profilo.get("tolleranza_pagine", 0.15))
    soglia_parole_tolleranza = math.ceil(obiettivo_parole * (1 - tolleranza_pagine))
    soglia_pagine_tolleranza = math.ceil(obiettivo_pagine * (1 - tolleranza_pagine))
    previsione_min_parole = parole_reali + parole_minime_rimanenti
    previsione_max_parole = parole_reali + parole_massime_rimanenti
    pagine_attuali = math.ceil(parole_reali / PAROLE_PER_PAGINA_6X9) if parole_reali else 0
    pagine_previste_min = math.ceil(previsione_min_parole / PAROLE_PER_PAGINA_6X9) if previsione_min_parole else 0
    pagine_previste_max = math.ceil(previsione_max_parole / PAROLE_PER_PAGINA_6X9) if previsione_max_parole else 0
    return {
        "parole_reali": parole_reali,
        "pagine_attuali": pagine_attuali,
        "sezioni_mancanti": sezioni_mancanti,
        "pagine_previste_min": pagine_previste_min,
        "pagine_previste_max": pagine_previste_max,
        "obiettivo_pagine": obiettivo_pagine,
        "obiettivo_parole": obiettivo_parole,
        "tolleranza_pagine": tolleranza_pagine,
        "soglia_pagine_tolleranza": soglia_pagine_tolleranza,
        "obiettivo_gia_raggiunto": parole_reali >= soglia_parole_tolleranza,
        "obiettivo_nominale_raggiunto": parole_reali >= obiettivo_parole,
        "obiettivo_realistico": previsione_max_parole >= soglia_parole_tolleranza,
    }


CHIAVE_MEMORIA_SEZIONI = "memoria_sezioni_editor"
# Archivio ridondante delle sezioni create dall'IA. È distinto dalla memoria
# dei widget Streamlit: se la sidebar si aggiorna durante una pausa, il testo
# generato resta disponibile e viene reidrato nell'editor e nell'anteprima.
# Viene cancellato insieme al progetto da RESET PROGETTO.
CHIAVE_MEMORIA_PROTETTA = "memoria_manoscritto_protetta"
# Archivio esclusivo della stesura completa. Non viene mai aggiornato da un
# widget vuoto o da un cambio della sezione selezionata: protegge soprattutto
# la prima sezione quando l'utente mette in pausa subito dopo il primo ciclo.
CHIAVE_ARCHIVIO_STESURA_COMPLETA = "archivio_stesura_completa"
# Elenco ordinato delle sezioni del manoscritto. Non dipende dall'indice né
# dai widget: mantiene visibili tutte le sezioni effettivamente create anche
# quando una pausa interrompe la stesura completa.
CHIAVE_REGISTRO_SEZIONI = "registro_sezioni_manoscritto"
CHIAVE_SEZIONI_DA_REIDRATARE = "sezioni_editor_da_reidratare"
CHIAVE_SEZIONE_EDITOR_ATTIVA = "sezione_editor_attiva"
CHIAVE_SELETTORE_EDITOR = "sezione_editor_selezionata"
# Ogni sezione ha una versione del relativo campo visuale. Quando l'AI crea
# o recupera un testo, la versione cambia: Streamlit non può quindi riusare
# una textarea vuota conservata dal browser prima della generazione.
CHIAVE_VERSIONI_WIDGET_SEZIONI = "versioni_widget_sezioni"


def chiave_widget_sezione(sezione):
    """Chiave esclusiva del campo visibile, distinta dalla memoria del testo."""
    versione = int(
        st.session_state.setdefault(CHIAVE_VERSIONI_WIDGET_SEZIONI, {}).get(sezione, 0)
    )
    return f"{chiave_sezione(sezione)}_v{versione}"


def leggi_sezione_memorizzata(sezione):
    """Legge una sezione dando priorità alla memoria stabile del progetto.

    Un widget Streamlit può conservare nel browser un testo precedente anche
    dopo una rigenerazione. Non deve mai sovrascrivere la nuova versione mentre
    vengono salvate o generate altre sezioni: le modifiche manuali passano
    invece dal callback ``sincronizza_modifica_manuale``.
    """
    chiave = chiave_sezione(sezione)
    # Prima della textarea e delle vecchie cache, leggiamo il manoscritto
    # unico. Una pausa o un rerun non può quindi rendere invisibile un testo
    # già concluso e registrato dal generatore.
    progetto = memoria_progetto_unica()
    testo_unico = str(progetto.get("contenuti", {}).get(sezione, "") or "")
    if testo_unico.strip():
        st.session_state.setdefault(CHIAVE_MEMORIA_SEZIONI, {})[sezione] = testo_unico
        st.session_state.setdefault(CHIAVE_MEMORIA_PROTETTA, {})[sezione] = testo_unico
        st.session_state[chiave] = testo_unico
        return testo_unico
    memoria = st.session_state.setdefault(CHIAVE_MEMORIA_SEZIONI, {})
    memoria_protetta = st.session_state.setdefault(CHIAVE_MEMORIA_PROTETTA, {})
    valore_widget = st.session_state.get(chiave, "") or st.session_state.get(chiave_sezione_precedente(sezione), "")
    if not str(memoria.get(sezione, "")).strip() and str(valore_widget).strip():
        memoria[sezione] = valore_widget
        memoria_protetta[sezione] = valore_widget
    # Una sezione generata non può sparire perché un widget si è ricreato
    # vuoto durante un rerun: la copia protetta ripristina quella editoriale.
    if not str(memoria.get(sezione, "")).strip() and str(memoria_protetta.get(sezione, "")).strip():
        memoria[sezione] = memoria_protetta[sezione]
    valore_memoria = memoria.get(sezione, "")
    if str(valore_memoria).strip():
        progetto["contenuti"][sezione] = valore_memoria
        return valore_memoria
    archivio_stesura = st.session_state.setdefault(CHIAVE_ARCHIVIO_STESURA_COMPLETA, {})
    testo = memoria_protetta.get(sezione, "") or archivio_stesura.get(sezione, "") or valore_widget or ""
    if str(testo).strip():
        progetto["contenuti"][sezione] = testo
    return testo


def scrivi_sezione_memorizzata(sezione, contenuto):
    """Scrive sempre sia nel widget sia nell'archivio del progetto.

    La doppia scrittura evita che il cambio di sezione o un rerun di Streamlit
    possa lasciare visibile solo l'ultimo testo generato.
    """
    testo = contenuto or ""
    progetto = memoria_progetto_unica()
    memoria = st.session_state.setdefault(CHIAVE_MEMORIA_SEZIONI, {})
    precedente = str(progetto.get("contenuti", {}).get(sezione, "") or "")
    progetto["contenuti"][sezione] = testo
    memoria[sezione] = testo
    st.session_state.setdefault(CHIAVE_MEMORIA_PROTETTA, {})[sezione] = testo
    # Questa chiave non è più usata da una textarea: può quindi essere
    # aggiornata in sicurezza anche mentre l'app effettua un rerun.
    st.session_state[chiave_sezione(sezione)] = testo
    if str(testo) != precedente:
        versioni = st.session_state.setdefault(CHIAVE_VERSIONI_WIDGET_SEZIONI, {})
        versioni[sezione] = int(versioni.get(sezione, 0)) + 1
    registro = st.session_state.setdefault(CHIAVE_REGISTRO_SEZIONI, [])
    if sezione not in registro:
        registro.append(sezione)
    # Non scriviamo direttamente nel widget: se il campo di testo è già stato
    # disegnato nel rerun corrente, Streamlit può ignorare o respingere la
    # modifica. Registriamo invece un ripristino certo al rerun successivo.
    da_reidratare = set(st.session_state.get(CHIAVE_SEZIONI_DA_REIDRATARE, []) or [])
    da_reidratare.add(sezione)
    st.session_state[CHIAVE_SEZIONI_DA_REIDRATARE] = list(da_reidratare)
    return testo


def scrivi_sezione_stesura_completa(sezione, contenuto):
    """Registra una sezione del job in un archivio ulteriore non modificabile dai widget."""
    testo = scrivi_sezione_memorizzata(sezione, contenuto)
    st.session_state.setdefault(CHIAVE_ARCHIVIO_STESURA_COMPLETA, {})[sezione] = testo
    return testo


def contenuto_memorizzato_puro(sezione):
    """Legge la copia stabile senza lasciarsi influenzare dal widget corrente."""
    contenuto_unico = memoria_progetto_unica().get("contenuti", {}).get(sezione, "")
    if str(contenuto_unico).strip():
        return contenuto_unico
    memoria = st.session_state.setdefault(CHIAVE_MEMORIA_SEZIONI, {})
    contenuto = memoria.get(sezione, "")
    if not str(contenuto).strip():
        contenuto = st.session_state.setdefault(CHIAVE_MEMORIA_PROTETTA, {}).get(sezione, "")
    if not str(contenuto).strip():
        contenuto = st.session_state.setdefault(CHIAVE_ARCHIVIO_STESURA_COMPLETA, {}).get(sezione, "")
    if not str(contenuto).strip():
        contenuto = st.session_state.get(chiave_sezione_precedente(sezione), "")
    return contenuto or ""


def elenco_sezioni_progetto(sezioni_base):
    """Unisce indice e memoria senza perdere testi già generati.

    Alcune sezioni possono esistere nella memoria dopo un ripristino CSV/cloud
    anche prima che il relativo widget venga ridisegnato. Anteprima, editor,
    controlli ed export devono quindi usare questo elenco, non soltanto le
    sezioni che Streamlit ha reso visibili nell'ultimo rerun.
    """
    risultato = []
    for sezione in [
        *(sezioni_base or []),
        *memoria_progetto_unica().get("contenuti", {}).keys(),
        *st.session_state.get(CHIAVE_REGISTRO_SEZIONI, []),
        *dict(st.session_state.get(CHIAVE_MEMORIA_SEZIONI, {}) or {}).keys(),
        *dict(st.session_state.get(CHIAVE_MEMORIA_PROTETTA, {}) or {}).keys(),
        *dict(st.session_state.get(CHIAVE_ARCHIVIO_STESURA_COMPLETA, {}) or {}).keys(),
    ]:
        if sezione and not sezione_dismessa(sezione) and sezione not in risultato:
            risultato.append(sezione)
    return risultato


def sincronizza_modifica_manuale(sezione, chiave_widget=None):
    """Callback dell'editor: conserva subito anche le modifiche digitate a mano."""
    chiave_da_leggere = chiave_widget or chiave_widget_sezione(sezione)
    contenuto = st.session_state.get(chiave_da_leggere, "")
    scrivi_sezione_memorizzata(sezione, contenuto)
    # L'archivio della stesura completa deve restare allineato anche dopo una
    # modifica manuale; altrimenti editor e anteprima potrebbero rileggere una
    # versione precedente del testo. La cancellazione esplicita resta possibile.
    archivio = st.session_state.setdefault(CHIAVE_ARCHIVIO_STESURA_COMPLETA, {})
    if sezione in archivio:
        if str(contenuto or "").strip():
            archivio[sezione] = contenuto
        else:
            archivio.pop(sezione, None)
    if not str(contenuto or "").strip():
        memoria_progetto_unica().get("contenuti", {}).pop(sezione, None)


def prepara_sezione_editor_selezionata():
    """Salva il campo lasciato e carica subito nel widget la sezione scelta.

    Streamlit conserva lo stato di ogni widget tra i rerun: questa funzione
    evita che un campo vuoto già presente nel browser prevalga sul testo che
    la memoria del progetto possiede davvero.
    """
    precedente = st.session_state.get(CHIAVE_SEZIONE_EDITOR_ATTIVA)
    if precedente:
        sincronizza_modifica_manuale(precedente)
    selezionata = st.session_state.get(CHIAVE_SELETTORE_EDITOR)
    if selezionata:
        st.session_state[chiave_sezione(selezionata)] = contenuto_memorizzato_puro(selezionata)
        st.session_state[CHIAVE_SEZIONE_EDITOR_ATTIVA] = selezionata


def reidrata_sezioni_memorizzate(sezioni):
    """Riporta nell'editor tutte le sezioni già salvate prima di renderizzare i widget."""
    progetto = memoria_progetto_unica()
    memoria_unica = progetto.get("contenuti", {}) or {}
    memoria = st.session_state.get(CHIAVE_MEMORIA_SEZIONI, {}) or {}
    memoria_protetta = st.session_state.get(CHIAVE_MEMORIA_PROTETTA, {}) or {}
    archivio_stesura = st.session_state.get(CHIAVE_ARCHIVIO_STESURA_COMPLETA, {}) or {}
    da_reidratare = set(st.session_state.get(CHIAVE_SEZIONI_DA_REIDRATARE, []) or [])
    for sezione in sezioni:
        contenuto = memoria_unica.get(sezione, "") or memoria.get(sezione)
        if not str(contenuto or "").strip():
            contenuto = memoria_protetta.get(sezione, "")
        if not str(contenuto or "").strip():
            contenuto = archivio_stesura.get(sezione, "")
        chiave = chiave_sezione(sezione)
        if not str(contenuto or "").strip():
            contenuto = st.session_state.get(chiave_sezione_precedente(sezione), "")
        if str(contenuto or "").strip() and (
            sezione in da_reidratare or not str(st.session_state.get(chiave, "")).strip()
        ):
            st.session_state[chiave] = contenuto
            memoria[sezione] = contenuto
            memoria_protetta[sezione] = contenuto
            progetto["contenuti"][sezione] = contenuto
            da_reidratare.discard(sezione)
    if da_reidratare:
        st.session_state[CHIAVE_SEZIONI_DA_REIDRATARE] = list(da_reidratare)
    else:
        st.session_state.pop(CHIAVE_SEZIONI_DA_REIDRATARE, None)


CAMPI_SALVATAGGIO_PROGETTO = {
    "titolo": "book_title",
    "autore": "book_author",
    "lingua": "editor_language",
    "genere": "book_genre",
    "tipologia_scrittura": "book_writing_style",
    "stile_racconto": "book_narrative_style",
    "punto_di_vista": "book_point_of_view",
    "obiettivo": "book_goal",
    "risultato_finale": "book_desired_result",
    "argomento": "book_plot",
    "approfondimenti": "book_further_details",
    # Personalizzazione facoltativa: sono dati editoriali veri e propri,
    # perciò vivono nella stessa memoria unica della sidebar e del manoscritto.
    "voce_personale": "book_personal_voice",
    "materiale_personale": "book_personal_material",
    "priorita_personali": "book_personal_priorities",
    "confini_personali": "book_personal_boundaries",
    "modalita_checkpoint": "book_personal_checkpoint_mode",
    "note_checkpoint": "book_personal_checkpoint_notes",
    "lunghezza": "profilo_lunghezza_stesura",
    "provider_ia": "provider_ia",
}
CHIAVE_MEMORIA_SIDEBAR = "memoria_sidebar_editor"


# La chat guidata non conserva un secondo progetto: quando l'utente approva
# la scheda, scrive direttamente nelle stesse chiavi della sidebar. Le funzioni
# sono fuori dalla UI perché i pulsanti usano callback Streamlit sicure anche
# dopo che i widget della sidebar sono stati creati.
CAMPI_CHAT_GUIDATA_SIDEBAR = {
    "TITOLO DEL LIBRO": "titolo",
    "NOME AUTORE": "autore",
    "LINGUA": "lingua",
    "GENERE LETTERARIO": "genere",
    "TIPOLOGIA SCRITTURA": "tipologia_scrittura",
    "STILE DI RACCONTO": "stile_racconto",
    "PUNTO DI VISTA": "punto_di_vista",
    "LUNGHEZZA DELLE SEZIONI": "lunghezza",
    "CERVELLO AI": "provider_ia",
    "OBIETTIVO DEL LIBRO": "obiettivo",
    "RISULTATO FINALE DESIDERATO": "risultato_finale",
    "TRAMA O ARGOMENTO": "argomento",
    "APPROFONDIMENTI": "approfondimenti",
    "VOCE O PROSPETTIVA PERSONALE": "voce_personale",
    "MATERIALI PERSONALI": "materiale_personale",
    "PRIORITÀ PERSONALI PER IL LETTORE": "priorita_personali",
    "CONFINI PERSONALI DA RISPETTARE": "confini_personali",
    "PAUSA GUIDATA DURANTE SCRIVI TUTTO IL LIBRO": "modalita_checkpoint",
}

CAMPI_CHAT_GUIDATA_OBBLIGATORI = {
    "titolo", "autore", "lingua", "genere", "tipologia_scrittura",
    "stile_racconto", "punto_di_vista", "lunghezza", "obiettivo",
    "risultato_finale", "argomento",
}

OPZIONI_CHAT_GUIDATA = {
    "genere": {
        "Saggio Scientifico", "Quiz Scientifico", "Manuale Tecnico", "Religioso / Teologico",
        "Spirituale / Esoterico", "Meditazione / Mindfulness", "Business & Marketing",
        "Economia e Finanza", "Romanzo Rosa", "Thriller / Noir", "Fantasy", "Fantascienza",
        "Manuale Psicologico", "Biografia", "Ricettario", "Test Prep (Preparazione Esami)",
        "Narrativo", "Romanzo Classico", "Contemporaneo", "Self-Help", "Manuale Pratico", "Storico",
    },
    "tipologia_scrittura": {
        "Standard", "Professionale Accademico", "Persuasivo (Neuromarketing Applicato)",
        "Conversazionale ed Empatico", "Scientifico Divulgativo", "Storytelling Immersivo",
        "Giornalistico d'Inchiesta", "Socratico (Dialogico / Riflessivo)", "Epico ed Evocativo",
        "Minimalista ed Essenziale",
    },
    "stile_racconto": {
        "Coinvolgente e Narrativo", "Tecnico e Analitico", "Ispirazionale e Motivante",
        "Socratico (Domanda/Risposta)", "Storytelling Emozionale",
        "Diretto e Pratico (Action-oriented)", "Storico e Documentale",
    },
    "punto_di_vista": {
        "Tu (Diretto, confidenziale e personale)", "Voi (Plurale, autorevole e rispettoso)",
        "Noi (Inclusivo, partecipativo e didattico)",
        "Impersonale / Terza Persona (Distaccato, analitico, oggettivo)",
    },
    "lunghezza": {"Compatto", "Standard KDP", "Approfondito"},
}


def _valore_scelta_chat_guidata(campo, valore):
    """Normalizza solo le selezioni tecniche; i testi editoriali restano liberi."""
    valore = re.sub(r"\s+", " ", str(valore or "")).strip()
    if campo == "lunghezza":
        valore_basso = valore.casefold()
        if valore_basso.startswith("compatto"):
            return "Compatto"
        if valore_basso.startswith("standard"):
            return "Standard KDP"
        if valore_basso.startswith("approfondito"):
            return "Approfondito"
    if campo == "modalita_checkpoint":
        valore_basso = valore.casefold()
        if not valore_basso or valore_basso.startswith("continua"):
            return ""
        if "ogni parte" in valore_basso:
            return "parti"
        if "parti" in valore_basso and "conclus" in valore_basso:
            return "parti_e_conclusione"
        if "conclus" in valore_basso:
            return "conclusione"
        return ""
    opzioni = OPZIONI_CHAT_GUIDATA.get(campo)
    if opzioni is not None and valore not in opzioni:
        return ""
    if campo == "lingua" and valore not in TRADUZIONI:
        return ""
    return valore


def estrai_scheda_chat_guidata(testo):
    """Legge soltanto la scheda finale con etichette note, mai testo libero."""
    testo = str(testo or "").replace("\r\n", "\n")
    risultati = {}
    corrispondenze = []
    for etichetta, campo in CAMPI_CHAT_GUIDATA_SIDEBAR.items():
        espressione = rf"(?im)^\s*{re.escape(etichetta)}\s*(?:\([^\n)]*\))?\s*:\s*"
        for match in re.finditer(espressione, testo):
            corrispondenze.append((match.start(), match.end(), campo))
    corrispondenze.sort(key=lambda dato: dato[0])
    for posizione, (_, fine, campo) in enumerate(corrispondenze):
        inizio_successivo = (
            corrispondenze[posizione + 1][0]
            if posizione + 1 < len(corrispondenze) else len(testo)
        )
        valore = testo[fine:inizio_successivo].strip(" \n:-")
        valore = _valore_scelta_chat_guidata(campo, valore)
        if valore:
            risultati[campo] = valore
    return risultati


def messaggio_benvenuto_chat_sidebar(lingua):
    """Primo messaggio gratuito, allineato alla lingua già scelta nel brief."""
    messaggi = {
        "Italiano": "Ciao! Che libro vuoi creare? Puoi descriverlo anche con una sola frase.",
        "English": "Hello! What book would you like to create? You can describe it in just one sentence.",
        "Español": "¡Hola! ¿Qué libro quieres crear? Puedes describirlo incluso con una sola frase.",
        "Français": "Bonjour ! Quel livre souhaitez-vous créer ? Vous pouvez le décrire en une seule phrase.",
        "Deutsch": "Hallo! Welches Buch möchten Sie erstellen? Sie können es auch mit nur einem Satz beschreiben.",
        "Română": "Bună! Ce carte vrei să creezi? O poți descrie și într-o singură propoziție.",
        "Русский": "Здравствуйте! Какую книгу вы хотите создать? Её можно описать даже одним предложением.",
        "العربية": "مرحبًا! ما الكتاب الذي تريد إنشاءه؟ يمكنك وصفه بجملة واحدة فقط.",
        "中文": "你好！你想创作什么书？即使只用一句话描述也可以。",
    }
    return messaggi.get(str(lingua or "").strip(), messaggi["Italiano"])


def avvia_chat_sidebar_guidata():
    lingua = st.session_state.get("editor_language", "Italiano") or "Italiano"
    st.session_state["chat_sidebar_attiva"] = True
    st.session_state["chat_sidebar_messaggi"] = [{
        "role": "assistant",
        "content": messaggio_benvenuto_chat_sidebar(lingua),
    }]
    st.session_state.pop("chat_sidebar_scheda_pronta", None)
    st.session_state.pop("chat_sidebar_esito_applicazione", None)
    st.session_state["chat_sidebar_input_nonce"] = 0


def avvia_chat_sidebar_da_guida():
    """Apre la chat esistente e chiede solo lo scorrimento al suo pannello."""
    avvia_chat_sidebar_guidata()
    st.session_state["chat_sidebar_scroll_request"] = True


def reimposta_chat_sidebar_guidata():
    for chiave in list(st.session_state):
        if chiave.startswith("chat_sidebar_"):
            st.session_state.pop(chiave, None)


def applica_scheda_chat_guidata(sovrascrivi=False):
    """Callback: compila le chiavi sidebar senza toccare il cervello scelto."""
    scheda = dict(st.session_state.get("chat_sidebar_scheda_pronta", {}) or {})
    applicati, mantenuti = [], []
    for campo, valore in scheda.items():
        if campo == "provider_ia":
            continue
        chiave_sidebar = CAMPI_SALVATAGGIO_PROGETTO.get(campo)
        if not chiave_sidebar or not str(valore).strip():
            continue
        valore_corrente = str(st.session_state.get(chiave_sidebar, "")).strip()
        if sovrascrivi or not valore_corrente:
            st.session_state[chiave_sidebar] = valore
            applicati.append(campo.replace("_", " "))
        else:
            mantenuti.append(campo.replace("_", " "))
    if applicati:
        # Mantiene la memoria unica coerente già nel callback, senza attendere
        # che il render successivo della sidebar ricostruisca i suoi widget.
        memoria = dict(st.session_state.get(CHIAVE_MEMORIA_SIDEBAR, {}) or {})
        for nome, chiave in CAMPI_SALVATAGGIO_PROGETTO.items():
            memoria[nome] = st.session_state.get(chiave, "")
        st.session_state[CHIAVE_MEMORIA_SIDEBAR] = memoria
        progetto = memoria_progetto_unica()
        progetto["sidebar"] = dict(memoria)
    messaggio = "Scheda applicata: " + (", ".join(applicati) if applicati else "nessun campo da aggiornare") + "."
    if mantenuti:
        messaggio += " Campi già compilati mantenuti: " + ", ".join(mantenuti) + "."
    st.session_state["chat_sidebar_esito_applicazione"] = messaggio

# Memoria unica del progetto. È la sola fonte autorevole per sidebar, indice,
# manoscritto, fonti e immagini. Le vecchie chiavi rimangono soltanto come
# copie di compatibilità con sessioni e CSV esportati in precedenza.
CHIAVE_PROGETTO_UNICO = "progetto_editoriale_unico"


MODALITA_CHECKPOINT_PERSONALE = ("", "parti", "conclusione", "parti_e_conclusione")


def etichette_personalizzazione(lingua):
    """Restituisce testi brevi per la personalizzazione, senza alterare il brief obbligatorio."""
    testi = {
        "Italiano": {
            "titolo": "✍️ Personalizza il tuo libro (facoltativo)",
            "intro": "Aggiungi ciò che rende il libro davvero tuo. Queste risposte non consumano crediti e vengono salvate con progetto, CSV e ripristino.",
            "voce": "Quale voce, esperienza o prospettiva deve distinguere questo libro?",
            "materiale": "Quali episodi, casi, esempi o materiali personali vuoi valorizzare?",
            "priorita": "Quali domande o risultati vuoi che il lettore trovi con particolare chiarezza?",
            "confini": "Cosa deve evitare il libro: tono, promesse, temi o esempi non adatti?",
            "check": "Durante Scrivi tutto il libro",
            "automatico": "Continua automaticamente (consigliato)",
            "parti": "Fermati prima di ogni Parte", "conclusione": "Fermati prima della conclusione",
            "parti_e_conclusione": "Fermati prima delle Parti e della conclusione",
            "nota": "Prima di proseguire, vuoi aggiungere un dettaglio personale per questa sezione?",
        },
        "English": {
            "titolo": "✍️ Personalize your book (optional)", "intro": "Add what makes the book truly yours. These answers use no credits and are saved with the project, CSV and restore.",
            "voce": "Which voice, experience or perspective should distinguish this book?", "materiale": "Which episodes, cases, examples or personal material should be used?",
            "priorita": "Which questions or outcomes must be especially clear to the reader?", "confini": "What should the book avoid: tone, claims, topics or unsuitable examples?",
            "check": "During Write the whole book", "automatico": "Continue automatically (recommended)", "parti": "Pause before each Part",
            "conclusione": "Pause before the conclusion", "parti_e_conclusione": "Pause before Parts and conclusion",
            "nota": "Before continuing, would you like to add a personal detail for this section?",
        },
        "Español": {
            "titolo": "✍️ Personaliza tu libro (opcional)", "intro": "Añade lo que hace que el libro sea realmente tuyo. Estas respuestas no consumen créditos y se guardan con el proyecto, CSV y restauración.",
            "voce": "¿Qué voz, experiencia o perspectiva debe distinguir este libro?", "materiale": "¿Qué episodios, casos, ejemplos o materiales personales deseas destacar?",
            "priorita": "¿Qué preguntas o resultados deben quedar especialmente claros para el lector?", "confini": "¿Qué debe evitar el libro: tono, promesas, temas o ejemplos inadecuados?",
            "check": "Durante Escribir todo el libro", "automatico": "Continuar automáticamente (recomendado)", "parti": "Pausar antes de cada Parte", "conclusione": "Pausar antes de la conclusión", "parti_e_conclusione": "Pausar antes de las Partes y de la conclusión", "nota": "Antes de continuar, ¿quieres añadir un detalle personal para esta sección?",
        },
        "Français": {
            "titolo": "✍️ Personnalisez votre livre (facultatif)", "intro": "Ajoutez ce qui rend le livre vraiment personnel. Ces réponses ne consomment pas de crédits et sont conservées avec le projet, le CSV et la restauration.",
            "voce": "Quelle voix, expérience ou perspective doit distinguer ce livre ?", "materiale": "Quels épisodes, cas, exemples ou matériaux personnels souhaitez-vous valoriser ?",
            "priorita": "Quelles questions ou quels résultats doivent être particulièrement clairs pour le lecteur ?", "confini": "Que doit éviter le livre : ton, promesses, sujets ou exemples inadaptés ?",
            "check": "Pendant Écrire tout le livre", "automatico": "Continuer automatiquement (recommandé)", "parti": "Pause avant chaque Partie", "conclusione": "Pause avant la conclusion", "parti_e_conclusione": "Pause avant les Parties et la conclusion", "nota": "Avant de continuer, souhaitez-vous ajouter un détail personnel pour cette section ?",
        },
        "Deutsch": {
            "titolo": "✍️ Personalisiere dein Buch (optional)", "intro": "Füge hinzu, was das Buch wirklich zu deinem macht. Diese Antworten verbrauchen keine Credits und werden mit Projekt, CSV und Wiederherstellung gespeichert.",
            "voce": "Welche Stimme, Erfahrung oder Perspektive soll dieses Buch prägen?", "materiale": "Welche Episoden, Fälle, Beispiele oder persönlichen Materialien sollen berücksichtigt werden?",
            "priorita": "Welche Fragen oder Ergebnisse sollen für Leser besonders klar sein?", "confini": "Was soll das Buch vermeiden: Ton, Versprechen, Themen oder ungeeignete Beispiele?",
            "check": "Während Ganzes Buch schreiben", "automatico": "Automatisch fortsetzen (empfohlen)", "parti": "Vor jedem Teil anhalten", "conclusione": "Vor dem Schluss anhalten", "parti_e_conclusione": "Vor Teilen und Schluss anhalten", "nota": "Möchtest du vor dem Fortsetzen ein persönliches Detail für diesen Abschnitt ergänzen?",
        },
        "Română": {
            "titolo": "✍️ Personalizează cartea (opțional)", "intro": "Adaugă ceea ce face cartea cu adevărat a ta. Aceste răspunsuri nu consumă credite și se salvează cu proiectul, CSV-ul și restaurarea.",
            "voce": "Ce voce, experiență sau perspectivă trebuie să distingă această carte?", "materiale": "Ce episoade, cazuri, exemple sau materiale personale vrei să valorifici?",
            "priorita": "Ce întrebări sau rezultate trebuie să fie deosebit de clare pentru cititor?", "confini": "Ce trebuie să evite cartea: ton, promisiuni, teme sau exemple nepotrivite?",
            "check": "În timpul Scrie toată cartea", "automatico": "Continuă automat (recomandat)", "parti": "Oprește înainte de fiecare Parte", "conclusione": "Oprește înainte de concluzie", "parti_e_conclusione": "Oprește înainte de părți și concluzie", "nota": "Înainte de a continua, vrei să adaugi un detaliu personal pentru această secțiune?",
        },
        "Русский": {
            "titolo": "✍️ Персонализируйте книгу (необязательно)", "intro": "Добавьте то, что сделает книгу по-настоящему вашей. Эти ответы не расходуют кредиты и сохраняются вместе с проектом, CSV и восстановлением.",
            "voce": "Какой голос, опыт или взгляд должен отличать эту книгу?", "materiale": "Какие эпизоды, случаи, примеры или личные материалы вы хотите использовать?",
            "priorita": "Какие вопросы или результаты должны быть особенно понятны читателю?", "confini": "Чего должна избегать книга: тона, обещаний, тем или неподходящих примеров?",
            "check": "Во время написания всей книги", "automatico": "Продолжать автоматически (рекомендуется)", "parti": "Остановиться перед каждой частью", "conclusione": "Остановиться перед заключением", "parti_e_conclusione": "Остановиться перед частями и заключением", "nota": "Перед продолжением хотите добавить личную деталь для этого раздела?",
        },
        "العربية": {
            "titolo": "✍️ خصص كتابك (اختياري)", "intro": "أضف ما يجعل الكتاب ملكك حقاً. هذه الإجابات لا تستهلك أرصدة وتُحفظ مع المشروع وملف CSV والاستعادة.",
            "voce": "ما الصوت أو الخبرة أو المنظور الذي يجب أن يميز هذا الكتاب؟", "materiale": "ما الأحداث أو الحالات أو الأمثلة أو المواد الشخصية التي تريد إبرازها؟",
            "priorita": "ما الأسئلة أو النتائج التي يجب أن تكون واضحة للقارئ بشكل خاص؟", "confini": "ما الذي يجب أن يتجنبه الكتاب: النبرة أو الوعود أو الموضوعات أو الأمثلة غير المناسبة؟",
            "check": "أثناء كتابة الكتاب كاملاً", "automatico": "المتابعة تلقائياً (موصى به)", "parti": "توقف قبل كل جزء", "conclusione": "توقف قبل الخاتمة", "parti_e_conclusione": "توقف قبل الأجزاء والخاتمة", "nota": "قبل المتابعة، هل تريد إضافة تفصيل شخصي لهذا القسم؟",
        },
        "中文": {
            "titolo": "✍️ 个性化你的图书（可选）", "intro": "添加真正能体现你特色的内容。这些回答不消耗积分，会随项目、CSV 和恢复功能一起保存。",
            "voce": "什么声音、经历或视角应使这本书与众不同？", "materiale": "你希望重点使用哪些事件、案例、示例或个人材料？",
            "priorita": "读者最需要清楚理解哪些问题或成果？", "confini": "本书应避免什么：语气、承诺、主题或不合适的示例？",
            "check": "在写完整本书期间", "automatico": "自动继续（推荐）", "parti": "每一部分前暂停", "conclusione": "结论前暂停", "parti_e_conclusione": "各部分和结论前暂停", "nota": "继续之前，是否要为本节添加个人细节？",
        },
    }
    return testi.get(lingua, testi["Italiano"])


def brief_personalizzazione_progetto(sezione=""):
    """Trasforma i dati facoltativi in istruzioni editoriali, mai in testo da copiare.

    Il blocco è usato da ricerca, indice e stesura. Limiti di lunghezza evitano
    che una nota personale renda il prompt sproporzionato o ripetitivo.
    """
    sidebar = sidebar_memorizzata_corrente()
    elementi = (
        ("Voce o prospettiva dell'autore", sidebar.get("voce_personale", "")),
        ("Esperienze, casi o materiali dell'autore", sidebar.get("materiale_personale", "")),
        ("Priorità per il lettore", sidebar.get("priorita_personali", "")),
        ("Confini da rispettare", sidebar.get("confini_personali", "")),
        ("Indicazioni aggiunte durante la stesura", sidebar.get("note_checkpoint", "")),
    )
    righe = [f"- {titolo}: {str(valore).strip()[:1200]}" for titolo, valore in elementi if str(valore).strip()]
    if not righe:
        return "Nessuna personalizzazione aggiuntiva: applica il brief editoriale ordinario."
    riferimento = f" per la sezione '{sezione}'" if str(sezione).strip() else ""
    return (
        f"PERSONALIZZAZIONE FACOLTATIVA{riferimento}\n" + "\n".join(righe) + "\n"
        "Usa queste indicazioni soltanto quando sono pertinenti. Rielabora tutto in modo originale: "
        "non presentare esempi personali come fatti non forniti e non ripetere la stessa nota in ogni sezione."
    )


def modalita_checkpoint_personale():
    valore = str(st.session_state.get("book_personal_checkpoint_mode", "") or "").strip()
    return valore if valore in MODALITA_CHECKPOINT_PERSONALE else ""


def sezione_conclusiva_personale(sezione):
    return bool(re.search(
        r"(?i)\b(conclusione|conclusion|conclusión|schluss|epilogo|epilogue|conclusie|заключ|итог|خاتمة|结论|总结)\b",
        str(sezione or ""),
    ))


def richiede_checkpoint_personale(sezione):
    """Decide se fermare una coda senza bloccare un checkpoint gia' superato."""
    normalizza = lambda valore: re.sub(r"\s+", " ", str(valore or "").strip()).casefold()
    sezione_norm = normalizza(sezione)

    # ``RIPRENDI GENERAZIONE`` registra anche questo lasciapassare esplicito.
    # E' separato dal registro storico per funzionare pure con vecchie sessioni
    # dove il titolo della Parte aveva spazi o segni diversi dopo un rerun.
    prosegui_forzato = normalizza(
        st.session_state.get("job_scrittura_checkpoint_da_superare", "")
    )
    if prosegui_forzato and prosegui_forzato == sezione_norm:
        st.session_state.pop("job_scrittura_checkpoint_da_superare", None)
        return False

    modalita = modalita_checkpoint_personale()
    if not modalita:
        return False
    gia_superati = {
        normalizza(voce)
        for voce in (st.session_state.get("job_scrittura_checkpoint_superati", []) or [])
    }
    if sezione_norm in gia_superati:
        return False
    if modalita in {"parti", "parti_e_conclusione"} and tipo_sezione_editoriale(sezione) == "parte":
        return True
    return modalita in {"conclusione", "parti_e_conclusione"} and sezione_conclusiva_personale(sezione)


def registra_nota_checkpoint(sezione, nota):
    """Conserva l'eventuale risposta senza sovrascrivere le indicazioni precedenti."""
    nota = str(nota or "").strip()
    if not nota:
        return
    precedente = str(st.session_state.get("book_personal_checkpoint_notes", "") or "").strip()
    nuova_nota = f"[{sezione}] {nota}"
    if nuova_nota not in precedente:
        st.session_state["book_personal_checkpoint_notes"] = (
            f"{precedente}\n{nuova_nota}".strip() if precedente else nuova_nota
        )
    sidebar_memorizzata_corrente()


def memoria_progetto_unica():
    """Restituisce la fotografia editoriale unica, riparando dati parziali."""
    progetto = st.session_state.setdefault(CHIAVE_PROGETTO_UNICO, {})
    if not isinstance(progetto, dict):
        progetto = {}
        st.session_state[CHIAVE_PROGETTO_UNICO] = progetto
    for nome in ("sidebar", "contenuti", "fonti", "immagini"):
        if not isinstance(progetto.get(nome), dict):
            progetto[nome] = {}
    progetto["indice"] = str(progetto.get("indice", "") or "")
    return progetto


def sidebar_memorizzata_corrente():
    """Restituisce tutti i campi editoriali, inclusi quelli non visibili dopo un rerun."""
    progetto = memoria_progetto_unica()
    memoria = dict(progetto.get("sidebar", {}) or {})
    memoria.update(dict(st.session_state.get(CHIAVE_MEMORIA_SIDEBAR, {}) or {}))
    for nome, chiave in CAMPI_SALVATAGGIO_PROGETTO.items():
        if chiave in st.session_state:
            memoria[nome] = st.session_state.get(chiave, "")
    st.session_state[CHIAVE_MEMORIA_SIDEBAR] = dict(memoria)
    progetto["sidebar"] = dict(memoria)
    return memoria


# ======================================================================================================================
# ADATTATORI DEL MODULO project_memory
# ======================================================================================================================
# Mantengono invariati i nomi usati dall'interfaccia e dalle versioni salvate,
# mentre la logica effettiva vive ora nel modulo dedicato. È una separazione
# intenzionalmente non distruttiva: nessun widget, CSV o progetto esistente
# cambia chiave o formato durante la migrazione.
def memoria_progetto_unica():
    return memoria_core.memoria_progetto_unica(st.session_state)


def sidebar_memorizzata_corrente():
    return memoria_core.sidebar_memorizzata_corrente(st.session_state)


def chiave_widget_sezione(sezione):
    return memoria_core.chiave_widget_sezione(st.session_state, sezione, chiave_sezione)


def leggi_sezione_memorizzata(sezione):
    return memoria_core.leggi_sezione_memorizzata(
        st.session_state, sezione, chiave_sezione, chiave_sezione_precedente
    )


def scrivi_sezione_memorizzata(sezione, contenuto):
    return memoria_core.scrivi_sezione_memorizzata(
        st.session_state, sezione, contenuto, chiave_sezione
    )


def scrivi_sezione_stesura_completa(sezione, contenuto):
    return memoria_core.scrivi_sezione_stesura_completa(
        st.session_state, sezione, contenuto, chiave_sezione
    )


def contenuto_memorizzato_puro(sezione):
    return memoria_core.contenuto_memorizzato_puro(
        st.session_state, sezione, chiave_sezione_precedente
    )


def elenco_sezioni_progetto(sezioni_base):
    return memoria_core.elenco_sezioni_progetto(
        st.session_state, sezioni_base, sezione_dismessa
    )


def sincronizza_modifica_manuale(sezione, chiave_widget=None):
    chiave_da_leggere = chiave_widget or chiave_widget_sezione(sezione)
    return memoria_core.sincronizza_modifica_manuale(
        st.session_state, sezione, chiave_da_leggere, chiave_sezione
    )


def prepara_sezione_editor_selezionata():
    return memoria_core.prepara_sezione_editor_selezionata(
        st.session_state, chiave_sezione, chiave_sezione_precedente
    )


def reidrata_sezioni_memorizzate(sezioni):
    return memoria_core.reidrata_sezioni_memorizzate(
        st.session_state, sezioni, chiave_sezione, chiave_sezione_precedente
    )


def immagini_per_snapshot_cloud(immagini):
    """Rende serializzabili le immagini caricate dall'utente per il cloud.

    Supabase salva JSON e non può ricevere direttamente byte. Codifichiamo solo
    le immagini già associate alle sezioni, senza crearne o modificarne alcuna.
    """
    risultato = {}
    for sezione, immagine in dict(immagini or {}).items():
        dati = dict(immagine or {})
        raw = dati.pop("bytes", None)
        if isinstance(raw, (bytes, bytearray)) and raw:
            dati["bytes_b64"] = base64.b64encode(bytes(raw)).decode("ascii")
        if dati:
            risultato[str(sezione)] = dati
    return risultato


def immagini_da_snapshot_cloud(immagini):
    """Ripristina in memoria le immagini esterne salvate in base64."""
    risultato = {}
    for sezione, immagine in dict(immagini or {}).items():
        dati = dict(immagine or {})
        raw_b64 = dati.pop("bytes_b64", "")
        if raw_b64:
            try:
                dati["bytes"] = base64.b64decode(raw_b64)
            except (ValueError, TypeError):
                continue
        if isinstance(dati.get("bytes"), bytearray):
            dati["bytes"] = bytes(dati["bytes"])
        if isinstance(dati.get("bytes"), bytes) and dati["bytes"]:
            risultato[str(sezione)] = dati
    return risultato


def esporta_progetto_editoriale_csv():
    """Esporta una fotografia completa e verificabile dell'intero progetto.

    Il CSV mantiene le righe leggibili in Excel, ma contiene anche una
    fotografia codificata completa. L'indice non dipende più da una sola riga
    lunga o dallo stato del widget: al reimport viene sempre recuperato dalla
    fotografia integra.
    """
    progetto = memoria_progetto_unica()
    sidebar = sidebar_memorizzata_corrente()
    indice = str(progetto.get("indice", "") or st.session_state.get("indice_raw", "") or "")
    contenuti = dict(progetto.get("contenuti", {}) or {})
    for sezione in elenco_sezioni_progetto(st.session_state.get("lista_capitoli", [])):
        testo = leggi_sezione_memorizzata(sezione)
        if str(testo).strip():
            contenuti[sezione] = str(testo)
    fonti = dict(progetto.get("fonti", {}) or {})
    for chiave in (
        "conoscenza_extra", "scheda_fonti", "dossier_fonti_ai", "brief_fonti_originale",
        "dossier_ricerca_preliminare", "registro_fonti_web", "firma_ricerca_preliminare",
    ):
        if chiave in st.session_state:
            valore_fonte = st.session_state.get(chiave, "")
            fonti[chiave] = (
                conserva_solo_fonti_web_selezionate(valore_fonte)
                if chiave == "registro_fonti_web" else valore_fonte
            )

    fotografia = {
        "sidebar": sidebar,
        "indice": indice,
        "contenuti": contenuti,
        "fonti": fonti,
        "immagini": st.session_state.get("immagini_capitoli", {}) or {},
    }
    return esporta_fotografia_csv(fotografia)


def importa_progetto_editoriale_csv(file_caricato):
    """Importa integralmente un CSV Scrittore Site, anche per libri lunghi.

    La fotografia v2 conserva in un unico campo CSV sidebar, indice, sezioni,
    fonti e immagini. Prima della lettura alziamo quindi il limite del parser,
    senza modificare il formato di esportazione o perdere compatibilità con i
    CSV già scaricati.
    """
    dati_grezzi = file_caricato.getvalue()
    return importa_fotografia_csv(dati_grezzi, CAMPI_SALVATAGGIO_PROGETTO)


def mostra_memoria_visiva_progetto():
    """Pannello leggibile che rende verificabile la memoria reale del progetto."""
    sidebar = sidebar_memorizzata_corrente()
    contenuti = dict(st.session_state.get(CHIAVE_ARCHIVIO_STESURA_COMPLETA, {}) or {})
    contenuti.update(dict(st.session_state.get(CHIAVE_MEMORIA_PROTETTA, {}) or {}))
    contenuti.update(dict(st.session_state.get(CHIAVE_MEMORIA_SEZIONI, {}) or {}))
    # Includiamo anche un testo presente nel widget ma non ancora confluito
    # nella cache, per esempio subito dopo una modifica manuale.
    for chiave, valore in st.session_state.items():
        # I campi visuali temporanei terminano in _vN: non sono sezioni
        # aggiuntive e non devono comparire come voci duplicate nella memoria.
        if chiave.startswith("txt_") and not re.search(r"_v\d+$", chiave) and str(valore).strip():
            titolo = next((s for s in contenuti if chiave_sezione(s) == chiave), chiave[4:].replace("_", " "))
            contenuti[titolo] = valore

    campi_compilati = sum(1 for valore in sidebar.values() if str(valore).strip())
    st.markdown("### 🧠 Memoria del progetto")
    registro_fonti = conserva_solo_fonti_web_selezionate(
        st.session_state.get("registro_fonti_web", "")
    )
    numero_fonti = len(re.findall(r"https?://\\S+", registro_fonti))
    if registro_fonti and not numero_fonti:
        numero_fonti = len([riga for riga in registro_fonti.splitlines() if riga.strip()])
    st.caption("Qui vedi ciò che il software conserva: campi della sidebar, indice, testi e fonti di ricerca, inclusi quelli modificati a mano.")
    col_a, col_b, col_c = st.columns(3)
    col_a.metric("Sidebar", f"{campi_compilati}/{len(CAMPI_SALVATAGGIO_PROGETTO)}")
    col_b.metric("Sezioni salvate", len(contenuti))
    col_c.metric("Fonti web", numero_fonti)
    st.caption("Indice: ✓ salvato" if st.session_state.get("indice_raw", "").strip() else "Indice: non ancora creato")

    with st.expander("Visualizza dati salvati", expanded=False):
        for nome in CAMPI_SALVATAGGIO_PROGETTO:
            valore = str(sidebar.get(nome, "")).strip() or "—"
            st.caption(f"**{nome.replace('_', ' ').capitalize()}:** {valore}")
        if st.session_state.get("indice_raw", "").strip():
            st.text_area(campo_ui("indice_salvato"), value=st.session_state["indice_raw"], height=160, disabled=True, key="memoria_indice")
        if registro_fonti:
            st.text_area(campo_ui("fonti_salvate"), value=registro_fonti, height=150, disabled=True, key="memoria_fonti_web")
        if contenuti:
            sezioni = list(contenuti)
            scelta = st.selectbox(campo_ui("sezione_salvata"), sezioni, key="memoria_sezione_scelta")
            testo = str(contenuti.get(scelta, ""))
            st.caption(f"{len(testo.split())} parole conservate")
            st.text_area(
                campo_ui("testo_salvato"), value=testo, height=230, disabled=True,
                key=f"memoria_testo_{hashlib.sha256(scelta.encode('utf-8')).hexdigest()[:12]}",
            )
        else:
            st.info("Non ci sono ancora sezioni salvate.")


def etichette_centro_progetto(lingua):
    """Etichette sintetiche del riepilogo, separate dalla logica editoriale.

    Il Centro del progetto e' soltanto una lettura della memoria condivisa:
    non crea richieste IA, non salva e non modifica l'indice o il manoscritto.
    """
    testi = {
        "Italiano": {
            "titolo": "🧭 Centro del progetto", "sidebar": "Brief", "indice": "Indice",
            "manoscritto": "Manoscritto", "fonti": "Fonti", "stato": "Stato operativo",
            "prossimo": "Prossimo passo", "dettaglio": "Dettaglio", "configura": "Completa il brief nella sidebar.",
            "crea_indice": "Genera l'indice professionale.", "continua": "Continua la stesura dalle sezioni rimaste.",
            "controlla": "Esegui il controllo finale prima dell'esportazione.",
            "in_corso": "Stesura automatica in corso", "in_pausa": "Stesura in pausa",
            "fermato": "Stesura interrotta", "indice_pronto": "Indice pronto", "nessun": "nessuna",
            "nota": "Riepilogo informativo: non genera né modifica contenuti.",
        },
        "English": {
            "titolo": "🧭 Project center", "sidebar": "Brief", "indice": "Outline",
            "manoscritto": "Manuscript", "fonti": "Sources", "stato": "Current status",
            "prossimo": "Next step", "dettaglio": "Detail", "configura": "Complete the brief in the sidebar.",
            "crea_indice": "Generate the professional outline.", "continua": "Continue writing the remaining sections.",
            "controlla": "Run the final review before exporting.",
            "in_corso": "Automatic writing in progress", "in_pausa": "Writing paused",
            "fermato": "Writing stopped", "indice_pronto": "Outline ready", "nessun": "none",
            "nota": "Informational summary: it does not generate or change content.",
        },
        "Español": {
            "titolo": "🧭 Centro del proyecto", "sidebar": "Brief", "indice": "Índice",
            "manoscritto": "Manuscrito", "fonti": "Fuentes", "stato": "Estado actual",
            "prossimo": "Siguiente paso", "dettaglio": "Detalle", "configura": "Completa el brief en la barra lateral.",
            "crea_indice": "Genera el índice profesional.", "continua": "Continúa la redacción de las secciones restantes.",
            "controlla": "Ejecuta la revisión final antes de exportar.",
            "in_corso": "Redacción automática en curso", "in_pausa": "Redacción en pausa",
            "fermato": "Redacción detenida", "indice_pronto": "Índice listo", "nessun": "ninguna",
            "nota": "Resumen informativo: no genera ni modifica contenido.",
        },
        "Français": {
            "titolo": "🧭 Centre du projet", "sidebar": "Brief", "indice": "Plan",
            "manoscritto": "Manuscrit", "fonti": "Sources", "stato": "État actuel",
            "prossimo": "Prochaine étape", "dettaglio": "Détail", "configura": "Complétez le brief dans la barre latérale.",
            "crea_indice": "Générez le plan professionnel.", "continua": "Continuez les sections restantes.",
            "controlla": "Lancez le contrôle final avant l’export.",
            "in_corso": "Rédaction automatique en cours", "in_pausa": "Rédaction en pause",
            "fermato": "Rédaction arrêtée", "indice_pronto": "Plan prêt", "nessun": "aucune",
            "nota": "Résumé informatif : il ne génère ni ne modifie le contenu.",
        },
        "Deutsch": {
            "titolo": "🧭 Projektzentrale", "sidebar": "Briefing", "indice": "Gliederung",
            "manoscritto": "Manuskript", "fonti": "Quellen", "stato": "Aktueller Status",
            "prossimo": "Nächster Schritt", "dettaglio": "Detail", "configura": "Briefing in der Seitenleiste vervollständigen.",
            "crea_indice": "Professionelle Gliederung erstellen.", "continua": "Verbleibende Abschnitte weiter schreiben.",
            "controlla": "Endkontrolle vor dem Export ausführen.",
            "in_corso": "Automatisches Schreiben läuft", "in_pausa": "Schreiben pausiert",
            "fermato": "Schreiben gestoppt", "indice_pronto": "Gliederung bereit", "nessun": "keine",
            "nota": "Nur Übersicht: Es werden keine Inhalte erstellt oder geändert.",
        },
        "Română": {
            "titolo": "🧭 Centrul proiectului", "sidebar": "Brief", "indice": "Cuprins",
            "manoscritto": "Manuscris", "fonti": "Surse", "stato": "Stare curentă",
            "prossimo": "Pasul următor", "dettaglio": "Detaliu", "configura": "Completează brief-ul din bara laterală.",
            "crea_indice": "Generează cuprinsul profesional.", "continua": "Continuă scrierea secțiunilor rămase.",
            "controlla": "Rulează controlul final înainte de export.",
            "in_corso": "Scriere automată în curs", "in_pausa": "Scriere întreruptă temporar",
            "fermato": "Scriere oprită", "indice_pronto": "Cuprins pregătit", "nessun": "niciuna",
            "nota": "Rezumat informativ: nu generează și nu modifică conținutul.",
        },
        "Русский": {
            "titolo": "🧭 Центр проекта", "sidebar": "Бриф", "indice": "Оглавление",
            "manoscritto": "Рукопись", "fonti": "Источники", "stato": "Текущий статус",
            "prossimo": "Следующий шаг", "dettaglio": "Детали", "configura": "Заполните бриф на боковой панели.",
            "crea_indice": "Создайте профессиональное оглавление.", "continua": "Продолжите создание оставшихся разделов.",
            "controlla": "Выполните итоговую проверку перед экспортом.",
            "in_corso": "Автоматическое написание идёт", "in_pausa": "Написание приостановлено",
            "fermato": "Написание остановлено", "indice_pronto": "Оглавление готово", "nessun": "нет",
            "nota": "Информационная сводка: она не создаёт и не изменяет содержимое.",
        },
        "العربية": {
            "titolo": "🧭 مركز المشروع", "sidebar": "الموجز", "indice": "الفهرس",
            "manoscritto": "المخطوطة", "fonti": "المصادر", "stato": "الحالة الحالية",
            "prossimo": "الخطوة التالية", "dettaglio": "التفاصيل", "configura": "أكمل الموجز في الشريط الجانبي.",
            "crea_indice": "أنشئ الفهرس الاحترافي.", "continua": "تابع كتابة الأقسام المتبقية.",
            "controlla": "شغّل الفحص النهائي قبل التصدير.",
            "in_corso": "الكتابة التلقائية جارية", "in_pausa": "الكتابة متوقفة مؤقتاً",
            "fermato": "تم إيقاف الكتابة", "indice_pronto": "الفهرس جاهز", "nessun": "لا يوجد",
            "nota": "ملخص معلوماتي: لا ينشئ المحتوى ولا يعدّله.",
        },
        "中文": {
            "titolo": "🧭 项目中心", "sidebar": "写作简报", "indice": "目录",
            "manoscritto": "手稿", "fonti": "来源", "stato": "当前状态",
            "prossimo": "下一步", "dettaglio": "详情", "configura": "请在侧边栏完成写作简报。",
            "crea_indice": "生成专业目录。", "continua": "继续撰写其余部分。",
            "controlla": "导出前运行最终检查。",
            "in_corso": "正在自动写作", "in_pausa": "写作已暂停",
            "fermato": "写作已停止", "indice_pronto": "目录已准备好", "nessun": "无",
            "nota": "信息摘要：不会生成或修改任何内容。",
        },
    }
    return testi.get(lingua, testi["English"])


# Il percorso è puramente visivo: riassume le funzioni già disponibili nelle
# tab, senza spostare l'utente né modificare i dati del progetto.
PASSI_PERCORSO_PROGETTO = {
    "Italiano": ("Idea", "Indice", "Scrittura", "Controllo", "Esporta"),
    "English": ("Idea", "Outline", "Writing", "Review", "Export"),
    "Español": ("Idea", "Índice", "Escritura", "Revisión", "Exportar"),
    "Français": ("Idée", "Plan", "Écriture", "Contrôle", "Exporter"),
    "Deutsch": ("Idee", "Gliederung", "Schreiben", "Prüfen", "Export"),
    "Română": ("Idee", "Cuprins", "Scriere", "Verificare", "Export"),
    "Русский": ("Идея", "План", "Написание", "Проверка", "Экспорт"),
    "العربية": ("الفكرة", "الفهرس", "الكتابة", "المراجعة", "التصدير"),
    "中文": ("想法", "目录", "写作", "检查", "导出"),
}

# Istruzioni concrete accanto al Centro del progetto. Descrivono soltanto
# pulsanti e schede già esistenti: non attivano alcuna funzione in autonomia.
ISTRUZIONI_PROSSIMO_PASSO = {
    "Italiano": {
        "brief": "Apri la sidebar a sinistra; su telefono tocca ☰. Completa i campi obbligatori del libro.",
        "indice": "Apri la scheda 1. Indice e premi «Genera indice professionale».",
        "scrittura": "Apri la scheda 2. Scrittura & Quiz e premi «Scrivi tutto il libro», oppure «Scrivi contenuto» per una sola sezione.",
        "pausa": "Resta nella scheda 2. Scrittura & Quiz e premi «Riprendi generazione» quando vuoi continuare. Il testo già creato resta salvato.",
        "controllo": "Apri la scheda 4. Importa / Esporta / Copyright, esegui «Controlla completezza del manoscritto», poi scarica Word o PDF.",
    },
    "English": {
        "brief": "Open the sidebar on the left; on a phone tap ☰. Complete the required book fields.",
        "indice": "Open tab 1. Outline and select “Generate professional outline”.",
        "scrittura": "Open tab 2. Writing & Quiz and select “Write the whole book”, or “Write content” for one section.",
        "pausa": "Stay in tab 2. Writing & Quiz and select “Resume generation” when you want to continue. Existing text stays saved.",
        "controllo": "Open tab 4. Import / Export / Copyright, run “Check manuscript completeness”, then download Word or PDF.",
    },
    "Español": {
        "brief": "Abre la barra lateral a la izquierda; en el teléfono toca ☰. Completa los campos obligatorios del libro.",
        "indice": "Abre la pestaña 1. Índice y pulsa «Generar índice profesional».",
        "scrittura": "Abre la pestaña 2. Escritura y Quiz y pulsa «Escribir todo el libro» o «Escribir contenido» para una sección.",
        "pausa": "Quédate en la pestaña 2. Escritura y Quiz y pulsa «Reanudar generación» cuando quieras continuar. El texto creado queda guardado.",
        "controllo": "Abre la pestaña 4. Importar / Exportar / Copyright, ejecuta «Comprobar la integridad del manuscrito» y descarga Word o PDF.",
    },
    "Français": {
        "brief": "Ouvrez la barre latérale à gauche ; sur téléphone, touchez ☰. Complétez les champs obligatoires du livre.",
        "indice": "Ouvrez l’onglet 1. Plan et choisissez « Générer le plan professionnel ».",
        "scrittura": "Ouvrez l’onglet 2. Écriture & Quiz et choisissez « Écrire tout le livre » ou « Écrire le contenu » pour une section.",
        "pausa": "Restez dans l’onglet 2. Écriture & Quiz et choisissez « Reprendre la génération » pour continuer. Le texte créé reste enregistré.",
        "controllo": "Ouvrez l’onglet 4. Importer / Exporter / Copyright, lancez « Vérifier la complétude du manuscrit », puis téléchargez Word ou PDF.",
    },
    "Deutsch": {
        "brief": "Öffnen Sie die Seitenleiste links; auf dem Telefon tippen Sie auf ☰. Füllen Sie die Pflichtfelder des Buchs aus.",
        "indice": "Öffnen Sie Tab 1. Gliederung und wählen Sie « Professionelle Gliederung erstellen ».",
        "scrittura": "Öffnen Sie Tab 2. Schreiben & Quiz und wählen Sie « Ganzes Buch schreiben » oder « Inhalt schreiben » für einen Abschnitt.",
        "pausa": "Bleiben Sie in Tab 2. Schreiben & Quiz und wählen Sie « Generierung fortsetzen », wenn Sie weiterarbeiten möchten. Bereits erstellter Text bleibt gespeichert.",
        "controllo": "Öffnen Sie Tab 4. Importieren / Exportieren / Copyright, führen Sie « Manuskriptvollständigkeit prüfen » aus und laden Sie dann Word oder PDF herunter.",
    },
    "Română": {
        "brief": "Deschide bara laterală din stânga; pe telefon atinge ☰. Completează câmpurile obligatorii ale cărții.",
        "indice": "Deschide fila 1. Cuprins și apasă « Generează cuprinsul profesional ».",
        "scrittura": "Deschide fila 2. Scriere & Quiz și apasă « Scrie toată cartea » sau « Scrie conținut » pentru o secțiune.",
        "pausa": "Rămâi în fila 2. Scriere & Quiz și apasă « Reia generarea » când vrei să continui. Textul creat rămâne salvat.",
        "controllo": "Deschide fila 4. Import / Export / Copyright, rulează « Verifică integritatea manuscrisului », apoi descarcă Word sau PDF.",
    },
    "Русский": {
        "brief": "Откройте боковую панель слева; на телефоне нажмите ☰. Заполните обязательные поля книги.",
        "indice": "Откройте вкладку 1. План и нажмите « Создать профессиональное оглавление ».",
        "scrittura": "Откройте вкладку 2. Написание и Quiz и нажмите « Написать всю книгу » или « Написать содержание » для одного раздела.",
        "pausa": "Оставайтесь на вкладке 2. Написание и Quiz и нажмите « Продолжить генерацию », когда захотите продолжить. Созданный текст сохранён.",
        "controllo": "Откройте вкладку 4. Импорт / Экспорт / Copyright, запустите « Проверить полноту рукописи », затем скачайте Word или PDF.",
    },
    "العربية": {
        "brief": "افتح الشريط الجانبي على اليسار؛ على الهاتف اضغط ☰. أكمل حقول الكتاب الإلزامية.",
        "indice": "افتح التبويب 1. الفهرس واضغط «إنشاء الفهرس الاحترافي».",
        "scrittura": "افتح التبويب 2. الكتابة والاختبار واضغط «اكتب الكتاب كاملاً» أو «اكتب المحتوى» لقسم واحد.",
        "pausa": "ابقَ في التبويب 2. الكتابة والاختبار واضغط «استئناف التوليد» عندما تريد المتابعة. النص الذي تم إنشاؤه يبقى محفوظاً.",
        "controllo": "افتح التبويب 4. استيراد / تصدير / حقوق النشر، نفّذ «فحص اكتمال المخطوطة» ثم نزّل Word أو PDF.",
    },
    "中文": {
        "brief": "打开左侧边栏；手机上请点击 ☰。完成图书的必填字段。",
        "indice": "打开第 1 个标签“目录”，点击“生成专业目录”。",
        "scrittura": "打开第 2 个标签“写作与测验”，点击“写完整本书”；如只写一节，点击“撰写内容”。",
        "pausa": "停留在第 2 个标签“写作与测验”，准备继续时点击“继续生成”。已生成的文本会被保存。",
        "controllo": "打开第 4 个标签“导入 / 导出 / 版权”，运行“检查手稿完整性”，然后下载 Word 或 PDF。",
    },
}


def riepilogo_operativo_progetto(campi_obbligatori=None):
    """Legge lo stato reale del progetto senza alterare sessione o memoria."""
    sezioni = elenco_sezioni_progetto(st.session_state.get("lista_capitoli", []))
    contenuti = dict(memoria_progetto_unica().get("contenuti", {}) or {})
    for sezione in sezioni:
        testo = str(leggi_sezione_memorizzata(sezione) or "").strip()
        if testo:
            contenuti[sezione] = testo
    sidebar = sidebar_memorizzata_corrente()
    # Il Centro usa gli stessi campi che abilitano realmente la generazione
    # dell'indice. Personalizzazione e checkpoint restano visibili in memoria,
    # ma non possono far apparire un brief completo come incompleto.
    campi_operativi = dict(campi_obbligatori or {})
    if campi_operativi:
        campi_compilati = sum(1 for valore in campi_operativi.values() if str(valore).strip())
        campi_totali = len(campi_operativi)
    else:
        campi_compilati = sum(1 for valore in sidebar.values() if str(valore).strip())
        campi_totali = len(CAMPI_SALVATAGGIO_PROGETTO)
    registro_fonti = conserva_solo_fonti_web_selezionate(
        st.session_state.get("registro_fonti_web", "")
    )
    numero_fonti = len(re.findall(r"https?://\\S+", registro_fonti))
    if registro_fonti and not numero_fonti:
        numero_fonti = len([riga for riga in registro_fonti.splitlines() if riga.strip()])
    coda = list(st.session_state.get("job_scrittura_coda", []) or [])
    totale_job = max(0, int(st.session_state.get("job_scrittura_totale", 0) or 0))
    completati_job = max(0, totale_job - len(coda)) if totale_job else 0
    return {
        "campi_compilati": campi_compilati,
        "campi_totali": campi_totali,
        "brief_pronto": bool(campi_totali) and campi_compilati == campi_totali,
        "indice_pronto": bool(str(st.session_state.get("indice_raw", "") or "").strip()),
        "sezioni_previste": len(sezioni),
        "sezioni_scritte": sum(1 for sezione in sezioni if str(contenuti.get(sezione, "") or "").strip()),
        "fonti": numero_fonti,
        "job_attivo": bool(st.session_state.get("job_scrittura_attivo")),
        "job_pausa": bool(st.session_state.get("job_scrittura_pausa")),
        "job_fermato": bool(st.session_state.get("job_scrittura_fermato")),
        "coda": coda,
        "totale_job": totale_job,
        "completati_job": completati_job,
        "ultima_sezione": str(st.session_state.get("job_scrittura_ultima_completata", "") or "").strip(),
        "errore": str(st.session_state.get("job_scrittura_errore", "") or "").strip(),
        "ultimo_indice": str(st.session_state.get("ultimo_controllo_indice", "") or "").strip(),
    }


def mostra_centro_progetto(lingua, campi_obbligatori=None):
    """Rende il punto del progetto come un percorso semplice e non interattivo.

    La vista legge soltanto lo stato già presente: non sostituisce tab,
    pulsanti, salvataggi o controlli e non avvia alcuna operazione.
    """
    etichette = etichette_centro_progetto(lingua)
    stato = riepilogo_operativo_progetto(campi_obbligatori)
    passi = PASSI_PERCORSO_PROGETTO.get(lingua, PASSI_PERCORSO_PROGETTO["English"])
    direzione = "rtl" if lingua == "العربية" else "ltr"

    if stato["job_attivo"]:
        fase_attiva, stato_visuale, icona, guida_chiave = 3, "", "✍️", "scrittura"
        sezione_corrente = stato["coda"][0] if stato["coda"] else stato["ultima_sezione"]
        dettaglio = (
            f"{stato['sezioni_scritte']}/{stato['sezioni_previste']} · {sezione_corrente}"
            if stato["sezioni_previste"] else sezione_corrente
        )
        messaggio, prossimo = f"{etichette['in_corso']} — {dettaglio}", etichette["in_corso"]
    elif stato["job_pausa"]:
        fase_attiva, stato_visuale, icona, guida_chiave = 3, "pause", "⏸", "pausa"
        dettagli_pausa = stato["coda"][0] if stato["coda"] else etichette["nessun"]
        messaggio, prossimo = f"{etichette['in_pausa']} — {dettagli_pausa}", etichette["continua"]
    elif stato["job_fermato"]:
        fase_attiva, stato_visuale, icona, guida_chiave = 3, "pause", "⏹", "pausa"
        messaggio, prossimo = etichette["fermato"], etichette["continua"]
    elif not stato["brief_pronto"]:
        fase_attiva, stato_visuale, icona, guida_chiave = 1, "", "🧭", "brief"
        messaggio, prossimo = etichette["configura"], etichette["configura"]
    elif not stato["indice_pronto"]:
        fase_attiva, stato_visuale, icona, guida_chiave = 2, "", "🧠", "indice"
        messaggio, prossimo = etichette["crea_indice"], etichette["crea_indice"]
    elif stato["sezioni_previste"] and stato["sezioni_scritte"] < stato["sezioni_previste"]:
        fase_attiva, stato_visuale, icona, guida_chiave = 3, "", "✍️", "scrittura"
        messaggio, prossimo = (
            f"{etichette['manoscritto']}: {stato['sezioni_scritte']}/{stato['sezioni_previste']}",
            etichette["continua"],
        )
    else:
        fase_attiva, stato_visuale, icona, guida_chiave = 4, "complete", "✓", "controllo"
        messaggio, prossimo = (
            f"{etichette['manoscritto']} ✓ — {stato['sezioni_scritte']}/{stato['sezioni_previste']}",
            etichette["controlla"],
        )

    istruzioni = ISTRUZIONI_PROSSIMO_PASSO.get(lingua, ISTRUZIONI_PROSSIMO_PASSO["English"])
    istruzione_operativa = istruzioni[guida_chiave]
    percorso_html = "".join(
        (
            f"<div class='ss-project-step {'done' if numero < fase_attiva else 'current' if numero == fase_attiva else ''}'>"
            f"<span>{'✓' if numero < fase_attiva else numero}</span><b>{html.escape(passo)}</b></div>"
        )
        for numero, passo in enumerate(passi, start=1)
    )
    dati_metriche = (
        (etichette["sidebar"], f"{stato['campi_compilati']}/{stato['campi_totali']}"),
        (etichette["indice"], "✓" if stato["indice_pronto"] else "—"),
        (etichette["manoscritto"], f"{stato['sezioni_scritte']}/{stato['sezioni_previste']}" if stato["sezioni_previste"] else "0"),
        (etichette["fonti"], str(stato["fonti"])),
    )
    metriche_html = "".join(
        f"<div class='ss-project-metric'><span>{html.escape(etichetta)}</span><b>{html.escape(valore)}</b></div>"
        for etichetta, valore in dati_metriche
    )
    st.markdown(
        f"""<section class='ss-project-panel' dir='{direzione}'>
          <h2 class='ss-project-title'>{html.escape(etichette['titolo'])}</h2>
          <div class='ss-project-journey'>{percorso_html}</div>
          <div class='ss-project-metrics'>{metriche_html}</div>
          <div class='ss-next-action {stato_visuale}'><span class='ss-next-icon'>{icona}</span>
            <div><small>{html.escape(etichette['prossimo'])}</small><strong>{html.escape(messaggio)}</strong>
            <p>{html.escape(istruzione_operativa)}</p></div>
          </div>
        </section>""",
        unsafe_allow_html=True,
    )
    if stato["job_attivo"] and stato["sezioni_previste"]:
        st.progress(
            min(100, int(stato["sezioni_scritte"] / stato["sezioni_previste"] * 100)),
            text=f"{etichette['in_corso']}: {dettaglio}",
        )
    st.caption(f"**{etichette['prossimo']}:** {prossimo}")
    if stato["errore"]:
        st.caption(f"**{etichette['dettaglio']}:** {stato['errore']}")
    elif stato["ultimo_indice"] and not stato["indice_pronto"]:
        st.caption(f"**{etichette['dettaglio']}:** {stato['ultimo_indice']}")
    if (
        stato["indice_pronto"]
        and stato["sezioni_previste"]
        and stato["sezioni_scritte"] >= stato["sezioni_previste"]
        and not stato["job_attivo"]
    ):
        st.caption("🛡️ " + nota_originalita_facoltativa(lingua))
    st.caption(etichette["nota"])


def applica_snapshot_progetto(snapshot):
    """Ripristina una fotografia completa, anche se il progetto e' ancora iniziale.

    Un utente puo' salvare correttamente la sidebar o le fonti prima di avere
    creato l'indice. Quel salvataggio e' comunque un progetto valido e il
    comando ``RIAGGIORNA ALL'ULTIMA STESURA`` deve poterlo recuperare, anziche'
    apparire inefficace per la sola assenza dell'indice.
    """
    if not isinstance(snapshot, dict):
        return False

    # La sidebar non viene più ricostruita da un dizionario generico: la
    # memoria centrale ripristina prima le esatte chiavi dei widget. Questo
    # avviene prima del rendering della sidebar, quindi Streamlit visualizza
    # l'ultima scelta dell'utente anziché i valori iniziali della nuova pagina.
    sidebar, campi_sidebar_ripristinati, campi_sidebar_totali = (
        memoria_core.ripristina_sidebar_integrale(st.session_state, snapshot)
    )
    contenuti = {
        str(nome): str(testo)
        for nome, testo in (snapshot.get("contenuti", {}) or {}).items()
        if str(nome).strip() and str(testo or "").strip()
    }
    indice = str(snapshot.get("indice_raw", "") or snapshot.get("indice_backup", "") or "")
    if not indice.strip() and contenuti:
        indice = "\n".join(contenuti.keys())
    fonti = dict(snapshot.get("fonti", {}) or {})
    immagini = immagini_da_snapshot_cloud(snapshot.get("immagini_capitoli", {}) or {})
    ha_dati_ripristinabili = bool(
        indice.strip()
        or contenuti
        or any(str(valore or "").strip() for valore in sidebar.values())
        or any(str(valore or "").strip() for valore in fonti.values())
        or immagini
    )
    if not ha_dati_ripristinabili:
        return False

    progetto = memoria_progetto_unica()
    progetto.clear()
    progetto.update({
        "sidebar": dict(sidebar),
        "indice": indice,
        "contenuti": dict(contenuti),
        "fonti": dict(fonti),
        "immagini": dict(immagini),
    })

    # Elimina soltanto le vecchie textarea, non i dati: l'indice e ogni
    # sezione verranno ridisegnati con valori nuovi dalla memoria unica.
    for chiave in list(st.session_state.keys()):
        if re.match(r"^txt_[0-9a-f]{20}_v\d+$", str(chiave)):
            del st.session_state[chiave]
    st.session_state[CHIAVE_VERSIONI_WIDGET_SEZIONI] = {}
    st.session_state[CHIAVE_MEMORIA_SEZIONI] = dict(contenuti)
    st.session_state[CHIAVE_MEMORIA_PROTETTA] = dict(contenuti)
    st.session_state[CHIAVE_ARCHIVIO_STESURA_COMPLETA] = dict(contenuti)
    st.session_state[CHIAVE_REGISTRO_SEZIONI] = list(contenuti.keys())
    st.session_state[CHIAVE_SEZIONI_DA_REIDRATARE] = list(contenuti.keys())
    st.session_state[CHIAVE_SEZIONE_EDITOR_ATTIVA] = None
    if contenuti:
        st.session_state[CHIAVE_SELETTORE_EDITOR] = next(iter(contenuti))

    # La tab Indice riceve sempre il testo completo importato. Se il progetto
    # e' stato salvato prima dell'indice, svuotiamo in modo esplicito soltanto
    # l'indice e la lista capitoli, senza annullare sidebar, fonti o sezioni.
    if indice.strip():
        imposta_indice_progetto(indice)
    else:
        st.session_state["indice_raw"] = ""
        st.session_state["indice_editoriale"] = ""
        st.session_state["indice_widget_version"] = int(
            st.session_state.get("indice_widget_version", 0)
        ) + 1
        sync_capitoli()
    st.session_state["immagini_capitoli"] = immagini
    for chiave in (
        "conoscenza_extra", "scheda_fonti", "dossier_fonti_ai", "brief_fonti_originale",
        "dossier_ricerca_preliminare", "registro_fonti_web", "firma_ricerca_preliminare",
    ):
        valore_fonte = fonti.get(chiave, "")
        st.session_state[chiave] = (
            conserva_solo_fonti_web_selezionate(valore_fonte)
            if chiave == "registro_fonti_web" else valore_fonte
        )

    riepilogo_sidebar = (
        f"Sidebar: {campi_sidebar_ripristinati}/{campi_sidebar_totali} campi recuperati."
    )
    if snapshot.get("_origine_importazione_csv"):
        st.session_state["autosave_stato"] = (
            "✓ CSV importato: " + riepilogo_sidebar + " Indice, sezioni, fonti e immagini sono stati ripristinati."
        )
    elif campi_sidebar_ripristinati == campi_sidebar_totali:
        st.session_state["autosave_stato"] = (
            "✓ Ultima stesura ripristinata integralmente: " + riepilogo_sidebar
            + " Indice, sezioni, fonti e immagini sono disponibili."
        )
    else:
        st.session_state["autosave_stato"] = (
            "⚠ Ultima stesura recuperata con una sidebar parziale ("
            + riepilogo_sidebar
            + "). I dati disponibili restano visibili; salva di nuovo il progetto per creare una fotografia completa."
        )
    return True


def ripristina_progetto_salvato():
    """Ripristina l'ultima bozza una sola volta, oppure una bozza richiesta dal pulsante manuale."""
    # RESET PROGETTO ha priorità assoluta: anche se una vecchia riga cloud non
    # fosse ancora stata eliminata, non deve ricomparire al rerun successivo.
    if st.session_state.get("commercial_project_reset_requested"):
        return False
    snapshot_richiesto = st.session_state.pop("autosave_snapshot_da_ripristinare", None)
    if snapshot_richiesto:
        st.session_state["autosave_ripristino_verificato"] = True
        return applica_snapshot_progetto(snapshot_richiesto)
    if st.session_state.get("autosave_ripristino_verificato"):
        return False
    st.session_state["autosave_ripristino_verificato"] = True
    return applica_snapshot_progetto(carica_progetto_automatico())


def prepara_ripristino_ultima_stesura():
    """Legge l'ultima bozza dell'account e la applica al rerun successivo in modo sicuro."""
    snapshot = carica_progetto_automatico()
    if not snapshot:
        return False
    st.session_state.pop("commercial_project_reset_requested", None)
    st.session_state["autosave_snapshot_da_ripristinare"] = snapshot
    return True


def salva_progetto_corrente(sidebar, sezioni):
    """Crea e conferma nel cloud una fotografia completa del progetto aperto.

    Sidebar, indice, testi, fonti e immagini vengono catturati nello stesso
    istante. Il logout usa questa funzione e si chiude soltanto dopo l'esito
    positivo, quindi non può lasciare una sidebar più vecchia del manoscritto.
    """
    if st.session_state.get("admin_test_mode"):
        # Il collaudo usa le stesse funzioni dell'app, ma è una sandbox: non
        # deve poter sostituire il progetto reale salvato dall'amministratore.
        st.session_state["autosave_stato"] = "🧪 Collaudo: dati mantenuti solo nella pagina di prova; il progetto reale non viene modificato."
        return True
    # La sidebar è parte essenziale del progetto quanto le sezioni. La
    # fotografia conserva sia i nomi editoriali sia le chiavi reali dei
    # widget, comprese le scelte volutamente vuote. In una nuova sessione non
    # dovremo quindi inferire né perdere alcun valore inserito dall'utente.
    for nome, chiave in CAMPI_SALVATAGGIO_PROGETTO.items():
        if chiave not in st.session_state and nome in (sidebar or {}):
            st.session_state[chiave] = sidebar.get(nome, "")
    # Se Esci è stato premuto, questa fotografia è stata acquisita prima del
    # rerun dal bottone stesso. Ha priorità sui widget che Streamlit può
    # ricreare durante il salvataggio finale.
    snapshot_logout = st.session_state.get("commercial_logout_sidebar_snapshot", {}) or {}
    owner_logout = str(snapshot_logout.get("owner", "") or "") if isinstance(snapshot_logout, dict) else ""
    owner_corrente = str((st.session_state.get("commercial_user_context") or {}).get("id", "") or "")
    valori_logout = (
        dict(snapshot_logout.get("widget_values", {}) or {})
        if owner_logout and owner_logout == owner_corrente and logout_sicuro_richiesto()
        else None
    )
    fotografia_sidebar = memoria_core.fotografia_sidebar_integrale(
        st.session_state,
        valori_widget_preferiti=valori_logout,
    )
    sidebar_completa = dict(fotografia_sidebar["valori"])
    sidebar_widget_values = dict(fotografia_sidebar["widget"])
    st.session_state[CHIAVE_MEMORIA_SIDEBAR] = dict(sidebar_completa)
    progetto = memoria_progetto_unica()
    progetto["sidebar"] = dict(sidebar_completa)
    # Il salvataggio manuale non deve mai trasformarsi in una fotografia
    # parziale. Prima conserva tutte le sezioni dell'ultima sessione cloud e
    # poi applica i testi più recenti presenti nella pagina.
    contenuti = dict(progetto.get("contenuti", {}) or {})
    precedente = {}
    if not st.session_state.get("commercial_project_reset_requested"):
        precedente = carica_progetto_automatico()
        contenuti.update((precedente.get("contenuti", {}) or {}))
    # La copia protetta ha precedenza sulla fotografia cloud precedente;
    # quella dell'editor può contenerne una revisione manuale più recente.
    contenuti.update(dict(st.session_state.get(CHIAVE_ARCHIVIO_STESURA_COMPLETA, {}) or {}))
    contenuti.update(dict(st.session_state.get(CHIAVE_MEMORIA_PROTETTA, {}) or {}))
    contenuti.update(dict(st.session_state.get(CHIAVE_MEMORIA_SEZIONI, {}) or {}))
    for sezione in sezioni:
        testo = leggi_sezione_memorizzata(sezione)
        if str(testo).strip():
            contenuti[sezione] = testo
    contenuti = {sezione: testo for sezione, testo in contenuti.items() if str(testo).strip()}
    progetto["contenuti"] = dict(contenuti)
    st.session_state[CHIAVE_MEMORIA_SEZIONI] = dict(contenuti)
    st.session_state[CHIAVE_MEMORIA_PROTETTA] = dict(contenuti)
    # Un salvataggio da una pagina appena riaperta non deve mai cancellare un
    # indice già presente nel cloud solo perché il relativo widget non è ancora
    # stato renderizzato. Il RESET PROGETTO è l'unica eccezione esplicita.
    indice_corrente = str(progetto.get("indice", "") or st.session_state.get("indice_raw", "") or "")
    indice_precedente = str(
        precedente.get("indice_raw", "") or precedente.get("indice_backup", "") or ""
    )
    if not indice_corrente.strip() and indice_precedente.strip():
        indice_corrente = indice_precedente
        imposta_indice_progetto(indice_corrente)
    progetto["indice"] = indice_corrente
    # Anche le immagini esterne fanno parte del progetto: prima le custodiamo
    # nella memoria unica e poi le rendiamo serializzabili per Supabase. Nessun
    # pulsante genera immagini; qui proteggiamo esclusivamente i file caricati.
    immagini = dict(progetto.get("immagini", {}) or {})
    immagini.update(dict(st.session_state.get("immagini_capitoli", {}) or {}))
    immagini = immagini_da_snapshot_cloud(immagini)
    progetto["immagini"] = dict(immagini)
    snapshot = {
        "snapshot_version": 2,
        "sidebar": sidebar_completa,
        "sidebar_widget_values": sidebar_widget_values,
        "sidebar_fields": list(fotografia_sidebar["campi"]),
        "sidebar_schema_version": fotografia_sidebar["versione"],
        "indice_raw": indice_corrente,
        "indice_backup": indice_corrente,
        "contenuti": contenuti,
        "immagini_capitoli": immagini_per_snapshot_cloud(immagini),
        # Conserviamo il dossier già elaborato: dopo logout o refresh l'AI può
        # continuare a usarlo senza richiedere nuovamente i file originali.
        "fonti": {
            "conoscenza_extra": st.session_state.get("conoscenza_extra", ""),
            "scheda_fonti": st.session_state.get("scheda_fonti", ""),
            "dossier_fonti_ai": st.session_state.get("dossier_fonti_ai", ""),
            "brief_fonti_originale": st.session_state.get("brief_fonti_originale", ""),
            "dossier_ricerca_preliminare": st.session_state.get("dossier_ricerca_preliminare", ""),
            "registro_fonti_web": conserva_solo_fonti_web_selezionate(
                st.session_state.get("registro_fonti_web", "")
            ),
            "firma_ricerca_preliminare": st.session_state.get("firma_ricerca_preliminare", ""),
        },
    }
    # Copia pronta per il logout: la sidebar commerciale viene eseguita prima
    # dell'editor, perciò non può ricostruire da sola il manoscritto. Qui
    # custodiamo una fotografia già completa, associata all'account corrente,
    # che il comando Esci potrà inviare al cloud prima di svuotare la pagina.
    proprietario_snapshot = str(
        (st.session_state.get("commercial_user_context") or {}).get("id", "")
    )
    if proprietario_snapshot:
        st.session_state["commercial_logout_snapshot"] = json.loads(
            json.dumps(snapshot, ensure_ascii=False)
        )
        st.session_state["commercial_logout_snapshot_owner"] = proprietario_snapshot
    progetto["fonti"] = dict(snapshot["fonti"])
    serializzato = json.dumps(snapshot, ensure_ascii=False, sort_keys=True)
    firma = hashlib.sha256(serializzato.encode("utf-8")).hexdigest()
    # Consideriamo concluso il salvataggio solo quando Supabase ha confermato
    # la fotografia. In questo modo un errore transitorio non lascia nel cloud
    # una vecchia bozza che poi potrebbe riapparire al rerun successivo.
    if st.session_state.get("autosave_firma_cloud") == firma:
        return True
    momento = datetime.datetime.now().strftime("%H:%M")
    if salva_progetto_automatico(snapshot):
        st.session_state.pop("commercial_project_reset_requested", None)
        st.session_state["autosave_firma"] = firma
        st.session_state["autosave_firma_cloud"] = firma
        st.session_state["autosave_stato"] = f"✓ Sessione salvata nel tuo account alle {momento}."
        return True
    else:
        st.session_state["autosave_firma"] = firma
        st.session_state["autosave_stato"] = f"⚠ Sessione mantenuta solo in questa pagina alle {momento}: il salvataggio nel tuo account non è riuscito."
        return False


def salva_stesura_immediata(sezioni):
    """Aggiorna solo la memoria della sessione, senza alcun salvataggio cloud."""
    for sezione in sezioni:
        leggi_sezione_memorizzata(sezione)
    return True


def salva_stesura_generata_in_cloud(sezioni, descrizione="contenuto generato"):
    """Protegge nel cloud il contenuto appena creato dall'IA.

    Il salvataggio automatico è volutamente limitato agli esiti di operazioni
    IA concluse con successo: le modifiche manuali restano sotto il controllo
    del pulsante ``SALVA SESSIONE``. Un errore di rete non annulla né nasconde
    mai il testo già ricevuto e conservato nella memoria della pagina.
    """
    salva_stesura_immediata(sezioni)
    if st.session_state.get("admin_test_mode"):
        st.session_state["autosave_stato"] = (
            f"🧪 {descrizione.capitalize()} conservato nel laboratorio di collaudo: nessuna bozza reale è stata sovrascritta."
        )
        return True
    sezioni_progetto = list(dict.fromkeys([
        *sezioni,
        *st.session_state.get("lista_capitoli", []),
        *st.session_state.get(CHIAVE_MEMORIA_SEZIONI, {}).keys(),
    ]))
    try:
        salvato = salva_progetto_corrente(sidebar_memorizzata_corrente(), sezioni_progetto)
    except Exception:
        salvato = False
    if salvato:
        momento = datetime.datetime.now().strftime("%H:%M")
        st.session_state["autosave_stato"] = (
            f"✓ {descrizione.capitalize()} salvato automaticamente nel tuo account alle {momento}."
        )
    else:
        st.session_state["autosave_stato"] = (
            f"⚠ {descrizione.capitalize()} conservato in questa pagina, ma il salvataggio automatico "
            "nel tuo account non è riuscito: premi SALVA SESSIONE appena possibile."
        )
    return salvato


def completa_logout_dopo_salvataggio_finale():
    """Salva l'intero progetto dopo il rendering e solo allora chiude l'account.

    Il comando Esci parte nella sidebar, che Streamlit esegue prima dell'editor.
    Questa fase conclusiva vede invece tutti i valori del rerun corrente: testi
    digitati, sidebar, indice, fonti e immagini caricate esternamente.
    """
    if not logout_sicuro_richiesto():
        return
    user = st.session_state.get("commercial_user_context") or {}
    if not str(user.get("id", "") or "").strip():
        annulla_logout_sicuro("La sessione non è più valida. Il progetto resta aperto in questa pagina: accedi di nuovo e riprova il salvataggio.")
        st.rerun()
    sezioni_complete = elenco_sezioni_progetto(
        list(st.session_state.get("lista_capitoli", []) or [])
    )
    try:
        salvato = salva_progetto_corrente(
            sidebar_memorizzata_corrente(),
            sezioni_complete,
        )
    except Exception:
        salvato = False
    if salvato:
        completa_logout_sicuro()
    annulla_logout_sicuro(
        "Non è stato possibile confermare il salvataggio finale nel tuo account. "
        "L’uscita è stata annullata: il progetto resta aperto e puoi riprovare senza perdere il lavoro."
    )
    st.rerun()


def minimo_parole_per_sezione_editoriale(sezione, genere):
    """Soglia unica usata da tutti i controlli prima dell'esportazione."""
    return minimo_parole_core(sezione, genere, tipo_sezione_editoriale)


def sezioni_mancanti_per_esportazione(sezioni, genere):
    """Non consente di esportare un libro se l'indice contiene sezioni non effettivamente redatte."""
    mancanti = []
    for sezione in sezioni:
        # La sezione può non essere il widget attualmente visibile, ma essere
        # già presente nella memoria stabile del progetto o dopo un ripristino.
        testo = pulisci_testo_editoriale(leggi_sezione_memorizzata(sezione)).strip()
        if len(testo.split()) < minimo_parole_per_sezione_editoriale(sezione, genere):
            mancanti.append(sezione)
    return mancanti


def stati_sezioni_editoriali(sezioni, genere, contenuti=None):
    """Distingue sezioni mancanti, deboli e complete usando la memoria reale.

    ``contenuti`` permette al controllo finale di usare una fotografia unica,
    senza dipendere dal solo widget della sezione che l'utente ha aperto per
    ultima nell'editor.
    """
    return stati_sezioni_core(
        sezioni,
        genere,
        dict(contenuti or {}),
        leggi_sezione_memorizzata,
        pulisci_testo_editoriale,
        tipo_sezione_editoriale,
    )


def controllo_completezza_testi_gratuito(sezioni, contenuti=None):
    """Controllo locale e gratuito contro testi assenti o evidentemente interrotti.

    Non usa IA e non valuta lo stile: segnala soltanto problemi tecnici che
    l'utente può risolvere prima del controllo editoriale finale.
    """
    return controllo_completezza_core(
        sezioni,
        dict(contenuti or {}),
        leggi_sezione_memorizzata,
        pulisci_testo_editoriale,
    )


def controllo_finale_pre_export(indice, sezioni, contenuti, titolo, trama, genere, obiettivo):
    """Controllo gratuito e non distruttivo: decide soltanto se il download è finale o BOZZA."""
    stati = stati_sezioni_editoriali(sezioni, genere, contenuti)
    problemi, prompt_correzione = [], []
    for voce in stati:
        if voce["Stato"] != "COMPLETA":
            problemi.append(f"{voce['Sezione']}: {voce['Stato'].lower()} ({voce['Dettaglio']})")
            prompt_correzione.append(
                f"SEZIONE: {voce['Sezione']}\n"
                f"PROBLEMA: {voce['Dettaglio']}.\n"
                "OBIETTIVO: completa o rafforza solo questa sezione, senza ripetere le altre.\n"
                "PROMPT DA INCOLLARE: Riscrivi questa sezione rispettando titolo, genere, stile, POV e argomento del libro. "
                "Aggiungi solo contenuto concreto necessario e non modificare le altre sezioni."
            )

    # La struttura dell'indice viene già validata quando è generata. Non la
    # riutilizziamo qui per non trasformare libri validi ma più brevi in bozze.

    if genere == "Test Prep (Preparazione Esami)":
        esiti_test = audit_simulazioni_test_prep(indice, contenuti, obiettivo, trama)
        problemi.extend(esito for esito in esiti_test if esito.startswith("ERRORE"))

    return {
        "pronto": not problemi,
        "problemi": problemi,
        "prompt_correzione": prompt_correzione,
        "stati": stati,
    }


def etichette_report_pubblicazione(lingua):
    """Testi compatti per l'esito a semaforo prima dell'esportazione."""
    testi = {
        "Italiano": ("🚦 Pronto per pubblicare?", "Struttura", "Manoscritto", "Frasi concluse", "Originalità", "Esportazione", "Interventi necessari", "Dettagli del controllo", "Da verificare"),
        "English": ("🚦 Ready to publish?", "Structure", "Manuscript", "Complete sentences", "Originality", "Export", "Required actions", "Review details", "To be checked"),
        "Español": ("🚦 ¿Listo para publicar?", "Estructura", "Manuscrito", "Frases completas", "Originalidad", "Exportación", "Acciones necesarias", "Detalles de la revisión", "Por verificar"),
        "Français": ("🚦 Prêt à publier ?", "Structure", "Manuscrit", "Phrases terminées", "Originalité", "Export", "Actions nécessaires", "Détails du contrôle", "À vérifier"),
        "Deutsch": ("🚦 Bereit zur Veröffentlichung?", "Struktur", "Manuskript", "Vollständige Sätze", "Originalität", "Export", "Erforderliche Schritte", "Prüfdetails", "Zu prüfen"),
        "Română": ("🚦 Gata de publicare?", "Structură", "Manuscris", "Fraze complete", "Originalitate", "Export", "Acțiuni necesare", "Detalii control", "De verificat"),
        "Русский": ("🚦 Готово к публикации?", "Структура", "Рукопись", "Завершённые фразы", "Оригинальность", "Экспорт", "Необходимые действия", "Детали проверки", "Нужно проверить"),
        "العربية": ("🚦 هل هو جاهز للنشر؟", "البنية", "المخطوطة", "اكتمال الجمل", "الأصالة", "التصدير", "الإجراءات المطلوبة", "تفاصيل الفحص", "بحاجة للتحقق"),
        "中文": ("🚦 可以发布了吗？", "结构", "手稿", "句子完整性", "原创性", "导出", "需要处理的事项", "检查详情", "待检查"),
    }
    return testi.get(lingua, testi["English"])


def nota_originalita_facoltativa(lingua):
    """Ricorda che il controllo è utile, ma non blocca l'esportazione."""
    testi = {
        "Italiano": "Controllo originalità: facoltativo e consigliato. Puoi comunque procedere al controllo finale e all'esportazione del manoscritto.",
        "English": "Originality check: optional and recommended. You can still proceed to the final check and export the manuscript.",
        "Español": "Control de originalidad: opcional y recomendado. Aun así puedes continuar con la revisión final y exportar el manuscrito.",
        "Français": "Contrôle d’originalité : facultatif et recommandé. Vous pouvez quand même passer au contrôle final et exporter le manuscrit.",
        "Deutsch": "Originalitätsprüfung: optional und empfohlen. Sie können dennoch die Endprüfung durchführen und das Manuskript exportieren.",
        "Română": "Controlul originalității: opțional și recomandat. Poți continua totuși cu verificarea finală și exportul manuscrisului.",
        "Русский": "Проверка оригинальности: необязательна, но рекомендуется. Вы всё равно можете выполнить финальную проверку и экспортировать рукопись.",
        "العربية": "فحص الأصالة اختياري وموصى به. يمكنك مع ذلك متابعة الفحص النهائي وتصدير المخطوطة.",
        "中文": "原创性检查为可选项，建议执行。你仍可继续最终检查并导出书稿。",
    }
    return testi.get(lingua, testi["English"])


def mostra_report_prontezza_pubblicazione(esito_finale, sezioni, contenuti, lingua):
    """Mostra un esito operativo, senza avviare controlli, IA o modifiche.

    Riusa gli esiti già disponibili nella sessione. I report estesi restano
    nell'expander, mentre in pagina compaiono soltanto gli interventi concreti.
    """
    (
        titolo, et_struttura, et_manoscritto, et_frasi, et_originalita,
        et_export, et_interventi, et_dettagli, et_da_verificare,
    ) = etichette_report_pubblicazione(lingua)
    stati = list(esito_finale.get("stati", []) or [])
    problemi = list(esito_finale.get("problemi", []) or [])
    stati_non_completi = [voce for voce in stati if voce.get("Stato") != "COMPLETA"]
    report_completezza = list(st.session_state.get("report_completezza_manoscritto", []) or [])
    completezza_non_superata = [voce for voce in report_completezza if voce.get("Esito") != "COMPLETA"]
    report_locale = st.session_state.get("report_originalita_fonti")
    report_web = str(st.session_state.get("report_originalita_web_completa", "") or st.session_state.get("report_originalita_web", "") or "").strip()
    ha_indice = bool(str(st.session_state.get("indice_raw", "") or "").strip())
    ha_testi = bool(sezioni) and any(str(contenuti.get(sezione, "") or "").strip() for sezione in sezioni)

    st.markdown(f"### {titolo}")
    col_a, col_b, col_c, col_d = st.columns(4)
    with col_a:
        if ha_indice and sezioni:
            st.success(f"**{et_struttura}**\n\n✓")
        else:
            st.error(f"**{et_struttura}**\n\n✕")
    with col_b:
        if not ha_testi:
            st.error(f"**{et_manoscritto}**\n\n✕")
        elif stati_non_completi:
            st.warning(f"**{et_manoscritto}**\n\n⚠ {len(stati_non_completi)}")
        else:
            st.success(f"**{et_manoscritto}**\n\n✓")
    with col_c:
        if not report_completezza:
            st.info(f"**{et_frasi}**\n\n• {et_da_verificare}")
        elif completezza_non_superata:
            st.warning(f"**{et_frasi}**\n\n⚠ {len(completezza_non_superata)}")
        else:
            st.success(f"**{et_frasi}**\n\n✓")
    with col_d:
        if isinstance(report_locale, dict) and report_locale.get("eseguito"):
            if report_locale.get("trovate"):
                st.error(f"**{et_originalita}**\n\n✕")
            elif report_web and richiede_revisione_copyright(report_web):
                st.warning(f"**{et_originalita}**\n\n⚠")
            else:
                st.success(f"**{et_originalita}**\n\n✓")
        else:
            st.info(f"**{et_originalita}**\n\n• {et_da_verificare}")

    if esito_finale.get("pronto"):
        st.success(f"**{et_export}:** ✓")
    else:
        st.warning(f"**{et_export}:** ⚠")

    azioni = list(problemi)
    azioni.extend(
        f"{voce.get('Sezione', '')}: {voce.get('Dettaglio', '')}"
        for voce in completezza_non_superata
        if voce.get("Sezione")
    )
    if isinstance(report_locale, dict) and report_locale.get("trovate"):
        azioni.append(report_locale.get("messaggio", "Controlla le somiglianze segnalate dalle fonti caricate."))
    if report_web and richiede_revisione_copyright(report_web):
        azioni.append("Il controllo web segnala elementi da rivedere: apri il dettaglio copyright e rielabora solo le sezioni indicate.")
    azioni = list(dict.fromkeys(azione for azione in azioni if str(azione).strip()))
    if azioni:
        st.markdown(f"**{et_interventi}**")
        for azione in azioni[:8]:
            st.write("- " + str(azione))
    else:
        st.success("✓ Nessun intervento tecnico richiesto dai controlli già eseguiti.")

    with st.expander(et_dettagli, expanded=False):
        st.caption("Questo riepilogo non sostituisce la rilettura editoriale né una certificazione legale.")
        if stati:
            st.dataframe(stati, hide_index=True, use_container_width=True)
        if report_completezza:
            st.dataframe(report_completezza, hide_index=True, use_container_width=True)
        if isinstance(report_locale, dict):
            st.write(report_locale.get("messaggio", "Nessun dettaglio disponibile."))
        if report_web:
            st.caption("Esito web disponibile nella sezione Controllo originalità e copyright.")


def suggerimento_editoriale_contestuale(sezione, sezioni, obiettivo_libro, argomento):
    """Crea una bussola locale per l'editor, senza IA e senza modificare testi."""
    sezioni = list(sezioni or [])
    posizione = sezioni.index(sezione) if sezione in sezioni else -1
    precedente = sezioni[posizione - 1] if posizione > 0 else ""
    successiva = sezioni[posizione + 1] if 0 <= posizione < len(sezioni) - 1 else ""
    titolo = str(sezione or "").strip()
    titolo_minuscolo = titolo.casefold()
    if sezione_prefazione(titolo):
        obiettivo = "Orientare il lettore: chiarisci promessa, contesto e utilità del libro senza anticiparne lo sviluppo."
        evita = "Non trasformarla in un capitolo completo e non ripetere le spiegazioni che seguiranno."
    elif titolo_minuscolo.startswith(("parte ", "part ", "parte", "teil ", "часть ", "الجزء", "部分")):
        obiettivo = "Presentare il filo conduttore della Parte e preparare il passaggio alle sezioni che la compongono."
        evita = "Non sostituirti ai capitoli o sottocapitoli successivi con spiegazioni troppo estese."
    elif titolo_minuscolo.startswith(("capitolo", "chapter", "chapitre", "kapitel", "capitol", "глава", "الفصل", "章节")):
        obiettivo = "Sviluppare il tema indicato dal titolo con un percorso chiaro, esempi pertinenti e una conclusione utile."
        evita = "Non ripetere definizioni o esempi già assegnati alle sezioni precedenti."
    else:
        obiettivo = "Trattare un aspetto specifico del percorso editoriale con contenuto concreto e autonomo."
        evita = "Non ampliare il tema fino a coprire sezioni vicine e non ripetere formule già usate."
    collegamento = (
        f"Riprende e approfondisce: {precedente}." if precedente else
        "È il primo punto del percorso: definisci subito il contesto necessario al lettore."
    )
    risultato = (
        f"Al termine il lettore dovrebbe fare un passo concreto verso: {obiettivo_libro}."
        if str(obiettivo_libro or "").strip() else
        f"Al termine il lettore deve comprendere o saper applicare il tema: {titolo}."
    )
    if successiva:
        risultato += f" Prepara senza anticiparla la sezione successiva: {successiva}."
    if str(argomento or "").strip():
        obiettivo += f" Mantieni la coerenza con l'argomento generale: {str(argomento).strip()[:220]}."
    return {
        "obiettivo": obiettivo,
        "collegamento": collegamento,
        "evita": evita,
        "risultato": risultato,
    }


def mostra_suggerimento_editoriale_contestuale(sezione, sezioni, obiettivo_libro, argomento):
    """Mostra la bussola soltanto su richiesta: nessun campo viene scritto."""
    suggerimento = suggerimento_editoriale_contestuale(sezione, sezioni, obiettivo_libro, argomento)
    with st.expander("💡 Bussola editoriale della sezione", expanded=False):
        st.caption("Suggerimento locale e gratuito: non genera testo e non modifica il manoscritto.")
        st.write(f"**Obiettivo:** {suggerimento['obiettivo']}")
        st.write(f"**Collegamento:** {suggerimento['collegamento']}")
        st.write(f"**Evita:** {suggerimento['evita']}")
        st.write(f"**Risultato per il lettore:** {suggerimento['risultato']}")


def genera_sezione_con_ripetizione(prompt, system_prompt, sezione, lingua, tentativi=2, amount=AI_REQUEST_CREDITS,
                                    max_completion_tokens=None, addebita=True):
    """Riprova una sezione senza perdere le precedenti; evita libri interrotti a metà dopo un errore transitorio."""
    ultimo_errore = None
    for tentativo in range(1, tentativi + 1):
        try:
            testo = chiedi_gpt(
                prompt, system_prompt, addebita=addebita, amount=amount,
                max_completion_tokens=max_completion_tokens,
            )
            if not testo or testo.startswith("ERRORE:"):
                raise RuntimeError(testo or "Risposta vuota")
            # La stesura non effettua automaticamente una ricerca: la verifica
            # viene attivata solo dal motore editoriale quando individua fatti variabili.
            return pulisci_testo_editoriale(testo)
        except Exception as exc:
            ultimo_errore = exc
    raise RuntimeError(f"Impossibile completare la sezione dopo {tentativi} tentativi: {ultimo_errore}")


MARCATORE_FINE_SEZIONE = "[[FINE_SEZIONE]]"


def genera_bozza_sezione_con_chiusura(prompt, system_prompt, sezione, lingua, *, max_completion_tokens,
                                      addebita=True):
    """Genera una sezione solo se il modello conferma una conclusione completa.

    La punteggiatura da sola non basta a distinguere una risposta conclusa da
    una risposta troncata dal limite di output. Il marcatore è tecnico, viene
    rimosso prima del salvataggio e non può comparire nel manoscritto.
    """
    richiesta = prompt + f"""

CONSEGNA TECNICA OBBLIGATORIA
Termina davvero il testo con una frase conclusa e pertinente. Solo dopo il punto,
il punto interrogativo o il punto esclamativo finale, scrivi su una riga separata
esattamente questo marcatore: {MARCATORE_FINE_SEZIONE}
Non scrivere altro dopo il marcatore. Il marcatore è interno al software e non fa
parte del libro.
"""
    risposta = genera_sezione_con_ripetizione(
        richiesta,
        system_prompt,
        sezione,
        lingua,
        tentativi=1,
        max_completion_tokens=max_completion_tokens,
        addebita=addebita,
    )
    finale = str(risposta or "").rstrip()
    consegna_confermata = finale.endswith(MARCATORE_FINE_SEZIONE)
    if consegna_confermata:
        finale = finale[:-len(MARCATORE_FINE_SEZIONE)].rstrip()
    return pulisci_testo_editoriale(finale), consegna_confermata


def criticita_consegna_sezione(testo, consegna_confermata, genere, sezione, profilo_lunghezza, indice):
    """Convalida una sezione senza confondere un dettaglio tecnico con la qualità.

    Il marcatore ``[[FINE_SEZIONE]]`` è utile per sapere che il modello ha
    seguito la consegna fino alla fine, ma GPT e DeepSeek possono ometterlo
    pur restituendo un testo integro. In quel caso non dobbiamo scartare la
    sezione, fermare la coda o far credere all'utente che il cervello non
    abbia scritto nulla. La decisione editoriale resta quindi basata su
    chiusura reale, lunghezza e specificità del testo; il marcatore è soltanto
    un segnale diagnostico non bloccante.
    """
    testo_pulito = pulisci_testo_editoriale(testo or "").strip()
    if not consegna_confermata:
        motivo_finale = motivo_chiusura_tecnica(testo_pulito)
        if motivo_finale:
            return f"ragionamento non concluso: {motivo_finale}"
        avvisi_tecnici = st.session_state.setdefault("avvisi_tecnici_stesura", {})
        avvisi_tecnici[sezione] = (
            "Il testo è stato accettato perché completo, ma il cervello non ha "
            "restituito il marcatore tecnico di fine sezione."
        )
    return criticita_specificita(testo_pulito, genere, sezione, profilo_lunghezza, indice)


def capitolo_padre(indice, sezione):
    """Restituisce il capitolo che contiene un sottocapitolo, utile per riconoscere le simulazioni."""
    padre = ""
    confronto = re.sub(r"\s+", " ", sezione.strip().lower())
    for riga in (indice or "").splitlines():
        pulita = riga.strip()
        if re.match(r"(?i)^capitolo\s+\d+", pulita):
            padre = pulita
        if re.sub(r"\s+", " ", pulita.lower()) == confronto:
            return padre
    return padre


def numero_domande_simulazione(indice, trama, obiettivo):
    """Ricava un conteggio soltanto quando il brief lo dichiara; evita numeri inventati."""
    testo = f"{indice}\n{trama}\n{obiettivo}"
    corrispondenze = re.findall(
        r"(?is)(?:simulazion\w*|test\s+completo|prova\s+completa).{0,90}?(\d{1,3})\s+(?:domand\w*|quesit\w*)"
        r"|(\d{1,3})\s+(?:domand\w*|quesit\w*).{0,90}?(?:simulazion\w*|test\s+completo|prova\s+completa)",
        testo
    )
    numeri = [int(a or b) for a, b in corrispondenze if (a or b)]
    return max(numeri) if numeri else 0


def sezione_simulazione_test_prep(sezione, indice, genere):
    """Attiva il generatore a blocchi solo nella sezione che deve contenere i quesiti della simulazione."""
    if genere != "Test Prep (Preparazione Esami)":
        return False
    titolo = sezione.lower()
    padre = capitolo_padre(indice, sezione).lower()
    è_simulazione = "simulazione" in titolo or "simulazione" in padre
    è_contenuto_test = any(parola in titolo for parola in ("domande", "quesiti", "quiz", "esecuzione", "test"))
    return è_simulazione and è_contenuto_test


def conta_domande_test_prep(testo):
    return len(re.findall(r"(?im)^\s*domanda\s+\d{1,3}\s*[:.-]", testo or ""))


def domande_normalizzate_test_prep(testo):
    domande = re.findall(r"(?im)^\s*domanda\s+\d{1,3}\s*[:.-]\s*(.+)$", testo or "")
    return [re.sub(r"[^a-z0-9àèéìòóù ]", "", domanda.lower()).strip() for domanda in domande]


def genera_simulazione_test_prep(prompt_base, system_prompt, sezione, indice, trama, obiettivo, lingua):
    """Genera prove lunghe in blocchi verificabili, evitando simulazioni promesse ma incomplete."""
    totale = numero_domande_simulazione(indice, trama, obiettivo)
    if totale < 10:
        return genera_sezione_con_ripetizione(prompt_base, system_prompt, sezione, lingua)

    dimensione_blocco = 10
    gruppi_domande = (totale + dimensione_blocco - 1) // dimensione_blocco
    blocchi_domande, domande_precedenti = [], []
    for inizio in range(1, totale + 1, dimensione_blocco):
        fine = min(inizio + dimensione_blocco - 1, totale)
        vincolo = f"""
Questa è la parte domande della simulazione '{sezione}', gruppo {inizio}-{fine} di {totale}.
Genera ESATTAMENTE {fine - inizio + 1} quesiti originali, numerati da DOMANDA {inizio:02d}: a DOMANDA {fine:02d}:.
Per ogni domanda usa quattro opzioni A), B), C), D). NON scrivere risposte, soluzioni, commenti,
punteggi o istruzioni su come prepararsi. Distribuisci i quesiti sui contenuti obbligatori dell'indice
e del brief; non ripetere le domande già prodotte qui sotto.
Domande già prodotte: {' | '.join(domande_precedenti) or 'nessuna'}
"""
        blocco, nuove_domande = "", []
        for tentativo_blocco in range(3):
            correzione = ""
            if tentativo_blocco:
                correzione = (
                    "\nCORREZIONE OBBLIGATORIA: la bozza precedente era incompleta o ripeteva quesiti già usati. "
                    "Sostituiscila interamente con domande nuove, numerate nel formato richiesto e senza testo introduttivo."
                )
            candidato = genera_sezione_con_ripetizione(prompt_base + vincolo + correzione, system_prompt, sezione, lingua)
            domande_candidate = domande_normalizzate_test_prep(candidato)
            conteggio_corretto = conta_domande_test_prep(candidato) == (fine - inizio + 1)
            senza_duplicati = not (set(domande_candidate) & set(domande_precedenti))
            if conteggio_corretto and senza_duplicati:
                blocco, nuove_domande = candidato, domande_candidate
                break
            blocco, nuove_domande = candidato, domande_candidate
        if conta_domande_test_prep(blocco) != (fine - inizio + 1) or (set(nuove_domande) & set(domande_precedenti)):
            st.session_state["avviso_simulazione_test_prep"] = (
                f"La simulazione '{sezione}' richiede una verifica: il blocco {inizio}-{fine} non ha superato "
                "il controllo automatico di quantità o unicità. Il controllo coerenza indicherà le correzioni necessarie."
            )
        blocchi_domande.append(blocco)
        domande_precedenti.extend(nuove_domande)

    corpo_domande = "\n\n".join(blocchi_domande)
    prompt_chiave = f"""Crea la chiave delle soluzioni per la simulazione '{sezione}' in lingua {lingua}.
Le domande seguenti sono già state redatte. Fornisci ESATTAMENTE una riga di soluzione per ogni
DOMANDA da 01 a {totale:02d}, con questo formato: SOLUZIONE 01: lettera corretta - spiegazione breve e concreta.
Non riscrivere le domande, non aggiungere nuove domande, non usare link o fonti e non omettere numeri.

DOMANDE DELLA SIMULAZIONE
{corpo_domande}
"""
    chiave = ""
    for tentativo_chiave in range(3):
        correzione = "" if not tentativo_chiave else "\nCORREZIONE OBBLIGATORIA: inserisci tutte e sole le soluzioni numerate richieste."
        chiave = genera_sezione_con_ripetizione(
            prompt_chiave + correzione, system_prompt, sezione, lingua,
            amount=gruppi_domande,
        )
        if len(re.findall(r"(?im)^\s*soluzione\s+\d{1,3}\s*[:.-]", chiave or "")) == totale:
            break
    if len(re.findall(r"(?im)^\s*soluzione\s+\d{1,3}\s*[:.-]", chiave or "")) != totale:
        st.session_state["avviso_simulazione_test_prep"] = (
            f"La simulazione '{sezione}' richiede una verifica: la chiave delle soluzioni non ha il conteggio previsto. "
            "Il controllo coerenza indicherà le correzioni necessarie."
        )
    return (
        f"SIMULAZIONE: DOMANDE\n\n{corpo_domande}\n\n"
        f"SOLUZIONI COMMENTATE - CONSULTALE SOLO DOPO AVER COMPLETATO LA PROVA\n\n{chiave}"
    )


def stima_massima_crediti_stesura(sezione, indice, trama, obiettivo, genere):
    """Stima prudente: include eventuali tentativi di recupero automatico."""
    if sezione_simulazione_test_prep(sezione, indice, genere):
        totale = numero_domande_simulazione(indice, trama, obiettivo)
        if totale >= 10:
            gruppi_domande = (totale + 9) // 10
            # Gruppi di domande + chiave soluzioni, fino a tre correzioni e due retry ciascuno.
            # Un credito ogni dieci domande e uno ogni dieci soluzioni.
            return gruppi_domande * 2
    return 1


def criticita_specificita(testo, genere, sezione, profilo_lunghezza=None, indice=""):
    """Individua bozze genericamente motivazionali prima che finiscano nel manoscritto."""
    pulito = pulisci_testo_editoriale(testo or "").strip()
    parole = pulito.split()
    basso = pulito.lower()
    tipo_sezione = tipo_sezione_editoriale(sezione)
    # Un finale senza chiusura è spesso il segnale di una risposta interrotta dal limite di output.
    # Viene rigenerata prima di essere salvata, senza accettare un ragionamento lasciato a metà.
    finale = pulito.rstrip()
    connettivi_finali = (
        " e", " o", " ma", " perché", " quindi", " inoltre", " come", " per", " con", " di", " da", " in", " su",
        " and", " or", " but", " because", " therefore", " with", " for", " to", " of", " in", " on",
        " y", " o", " pero", " porque", " por", " para", " con", " de", " en",
        " et", " ou", " mais", " parce que", " pour", " avec", " de", " dans",
        " und", " oder", " aber", " weil", " für", " mit", " von", " im",
        " și", " sau", " dar", " deoarece", " pentru", " cu", " din", " în",
    )
    chiusura_sospesa = finale.endswith((",", ";", ":", "—", "–", "-", "…", "...", "(", "[", "{"))
    if finale and (
        chiusura_sospesa
        or finale[-1] not in ".!?。！？؟»”)]}"
        or any(finale.lower().endswith(connettivo) for connettivo in connettivi_finali)
    ):
        return "ragionamento non concluso: chiudi l'ultima idea con una frase completa e utile"
    # Le Parti sono cornici editoriali più brevi, ma devono comunque finire
    # con una frase completa.
    if tipo_sezione == "parte":
        return ""
    # I capitoli che possiedono sottocapitoli restano cornici intenzionalmente brevi:
    # la trattazione completa è affidata alle sezioni figlie e non va duplicata.
    capitolo_cornice = (
        tipo_sezione == "capitolo"
        and bool(individua_sottocapitoli_del_capitolo(sezione, (indice or "").splitlines()))
    )
    if profilo_lunghezza and not capitolo_cornice:
        minimo_parole, _ = vincolo_parole_con_tolleranza(profilo_lunghezza)
        if len(parole) < minimo_parole:
            return (
                f"testo troppo breve: servono almeno {minimo_parole} parole per il profilo "
                f"'{profilo_lunghezza}', senza aggiungere ripetizioni"
            )
    elif len(parole) < 150:
        return "testo troppo breve per sviluppare l'argomento assegnato"

    # Un capitolo che possiede sottocapitoli introduce e collega il percorso:
    # l'operatività dettagliata appartiene alle sezioni figlie. Un capitolo
    # autonomo e ogni sottocapitolo, invece, devono restare concretamente utili.
    # Le istruzioni per i Manuali chiedono già esempi, passaggi e verifiche
    # quando pertinenti. Non imponiamo però più un criterio meccanico di
    # procedura numerata come blocco finale: espressioni come “applicazione
    # pratica” possono richiedere casi, checklist o spiegazioni, non sempre
    # una sequenza artificiale. Restano obbligatori completezza, lunghezza,
    # coerenza e assenza di frasi spezzate per ogni sezione.

    formule_generiche = (
        "è fondamentale", "e fondamentale", "è cruciale", "e cruciale", "in modo efficace",
        "è importante", "e importante", "con sicurezza", "molto utile", "potente strumento"
    )
    genericita = sum(basso.count(formula) for formula in formule_generiche)
    segnali_per_genere = {
        "Saggio Scientifico": ("definiz", "evidenz", "limite", "esempio"),
        "Quiz Scientifico": ("domanda", "risposta", "spiegazione", "errore"),
        "Manuale Tecnico": ("passo", "verifica", "errore", "esempio"),
        "Religioso / Teologico": ("testo", "tradizion", "interpret", "contesto"),
        "Spirituale / Esoterico": ("pratica", "esperienza", "limite", "esercizio"),
        "Meditazione / Mindfulness": ("esercizio", "respiro", "osserv", "durata"),
        "Business & Marketing": ("caso", "metrica", "azione", "cliente"),
        "Economia e Finanza": ("dato", "rischio", "scenario", "esempio"),
        "Romanzo Rosa": ("dialog", "scena", "personagg", "relazione"),
        "Thriller / Noir": ("scena", "indizio", "conflitto", "personagg"),
        "Fantasy": ("scena", "personagg", "conflitto", "mondo"),
        "Fantascienza": ("scena", "personagg", "conseguenz", "tecnolog"),
        "Manuale Psicologico": ("esercizio", "esempio", "limite", "pratica"),
        "Biografia": ("evento", "contesto", "periodo", "scelta"),
        "Ricettario": ("ingredient", "procedimento", "cottura", "porzion"),
        "Test Prep (Preparazione Esami)": ("domanda", "risposta", "esercizio", "errore"),
        "Narrativo": ("scena", "personagg", "azione", "dialog"),
        "Romanzo Classico": ("scena", "personagg", "azione", "dialog"),
        "Contemporaneo": ("scena", "personagg", "azione", "dialog"),
        "Self-Help": ("esercizio", "passo", "esempio", "verifica"),
        "Manuale Pratico": ("passo", "material", "errore", "risultato"),
        "Storico": ("evento", "contesto", "fonte", "periodo")
    }
    segnali = segnali_per_genere.get(genere, ())
    trovati = sum(1 for segnale in segnali if segnale in basso)
    if genericita >= 5 and trovati < 2:
        return "eccesso di formule generiche senza esempi, scene, dati, procedure o strumenti specifici del genere"
    return ""


def genera_apertura_di_parte(prompt, system_prompt, sezione, genere, lingua, limite_output):
    """Genera e valida integralmente le aperture di Parte.

    Una Parte è una cornice editoriale breve, non un capitolo normale. Perciò
    ha una consegna dedicata: il testo deve essere leggibile, concluso e
    sufficiente a orientare il lettore, ma non viene scartato soltanto perché
    il modello omette il marcatore tecnico di fine risposta.
    """
    istruzione = prompt + f"""

APERTURA EDITORIALE DELLA PARTE
Questa è la Parte '{sezione}'. Scrivi un'apertura autonoma di 70-120 parole.
Spiega quale passaggio del percorso affronta questa Parte, come il lettore può
usarla e quale cambiamento concreto prepara. Non elencare né anticipare i
capitoli successivi. Chiudi con una frase piena e definitiva.
"""
    ultimo_testo = ""
    for tentativo in range(3):
        correzione = "" if tentativo == 0 else f"""

RISCRITTURA OBBLIGATORIA
La bozza precedente non era una vera apertura conclusa. Riscrivi integralmente
la Parte '{sezione}' in 70-120 parole. Non usare elenchi, titoli interni,
puntini di sospensione o frasi sospese. L'ultima frase deve chiudere davvero
l'idea, senza rimandare la conclusione a una sezione successiva.
"""
        testo, marcatore = genera_bozza_sezione_con_chiusura(
            istruzione + correzione,
            system_prompt,
            sezione,
            lingua,
            max_completion_tokens=min(max(320, limite_output), 700),
            addebita=(tentativo == 0),
        )
        ultimo_testo = pulisci_testo_editoriale(testo)
        testo_concluso = not criticita_specificita(
            ultimo_testo, genere, sezione, "Compatto", ""
        )
        if len(ultimo_testo.split()) >= 40 and testo_concluso:
            # Per le Parti il controllo reale è il finale completo; il
            # marcatore tecnico resta preferibile ma non può bloccare il libro.
            return ultimo_testo
    raise RuntimeError(
        f"L'apertura '{sezione}' non ha prodotto un testo concluso dopo tre tentativi. "
        "Le altre sezioni non sono state modificate."
    )


def genera_prefazione_del_libro(prompt, system_prompt, sezione, lingua, limite_output):
    """Genera in modo affidabile la prima sezione della coda editoriale.

    La Prefazione e' più breve di un capitolo, ma non deve mai essere la causa
    di un libro bloccato al primo passo. Il primo tentativo è quello a credito;
    tutti i recuperi sono gratuiti. Un testo già completo e concluso viene
    mantenuto come base di sicurezza, invece di essere scartato solo perché
    non ha raggiunto una soglia ideale al primo tentativo.
    """
    istruzione_base = prompt + f"""

CONSEGNA FINALE — PREFAZIONE
Restituisci esclusivamente il corpo della Prefazione iniziale in lingua
{lingua}, senza titolo, Markdown, elenco, URL, fonti, citazioni o marcatore
tecnico. Scrivi indicativamente 140-220 parole. Accogli il lettore, chiarisci
il bisogno da cui nasce il libro e il percorso che trovera', senza riassumere
o anticipare i capitoli. Mantieni con precisione genere, tono e punto di vista
gia' indicati nel brief. Termina con una frase piena, definitiva e pertinente.
"""
    minimo_ideale = 100
    minimo_salvabile = 60
    migliore_prefazione_completa = ""
    ultimo_errore = ""
    # Il primo tentativo è l'unica nuova operazione addebitata. I due recuperi
    # successivi intervengono esclusivamente se la stessa prima risposta non
    # era utilizzabile: non sono nuove richieste dell'utente.
    for tentativo in range(3):
        recupero = "" if tentativo == 0 else """

RISCRITTURA DI RECUPERO
La Prefazione precedente era troppo breve oppure non risultava chiaramente
conclusa. Riscrivi da zero una Prefazione autonoma di almeno 100 parole, senza
aggiungere indice, titoli interni o parti del libro. L'ultimo periodo deve
chiudere un'idea concreta e non può terminare con virgola, due punti, elenco o
parola di collegamento.
"""
        try:
            candidato = genera_sezione_con_ripetizione(
                istruzione_base + recupero,
                system_prompt,
                sezione,
                lingua,
                tentativi=1,
                max_completion_tokens=min(max(420, limite_output), 760),
                addebita=(tentativo == 0),
            )
        except Exception as exc:
            ultimo_errore = str(exc)
            continue

        testo = pulisci_testo_editoriale(candidato).strip()
        if not testo or testo.upper().startswith("ERRORE:"):
            ultimo_errore = "risposta vuota o non utilizzabile"
            continue
        if motivo_chiusura_tecnica(testo):
            ultimo_errore = "ultima frase non conclusa"
            continue
        # Conserviamo ogni versione realmente conclusa: se il motore risponde
        # con una Prefazione valida ma un po' breve, il recupero può ampliarla
        # senza fare sparire il testo già buono.
        if len(testo.split()) > len(migliore_prefazione_completa.split()):
            migliore_prefazione_completa = testo
        if len(testo.split()) >= minimo_ideale:
            return testo
        ultimo_errore = "prefazione completa ma troppo breve"

    if migliore_prefazione_completa:
        integrazione_prompt = f"""
Completa la Prefazione qui sotto senza riscriverla e senza aggiungere titoli,
elenchi, fonti o commenti. Aggiungi un solo breve paragrafo coerente sul
percorso del lettore, poi chiudi con una frase completa. Restituisci solo
l'integrazione finale in lingua {lingua}.

PREFAZIONE DA COMPLETARE:
{migliore_prefazione_completa}
"""
        try:
            integrazione = genera_sezione_con_ripetizione(
                integrazione_prompt, system_prompt, sezione, lingua,
                tentativi=2,
                max_completion_tokens=min(max(320, limite_output), 620),
                addebita=False,
            )
            testo_integrato = pulisci_testo_editoriale(
                f"{migliore_prefazione_completa}\n\n{integrazione}"
            ).strip()
            if (
                len(testo_integrato.split()) >= minimo_salvabile
                and not motivo_chiusura_tecnica(testo_integrato)
            ):
                return testo_integrato
        except Exception as exc:
            ultimo_errore = str(exc)

        # Una Prefazione compiuta non deve bloccare l'intero libro soltanto
        # perché il modello non ha ampliato un testo già valido. Il controllo
        # finale potrà comunque segnalarla come breve, ma memoria e anteprima
        # riceveranno sempre la prima sezione completa.
        if len(migliore_prefazione_completa.split()) >= minimo_salvabile:
            return migliore_prefazione_completa

    raise RuntimeError(
        "La Prefazione non ha prodotto un testo completo e salvabile. "
        f"Dettaglio: {ultimo_errore or 'risposta non disponibile'}."
    )


def genera_contenuto_editoriale(prompt, system_prompt, sezione, indice, trama, genere, obiettivo, lingua, profilo_lunghezza="Standard KDP"):
    """Mantiene il flusso comune per tutti i generi e applica la logica speciale solo quando serve."""
    if sezione_simulazione_test_prep(sezione, indice, genere):
        return genera_simulazione_test_prep(prompt, system_prompt, sezione, indice, trama, obiettivo, lingua)
    limite_output = PROFILI_LUNGHEZZA_STESURA.get(profilo_lunghezza, PROFILI_LUNGHEZZA_STESURA["Standard KDP"])["max_completion_tokens"]
    if tipo_sezione_editoriale(sezione) == "prefazione":
        # Il system prompt globale contiene il profilo di lunghezza del
        # manoscritto. Per la Prefazione, più breve per definizione, questa
        # direttiva locale ha priorità assoluta in entrambi i cervelli.
        system_prefazione = system_prompt + """

ECCEZIONE VINCOLANTE — PREFAZIONE
Stai scrivendo la Prefazione, non un capitolo né un sottocapitolo. Ignora
qualsiasi istruzione generale che chieda la lunghezza delle sezioni standard:
per questa sola risposta il limite corretto è 140-220 parole. Restituisci
soltanto una Prefazione completa, con ultima frase conclusa.
"""
        return genera_prefazione_del_libro(prompt, system_prefazione, sezione, lingua, limite_output)
    if tipo_sezione_editoriale(sezione) == "parte":
        return genera_apertura_di_parte(
            prompt, system_prompt, sezione, genere, lingua, limite_output
        )
    testo, consegna_confermata = genera_bozza_sezione_con_chiusura(
        prompt, system_prompt, sezione, lingua, max_completion_tokens=limite_output
    )
    criticita = criticita_consegna_sezione(
        testo, consegna_confermata, genere, sezione, profilo_lunghezza, indice
    )
    # La completezza non è un controllo da lasciare all'utente dopo l'export:
    # una bozza che termina a metà non viene mai salvata. I tentativi di
    # riparazione sono a carico del servizio, perché correggono un difetto di
    # generazione e non una nuova richiesta editoriale dell'utente.
    for tentativo_riparazione in range(2):
        if not criticita:
            return pulisci_testo_editoriale(testo)
        minimo_richiesto, massimo_richiesto = vincolo_parole_con_tolleranza(profilo_lunghezza)
        istruzione_lunghezza = (
            f"Il conteggio automatico della bozza è {len(testo.split())} parole. "
            f"La nuova sezione deve contenere tra {minimo_richiesto} e {massimo_richiesto} parole, "
            "senza riempitivi. Non interrompere il testo prima di aver raggiunto il minimo."
            if "testo troppo breve" in criticita else
            "Rispetta con precisione la lunghezza indicata nel prompt originario."
        )
        testo, consegna_confermata = genera_bozza_sezione_con_chiusura(
            prompt + f"""

REVISIONE OBBLIGATORIA DI QUALITÀ
La prima bozza è stata rifiutata perché presenta: {criticita}.
{istruzione_lunghezza}
Riscrivi integralmente la sezione. Ogni paragrafo deve aggiungere un fatto, una scena, una procedura,
un esempio, un caso, un esercizio, un dato o una conseguenza specifica del genere '{genere}'.
Elimina frasi motivazionali, definizioni vaghe e ripetizioni. Non descrivere ciò che il lettore potrebbe fare:
            mostra il contenuto concreto richiesto dal titolo della sezione.
Chiudi l'ultima idea con una frase completa e significativa, rispettando rigorosamente il limite di parole.
Prima di inviare il testo, rileggi le ultime due frasi: non possono terminare con una parola di collegamento,
una virgola, due punti, un trattino, un elenco incompleto o un ragionamento lasciato a metà. Se lo spazio non
basta, elimina l'ultimo dettaglio secondario e termina con una conclusione breve ma pienamente compiuta.
""",
            system_prompt, sezione, lingua, max_completion_tokens=limite_output,
            addebita=False,
        )
        criticita = criticita_consegna_sezione(
            testo, consegna_confermata, genere, sezione, profilo_lunghezza, indice
        )
    # Una frase tronca rende inaffidabile l'intera bozza: la scartiamo e
    # rigeneriamo integralmente la sezione con un margine tecnico aggiuntivo.
    # Non aggiungiamo mai una semplice coda a un testo incompleto.
    if criticita and criticita.startswith("ragionamento non concluso"):
        limite_riscrittura = int(limite_output * 1.20)
        profilo_riscrittura = PROFILI_LUNGHEZZA_STESURA.get(
            profilo_lunghezza, PROFILI_LUNGHEZZA_STESURA["Standard KDP"]
        )
        massimo_ordinario = vincolo_parole_con_tolleranza(profilo_lunghezza)[1]
        massimo_eccezionale = math.ceil(profilo_riscrittura["max_parole"] * 1.10)
        testo, consegna_confermata = genera_bozza_sezione_con_chiusura(
            prompt + f"""

RISCRITTURA INTEGRALE OBBLIGATORIA DI QUALITÀ
La bozza precedente viene scartata perché termina con un ragionamento incompleto.
Riscrivi da zero l'intera sezione '{sezione}', senza recuperare né proseguire la bozza precedente.
Mantieni tutti i vincoli editoriali, sviluppa il contenuto in modo originale e concludi l'ultima idea
con una frase piena, definitiva e pertinente. Dedica l'ultimo paragrafo alla conclusione; non iniziare
un nuovo esempio o elenco negli ultimi 80-100 vocaboli. Prima di inviare, verifica che il carattere
finale sia una chiusura di frase e che non rimangano elenchi, passaggi o ragionamenti sospesi.
Mantieni di norma il limite di {massimo_ordinario} parole. Solo se indispensabile per concludere bene
un ragionamento complesso, questa specifica sezione può arrivare eccezionalmente a {massimo_eccezionale}
parole: usa tale margine solo per contenuto utile, mai per ripetizioni o riempitivi.
""",
            system_prompt, sezione, lingua,
            max_completion_tokens=limite_riscrittura,
            addebita=False,
        )
        criticita = criticita_consegna_sezione(
            testo, consegna_confermata, genere, sezione, profilo_lunghezza, indice
        )
    # Se le riscritture hanno prodotto un testo completo ma ancora corto, non
    # scartiamo una buona base e non blocchiamo il libro: chiediamo soltanto
    # l'integrazione strettamente necessaria. È un recupero interno gratuito,
    # usato da GPT e DeepSeek, che mantiene coerente il ragionamento già
    # scritto e porta davvero la sezione nel margine del 5%.
    if criticita and "testo troppo breve" in criticita:
        minimo_richiesto, massimo_richiesto = vincolo_parole_con_tolleranza(profilo_lunghezza)
        for _ in range(2):
            parole_attuali = len(testo.split())
            parole_da_aggiungere = max(1, minimo_richiesto - parole_attuali)
            if parole_da_aggiungere <= 0:
                break
            integrazione, integrazione_conclusa = genera_bozza_sezione_con_chiusura(
                f"""
Completa la sezione qui sotto senza riscriverla e senza ripeterne i concetti.
Scrivi SOLO l'integrazione finale, pronta da aggiungere dopo l'ultimo paragrafo.
Servono almeno {parole_da_aggiungere} parole nuove e non devi superare {massimo_richiesto} parole totali.
Approfondisci un solo aspetto concreto e pertinente al titolo, poi chiudi con una frase completa.
Non introdurre titoli, fonti, note, promesse o elenchi lasciati a metà.

TITOLO DELLA SEZIONE: {sezione}
TESTO GIÀ APPROVATO NELLA FORMA:
{testo}
""",
                system_prompt, sezione, lingua, max_completion_tokens=limite_output,
                addebita=False,
            )
            if not integrazione or not integrazione_conclusa:
                continue
            testo = pulisci_testo_editoriale(f"{testo}\n\n{integrazione}")
            consegna_confermata = True
            criticita = criticita_consegna_sezione(
                testo, consegna_confermata, genere, sezione, profilo_lunghezza, indice
            )
            if not criticita:
                return testo
    if criticita:
        # Un testo realmente troncato non è pubblicabile e deve essere
        # rigenerato. Gli altri rilievi (densità o lunghezza) non possono però
        # far sparire una sezione già completa: viene conservata, segnalata nel
        # controllo finale e potrà essere migliorata con RIELABORA CON IA.
        if criticita.startswith("ragionamento non concluso"):
            raise RuntimeError(
                f"La sezione '{sezione}' termina in modo incompleto ({criticita}). "
                "Non è stata salvata: riprova senza perdere le altre sezioni."
            )
        avvisi = st.session_state.setdefault("avvisi_qualita_sezioni", {})
        avvisi[sezione] = criticita
    # Le ricerche web sono riservate a leggi, prezzi, versioni, requisiti e altri
    # dati soggetti a cambiamento; i contenuti didattici stabili non consumano credito web.
    generi_con_verifica_estesa = {
        "Saggio Scientifico", "Quiz Scientifico", "Manuale Tecnico", "Economia e Finanza",
        "Biografia", "Test Prep (Preparazione Esami)", "Storico"
    }
    # La verifica online è una scelta esplicita dell'utente: costa 2 crediti e
    # viene proposta nel relativo pulsante del capitolo, senza addebiti nascosti
    # durante la normale scrittura di una sezione.
    return pulisci_testo_editoriale(testo)


def scrivi_contenuto_dettagliato(sezione, indice, trama, genere, tipologia, stile, punto_di_vista,
                                 obiettivo, lingua, approfondimenti, profilo_lunghezza):
    """Scrive e salva una sezione con un recupero indipendente dai controlli.

    Prefazione, Parte e prima sezione concreta sono le voci piu' esposte a
    conflitti fra vincoli di lunghezza e controlli automatici. Se il motore
    editoriale non riesce a convalidarle, facciamo un ultimo tentativo diretto
    usando lo stesso brief: non saltiamo la voce, non passiamo alla successiva
    e non perdiamo il resto del libro.
    """
    prompt = crea_prompt_stesura_sezione(
        sezione, indice, trama, genere, tipologia, stile, punto_di_vista,
        obiettivo, lingua, approfondimenti, profilo_lunghezza,
    )
    limite_output = PROFILI_LUNGHEZZA_STESURA.get(
        profilo_lunghezza,
        PROFILI_LUNGHEZZA_STESURA["Standard KDP"],
    )["max_completion_tokens"]
    errore_principale = None
    try:
        contenuto_generato = genera_contenuto_editoriale(
            prompt, S_PROMPT, sezione, indice, trama, genere, obiettivo,
            lingua, profilo_lunghezza,
        )
    except Exception as exc:
        errore_principale = exc
        contenuto_generato = ""

    testo_pulito = pulisci_testo_editoriale(contenuto_generato).strip()
    testo_valido = (
        len(testo_pulito.split()) >= 30
        and not testo_pulito.upper().startswith("ERRORE:")
        and not motivo_chiusura_tecnica(testo_pulito)
    )

    if not testo_valido:
        # Recupero neutro: non chiede il marcatore interno e non reintroduce
        # la soglia dei capitoli per le aperture del libro. E' gratuito per
        # l'utente, poiche' sostituisce una risposta che non era salvabile.
        prompt_recupero = prompt + f"""

RECUPERO DIRETTO DELLA SEZIONE
Scrivi ora esclusivamente il testo completo e leggibile della sezione
'{sezione}'. Non usare marcatori, titoli interni, commenti sul processo,
Markdown, URL o fonti. Mantieni tutti i parametri gia' presenti nel brief.
Completa ogni frase e chiudi l'ultima idea con un punto, punto interrogativo
o punto esclamativo. Non passare alla sezione successiva e non riassumere
l'indice. Restituisci un testo autonomo di almeno 80 parole.
"""
        try:
            recupero = genera_sezione_con_ripetizione(
                prompt_recupero,
                S_PROMPT,
                sezione,
                lingua,
                tentativi=2,
                max_completion_tokens=max(420, limite_output),
                addebita=False,
            )
            testo_recupero = pulisci_testo_editoriale(recupero).strip()
            if (
                len(testo_recupero.split()) >= 30
                and not testo_recupero.upper().startswith("ERRORE:")
                and not motivo_chiusura_tecnica(testo_recupero)
            ):
                testo_pulito = testo_recupero
                testo_valido = True
        except Exception as exc:
            if errore_principale is None:
                errore_principale = exc

    if not testo_valido:
        dettaglio = str(errore_principale or "risposta non valida")
        raise RuntimeError(f"nessun testo completo restituito per '{sezione}': {dettaglio}")

    # Ogni sezione, inclusa la prima Prefazione, entra subito nella memoria
    # stabile della stesura completa prima del rerun della coda automatica.
    return scrivi_sezione_stesura_completa(sezione, testo_pulito)


def genera_e_conferma_sezione_del_libro(
    sezione, indice, trama, genere, tipologia, stile, punto_di_vista,
    obiettivo, lingua, approfondimenti, profilo_lunghezza,
):
    """Usa il percorso del pulsante dettagliato e verifica subito la memoria.

    La stesura di una sezione singola e quella del libro intero devono essere
    indistinguibili: stesso generatore, stessa validazione, stesso archivio
    editoriale. La verifica finale impedisce alla coda di avanzare se il testo
    non è realmente recuperabile da editor, anteprima ed export.
    """
    testo_generato = scrivi_contenuto_dettagliato(
        sezione, indice, trama, genere, tipologia, stile, punto_di_vista,
        obiettivo, lingua, approfondimenti, profilo_lunghezza,
    )
    testo_salvato = pulisci_testo_editoriale(
        contenuto_memorizzato_puro(sezione)
    ).strip()
    if not testo_salvato:
        raise RuntimeError(
            f"'{sezione}' è stata generata ma non risulta nella memoria del progetto"
        )
    # Conserva una seconda copia coerente con l'archivio della stesura completa
    # senza dipendere da widget o cache del browser.
    scrivi_sezione_stesura_completa(sezione, testo_salvato or testo_generato)
    testo_verificato = pulisci_testo_editoriale(
        contenuto_memorizzato_puro(sezione)
    ).strip()
    if not testo_verificato:
        raise RuntimeError(
            f"'{sezione}' non è disponibile dopo il salvataggio verificato"
        )
    # La Prefazione della stesura completa possiede una ricevuta propria.
    # Non dipende dal nome del widget, dalla tab aperta o dalla ricostruzione
    # della coda: dopo il primo salvataggio non potrà più essere generata una
    # seconda volta soltanto perché Streamlit ha ridisegnato la pagina.
    if sezione_prefazione(sezione) and "job_scrittura_prefazione_confermata" in st.session_state:
        st.session_state["job_scrittura_prefazione_confermata"] = testo_verificato
    return testo_verificato

# NUOVA FUNZIONE: Motore Decisionale per attivare i 3 Cervelli in base alla Sidebar
def valuta_approccio_neurologico(genere, stile, narrativa):
    """
    Decide se l'argomento e lo stile richiedono la manipolazione dei 3 cervelli
    o un approccio più analitico/oggettivo.
    """
    trigger_neuro_stile = ["Persuasivo (Neuromarketing Applicato)", "Conversazionale ed Empatico", "Storytelling Immersivo", "Epico ed Evocativo"]
    trigger_neuro_narrativa = ["Coinvolgente e Narrativo", "Ispirazionale e Motivante", "Storytelling Emozionale", "Diretto e Pratico (Action-oriented)"]
    trigger_neuro_genere = ["Business & Marketing", "Economia e Finanza", "Manuale Psicologico", "Romanzo Rosa", "Thriller / Noir", "Spirituale / Esoterico"]
    
    if stile in trigger_neuro_stile or narrativa in trigger_neuro_narrativa or genere in trigger_neuro_genere:
        return True
    return False

# ======================================================================================================================
# 6. SIDEBAR: SETUP EDITORIALE AVANZATO E CARICAMENTO FONTI
# ======================================================================================================================
# Il progetto cloud non viene più caricato automaticamente: l'utente sceglie
# esplicitamente quando recuperarlo con “RIAGGIORNA ALL'ULTIMA STESURA”. Il
# flag esiste soltanto nel rerun immediatamente successivo a quel pulsante.
if st.session_state.get("autosave_snapshot_da_ripristinare"):
    ripristina_progetto_salvato()

with st.sidebar:
    lingua_scelta = st.selectbox(testo_ui("lingua"), [""] + list(TRADUZIONI.keys()), key="editor_language", format_func=lambda valore: valore or testo_ui("seleziona"))
    lingua_sel = lingua_scelta or "Italiano"
    L = TRADUZIONI.get(lingua_sel, TRADUZIONI["Italiano"])
    intestazioni_sidebar = {
        "Italiano": ("Percorso: configura → indice → scrivi → controlla ed esporta", "Configura il progetto", "Fonti e ricerca (opzionale)", "Dettagli editoriali", "Sessione e memoria", "Salva per conservare anche le modifiche manuali; il ripristino è sempre volontario."),
        "English": ("Path: configure → outline → write → review and export", "Configure your project", "Sources and research (optional)", "Editorial details", "Session and memory", "Save to keep manual edits too; restoring is always voluntary."),
        "Español": ("Ruta: configura → índice → escribe → revisa y exporta", "Configura el proyecto", "Fuentes y búsqueda (opcional)", "Detalles editoriales", "Sesión y memoria", "Guarda también los cambios manuales; restaurar siempre es voluntario."),
        "Français": ("Parcours : configurez → planifiez → rédigez → contrôlez et exportez", "Configurer le projet", "Sources et recherche (facultatif)", "Détails éditoriaux", "Session et mémoire", "Enregistrez aussi les modifications manuelles ; la restauration reste volontaire."),
        "Deutsch": ("Ablauf: konfigurieren → Gliederung → schreiben → prüfen und exportieren", "Projekt konfigurieren", "Quellen und Recherche (optional)", "Redaktionelle Details", "Sitzung und Speicher", "Speichern Sie auch manuelle Änderungen; eine Wiederherstellung bleibt freiwillig."),
        "Română": ("Parcurs: configurează → cuprins → scrie → verifică și exportă", "Configurează proiectul", "Surse și cercetare (opțional)", "Detalii editoriale", "Sesiune și memorie", "Salvează și modificările manuale; restaurarea rămâne voluntară."),
        "Русский": ("Порядок: настройка → оглавление → текст → проверка и экспорт", "Настройка проекта", "Источники и поиск (необязательно)", "Редакционные параметры", "Сессия и память", "Сохраните и ручные правки; восстановление всегда выполняется по вашему выбору."),
        "العربية": ("المسار: الإعداد ← الفهرس ← الكتابة ← المراجعة والتصدير", "إعداد المشروع", "المصادر والبحث (اختياري)", "التفاصيل التحريرية", "الجلسة والذاكرة", "احفظ أيضاً التعديلات اليدوية؛ والاستعادة اختيارية دائماً."),
        "中文": ("流程：配置 → 目录 → 写作 → 检查并导出", "配置项目", "资料与研究（可选）", "编辑详情", "会话与记忆", "保存可保留手动修改；恢复始终由你决定。"),
    }.get(lingua_sel, ("Path: configure → outline → write → review and export", "Configure your project", "Sources and research (optional)", "Editorial details", "Session and memory", "Save to keep manual edits too; restoring is always voluntary."))
    st.caption(intestazioni_sidebar[0])
    st.title(L["side_tit"])
    st.markdown(f"#### 1 · {intestazioni_sidebar[1]}")
    provider_ia = st.selectbox(
        testo_ui("cervello", lingua_sel),
        ["GPT-5.4 (OpenAI)", "DeepSeek V4 Pro"],
        key="provider_ia",
        help="GPT conserva ricerca web e verifica copyright web. Le immagini vengono esclusivamente caricate dall'utente; DeepSeek Pro usa un motore separato per ricerca fonti con registro visibile, indice, fonti caricate, scrittura e controlli editoriali, con consumi più leggeri.",
    )
    tariffari_sidebar = {
        "Italiano": {
            "gpt_info": "GPT-5.4 attivo: usa il cervello completo, incluse ricerca web, verifica copyright web e generazione immagini.",
            "gpt_title": "Tariffario GPT-5.4",
            "gpt": ["Scrittura, rigenerazione, quiz ed esempi: 1 credito per ogni operazione.", "Indice completo: 5 crediti (2 ricerca web + 3 progettazione editoriale).", "Voto indice: 1 credito; rigenerazione indice: 3 crediti.", "Verifica fatti: 2 crediti; report sintattico e metadati KDP: 1 credito ciascuno.", "Coerenza completa: 10 crediti; controllo successivo: 1 credito per blocco modificato.", "10 ricette: 10 crediti; immagine di capitolo: 5 crediti; copyright web: da 2 crediti."],
            "ds_info": "DeepSeek Pro attivo: scrittura, indice, ricerca fonti e controlli editoriali usano DeepSeek. Il copyright web e le immagini richiedono GPT e restano disattivati.",
            "ds_title": "Tariffario DeepSeek Pro",
            "ds": ["Rapporto DeepSeek/GPT: 1 a 3. Tre operazioni equivalenti da 1 credito GPT consumano 1 credito DeepSeek.", "Scrittura, rigenerazione, quiz ed esempi: 1 credito ogni 3 operazioni.", "Ricerca fonti web nativa + indice completo: circa 2 crediti in totale.", "Voto indice, report sintattico e metadati: 1 credito ogni 3 controlli equivalenti.", "Coerenza completa: circa 4 crediti; aggiornamenti: 1 credito ogni 3 blocchi.", "10 ricette: circa 4 crediti. Copyright web e immagini sono disponibili solo con GPT."],
        },
        "English": {
            "gpt_info": "GPT-5.4 is active: it uses the complete engine, including web research, web copyright checks and image generation.", "gpt_title": "GPT-5.4 pricing", "gpt": ["Writing, rewriting, quizzes and examples: 1 credit per action.", "Complete outline: 5 credits (2 web research + 3 editorial design).", "Outline review: 1 credit; outline rewrite: 3 credits.", "Fact check: 2 credits; syntax report and KDP metadata: 1 credit each.", "Full consistency check: 10 credits; later checks: 1 credit per changed block.", "10 recipes: 10 credits; chapter image: 5 credits; web copyright: from 2 credits."],
            "ds_info": "DeepSeek Pro is active: writing, outlines, source research and editorial checks use DeepSeek. Web copyright and images require GPT and are disabled.", "ds_title": "DeepSeek Pro pricing", "ds": ["DeepSeek/GPT ratio: 1 to 3. Three equivalent 1-credit GPT actions use 1 DeepSeek credit.", "Writing, rewriting, quizzes and examples: 1 credit every 3 actions.", "Native web source research + complete outline: about 2 credits total.", "Outline review, syntax report and metadata: 1 credit every 3 equivalent checks.", "Full consistency: about 4 credits; updates: 1 credit every 3 blocks.", "10 recipes: about 4 credits. Web copyright and images are available with GPT only."],
        },
        "Español": {
            "gpt_info": "GPT-5.4 está activo: utiliza el motor completo, incluida la investigación web, la comprobación de copyright web y la generación de imágenes.", "gpt_title": "Tarifas de GPT-5.4", "gpt": ["Escritura, reescritura, cuestionarios y ejemplos: 1 crédito por acción.", "Índice completo: 5 créditos (2 de investigación web + 3 de diseño editorial).", "Evaluación del índice: 1 crédito; regeneración: 3 créditos.", "Verificación de hechos: 2 créditos; informe sintáctico y metadatos KDP: 1 crédito cada uno.", "Coherencia completa: 10 créditos; controles posteriores: 1 crédito por bloque modificado.", "10 recetas: 10 créditos; imagen de capítulo: 5 créditos; copyright web: desde 2 créditos."],
            "ds_info": "DeepSeek Pro está activo: escritura, índice, investigación de fuentes y controles editoriales usan DeepSeek. El copyright web y las imágenes requieren GPT y están desactivados.", "ds_title": "Tarifas de DeepSeek Pro", "ds": ["Relación DeepSeek/GPT: 1 a 3. Tres operaciones GPT equivalentes de 1 crédito consumen 1 crédito DeepSeek.", "Escritura, reescritura, cuestionarios y ejemplos: 1 crédito cada 3 operaciones.", "Investigación nativa de fuentes web + índice completo: unos 2 créditos en total.", "Evaluación del índice, informe sintáctico y metadatos: 1 crédito cada 3 controles equivalentes.", "Coherencia completa: unos 4 créditos; actualizaciones: 1 crédito cada 3 bloques.", "10 recetas: unos 4 créditos. Copyright web e imágenes solo con GPT."],
        },
        "Français": {
            "gpt_info": "GPT-5.4 est actif : il utilise le moteur complet, avec recherche web, contrôle de copyright web et génération d’images.", "gpt_title": "Tarifs GPT-5.4", "gpt": ["Rédaction, réécriture, quiz et exemples : 1 crédit par action.", "Plan complet : 5 crédits (2 recherche web + 3 conception éditoriale).", "Évaluation du plan : 1 crédit ; régénération : 3 crédits.", "Vérification des faits : 2 crédits ; rapport syntaxique et métadonnées KDP : 1 crédit chacun.", "Cohérence complète : 10 crédits ; contrôles suivants : 1 crédit par bloc modifié.", "10 recettes : 10 crédits ; image de chapitre : 5 crédits ; copyright web : dès 2 crédits."],
            "ds_info": "DeepSeek Pro est actif : rédaction, plan, recherche de sources et contrôles éditoriaux utilisent DeepSeek. Le copyright web et les images nécessitent GPT et sont désactivés.", "ds_title": "Tarifs DeepSeek Pro", "ds": ["Rapport DeepSeek/GPT : 1 pour 3. Trois opérations GPT équivalentes à 1 crédit consomment 1 crédit DeepSeek.", "Rédaction, réécriture, quiz et exemples : 1 crédit toutes les 3 opérations.", "Recherche native de sources web + plan complet : environ 2 crédits au total.", "Évaluation du plan, rapport syntaxique et métadonnées : 1 crédit tous les 3 contrôles équivalents.", "Cohérence complète : environ 4 crédits ; mises à jour : 1 crédit tous les 3 blocs.", "10 recettes : environ 4 crédits. Copyright web et images uniquement avec GPT."],
        },
        "Deutsch": {
            "gpt_info": "GPT-5.4 ist aktiv: Es nutzt die vollständige Engine einschließlich Webrecherche, Web-Copyright-Prüfung und Bildgenerierung.", "gpt_title": "GPT-5.4 Preise", "gpt": ["Schreiben, Überarbeiten, Quiz und Beispiele: 1 Credit pro Vorgang.", "Vollständige Gliederung: 5 Credits (2 Webrecherche + 3 redaktionelle Planung).", "Gliederungsbewertung: 1 Credit; Neugenerierung: 3 Credits.", "Faktenprüfung: 2 Credits; Syntaxbericht und KDP-Metadaten: je 1 Credit.", "Vollständige Kohärenzprüfung: 10 Credits; weitere Prüfungen: 1 Credit pro geändertem Block.", "10 Rezepte: 10 Credits; Kapitelbild: 5 Credits; Web-Copyright: ab 2 Credits."],
            "ds_info": "DeepSeek Pro ist aktiv: Schreiben, Gliederung, Quellenrecherche und redaktionelle Kontrollen nutzen DeepSeek. Web-Copyright und Bilder benötigen GPT und sind deaktiviert.", "ds_title": "DeepSeek Pro Preise", "ds": ["DeepSeek/GPT-Verhältnis: 1 zu 3. Drei gleichwertige GPT-Vorgänge zu 1 Credit verbrauchen 1 DeepSeek-Credit.", "Schreiben, Überarbeiten, Quiz und Beispiele: 1 Credit je 3 Vorgänge.", "Native Webquellen-Recherche + vollständige Gliederung: insgesamt etwa 2 Credits.", "Gliederungsbewertung, Syntaxbericht und Metadaten: 1 Credit je 3 gleichwertige Prüfungen.", "Vollständige Kohärenz: etwa 4 Credits; Updates: 1 Credit je 3 Blöcke.", "10 Rezepte: etwa 4 Credits. Web-Copyright und Bilder nur mit GPT."],
        },
        "Română": {
            "gpt_info": "GPT-5.4 este activ: folosește motorul complet, inclusiv cercetare web, verificare copyright web și generare de imagini.", "gpt_title": "Tarife GPT-5.4", "gpt": ["Scriere, rescriere, quiz-uri și exemple: 1 credit per acțiune.", "Cuprins complet: 5 credite (2 cercetare web + 3 proiectare editorială).", "Evaluarea cuprinsului: 1 credit; regenerare: 3 credite.", "Verificarea faptelor: 2 credite; raport sintactic și metadate KDP: câte 1 credit.", "Coerență completă: 10 credite; controale ulterioare: 1 credit per bloc modificat.", "10 rețete: 10 credite; imagine de capitol: 5 credite; copyright web: de la 2 credite."],
            "ds_info": "DeepSeek Pro este activ: scrierea, cuprinsul, cercetarea surselor și controalele editoriale folosesc DeepSeek. Copyright-ul web și imaginile necesită GPT și sunt dezactivate.", "ds_title": "Tarife DeepSeek Pro", "ds": ["Raport DeepSeek/GPT: 1 la 3. Trei operațiuni GPT echivalente de 1 credit consumă 1 credit DeepSeek.", "Scriere, rescriere, quiz-uri și exemple: 1 credit la 3 operațiuni.", "Cercetare nativă de surse web + cuprins complet: circa 2 credite în total.", "Evaluare cuprins, raport sintactic și metadate: 1 credit la 3 controale echivalente.", "Coerență completă: circa 4 credite; actualizări: 1 credit la 3 blocuri.", "10 rețete: circa 4 credite. Copyright web și imagini doar cu GPT."],
        },
        "Русский": {
            "gpt_info": "GPT-5.4 активен: используется полный движок, включая веб-поиск, проверку авторских прав в сети и генерацию изображений.", "gpt_title": "Тарифы GPT-5.4", "gpt": ["Написание, переработка, тесты и примеры: 1 кредит за действие.", "Полная структура: 5 кредитов (2 веб-поиск + 3 редакционное проектирование).", "Оценка структуры: 1 кредит; регенерация: 3 кредита.", "Проверка фактов: 2 кредита; синтаксический отчёт и метаданные KDP: по 1 кредиту.", "Полная проверка связности: 10 кредитов; следующие проверки: 1 кредит за изменённый блок.", "10 рецептов: 10 кредитов; изображение главы: 5 кредитов; веб-copyright: от 2 кредитов."],
            "ds_info": "DeepSeek Pro активен: написание, структура, поиск источников и редакционные проверки используют DeepSeek. Веб-copyright и изображения требуют GPT и отключены.", "ds_title": "Тарифы DeepSeek Pro", "ds": ["Соотношение DeepSeek/GPT: 1 к 3. Три эквивалентных действия GPT по 1 кредиту используют 1 кредит DeepSeek.", "Написание, переработка, тесты и примеры: 1 кредит за 3 действия.", "Нативный поиск веб-источников + полная структура: около 2 кредитов всего.", "Оценка структуры, синтаксический отчёт и метаданные: 1 кредит за 3 эквивалентные проверки.", "Полная связность: около 4 кредитов; обновления: 1 кредит за 3 блока.", "10 рецептов: около 4 кредитов. Веб-copyright и изображения доступны только с GPT."],
        },
        "العربية": {
            "gpt_info": "GPT-5.4 نشط: يستخدم المحرك الكامل، بما في ذلك البحث على الويب وفحص حقوق النشر وتوليد الصور.", "gpt_title": "أسعار GPT-5.4", "gpt": ["الكتابة وإعادة الصياغة والاختبارات والأمثلة: رصيد واحد لكل عملية.", "فهرس كامل: 5 أرصدة (2 للبحث على الويب + 3 للتخطيط التحريري).", "تقييم الفهرس: رصيد واحد؛ إعادة التوليد: 3 أرصدة.", "التحقق من الحقائق: رصيدان؛ التقرير النحوي وبيانات KDP: رصيد واحد لكل منهما.", "فحص الاتساق الكامل: 10 أرصدة؛ الفحوصات اللاحقة: رصيد لكل كتلة معدلة.", "10 وصفات: 10 أرصدة؛ صورة الفصل: 5 أرصدة؛ حقوق النشر على الويب: من رصيدين."],
            "ds_info": "DeepSeek Pro نشط: الكتابة والفهرس وبحث المصادر والفحوصات التحريرية تستخدم DeepSeek. حقوق النشر على الويب والصور تحتاج GPT وهي معطلة.", "ds_title": "أسعار DeepSeek Pro", "ds": ["نسبة DeepSeek/GPT هي 1 إلى 3. ثلاث عمليات GPT مكافئة لرصيد واحد تستهلك رصيد DeepSeek واحداً.", "الكتابة وإعادة الصياغة والاختبارات والأمثلة: رصيد واحد كل 3 عمليات.", "بحث أصلي في مصادر الويب + فهرس كامل: نحو رصيدين إجمالاً.", "تقييم الفهرس والتقرير النحوي والبيانات: رصيد واحد كل 3 فحوصات مكافئة.", "اتساق كامل: نحو 4 أرصدة؛ التحديثات: رصيد واحد كل 3 كتل.", "10 وصفات: نحو 4 أرصدة. حقوق النشر على الويب والصور متاحة مع GPT فقط."],
        },
        "中文": {
            "gpt_info": "GPT-5.4 已启用：使用完整引擎，包括网页研究、网页版权检查和图片生成。", "gpt_title": "GPT-5.4 价格", "gpt": ["写作、改写、测验和示例：每项操作 1 积分。", "完整目录：5 积分（2 积分网页研究 + 3 积分编辑规划）。", "目录评估：1 积分；重新生成：3 积分。", "事实核查：2 积分；句法报告和 KDP 元数据：各 1 积分。", "完整一致性检查：10 积分；后续检查：每个修改区块 1 积分。", "10 道食谱：10 积分；章节图片：5 积分；网页版权：2 积分起。"],
            "ds_info": "DeepSeek Pro 已启用：写作、目录、来源研究和编辑检查使用 DeepSeek。网页版权和图片需要 GPT，现已禁用。", "ds_title": "DeepSeek Pro 价格", "ds": ["DeepSeek/GPT 比例为 1:3。三项相当于 1 GPT 积分的操作消耗 1 DeepSeek 积分。", "写作、改写、测验和示例：每 3 项操作 1 积分。", "原生网页来源研究 + 完整目录：总计约 2 积分。", "目录评估、句法报告和元数据：每 3 项等效检查 1 积分。", "完整一致性：约 4 积分；更新：每 3 个区块 1 积分。", "10 道食谱：约 4 积分。网页版权和图片仅可使用 GPT。"],
        },
    }
    # Questa seconda definizione è il tariffario effettivo dell'app. Mantiene
    # separati i due cervelli, mostra il costo reale dell'indice e non propone
    # più immagini generate: il software accetta soltanto file esterni.
    tariffari_sidebar.update({
        "Italiano": {
            "gpt_info": "GPT-5.4 attivo: ricerca web, controlli editoriali e verifica copyright web. Le immagini si caricano dall'esterno.",
            "gpt_title": "Tariffario GPT-5.4",
            "gpt": ["Scrittura, rigenerazione, quiz ed esempi: 1 credito per operazione.", "Indice completo: 10 crediti (4 ricerca fonti + 6 progettazione editoriale).", "Voto indice: 2 crediti; rigenerazione indice: 6 crediti.", "Verifica fatti: 2 crediti; report sintattico e metadati KDP: 1 credito ciascuno.", "Coerenza completa: 10 crediti; controllo successivo: 1 credito per blocco modificato.", "10 ricette: 10 crediti; immagini esterne: gratuite da caricare; copyright web: da 2 crediti."],
            "ds_info": "DeepSeek Pro attivo: scrittura, indice, ricerca fonti e controlli editoriali usano DeepSeek. Le immagini restano esclusivamente esterne.",
            "ds_title": "Tariffario DeepSeek Pro",
            "ds": ["Rapporto DeepSeek/GPT: 1 a 3.", "Scrittura, rigenerazione, quiz ed esempi: 1 credito ogni 3 operazioni.", "Ricerca fonti nativa + indice completo: circa 3⅓ crediti.", "Voto indice: 1 credito; report sintattico e metadati: 1 credito ogni 3 controlli equivalenti.", "Coerenza completa: circa 4 crediti; aggiornamenti: 1 credito ogni 3 blocchi.", "10 ricette: circa 4 crediti. Le immagini esterne non consumano crediti."],
        },
        "English": {
            "gpt_info": "GPT-5.4 is active: web research, editorial checks and web copyright screening. Images are uploaded externally.", "gpt_title": "GPT-5.4 pricing",
            "gpt": ["Writing, rewriting, quizzes and examples: 1 credit per action.", "Complete outline: 10 credits (4 source research + 6 editorial design).", "Outline review: 2 credits; outline rewrite: 6 credits.", "Fact check: 2 credits; syntax report and KDP metadata: 1 credit each.", "Full consistency check: 10 credits; later checks: 1 credit per changed block.", "10 recipes: 10 credits; external image upload is free; web copyright: from 2 credits."],
            "ds_info": "DeepSeek Pro is active for writing, outlines, source research and editorial checks. Images remain external uploads only.", "ds_title": "DeepSeek Pro pricing",
            "ds": ["DeepSeek/GPT ratio: 1 to 3.", "Writing, rewriting, quizzes and examples: 1 credit every 3 actions.", "Native source research + complete outline: about 3⅓ credits.", "Outline review: 1 credit; syntax report and metadata: 1 credit every 3 equivalent checks.", "Full consistency: about 4 credits; updates: 1 credit every 3 blocks.", "10 recipes: about 4 credits. External images use no credits."],
        },
        "Español": {"gpt_info": "GPT-5.4 activo: investigación web, controles editoriales y copyright web. Las imágenes se cargan externamente.", "gpt_title": "Tarifas de GPT-5.4", "gpt": ["Escritura, reescritura, cuestionarios y ejemplos: 1 crédito por acción.", "Índice completo: 10 créditos (4 investigación + 6 diseño editorial).", "Evaluación del índice: 2 créditos; regeneración: 6 créditos.", "Verificación de hechos: 2 créditos; informe sintáctico y metadatos KDP: 1 crédito cada uno.", "Coherencia completa: 10 créditos; controles posteriores: 1 crédito por bloque modificado.", "10 recetas: 10 créditos; carga externa de imágenes gratuita; copyright web: desde 2 créditos."], "ds_info": "DeepSeek Pro activo para escritura, índice, fuentes y controles. Las imágenes son solo cargas externas.", "ds_title": "Tarifas de DeepSeek Pro", "ds": ["Relación DeepSeek/GPT: 1 a 3.", "Escritura, reescritura, cuestionarios y ejemplos: 1 crédito cada 3 operaciones.", "Fuentes nativas + índice completo: unos 3⅓ créditos.", "Evaluación del índice: 1 crédito; informe sintáctico y metadatos: 1 crédito cada 3 controles.", "Coherencia completa: unos 4 créditos; actualizaciones: 1 crédito cada 3 bloques.", "Las imágenes externas no consumen créditos."]},
        "Français": {"gpt_info": "GPT-5.4 actif : recherche web, contrôles éditoriaux et copyright web. Les images sont importées de l’extérieur.", "gpt_title": "Tarifs GPT-5.4", "gpt": ["Rédaction, réécriture, quiz et exemples : 1 crédit par action.", "Plan complet : 10 crédits (4 recherche + 6 conception éditoriale).", "Évaluation du plan : 2 crédits ; régénération : 6 crédits.", "Vérification des faits : 2 crédits ; rapport syntaxique et métadonnées KDP : 1 crédit chacun.", "Cohérence complète : 10 crédits ; contrôles suivants : 1 crédit par bloc modifié.", "Import d’images externes gratuit ; copyright web : dès 2 crédits."], "ds_info": "DeepSeek Pro actif pour rédaction, plan, sources et contrôles. Images externes uniquement.", "ds_title": "Tarifs DeepSeek Pro", "ds": ["Rapport DeepSeek/GPT : 1 pour 3.", "Rédaction, réécriture, quiz et exemples : 1 crédit toutes les 3 actions.", "Sources natives + plan complet : environ 3⅓ crédits.", "Évaluation du plan : 1 crédit ; rapport et métadonnées : 1 crédit tous les 3 contrôles.", "Cohérence complète : environ 4 crédits.", "Les images externes ne consomment aucun crédit."]},
        "Deutsch": {"gpt_info": "GPT-5.4 aktiv: Webrecherche, redaktionelle Prüfungen und Web-Copyright. Bilder werden extern hochgeladen.", "gpt_title": "GPT-5.4 Preise", "gpt": ["Schreiben, Überarbeiten, Quiz und Beispiele: 1 Credit pro Vorgang.", "Vollständige Gliederung: 10 Credits (4 Recherche + 6 Redaktion).", "Gliederungsbewertung: 2 Credits; Neugenerierung: 6 Credits.", "Faktenprüfung: 2 Credits; Syntaxbericht und KDP-Metadaten: je 1 Credit.", "Vollständige Kohärenzprüfung: 10 Credits; weitere Prüfungen: 1 Credit pro geändertem Block.", "Externe Bilder kostenlos hochladen; Web-Copyright: ab 2 Credits."], "ds_info": "DeepSeek Pro aktiv für Text, Gliederung, Quellen und Prüfungen. Bilder sind ausschließlich externe Uploads.", "ds_title": "DeepSeek Pro Preise", "ds": ["DeepSeek/GPT-Verhältnis: 1 zu 3.", "Schreiben, Überarbeiten, Quiz und Beispiele: 1 Credit je 3 Vorgänge.", "Native Quellenrecherche + Gliederung: etwa 3⅓ Credits.", "Gliederungsbewertung: 1 Credit; Bericht und Metadaten: 1 Credit je 3 Prüfungen.", "Vollständige Kohärenz: etwa 4 Credits.", "Externe Bilder verbrauchen keine Credits."]},
        "Română": {"gpt_info": "GPT-5.4 activ: cercetare web, controale editoriale și copyright web. Imaginile sunt încărcate extern.", "gpt_title": "Tarife GPT-5.4", "gpt": ["Scriere, rescriere, quiz-uri și exemple: 1 credit per operațiune.", "Cuprins complet: 10 credite (4 cercetare + 6 proiectare editorială).", "Evaluare cuprins: 2 credite; regenerare: 6 credite.", "Verificare fapte: 2 credite; raport sintactic și metadate KDP: câte 1 credit.", "Coerență completă: 10 credite; ulterior 1 credit per bloc modificat.", "Încărcarea imaginilor externe este gratuită; copyright web: de la 2 credite."], "ds_info": "DeepSeek Pro activ pentru scriere, cuprins, surse și controale. Imagini externe numai.", "ds_title": "Tarife DeepSeek Pro", "ds": ["Raport DeepSeek/GPT: 1 la 3.", "Scriere, rescriere, quiz-uri și exemple: 1 credit la 3 operațiuni.", "Surse native + cuprins complet: circa 3⅓ credite.", "Evaluare cuprins: 1 credit; raport și metadate: 1 credit la 3 controale.", "Coerență completă: circa 4 credite.", "Imaginile externe nu consumă credite."]},
        "Русский": {"gpt_info": "GPT-5.4 активен: веб-поиск, редакторские проверки и веб-copyright. Изображения загружаются извне.", "gpt_title": "Тарифы GPT-5.4", "gpt": ["Текст, переработка, тесты и примеры: 1 кредит за действие.", "Полная структура: 10 кредитов (4 поиск + 6 редактура).", "Оценка структуры: 2 кредита; регенерация: 6 кредитов.", "Проверка фактов: 2 кредита; синтаксический отчёт и KDP-метаданные: по 1 кредиту.", "Полная проверка связности: 10 кредитов; далее 1 кредит за блок.", "Внешняя загрузка изображений бесплатна; веб-copyright: от 2 кредитов."], "ds_info": "DeepSeek Pro активен для текста, структуры, источников и проверок. Только внешние изображения.", "ds_title": "Тарифы DeepSeek Pro", "ds": ["Соотношение DeepSeek/GPT: 1 к 3.", "Текст, переработка, тесты и примеры: 1 кредит за 3 действия.", "Источники + структура: около 3⅓ кредита.", "Оценка структуры: 1 кредит; отчёт и метаданные: 1 кредит за 3 проверки.", "Полная связность: около 4 кредитов.", "Внешние изображения не расходуют кредиты."]},
        "العربية": {"gpt_info": "GPT-5.4 نشط: بحث ويب وفحوص تحريرية وحقوق نشر على الويب. تُرفع الصور من الخارج.", "gpt_title": "أسعار GPT-5.4", "gpt": ["الكتابة وإعادة الصياغة والاختبارات والأمثلة: رصيد لكل عملية.", "فهرس كامل: 10 أرصدة (4 بحث + 6 تخطيط تحريري).", "تقييم الفهرس: رصيدان؛ إعادة التوليد: 6 أرصدة.", "التحقق من الحقائق: رصيدان؛ التقرير النحوي وبيانات KDP: رصيد لكل منهما.", "فحص الاتساق الكامل: 10 أرصدة؛ ثم رصيد لكل كتلة معدلة.", "رفع الصور الخارجية مجاني؛ حقوق النشر على الويب: من رصيدين."], "ds_info": "DeepSeek Pro نشط للكتابة والفهرس والمصادر والفحوص. الصور خارجية فقط.", "ds_title": "أسعار DeepSeek Pro", "ds": ["نسبة DeepSeek/GPT هي 1 إلى 3.", "الكتابة وإعادة الصياغة والاختبارات والأمثلة: رصيد كل 3 عمليات.", "مصادر أصلية + فهرس كامل: نحو 3⅓ أرصدة.", "تقييم الفهرس: رصيد واحد؛ التقرير والبيانات: رصيد كل 3 فحوص.", "اتساق كامل: نحو 4 أرصدة.", "الصور الخارجية لا تستهلك أرصدة."]},
        "中文": {"gpt_info": "GPT-5.4 已启用：网页研究、编辑检查和网页版权检查。图片从外部上传。", "gpt_title": "GPT-5.4 价格", "gpt": ["写作、改写、测验和示例：每项操作 1 积分。", "完整目录：10 积分（4 积分研究 + 6 积分编辑规划）。", "目录评估：2 积分；重新生成：6 积分。", "事实核查：2 积分；句法报告和 KDP 元数据：各 1 积分。", "完整一致性检查：10 积分；后续每个修改区块 1 积分。", "外部图片上传免费；网页版权检查：2 积分起。"], "ds_info": "DeepSeek Pro 用于写作、目录、资料和检查。图片仅可从外部上传。", "ds_title": "DeepSeek Pro 价格", "ds": ["DeepSeek/GPT 比例为 1:3。", "写作、改写、测验和示例：每 3 项操作 1 积分。", "原生资料研究 + 完整目录：约 3⅓ 积分。", "目录评估：1 积分；报告和元数据：每 3 次检查 1 积分。", "完整一致性检查：约 4 积分。", "外部图片不消耗积分。"]},
    })
    # Il tariffario viene scritto in nove lingue. Questo passaggio conclusivo
    # mantiene coerenti in ogni traduzione le due stime DeepSeek ricalibrate:
    # 10 unità interne corrispondono a 3⅓ crediti, non a 4.
    sostituzioni_tariffario_ds = (
        ("circa 4", "circa 3⅓"), ("about 4", "about 3⅓"),
        ("unos 4", "unos 3⅓"), ("environ 4", "environ 3⅓"),
        ("etwa 4", "etwa 3⅓"), ("circa 4", "circa 3⅓"),
        ("около 4", "около 3⅓"), ("نحو 4", "نحو 3⅓"),
        ("约 4", "约 3⅓"),
    )
    for tariffario_tradotto in tariffari_sidebar.values():
        tariffario_tradotto["ds"] = [
            applica_sostituzioni_tariffario(voce, sostituzioni_tariffario_ds)
            for voce in tariffario_tradotto.get("ds", [])
        ]
    tariffario = tariffari_sidebar.get(lingua_sel, tariffari_sidebar["Italiano"])
    motore_tariffario = "ds" if usa_deepseek_pro() else "gpt"
    st.info(tariffario[f"{motore_tariffario}_info"])
    with st.expander(tariffario[f"{motore_tariffario}_title"], expanded=False):
        for voce_tariffario in tariffario[motore_tariffario]:
            st.write(f"• {voce_tariffario}")
    val_titolo = st.text_input(L["lbl_tit"], key="book_title")
    val_autore = st.text_input(L["lbl_auth"], key="book_author")
    
    # --- NUOVA SEZIONE CARICAMENTO FONTI ---
    st.divider()
    st.markdown(f"#### 2 · {intestazioni_sidebar[2]}")
    st.markdown("<small>Carica PDF o DOCX: l'IA ne ricava una mappa concettuale interna e scrive un testo autonomo, senza riprendere formulazioni delle fonti.</small>", unsafe_allow_html=True)
    file_caricati = st.file_uploader(testo_ui("fonti_carica", lingua_sel), type=['pdf', 'docx'], accept_multiple_files=True, label_visibility="collapsed")
    if file_caricati:
        if len(file_caricati) > 10:
            st.warning("Hai superato il limite di 10 file. Verranno analizzati i primi 10.")
            file_caricati = file_caricati[:10]
        firma_fonti = firma_fonti_esterne(file_caricati)
        if st.session_state.get("firma_fonti") != firma_fonti:
            with st.spinner("Lettura, studio editoriale e preparazione delle fonti in corso..."):
                st.session_state["conoscenza_extra"] = estrai_testo_da_files(file_caricati)
                st.session_state["scheda_fonti"] = crea_scheda_fonti(st.session_state["conoscenza_extra"])
                mappa_fonti = studia_fonti_con_ai(st.session_state["conoscenza_extra"])
                st.session_state["brief_fonti_originale"] = mappa_fonti
                # Compatibilità con salvataggi e CSV precedenti: il dossier ora
                # contiene esclusivamente la mappa concettuale, non estratti.
                st.session_state["dossier_fonti_ai"] = mappa_fonti
                st.session_state["firma_fonti"] = firma_fonti
        if st.session_state.get("conoscenza_extra"):
            st.success(f"Studiati {len(file_caricati)} documenti. La mappa concettuale originale guiderà indice e stesura.")
            st.caption(f"Analisi fonti: {'DeepSeek V4 Pro' if usa_deepseek_pro() else MODELLO_ANALISI_FONTI}. I brani caricati non vengono passati alla stesura: l'IA usa solo concetti rielaborati e il controllo di originalità confronta il testo prima dell'esportazione.")
    elif st.session_state.get("conoscenza_extra"):
        st.caption("Fonti già elaborate e conservate nel progetto. Per sostituirle, carica nuovi file oppure usa RESET PROGETTO.")

    etichette_fonti_web = {
        "Italiano": ("🔄 RIGENERA FONTI WEB", "Cerca di nuovo fonti aggiornate e sostituisce soltanto la mappa e il registro delle fonti web. Indice e sezioni già scritti non cambiano."),
        "English": ("🔄 REFRESH WEB SOURCES", "Searches for updated sources again and replaces only the web-source map and register. Your index and written sections will not change."),
        "Español": ("🔄 REGENERAR FUENTES WEB", "Busca de nuevo fuentes actualizadas y sustituye solo el mapa y el registro web. El índice y las secciones escritas no cambian."),
        "Français": ("🔄 ACTUALISER LES SOURCES WEB", "Recherche à nouveau des sources actualisées et remplace uniquement la carte et le registre web. L’index et les sections déjà écrites ne changent pas."),
        "Deutsch": ("🔄 WEBQUELLEN AKTUALISIEREN", "Sucht erneut aktuelle Quellen und ersetzt nur Karte und Webquellenregister. Inhaltsverzeichnis und geschriebene Abschnitte bleiben unverändert."),
    }
    etichetta_fonti, descrizione_fonti = etichette_fonti_web.get(lingua_sel, etichette_fonti_web["Italiano"])
    dati_fonti_pronti = all(str(st.session_state.get(chiave, "")).strip() for chiave in (
        "book_title", "book_genre", "book_goal", "book_plot",
    ))
    if pulsante_con_preventivo(
        "rigenera_fonti_web", etichetta_fonti, CREDIT_COSTS["indice_ricerca_web"], descrizione_fonti,
        use_container_width=True, disabled=not dati_fonti_pronti,
    ):
        with st.spinner("Ricerca e aggiornamento delle fonti web in corso..."):
            dossier_aggiornato = ricerca_preliminare_per_indice(
                st.session_state.get("book_title", ""),
                st.session_state.get("book_genre", ""),
                st.session_state.get("book_plot", ""),
                st.session_state.get("book_goal", ""),
                lingua_sel,
                f"{st.session_state.get('book_further_details', '')}\n\n{brief_personalizzazione_progetto()}",
                forza=True,
            )
        if dossier_aggiornato:
            st.success("Fonti web aggiornate e salvate nel progetto. Indice e sezioni esistenti sono rimasti invariati.")
        else:
            st.error("Non è stato possibile aggiornare le fonti web. Nessun credito viene trattenuto se la ricerca non restituisce un dossier valido.")

    registro_fonti_web = conserva_solo_fonti_web_selezionate(
        st.session_state.get("registro_fonti_web", "")
    )
    if registro_fonti_web != str(st.session_state.get("registro_fonti_web", "") or "").strip():
        st.session_state["registro_fonti_web"] = registro_fonti_web
    if registro_fonti_web:
        with st.expander("🌐 Fonti web selezionate per il progetto", expanded=False):
            st.caption("Sono visualizzate solo le fonti più pertinenti e autorevoli selezionate automaticamente per questo brief. Non vengono inserite nel libro.")
            st.markdown(registro_fonti_web)
    
    st.divider()
    st.markdown(f"#### 3 · {intestazioni_sidebar[3]}")
    # --- AGGIUNTA "STORICO" AI GENERI ---
    lista_gen = ["Saggio Scientifico", "Quiz Scientifico", "Manuale Tecnico", "Religioso / Teologico", "Spirituale / Esoterico", "Meditazione / Mindfulness", "Business & Marketing", "Economia e Finanza", "Romanzo Rosa", "Thriller / Noir", "Fantasy", "Fantascienza", "Manuale Psicologico", "Biografia", "Ricettario", "Test Prep (Preparazione Esami)", "Narrativo", "Romanzo Classico", "Contemporaneo", "Self-Help", "Manuale Pratico", "Storico"]
    val_genere = st.selectbox(L["lbl_gen"], [""] + lista_gen, key="book_genre", format_func=lambda valore: valore or testo_ui("seleziona", lingua_sel))
    
    stili_estesi = [
        "Standard", 
        "Professionale Accademico", 
        "Persuasivo (Neuromarketing Applicato)", 
        "Conversazionale ed Empatico", 
        "Scientifico Divulgativo", 
        "Storytelling Immersivo", 
        "Giornalistico d'Inchiesta", 
        "Socratico (Dialogico / Riflessivo)", 
        "Epico ed Evocativo", 
        "Minimalista ed Essenziale"
    ]
    val_stile = st.selectbox(L["lbl_style"], [""] + stili_estesi, key="book_writing_style", format_func=lambda valore: valore or testo_ui("seleziona", lingua_sel))

    direttive_indice_tipologia = {
        "Standard": "Crea un percorso lineare da basi a sviluppo, applicazione, verifica e sintesi. Rispetta il budget di sezioni indicato nel prompt: non espandere l'indice con capitoli o sottocapitoli ripetitivi. Ogni sottocapitolo deve avere un obiettivo concreto e un risultato leggibile.",
        "Professionale Accademico": "Organizza l'indice in contesto, definizioni, quadro teorico, metodologia, analisi, evidenze, limiti, implicazioni e riferimenti. Separa chiaramente ipotesi, dati, metodo e risultati; prevedi criteri di valutazione e fonti.",
        "Persuasivo (Neuromarketing Applicato)": "Costruisci il percorso da problema e consapevolezza a soluzione, prove, obiezioni, benefici, applicazione e azione. Inserisci casi, comparazioni e piani d'azione senza promesse garantite o claim non verificabili.",
        "Conversazionale ed Empatico": "Sequenzia l'indice come un accompagnamento: situazione del lettore, ostacoli, spiegazione semplice, esercitazione guidata, verifica e autonomia. Usa domande guida, riepiloghi e passaggi graduali senza infantilizzare.",
        "Scientifico Divulgativo": "Procedi da basi e contesto a meccanismi, evidenze, esempi, applicazioni e limiti. Indica dove servono fonti aggiornate, distingui fatti, ipotesi e analogie e inserisci esperimenti mentali o verifiche pratiche quando pertinenti.",
        "Storytelling Immersivo": "Progetta un arco narrativo completo con situazione iniziale, personaggi, desiderio, conflitto, ostacoli, svolte, conseguenze, climax e risoluzione. Ogni capitolo deve modificare la situazione o approfondire un personaggio, evitando capitoli di riempimento.",
        "Giornalistico d'Inchiesta": "Organizza il percorso da domanda iniziale a contesto, fonti primarie, testimonianze, verifiche indipendenti, contraddizioni, prove, responsabilità e conclusioni. Specifica quali fatti devono essere documentati e separa dati accertati da ipotesi.",
        "Socratico (Dialogico / Riflessivo)": "Costruisci l'indice attraverso domande progressive: presupposti, definizioni, dubbi, obiezioni, esempi, conseguenze e sintesi. Ogni capitolo deve porre una domanda centrale e chiuderla con una risposta argomentata o un esercizio di riflessione.",
        "Epico ed Evocativo": "Crea una progressione ampia con origine, chiamata, prove, alleati, opposizioni, trasformazione, crisi, compimento e significato finale. Mantieni immagini evocative nei titoli ma indica sempre un contenuto concreto e coerente con il genere.",
        "Minimalista ed Essenziale": "Riduci il libro a 15-18 capitoli indispensabili, con un solo obiettivo per capitolo e sottocapitoli non sovrapposti. Usa titoli brevi e operativi, elimina digressioni e assegna a ogni sezione un risultato verificabile."
    }
    direttiva_indice_selezionata = direttive_indice_tipologia.get(val_stile, direttive_indice_tipologia["Standard"])
    
    st.markdown("---")
    # --- AGGIUNTA "STORICO E DOCUMENTALE" AGLI STILI DI RACCONTO ---
    val_narrativa = st.selectbox(L["lbl_narrative"], [""] + [
        "Coinvolgente e Narrativo", "Tecnico e Analitico", "Ispirazionale e Motivante", 
        "Socratico (Domanda/Risposta)", "Storytelling Emozionale", "Diretto e Pratico (Action-oriented)", "Storico e Documentale"
    ], key="book_narrative_style", format_func=lambda valore: valore or testo_ui("seleziona", lingua_sel))
    
    # NUOVO BLOCCO: Punto di Vista (POV)
    lista_pov = [
        "Tu (Diretto, confidenziale e personale)",
        "Voi (Plurale, autorevole e rispettoso)",
        "Noi (Inclusivo, partecipativo e didattico)",
        "Impersonale / Terza Persona (Distaccato, analitico, oggettivo)"
    ]
    val_pov = st.selectbox(L.get("lbl_pov", "Punto di Vista (Pronome)"), [""] + lista_pov, key="book_point_of_view", format_func=lambda valore: valore or testo_ui("seleziona", lingua_sel))
    
    # Definizioni disponibili prima del loro primo utilizzo nella UI.
    # Restano presenti anche nel modulo di memoria sottostante per compatibilità.
    def costruisci_specifica_editoriale(titolo, genere, stile, narrativa, pov, obiettivo, argomento, risultato_finale="", approfondimenti=""):
        return f"""=== SPECIFICA EDITORIALE STRUTTURATA ===
Titolo: {titolo}
Genere: {genere}
Tipologia di scrittura: {stile}
Stile di racconto: {narrativa}
Punto di vista: {pov}

OBIETTIVO OPERATIVO:
{obiettivo}

RISULTATO FINALE DESIDERATO E VERIFICABILE:
{risultato_finale or "Non dichiarato: definisci cosa il lettore deve saper fare o ottenere alla fine."}

ARGOMENTO E CONFINI:
{argomento}

APPROFONDIMENTI PRIORITARI (FACOLTATIVI):
{approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}

Per ogni sezione ricava un risultato concreto, il livello del lettore, i concetti necessari,
gli esempi o le procedure da produrre e ciò che deve restare fuori per evitare ripetizioni.
"""

    def analizza_coerenza_libro(indice, contenuti, obiettivo, argomento):
        risultati = ["REPORT CONTROLLO COERENZA DEL LIBRO"]
        testo = "\n".join(contenuti.values()) if contenuti else ""
        capitoli = re.findall(r"(?im)^(?:Capitolo|Chapter|CAPITOLO)\\s+\\d+", indice or "")
        sottocapitoli = re.findall(r"(?m)^\\d+\\.\\d+\\s+", indice or "")
        risultati.append(f"Capitoli rilevati: {len(capitoli)}")
        risultati.append(f"Sottocapitoli rilevati: {len(sottocapitoli)}")
        if not indice.strip(): risultati.append("ERRORE: indice assente")
        if not obiettivo.strip(): risultati.append("AVVISO: obiettivo assente")
        if not argomento.strip(): risultati.append("AVVISO: argomento assente")
        if len(testo.strip()) < 1000: risultati.append("AVVISO: contenuto ancora troppo breve per una verifica completa")
        frasi = [f.strip().lower() for f in re.split(r"[.!?]+", testo) if len(f.strip()) > 40]
        duplicati = len(frasi) - len(set(frasi))
        risultati.append(f"Frasi duplicate identiche rilevate: {max(0, duplicati)}")
        if duplicati == 0: risultati.append("OK: nessuna duplicazione identica rilevata nel testo disponibile")
        return "\n".join(risultati)

    val_goal = st.text_input(L["lbl_goal"], placeholder="Es: Mantenere l'attenzione alta, far emozionare...", key="book_goal")
    etichette_risultato = {
        "Italiano": "Risultato finale desiderato",
        "English": "Desired final result",
        "Español": "Resultado final deseado",
        "Français": "Résultat final souhaité",
        "Deutsch": "Gewünschtes Endergebnis",
        "Română": "Rezultatul final dorit",
        "Русский": "Желаемый итоговый результат",
        "العربية": "النتيجة النهائية المطلوبة",
        "中文": "期望的最终结果",
    }
    val_risultato = st.text_area(
        testo_ui("risultato", lingua_sel),
        height=100,
        placeholder=testo_ui("risultato_placeholder", lingua_sel),
        key="book_desired_result"
    )
    val_trama = st.text_area(L["lbl_plot"], height=150, key="book_plot")
    val_approfondimenti = st.text_area(
        testo_ui("approfondimenti", lingua_sel),
        height=130,
        placeholder=testo_ui("approfondimenti_placeholder", lingua_sel),
        key="book_further_details"
    )
    # Blocco autonomo: non rende obbligatoria la personalizzazione e non
    # modifica i campi editoriali già richiesti per generare indice e testo.
    testi_personalizzazione = etichette_personalizzazione(lingua_sel)
    with st.expander(testi_personalizzazione["titolo"], expanded=False):
        st.caption(testi_personalizzazione["intro"])
        val_voce_personale = st.text_area(
            testi_personalizzazione["voce"], height=92,
            key="book_personal_voice",
        )
        val_materiale_personale = st.text_area(
            testi_personalizzazione["materiale"], height=92,
            key="book_personal_material",
        )
        val_priorita_personali = st.text_area(
            testi_personalizzazione["priorita"], height=92,
            key="book_personal_priorities",
        )
        val_confini_personali = st.text_area(
            testi_personalizzazione["confini"], height=92,
            key="book_personal_boundaries",
        )
        val_modalita_checkpoint = st.selectbox(
            testi_personalizzazione["check"],
            list(MODALITA_CHECKPOINT_PERSONALE),
            key="book_personal_checkpoint_mode",
            format_func=lambda valore: testi_personalizzazione.get(valore or "automatico", testi_personalizzazione["automatico"]),
            help="Le pause guidate non generano testo e non consumano crediti. Se non scegli nulla, la stesura prosegue automaticamente.",
        )
    # Le variabili sono lette qui per rendere esplicita la loro appartenenza
    # alla stessa fotografia della sidebar, salvata anche prima di un rerun.
    val_note_checkpoint = st.session_state.get("book_personal_checkpoint_notes", "")
    opzioni_lunghezza = list(PROFILI_LUNGHEZZA_STESURA.keys())
    val_lunghezza_scelta = st.selectbox(
        testo_ui("lunghezza", lingua_sel),
        [""] + opzioni_lunghezza,
        index=0,
        key="profilo_lunghezza_stesura",
        format_func=lambda valore: valore or testo_ui("seleziona", lingua_sel),
        help=testo_ui("lunghezza_help", lingua_sel),
    )
    # Il valore visibile può provenire da una sessione/CSV precedente. Tutta la
    # logica usa esclusivamente una chiave tecnica presente nel profilo: un
    # valore non valido non può più interrompere la sidebar con un KeyError.
    val_lunghezza_normalizzata = _valore_scelta_chat_guidata(
        "lunghezza", val_lunghezza_scelta
    )
    if val_lunghezza_normalizzata not in PROFILI_LUNGHEZZA_STESURA:
        val_lunghezza_normalizzata = ""
    val_lunghezza = val_lunghezza_normalizzata or "Standard KDP"
    # Memorizza la sidebar in ogni esecuzione, prima di qualsiasi pulsante che
    # possa avviare una generazione o un rerun.
    st.session_state[CHIAVE_MEMORIA_SIDEBAR] = {
        "titolo": val_titolo,
        "autore": val_autore,
        "lingua": lingua_scelta,
        "genere": val_genere,
        "tipologia_scrittura": val_stile,
        "stile_racconto": val_narrativa,
        "punto_di_vista": val_pov,
        "obiettivo": val_goal,
        "risultato_finale": val_risultato,
        "argomento": val_trama,
        "approfondimenti": val_approfondimenti,
        "voce_personale": val_voce_personale,
        "materiale_personale": val_materiale_personale,
        "priorita_personali": val_priorita_personali,
        "confini_personali": val_confini_personali,
        "modalita_checkpoint": val_modalita_checkpoint,
        "note_checkpoint": val_note_checkpoint,
        "lunghezza": val_lunghezza_scelta,
        "provider_ia": provider_ia,
    }
    if val_lunghezza_scelta:
        st.caption(
            f"{val_lunghezza}: {PROFILI_LUNGHEZZA_STESURA[val_lunghezza]['parole']} per sezione — "
            f"{PROFILI_LUNGHEZZA_STESURA[val_lunghezza]['descrizione']}. "
            f"Massimo {PROFILI_LUNGHEZZA_STESURA[val_lunghezza]['max_sezioni']} sezioni totali, "
            f"tutte dedicate ai contenuti dell'indice. Obiettivo indicativo: circa "
            f"{PROFILI_LUNGHEZZA_STESURA[val_lunghezza]['pagine_minime']} pagine nel manoscritto 6×9. "
            "La singola sezione mantiene una tolleranza del 5%; l'obiettivo di pagine ha una tolleranza del 15%."
        )
    else:
        st.caption("Scegli una lunghezza delle sezioni per completare la sidebar.")
    profilo_indice = PROFILI_LUNGHEZZA_STESURA[val_lunghezza]
    limite_sezioni_totali = profilo_indice["max_sezioni"]
    # La Prefazione viene aggiunta dal software dopo l'output del modello:
    # il limite passato al generatore le riserva sempre un posto.
    limite_voci_indice = max(1, limite_sezioni_totali - 1)
    budget_struttura_indice = {
        "Compatto": "massimo 65 sezioni totali, Prefazione inclusa; usa soltanto Parti, Capitoli e sottocapitoli indispensabili al brief",
        "Standard KDP": "massimo 110 sezioni totali, Prefazione inclusa; usa soltanto Parti, Capitoli e sottocapitoli indispensabili al brief",
        "Approfondito": "massimo 160 sezioni totali, Prefazione inclusa; amplia soltanto argomenti davvero distinti e verificabili",
    }[val_lunghezza]
    minimi_struttura_indice = {
        "Compatto": (0, 0),
        "Standard KDP": (0, 0),
        "Approfondito": (0, 0),
    }[val_lunghezza]
    budget_editoriale_indice = {
        "Compatto": {"parti": 3, "capitoli": 10, "sottocapitoli": 3},
        "Standard KDP": {"parti": 4, "capitoli": 15, "sottocapitoli": 4},
        "Approfondito": {"parti": 5, "capitoli": 20, "sottocapitoli": 5},
    }[val_lunghezza]
    massimo_parti_editoriali = budget_editoriale_indice["parti"]
    massimo_capitoli_editoriali = budget_editoriale_indice["capitoli"]
    massimo_sottocapitoli_per_capitolo = budget_editoriale_indice["sottocapitoli"]
    specifica_editoriale = costruisci_specifica_editoriale(
        val_titolo, val_genere, val_stile, val_narrativa, val_pov, val_goal, val_trama, val_risultato, val_approfondimenti
    )
    campi_obbligatori_sidebar = {
        L["lbl_tit"]: val_titolo,
        L["lbl_auth"]: val_autore,
        L["lbl_gen"]: val_genere,
        L["lbl_style"]: val_stile,
        L["lbl_narrative"]: val_narrativa,
        L["lbl_pov"]: val_pov,
        L["lbl_goal"]: val_goal,
        testo_ui("risultato", lingua_sel): val_risultato,
        L["lbl_plot"]: val_trama,
        testo_ui("lunghezza", lingua_sel): val_lunghezza_scelta,
    }
    campi_sidebar_mancanti = [etichetta for etichetta, valore in campi_obbligatori_sidebar.items() if not str(valore).strip()]
    sidebar_pronta = not campi_sidebar_mancanti
    firma_sidebar_pronta = "|".join(str(valore).strip() for valore in campi_obbligatori_sidebar.values())
    if sidebar_pronta and st.session_state.get("firma_notifica_sidebar") != firma_sidebar_pronta:
        notifica_sonora("sidebar_pronta", lingua_sel)
        st.session_state["firma_notifica_sidebar"] = firma_sidebar_pronta
    elif not sidebar_pronta:
        st.session_state.pop("firma_notifica_sidebar", None)

    if st.session_state.get("autosave_stato"):
        st.caption(st.session_state["autosave_stato"])

    st.divider()
    st.markdown(f"#### 4 · {intestazioni_sidebar[4]}")
    st.caption(intestazioni_sidebar[5])
    
    # Reset del solo progetto: l'accesso commerciale e il saldo crediti restano attivi.
    if st.button(
        L["btn_res"],
        type="primary",
        use_container_width=True,
        key="reset_progetto_distruttivo",
        help="Cancella definitivamente il progetto aperto, inclusi campi, indice, testi, fonti e immagini.",
    ):
        reset_eseguito = False
        if st.session_state.get("admin_test_mode"):
            # Nel laboratorio RESET non deve mai cancellare la bozza cloud
            # dell'amministratore: chiude soltanto il progetto di prova.
            termina_collaudo_amministratore()
            st.session_state["messaggio_aggiornamento_pagina"] = "Collaudo chiuso: la sessione precedente è stata ripristinata."
            reset_eseguito = True
        else:
            # Reset significa eliminazione integrale: prima il cloud, poi la
            # memoria della pagina. Se il cloud non conferma, non svuotiamo la
            # UI e non rischiamo che la bozza ricompaia al login successivo.
            if elimina_progetto_automatico():
                st.session_state.pop("commercial_logout_snapshot", None)
                st.session_state.pop("commercial_logout_snapshot_owner", None)
                st.session_state.pop("commercial_logout_requested_for", None)
                st.session_state["commercial_project_reset_requested"] = True
                for key in list(st.session_state.keys()):
                    if not key.startswith("commercial_"):
                        del st.session_state[key]
                reset_eseguito = True
            else:
                st.error(
                    "Il progetto non è stato cancellato perché l'archivio cloud non ha confermato l'operazione. "
                    "Riprova tra un momento: i dati restano visibili e al sicuro."
                )
        if reset_eseguito:
            st.rerun()

    etichette_aggiorna = {
        "Italiano": ("🔄 AGGIORNA PAGINA", "Pagina aggiornata: cache e messaggi temporanei ripuliti. Il progetto resta invariato."),
        "English": ("🔄 REFRESH PAGE", "Page refreshed: cache and temporary messages cleared. Your project is unchanged."),
        "Español": ("🔄 ACTUALIZAR PÁGINA", "Página actualizada: caché y mensajes temporales eliminados. El proyecto no cambia."),
        "Français": ("🔄 ACTUALISER LA PAGE", "Page actualisée : cache et messages temporaires effacés. Le projet reste inchangé."),
        "Deutsch": ("🔄 SEITE AKTUALISIEREN", "Seite aktualisiert: Cache und temporäre Meldungen wurden gelöscht. Das Projekt bleibt unverändert."),
        "Română": ("🔄 REÎNCARCĂ PAGINA", "Pagina a fost reîncărcată: cache-ul și mesajele temporare au fost șterse. Proiectul rămâne neschimbat."),
    }
    etichetta_aggiorna, messaggio_aggiorna = etichette_aggiorna.get(
        lingua_sel, etichette_aggiorna["Italiano"]
    )
    if st.button(etichetta_aggiorna, use_container_width=True, key="aggiorna_pagina_sicura"):
        # Pulizia sicura: non tocca sidebar, indice, memoria delle sezioni,
        # crediti, fonti o coda del libro. Rimuove soltanto cache e messaggi
        # transitori, così un errore già visualizzato non resta nella pagina.
        try:
            st.cache_data.clear()
        except Exception:
            pass
        for chiave_temporanea in (
            "job_scrittura_errore", "messaggio_stesura_sezione",
            "messaggio_stesura_sottocapitoli", "dettaglio_errori_sottocapitoli",
            "preventivo_crediti_attesa", "azione_crediti_confermata",
        ):
            st.session_state.pop(chiave_temporanea, None)
        st.session_state["messaggio_aggiornamento_pagina"] = messaggio_aggiorna
        st.rerun()

    if st.session_state.pop("messaggio_aggiornamento_pagina", ""):
        st.success(messaggio_aggiorna)

    if st.button(testo_ui("salva", lingua_sel), type="primary", use_container_width=True, key="salva_sessione_manuale"):
        # Il cloud viene aggiornato esclusivamente con questo comando: durante
        # la stesura normale i testi restano nella memoria della pagina.
        sezioni_da_salvare = list(
            set(st.session_state.get("lista_capitoli", []))
            | set((st.session_state.get(CHIAVE_MEMORIA_SEZIONI, {}) or {}).keys())
        )
        if st.session_state.get("admin_test_mode"):
            salva_stesura_immediata(sezioni_da_salvare)
            st.info("🧪 Collaudo salvato solo nel laboratorio corrente: il progetto reale resta invariato.")
        elif salva_progetto_corrente(sidebar_memorizzata_corrente(), sezioni_da_salvare):
            st.success("Sessione salvata: sidebar, indice, sezioni e fonti web sono stati memorizzati nel tuo account.")
        else:
            st.error("Non è stato possibile salvare la sessione nel tuo account. Il lavoro resta aperto in questa pagina.")

    if st.button(testo_ui("ripristina", lingua_sel), use_container_width=True, key="ripristina_ultima_stesura"):
        if prepara_ripristino_ultima_stesura():
            st.rerun()
        else:
            st.warning("Non è stata trovata una stesura salvata nel tuo account. Verifica che il salvataggio automatico indichi “nel tuo account”.")

    mostra_memoria_visiva_progetto()

# ======================================================================================================================
# 7. LOGICA DI MEMORIA E COERENZA (EVITA RIPETIZIONI GLOBALI) E INTEGRAZIONE FONTI
# ======================================================================================================================
def genera_contesto_avanzato(sezione_corrente, argomento=""):
    contesto = ""
    estratti_fonti = estratti_fonti_pertinenti(sezione_corrente, argomento)
    if estratti_fonti:
        contesto += f"=== ESTRATTI PERTINENTI DELLE FONTI ESTERNE (SOLO PER RAGIONAMENTO) ===\n{estratti_fonti}\n\n"
        
    for s in st.session_state.get("lista_capitoli", []):
        if s == sezione_corrente: break
        testo_precedente = leggi_sezione_memorizzata(s)
        if str(testo_precedente).strip():
            # Memoria estesa: il riepilogo breve da 150 caratteri non era sufficiente
            # per distinguere concetti, esempi e procedure già utilizzati.
            contesto += f"- Trattato in {s}:\n{testo_precedente[:1200]}\n"
    return contesto

def individua_sottocapitoli_del_capitolo(capitolo, sezioni):
    """Restituisce soltanto i sottocapitoli numerati appartenenti al capitolo selezionato."""
    match = re.match(r"(?i)^(?:capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+(\d+)", capitolo.strip())
    if not match:
        return []
    numero = match.group(1)
    return [s for s in sezioni if re.match(rf"^{re.escape(numero)}\.\d+\s", s.strip())]

def individua_sezioni_da_stendere(sezioni):
    """Restituisce tutte le sezioni che devono apparire nel manoscritto.

    Le Parti sono aperture editoriali brevi, ma sono vere voci dell'indice:
    escluderle dalla coda di “Scrivi tutto il libro” lasciava un falso
    “MANCANTE” nel controllo di completezza.
    """
    regex_capitolo = r'(?i)^(?:capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+'
    regex_sottocapitolo = r'^\d+\.\d+\s+'
    risultato = []
    for sezione in sezioni:
        pulita = sezione.strip()
        if tipo_sezione_editoriale(pulita) == "parte":
            risultato.append(sezione)
        elif re.match(regex_sottocapitolo, pulita):
            risultato.append(sezione)
        elif re.match(regex_capitolo, pulita) and not individua_sottocapitoli_del_capitolo(sezione, sezioni):
            # Un capitolo senza figli ha contenuto autonomo e viene scritto normalmente.
            risultato.append(sezione)
    return risultato

def crea_prompt_stesura_sezione(sezione, indice, trama, genere, stile, narrativa, pov, obiettivo, lingua, approfondimenti="", profilo_lunghezza="Standard KDP"):
    """Costruisce il prompt comune usato dalla stesura singola e dalla stesura di un capitolo intero."""
    memoria = genera_contesto_avanzato(sezione, trama)
    tipo_sezione = tipo_sezione_editoriale(sezione)
    brief_personale = brief_personalizzazione_progetto(sezione)
    profilo_genere = profilo_genere_stesura(genere)
    profilo_tipologia = profilo_tipologia_stesura(stile)
    regola_struttura = profilo_struttura_indice(genere, "", trama, obiettivo)
    profilo_lunghezza_dati = PROFILI_LUNGHEZZA_STESURA.get(
        profilo_lunghezza, PROFILI_LUNGHEZZA_STESURA["Standard KDP"]
    )
    minimo_parole_tolleranza, massimo_parole_tolleranza = vincolo_parole_con_tolleranza(profilo_lunghezza)
    riserva_chiusura = max(55, round(profilo_lunghezza_dati["max_parole"] * 0.12))
    soglia_sviluppo = profilo_lunghezza_dati["max_parole"] - riserva_chiusura
    # La Prefazione non è un capitolo breve: è una soglia editoriale. Questo
    # blocco sostituisce integralmente per lei le istruzioni generiche sotto,
    # evitando il conflitto fra 140-220 parole e il profilo standard del libro.
    if tipo_sezione == "prefazione":
        return f"""
CONTESTO DEL LIBRO (usa solo concetti, mai frasi o struttura delle fonti):
{memoria}

DATI EDITORIALI DELLA PREFAZIONE
- Titolo del libro: {val_titolo}
- Argomento centrale: {trama}
- Genere: {genere}
- Tipologia di scrittura: {stile}
- Stile di racconto: {narrativa}
- Punto di vista: {pov}
- Obiettivo del libro: {obiettivo}
- Risultato finale desiderato: {val_risultato or "Non dichiarato"}
- Approfondimenti prioritari: {approfondimenti.strip() or "Nessuno"}

{brief_personale}

AZIONE
Redigi la sezione iniziale '{sezione}', rigorosamente in {lingua}. Questa
Prefazione deve avere 140-220 parole; non applicare il profilo di lunghezza
delle sezioni ordinarie. Accogli il lettore, chiarisci il bisogno a cui il
libro risponde e descrivi il percorso senza elencare, riassumere o anticipare
i capitoli. Mantieni il POV richiesto, crea un testo originale, non ripetere
il titolo come intestazione e non usare Markdown, URL, citazioni, fonti,
elenchi o promesse garantite. Concludi con una frase completa e definitiva.
"""
    ha_sottocapitoli = bool(individua_sottocapitoli_del_capitolo(sezione, indice.splitlines()))
    if tipo_sezione == "capitolo" and ha_sottocapitoli:
        istruzione_lunghezza = (
            "Questo capitolo ha sottocapitoli: scrivi SOLO una cornice iniziale di 70-100 parole. "
            "Non svolgere argomenti, procedure, esempi o conclusioni riservati ai sottocapitoli."
        )
    else:
        istruzione_lunghezza = (
            f"Obiettivo: {profilo_lunghezza_dati['parole']}. Tolleranza massima del 5%: non scrivere meno di "
            f"{minimo_parole_tolleranza} né più di {massimo_parole_tolleranza} parole. La qualità dipende dalla densità "
            "delle informazioni, non dalla quantità di parole: scegli solo gli esempi, passaggi e dettagli indispensabili al titolo. "
            f"Completa lo sviluppo principale entro circa {soglia_sviluppo} parole e conserva le ultime circa "
            f"{riserva_chiusura} parole per chiudere con calma l'ultima idea."
        )
    direttiva_test_prep = ""
    if genere == "Test Prep (Preparazione Esami)":
        direttiva_test_prep = """
- Se il titolo della sezione contiene quiz, test, domande, autovalutazione o esercizi, crea il materiale promesso: quesiti originali con quattro opzioni, risposta corretta e spiegazione della scelta. Non limitarti a spiegare come affrontare un quiz.
- Se il titolo contiene simulazione, genera solo ciò che il titolo assegna alla sezione; la sezione che contiene domande o esecuzione sarà completata dal generatore in blocchi con quesiti e soluzioni separati.
- Evita formule generiche sullo studio: ogni paragrafo deve introdurre una competenza d'esame, un errore concreto, una procedura, un quesito o una decisione verificabile.
"""
    return f"""
INDICE GENERALE (STUDIALO PER CAPIRE COSA NON DEVI ANTICIPARE):
{indice}

MEMORIA CONTENUTI PRECEDENTI (Per non ripetersi):
{memoria}

=== PARAMETRI EDITORIALI SARTORIALI (DA APPLICARE TASSATIVAMENTE IN QUESTA SEZIONE) ===
- Argomento Centrale / Trama: {trama}
- Genere Letterario: {genere}
- Tipologia di Scrittura: {stile}
- Stile di Racconto: {narrativa}
- Punto di Vista (POV): {pov}
- Obiettivo Emozionale/Pratico: {obiettivo}
- Risultato finale desiderato: {val_risultato or "Non dichiarato"}
- Approfondimenti prioritari: {approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}

{brief_personale}

=== PROFILO EDITORIALE DA RISPETTARE ===
- Regole del genere: {profilo_genere}
- Regole della tipologia di scrittura: {profilo_tipologia}
- Regole della struttura: {regola_struttura}
- Tipo della sezione corrente: {tipo_sezione}
- Profilo di lunghezza scelto: {profilo_lunghezza}

AZIONE:
Scrivi ora la sezione ESATTA: '{sezione}'. Il testo deve essere rigorosamente in lingua {lingua}.
- Se la sezione è una Parte, scrivi soltanto una breve apertura che spiega lo scopo della Parte e come usarla: non sviluppare o riassumere i capitoli che seguono.
- Se la sezione è un Capitolo con sottocapitoli nell'indice, non anticipare né risolvere gli argomenti assegnati ai relativi sottocapitoli.
- Se il genere è Ricettario, ogni Capitolo che porta il nome di un piatto è una sola ricetta completa: non aggiungere sottocapitoli autonomi e non ripetere una ricetta già presente nella Parte o in altre sezioni.
{direttiva_test_prep}
- Rispetta integralmente i parametri editoriali e usa tassativamente il POV richiesto ({pov}).
- Tratta con priorità gli approfondimenti forniti, ma soltanto nelle sezioni cui sono pertinenti; non ripeterli artificialmente e non anticipare contenuti assegnati a sezioni successive.
- Sii profondo ed esaustivo nell'ambito della sezione, senza rubare materiale alle altre.
- {istruzione_lunghezza}
- Pianifica internamente prima di scrivere, senza mostrare il piano: apri il tema, sviluppa solo i passaggi indispensabili e riserva già lo spazio finale indicato per concludere.
- Non iniziare mai un nuovo esempio, elenco, periodo o ragionamento quando non hai spazio per terminarlo. Raggiunta la soglia di sviluppo, smetti di aggiungere dettagli secondari e completa con ordine il concetto già avviato.
- L'ultimo paragrafo deve concludere il ragionamento, la procedura, l'esempio o la scena con una frase piena e definitiva. Non usare ellissi, frasi sospese, elenchi lasciati aperti o collegamenti rimandati alla sezione successiva.
- Redigi contenuto concreto suggerito dal titolo, senza preamboli inutili.
- Non scrivere e non ripetere mai '{sezione}' come intestazione. Inizia direttamente con il contenuto.
- Usa formattazione editoriale pulita: non usare Markdown, simboli ###, ##, **, __, ``` o intestazioni tecniche. Se servono elenchi, usa semplici punti o numeri senza caratteri decorativi.
- Non inserire URL, link, citazioni, note bibliografiche o sezioni fonti.
- Se sono disponibili fonti esterne, la memoria contiene soltanto una mappa concettuale interna: usa i concetti come conoscenza da rielaborare, mai come testo da parafrasare riga per riga.
- ORIGINALITÀ OBBLIGATORIA: costruisci una spiegazione nuova con ordine, esempi, collegamenti, frasi e sviluppo propri. Non imitare la struttura della fonte, non riprendere formulazioni caratteristiche e non riprodurre sequenze di sei o più parole che potrebbero provenire dalle fonti. Se un fatto è necessario, esprimilo in modo autonomo e contestualizzalo per questo specifico libro.
- Prima di consegnare, verifica internamente che il contenuto sia completo per la sezione assegnata, che non sia una bozza o un frammento, che l'ultima frase sia completa e che il ragionamento abbia una chiusura utile entro il limite di parole.

=== PROFONDITÀ ADATTIVA E SPIEGAZIONE PASSO PASSO ===
- Scrivi un testo professionale, completo e proporzionato alla complessità della sezione e al profilo di lunghezza scelto. Una spiegazione accurata non deve diventare una spiegazione lunga: evita frasi motivazionali, ripetizioni, riassunti e varianti dello stesso esempio.
- Per una procedura o una funzione software, segui questa sequenza: scopo pratico; prerequisiti e strumenti; passaggi numerati nell'ordine esatto; cosa osservare dopo ogni passaggio; controllo del risultato; errori frequenti e correzioni; esempio realistico di applicazione.
- Per un concetto teorico, segui questa sequenza: definizione chiara; perché è importante; come si applica; esempio concreto; limiti, eccezioni o errori da evitare.
- Non saltare passaggi impliciti. Spiega ogni azione in modo che un lettore del livello dichiarato possa ripeterla autonomamente.
- Quando è utile, includi checklist, tabella, caso pratico o criterio verificabile all'interno della spiegazione, senza trasformarli in capitoli separati.
- Mantieni paragrafi leggibili e sottotitoli brevi solo quando migliorano la consultazione; non usare formule generiche come "semplice", "intuitivo" o "fondamentale" senza spiegare concretamente il perché.
"""

def valuta_indice_editoriale(indice, titolo, trama, genere, stile, narrativa, pov, obiettivo, lingua, approfondimenti="", *, addebita=True):
    """Valuta l'indice rispetto al brief compilato nella sidebar."""
    prompt = f"""Valuta professionalmente il seguente indice editoriale in lingua {lingua}.

DATI DEL BRIEF
Titolo: {titolo}
Trama/argomento: {trama}
Genere: {genere}
Tipologia di scrittura: {stile}
Stile di racconto: {narrativa}
Punto di vista: {pov}
Obiettivo del libro: {obiettivo}
Risultato finale desiderato: {val_risultato or "Non dichiarato"}
Approfondimenti prioritari: {approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}
{brief_personalizzazione_progetto()}

INDICE DA VALUTARE
{indice}

Esamina esclusivamente: attinenza al brief, copertura degli approfondimenti prioritari, raggiungibilità del risultato finale desiderato, ordine logico, completezza, progressione del lettore, assenza di ripetizioni, distinzione tra capitoli e sottocapitoli, concretezza dei titoli e capacità di sostenere un libro completo.
Non valutare il libro non ancora scritto e non inventare informazioni mancanti.

Restituisci testo semplice, senza Markdown e senza URL, in questo formato:
VOTO COMPLESSIVO: X/10
VERDETTO: una frase chiara.
PUNTI DI FORZA: 3-5 osservazioni concrete.
MIGLIORAMENTI CONSIGLIATI: soltanto modifiche necessarie e direttamente applicabili; se l'indice è già valido, scrivi "Nessuna modifica strutturale necessaria".
COERENZA CON IL BRIEF: breve verifica di titolo, pubblico, obiettivo, genere e stile.
"""
    return pulisci_testo_editoriale(chiedi_gpt(
        prompt,
        "Sei un editor senior specializzato in architettura di libri. Sei rigoroso, concreto e non usi valutazioni vaghe.",
        addebita=addebita,
        amount=CREDIT_COSTS["voto_indice"],
        model=MODELLO_EDITORIALE,
        reason="voto_indice",
    ))

def firma_controllo_coerenza(indice, contenuti, titolo, trama, genere, stile, narrativa, pov, obiettivo, risultato_finale, approfondimenti):
    """Identifica con certezza la versione del manoscritto su cui è stato prodotto il report."""
    parti = [indice, titolo, trama, genere, stile, narrativa, pov, obiettivo, risultato_finale, approfondimenti, brief_personalizzazione_progetto()]
    parti.extend(f"{sezione}\n{contenuto}" for sezione, contenuto in contenuti.items())
    return hashlib.sha256("\n\u241e\n".join(str(p or "") for p in parti).encode("utf-8")).hexdigest()


def blocchi_per_audit_manoscritto(contenuti, limite_caratteri=18000):
    """Divide il testo intero in blocchi consecutivi, senza omettere la parte centrale delle sezioni."""
    blocchi, corrente = [], ""
    for sezione, contenuto in contenuti.items():
        testo = pulisci_testo_editoriale(contenuto).strip()
        if not testo:
            continue
        unita = f"SEZIONE: {sezione}\nTESTO:\n{testo}\n\n"
        while unita:
            spazio = limite_caratteri - len(corrente)
            if spazio <= 300:
                blocchi.append(corrente)
                corrente, spazio = "", limite_caratteri
            if len(unita) <= spazio:
                corrente += unita
                unita = ""
            else:
                punto_taglio = unita.rfind("\n", 0, spazio)
                if punto_taglio < max(500, spazio // 2):
                    punto_taglio = spazio
                corrente += unita[:punto_taglio]
                blocchi.append(corrente)
                corrente, unita = "", unita[punto_taglio:]
    if corrente.strip():
        blocchi.append(corrente)
    return blocchi


def chiedi_audit_editoriale(prompt, *, addebita=True):
    """Esegue l'audit editoriale completo con il modello rapido dedicato ai controlli."""
    return pulisci_testo_editoriale(chiedi_gpt(
        prompt,
        "Sei un direttore editoriale rigoroso. Restituisci solo l'audit richiesto, senza ragionamento interno.",
        addebita=addebita,
        amount=CREDIT_COSTS["voto_indice"],
        model=MODELLO_EDITORIALE,
        reason="audit_editoriale",
    )).strip()


def mappa_capitoli_e_sottocapitoli(indice):
    """Restituisce una mappa leggibile sottocapitolo -> capitolo padre per rendere il report azionabile."""
    capitolo_corrente = ""
    righe_mappa = []
    pattern_capitolo = re.compile(
        r"(?i)^(?:capitolo|chapter|kapitel|cap[ií]tulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+.*"
    )
    pattern_sottocapitolo = re.compile(r"^\d+\.\d+(?:\.\d+)?\s+.+")
    for riga in (indice or "").splitlines():
        voce = riga.strip()
        if not voce:
            continue
        if pattern_capitolo.match(voce):
            capitolo_corrente = voce
        elif pattern_sottocapitolo.match(voce):
            if capitolo_corrente:
                righe_mappa.append(f"{voce}  →  {capitolo_corrente}")
    return "\n".join(righe_mappa) or "Nessun sottocapitolo mappabile nell'indice."


def valuta_manoscritto_completo(indice, contenuti, titolo, trama, genere, stile, narrativa, pov, obiettivo, lingua, approfondimenti="", avanzamento=None):
    """Valuta ogni porzione del manoscritto, poi crea una sintesi basata sugli esiti effettivi."""
    sezioni_scritte = [s for s, c in contenuti.items() if pulisci_testo_editoriale(c).strip()]
    sezioni_vuote = [s for s, c in contenuti.items() if not pulisci_testo_editoriale(c).strip()]
    if not sezioni_scritte:
        return "VALUTAZIONE NON ESEGUIBILE: non è ancora presente testo da analizzare. Genera almeno una sezione e riprova."

    brief = f"""Titolo: {titolo}
Trama/argomento: {trama}
Genere: {genere}
Tipologia di scrittura: {stile}
Stile di racconto: {narrativa}
Punto di vista: {pov}
Obiettivo: {obiettivo}
Risultato finale desiderato: {val_risultato or "Non dichiarato"}
Approfondimenti prioritari: {approfondimenti.strip() or "Nessuno"}
{brief_personalizzazione_progetto()}"""
    blocchi = blocchi_per_audit_manoscritto(contenuti)
    totale_blocchi = len(blocchi)
    mappa_indice = mappa_capitoli_e_sottocapitoli(indice)
    # La versione dell'audit invalida in modo sicuro i vecchi report quando cambiano le regole di valutazione.
    firma_base = hashlib.sha256(
        f"{VERSIONE_AUDIT_COHERENZA}\n{brief}\n{indice}".encode("utf-8")
    ).hexdigest()
    cache_blocchi = st.session_state.setdefault("cache_audit_blocchi", {})

    def aggiorna_avanzamento(completati, fase):
        if avanzamento:
            avanzamento(completati, totale_blocchi, fase)

    def prompt_audit_blocco(numero, blocco):
        return f"""Sei un editor rigoroso. Analizza il BLOCCO {numero} di {len(blocchi)} di un manoscritto in lingua {lingua}.

BRIEF
{brief}

INDICE
{indice}

BLOCCO DA VALUTARE
{blocco}

Valuta soltanto le prove contenute nel blocco: aderenza al titolo e al brief, contributo al risultato finale desiderato, pertinenza, profondità, chiarezza, progressione locale, eventuali ripetizioni e istruzioni mancanti.

IDENTIFICAZIONE OBBLIGATORIA DELLE CORREZIONI
- Usa sempre il titolo esatto presente nell'indice o nell'intestazione del testo, compreso il numero (per esempio "Capitolo 4: ..." oppure "4.2 ...").
- Non usare mai formule vaghe come "questa parte", "il capitolo precedente" o "la sezione tecnica".
- Se il problema riguarda più sottocapitoli, crea una riga distinta per ciascun sottocapitolo interessato.
- Se non esistono problemi osservabili, scrivi "NESSUN PROBLEMA SPECIFICO RILEVATO"; non inventare difetti.

Non fare verifiche online e non citare fonti. Restituisci testo semplice, senza Markdown, con queste etichette:
SEZIONI ESAMINATE:
PUNTI FORTI:
PROBLEMI SPECIFICI:
Per ogni problema: SEZIONE ESATTA | PRIORITÀ (alta/media/bassa) | DIFETTO OSSERVATO | INTERVENTO RICHIESTO.
INTERVENTI PROPOSTI:"""

    esiti_per_numero, da_analizzare = {}, []
    completati = 0
    for numero, blocco in enumerate(blocchi, 1):
        firma_blocco = hashlib.sha256(f"{firma_base}\n{blocco}".encode("utf-8")).hexdigest()
        if firma_blocco in cache_blocchi:
            esiti_per_numero[numero] = cache_blocchi[firma_blocco]
            completati += 1
            aggiorna_avanzamento(completati, "Riutilizzo delle analisi già disponibili")
        else:
            da_analizzare.append((numero, blocco, firma_blocco))

    # Primo controllo: prezzo fisso. In seguito si pagano solo i blocchi modificati.
    if da_analizzare:
        # Se nessun blocco del manoscritto corrente è in cache è un primo
        # controllo, anche se la sessione contiene cache di un progetto passato.
        costo_controllo = (
            CREDIT_COSTS["controllo_coerenza_iniziale"]
            if completati == 0
            else len(da_analizzare) * CREDIT_COSTS["controllo_coerenza_blocco_modificato"]
        )
        addebita_azione_diretta("controllo_coerenza", amount=costo_controllo)
        with ThreadPoolExecutor(max_workers=3) as esecutore:
            lavori = {
                esecutore.submit(chiedi_audit_editoriale, prompt_audit_blocco(numero, blocco), addebita=False): (numero, firma_blocco)
                for numero, blocco, firma_blocco in da_analizzare
            }
            for lavoro in as_completed(lavori):
                numero, firma_blocco = lavori[lavoro]
                try:
                    esito = lavoro.result()
                except Exception as errore:
                    esito = f"ERRORE AUDIT: {errore}"
                cache_blocchi[firma_blocco] = esito
                esiti_per_numero[numero] = esito
                completati += 1
                aggiorna_avanzamento(completati, "Analisi editoriale dei blocchi del manoscritto")

    esiti_blocchi = [f"AUDIT BLOCCO {numero}\n{esiti_per_numero[numero]}" for numero in range(1, len(blocchi) + 1)]

    audit_compilati = "\n\n".join(esiti_blocchi)
    firma_completa = hashlib.sha256(f"{firma_base}\n{audit_compilati}".encode("utf-8")).hexdigest()
    cache_sintesi = st.session_state.setdefault("cache_sintesi_audit", {})
    if firma_completa in cache_sintesi:
        aggiorna_avanzamento(totale_blocchi, "Sintesi già disponibile")
        return cache_sintesi[firma_completa]
    aggiorna_avanzamento(totale_blocchi, "Preparazione della valutazione finale")
    sintesi = chiedi_audit_editoriale(f"""Sei un direttore editoriale. Prepara la valutazione finale di un manoscritto in lingua {lingua}, basandoti esclusivamente sul brief, sui dati oggettivi e sugli audit qui sotto. Non inventare contenuti non riportati.

BRIEF
{brief}

DATI OGGETTIVI
Sezioni scritte: {len(sezioni_scritte)}
Sezioni non scritte: {len(sezioni_vuote)}
Elenco sezioni non scritte: {', '.join(sezioni_vuote) if sezioni_vuote else 'nessuna'}

AUDIT DI TUTTI I BLOCCHI DEL MANOSCRITTO
{audit_compilati}

MAPPA UFFICIALE CAPITOLI E SOTTOCAPITOLI
{mappa_indice}

Se ci sono sezioni non scritte, dichiara con chiarezza che la valutazione riguarda un manoscritto parziale e non assegnare un voto finale al libro completo. Altrimenti assegna un voto da 1 a 10.
Non inserire URL, fonti, Markdown o ragionamenti interni. Restituisci esattamente queste voci:
STATO DEL MANOSCRITTO:
VOTO COMPLESSIVO DEL LIBRO:
VERDETTO EDITORIALE:
ADERENZA ALLA SIDEBAR:
PROFONDITÀ DELLE TEMATICHE:
COERENZA E PROGRESSIONE:
SEZIONI DA MIGLIORARE:
AZIONI PRIORITARIE:
PROMPT PRONTI PER RIGENERA CON AI:

REGOLE OBBLIGATORIE PER "SEZIONI DA MIGLIORARE" E PER I PROMPT:
- Elenca soltanto problemi realmente emersi dagli audit.
- Ogni voce deve identificare la destinazione con il titolo esatto dell'indice: prima il capitolo, poi il sottocapitolo specifico quando il difetto è locale.
- Usa la MAPPA UFFICIALE: un sottocapitolo (es. 4.2) non può mai essere indicato come capitolo di riferimento.
- Non scrivere mai indicazioni generiche quali "migliorare il capitolo" senza identificare il punto esatto.
- Riunisci tutti i difetti della stessa sezione in un unico blocco e in un solo prompt completo. Non creare cinque prompt separati per lo stesso sottocapitolo.
- Non accorpare sottocapitoli diversi nella stessa voce se richiedono interventi differenti.
- Se non rilevi criticità, scrivi "NESSUNA SEZIONE DA RIGENERARE" in entrambe le voci.

Per ogni sezione da migliorare, usa obbligatoriamente questo blocco separato:
CAPITOLO DI RIFERIMENTO: numero e titolo esatti del capitolo
SOTTOCAPITOLO DA SISTEMARE: numero e titolo esatti; scrivi "INTERO CAPITOLO" solo se il problema riguarda davvero ogni sua parte
PRIORITÀ: alta, media o bassa
PROBLEMA: difetto concreto osservato nel testo
OBIETTIVO DELLA CORREZIONE: risultato verificabile da ottenere
PROMPT DA INCOLLARE: istruzione autonoma, pronta da copiare nel campo "Rigenera con AI" della sezione indicata. Indica cosa mantenere, cosa aggiungere, cosa eliminare, genere, stile, POV e divieto di ripetere altre sezioni. Non proporre alcuna riscrittura automatica.""", addebita=False)
    cache_sintesi[firma_completa] = sintesi
    return sintesi

def costruisci_specifica_editoriale(titolo, genere, stile, narrativa, pov, obiettivo, argomento, risultato_finale="", approfondimenti=""):
    """Crea una specifica strutturata che indice e capitoli possono applicare in modo coerente."""
    return f"""=== SPECIFICA EDITORIALE STRUTTURATA ===
Titolo: {titolo}
Genere: {genere}
Tipologia di scrittura: {stile}
Stile di racconto: {narrativa}
Punto di vista: {pov}

OBIETTIVO OPERATIVO:
{obiettivo}

RISULTATO FINALE DESIDERATO E VERIFICABILE:
{risultato_finale or "Non dichiarato: definisci cosa il lettore deve saper fare o ottenere alla fine."}

ARGOMENTO E CONFINI:
{argomento}

APPROFONDIMENTI PRIORITARI (FACOLTATIVI):
{approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}

Per ogni sezione ricava un risultato concreto, il livello del lettore, i concetti necessari,
gli esempi o le procedure da produrre e ciò che deve restare fuori per evitare ripetizioni.
"""


def audit_simulazioni_test_prep(indice, contenuti, obiettivo, argomento):
    """Controlla che le simulazioni di un test prep contengano davvero quesiti e soluzioni completi."""
    risultati = []
    attese = numero_domande_simulazione(indice, argomento, obiettivo)
    sezioni = [
        sezione for sezione in contenuti
        if sezione_simulazione_test_prep(sezione, indice, "Test Prep (Preparazione Esami)")
    ]
    if not sezioni:
        return ["ERRORE TEST PREP: nell'indice non è stata individuata una sezione che contenga i quesiti di una simulazione."]
    if not attese:
        risultati.append("AVVISO TEST PREP: il brief non dichiara quante domande deve contenere ogni simulazione; il controllo non può validarne la completezza numerica.")

    domande_fuori_simulazione = []
    for sezione, testo in contenuti.items():
        if sezione not in sezioni:
            domande_fuori_simulazione.extend(domande_normalizzate_test_prep(testo))

    for sezione in sezioni:
        testo = contenuti.get(sezione, "")
        domande = domande_normalizzate_test_prep(testo)
        soluzioni = len(re.findall(r"(?im)^\s*soluzione\s+\d{1,3}\s*[:.-]", testo or ""))
        risultati.append(f"Simulazione '{sezione}': domande rilevate {len(domande)}; soluzioni rilevate {soluzioni}.")
        if attese and len(domande) != attese:
            risultati.append(f"ERRORE TEST PREP: '{sezione}' richiede {attese} domande ma ne contiene {len(domande)}.")
        if attese and soluzioni != attese:
            risultati.append(f"ERRORE TEST PREP: '{sezione}' richiede {attese} soluzioni commentate ma ne contiene {soluzioni}.")
        ripetute_interne = len(domande) - len(set(domande))
        ripetute_esterne = len(set(domande) & set(domande_fuori_simulazione))
        if ripetute_interne:
            risultati.append(f"ERRORE TEST PREP: '{sezione}' contiene {ripetute_interne} domande duplicate al suo interno.")
        if ripetute_esterne:
            risultati.append(f"AVVISO TEST PREP: '{sezione}' riutilizza {ripetute_esterne} domande già presenti in quiz o altre sezioni.")
        if (not attese or (len(domande) == attese and soluzioni == attese)) and not ripetute_interne:
            risultati.append(f"OK TEST PREP: struttura numerica della simulazione '{sezione}' valida.")
    return risultati

def analizza_coerenza_libro(indice, contenuti, obiettivo, argomento, genere="", risultato_finale=""):
    """Controllo deterministico preliminare su struttura, copertura e ripetizioni."""
    risultati = ["REPORT CONTROLLO COERENZA DEL LIBRO"]
    testo = "\n".join(contenuti.values()) if contenuti else ""
    capitoli = re.findall(r"(?im)^(?:Capitolo|Chapter|CAPITOLO)\\s+\\d+", indice or "")
    sottocapitoli = re.findall(r"(?m)^\\d+\\.\\d+\\s+", indice or "")
    risultati.append(f"Capitoli rilevati: {len(capitoli)}")
    risultati.append(f"Sottocapitoli rilevati: {len(sottocapitoli)}")
    if not indice.strip(): risultati.append("ERRORE: indice assente")
    if not obiettivo.strip(): risultati.append("AVVISO: obiettivo assente")
    if not risultato_finale.strip(): risultati.append("AVVISO: risultato finale desiderato assente")
    if not argomento.strip(): risultati.append("AVVISO: argomento assente")
    sezioni_indice = st.session_state.get("lista_capitoli", [])
    if sezioni_indice:
        mancanti = sezioni_mancanti_per_esportazione(sezioni_indice, genere)
        risultati.append(f"Sezioni dell'indice non ancora complete: {len(mancanti)}")
        if mancanti:
            anteprima = "; ".join(mancanti[:8])
            risultati.append(f"ERRORE: non esportare il libro. Sezioni incomplete: {anteprima}" + (" ..." if len(mancanti) > 8 else ""))
        else:
            risultati.append("OK: tutte le sezioni previste dall'indice risultano compilate.")
    if genere == "Ricettario":
        numero_richiesto = estrai_numero_ricette("", argomento, obiettivo)
        capitoli_ricetta = [s for s in sezioni_indice if tipo_sezione_editoriale(s) == "capitolo"]
        if numero_richiesto:
            risultati.append(f"Ricette richieste dal brief: {numero_richiesto}; capitoli-ricetta nell'indice: {len(capitoli_ricetta)}")
            if len(capitoli_ricetta) != numero_richiesto:
                risultati.append("ERRORE RICETTARIO: il numero delle ricette nell'indice non coincide con la richiesta.")
        termini_non_vegani = r"\b(miele|latte vaccino|burro|uova|formaggio|panna|yogurt greco)\b"
        non_vegani = sorted(set(re.findall(termini_non_vegani, testo.lower())))
        if non_vegani:
            risultati.append("AVVISO RICETTARIO: verificare ingredienti non vegani rilevati: " + ", ".join(non_vegani))
    if genere == "Test Prep (Preparazione Esami)":
        risultati.extend(audit_simulazioni_test_prep(indice, contenuti, obiettivo, argomento))
    if len(testo.strip()) < 1000: risultati.append("AVVISO: contenuto ancora troppo breve per una verifica completa")
    frasi = [f.strip().lower() for f in re.split(r"[.!?]+", testo) if len(f.strip()) > 40]
    duplicati = len(frasi) - len(set(frasi))
    risultati.append(f"Frasi duplicate identiche rilevate: {max(0, duplicati)}")
    if duplicati == 0: risultati.append("OK: nessuna duplicazione identica rilevata nel testo disponibile")
    return "\n".join(risultati)

# ======================================================================================================================
# 8. UI PRINCIPALE & GENERAZIONE PROMPT DINAMICO
# ======================================================================================================================
sync_capitoli()
lista_cap_base = st.session_state.get("lista_capitoli", [])
# Il manoscritto segue esclusivamente le sezioni definite nell'indice.
sezioni_struttura_corrente = [
    sezione for sezione in lista_cap_base
    if not sezione_dismessa(sezione)
]
# Il registro editoriale mantiene l'ordine e la visibilità delle sezioni già
# elaborate anche dopo una pausa o un aggiornamento della pagina.
if lista_cap_base:
    registro_editoriale = st.session_state.setdefault(CHIAVE_REGISTRO_SEZIONI, [])
    for sezione_editoriale in sezioni_struttura_corrente:
        if sezione_editoriale not in registro_editoriale:
            registro_editoriale.append(sezione_editoriale)
# Durante “Scrivi tutto il libro” l'elenco viene fissato anche nella sessione.
# In questo modo una pausa o un rerun non può ridurre l'elenco dell'editor alle
# sole voci rileggibili dall'indice.
sezioni_job_protette = [
    sezione for sezione in (st.session_state.get("job_scrittura_sezioni", []) or [])
    if not sezione_dismessa(sezione)
]
opzioni_editor = elenco_sezioni_progetto([
    *sezioni_struttura_corrente,
    *sezioni_job_protette,
])
# Prima di disegnare l'editor, ripristina ogni testo già protetto nella memoria
# del progetto. Il cambio della sezione selezionata non può quindi svuotare le
# altre caselle o nasconderle dall'anteprima/esportazione.
reidrata_sezioni_memorizzate(opzioni_editor)

# Intestazione di lavoro: riassume il progetto senza occupare lo spazio della
# scrittura. I pulsanti operativi restano nelle rispettive sezioni, così non
# cambiano il flusso o la logica esistente dell'editor.
titolo_workspace = html.escape((val_titolo or "Editor professionale").strip())
motore_workspace = html.escape(provider_ia.replace(" (OpenAI)", ""))
sezioni_con_testo_workspace = sum(
    1 for contenuto in (st.session_state.get(CHIAVE_MEMORIA_SEZIONI, {}) or {}).values()
    if str(contenuto).strip()
)
stato_workspace = "Progetto pronto" if sidebar_pronta else "Completa la configurazione"
classe_stato_workspace = "ok" if sidebar_pronta else ""
st.markdown(
    f'''<div class="ss-workspace-header">
        <div class="ss-workspace-brand">
            <div class="ss-workspace-mark">✒</div>
            <div>
                <div class="ss-workspace-title">{titolo_workspace}</div>
                <div class="ss-workspace-subtitle">Scrittore Site · spazio di lavoro editoriale</div>
            </div>
        </div>
        <div class="ss-workspace-meta">
            <span class="ss-workspace-chip {classe_stato_workspace}">{stato_workspace}</span>
            <span class="ss-workspace-chip ai">🧠 {motore_workspace}</span>
            <span class="ss-workspace-chip">✍ {sezioni_con_testo_workspace} sezioni scritte</span>
            <span class="ss-workspace-chip credit">💾 Salvataggio disponibile</span>
        </div>
    </div>''',
    unsafe_allow_html=True,
)
st.caption(VERSIONE_DEPLOY)
mostra_centro_progetto(lingua_sel, campi_obbligatori_sidebar)

if st.session_state.get("admin_test_mode"):
    # Pannello visibile soltanto dopo l'avvio dal comando protetto della
    # sidebar amministratore. Non avvia chiamate AI da solo: i test usano gli
    # stessi pulsanti e gli stessi prompt dell'app pubblica.
    st.warning(
        f"🧪 **Laboratorio amministratore attivo** — {st.session_state.get('admin_test_profile', 'Collaudo breve')} "
        f"con **{st.session_state.get('admin_test_provider', provider_ia)}**. "
        "I crediti utente non vengono addebitati; le chiamate AI sono reali. Il progetto è isolato e non sovrascrive la bozza cloud."
    )
    indice_test = str(st.session_state.get("indice_raw", "") or "").strip()
    fonti_test = str(st.session_state.get("registro_fonti_web", "") or "").strip()
    testi_test = {nome: str(testo).strip() for nome, testo in (st.session_state.get(CHIAVE_MEMORIA_SEZIONI, {}) or {}).items() if str(testo).strip()}
    frasi_non_concluse = [nome for nome, testo in testi_test.items() if testo and testo.rstrip()[-1:] not in ".!?…»”\"'"]
    massimo_voci_visualizzato = limite_voci_indice
    indice_nei_limiti = bool(indice_test) and conta_sezioni_indice(indice_test) <= limite_voci_indice
    verifica_visiva_test = st.checkbox(
        campo_controllo_ui(0),
        key="admin_test_verifica_visiva",
        help="Spunta soltanto dopo avere provato davvero le funzioni nel browser.",
    )
    fasi_collaudo = [
        ("Sidebar e profilo", True),
        ("Personalizzazione", bool(st.session_state.get("admin_test_personalizzazione_verificata"))),
        ("Fonti web", bool(fonti_test)),
        ("Indice nei limiti", indice_nei_limiti),
        ("Scrittura sezione", bool(testi_test)),
        ("Pausa e recupero", bool(st.session_state.get("admin_test_pausa_ripresa_verificata"))),
        ("Completezza testi", bool(testi_test) and not frasi_non_concluse),
        ("Anteprima, voce, CSV e PDF", verifica_visiva_test),
    ]
    fasi_completate = sum(1 for _nome, completata in fasi_collaudo if completata)
    percentuale_collaudo = int(fasi_completate / len(fasi_collaudo) * 100)
    prossima_fase = next((nome for nome, completata in fasi_collaudo if not completata), "Collaudo completato")
    st.progress(
        percentuale_collaudo,
        text=f"Avanzamento collaudo: {fasi_completate}/{len(fasi_collaudo)} fasi completate — prossima: {prossima_fase}.",
    )
    st.caption("Avvio automatico: ricerca fonti, indice breve con controllo editoriale e due sezioni reali. Include una simulazione di salvataggio e recupero della stesura. Non può verificare da solo funzioni del browser come voce, download CSV e PDF.")
    if st.button("🚀 AVVIA COLLAUDO AUTOMATICO", type="primary", key="avvia_collaudo_automatico", use_container_width=True):
        st.session_state["admin_test_run_requested"] = True
    esito_collaudo_automatico = st.session_state.get("admin_test_run_report", "")
    if esito_collaudo_automatico:
        if esito_collaudo_automatico.startswith("ERRORE:"):
            st.error(esito_collaudo_automatico)
        else:
            st.success(esito_collaudo_automatico)
    risultati_test = [
        ("Sidebar e profilo", True, "Brief breve completo e cervello selezionato."),
        ("Personalizzazione", bool(st.session_state.get("admin_test_personalizzazione_verificata")), "Il collaudo verifica che le indicazioni personali entrino realmente nel brief di indice e stesura." if not st.session_state.get("admin_test_personalizzazione_verificata") else "Indicazioni personali presenti nel brief reale, senza modificare i parametri ordinari del collaudo."),
        ("Fonti web", bool(fonti_test), "Genera l'indice per verificare la ricerca e il registro fonti." if not fonti_test else "Registro fonti disponibile."),
        ("Indice e prompt", indice_nei_limiti, "Apri Indice e premi Genera indice professionale." if not indice_test else (f"{conta_sezioni_indice(indice_test)} voci rilevate (massimo test: {massimo_voci_visualizzato})." if indice_nei_limiti else f"Indice fuori limite: {conta_sezioni_indice(indice_test)} voci su massimo {massimo_voci_visualizzato}.")),
        ("Scrittura sezione", bool(testi_test), "Apri Scrittura e genera almeno una sezione." if not testi_test else f"{len(testi_test)} sezioni create."),
        ("Pausa e recupero", bool(st.session_state.get("admin_test_pausa_ripresa_verificata")), "Avvia il collaudo automatico: genera due sezioni dell'indice, simula una pausa e verifica che restino leggibili nella memoria editoriale." if not st.session_state.get("admin_test_pausa_ripresa_verificata") else "Sezioni conservate e rilette correttamente dopo la simulazione della pausa."),
        ("Completezza frasi", bool(testi_test) and not frasi_non_concluse, "Da verificare dopo la prima sezione." if not testi_test else ("Nessuna chiusura tronca rilevata." if not frasi_non_concluse else "Da rielaborare: " + ", ".join(frasi_non_concluse[:3]))),
        ("Anteprima / voce / CSV", verifica_visiva_test, "Verifica manualmente Anteprima, lettore vocale e Importa / Esporta / Copyright: sono funzioni del browser e richiedono un controllo visivo." if not verifica_visiva_test else "Verifica visiva confermata dall'amministratore."),
    ]
    colonne_test = st.columns([2, 2, 2])
    colonne_test[0].metric("Indice", "✓" if indice_test else "—")
    colonne_test[1].metric("Sezioni", len(testi_test))
    colonne_test[2].metric("Fonti web", len(re.findall(r"https?://\\S+", fonti_test)))
    with st.expander("Checklist del collaudo", expanded=True):
        for nome, superato, dettaglio in risultati_test:
            st.write(f"{'✅' if superato else '⬜'} **{nome}** — {dettaglio}")
        st.caption("Sequenza: 1) genera l'indice; 2) scrivi una sezione; 3) prova rielaborazione e controllo; 4) verifica anteprima, voce, CSV e PDF; 5) ripeti con l'altro cervello e un'altra modalità.")
    if st.button("↩ TERMINA COLLAUDO E RIPRISTINA IL PROGETTO PRECEDENTE", key="termina_collaudo_amministratore", use_container_width=True):
        termina_collaudo_amministratore()
        st.rerun()

# La guida deve essere disponibile anche al primo avvio, prima che l'utente compili il brief.
interfaccia_editor_disponibile = True
if interfaccia_editor_disponibile:
    
    # VALUTAZIONE DINAMICA: L'IA decide se usare o meno la manipolazione cerebrale
    usa_tre_cervelli = valuta_approccio_neurologico(val_genere, val_stile, val_narrativa)
    
    if usa_tre_cervelli:
        modulo_stilistico = """
=== METODOLOGIA DEI 3 CERVELLI (NEUROMARKETING) ===
Devi strutturare il testo per comunicare simultaneamente con i 3 livelli cerebrali del lettore, iniettando la giusta chimica:
1. CERVELLO RETTILE (Sopravvivenza & Istinto): Usa un linguaggio netto, tangibile e basato sui contrasti (prima/dopo, problema/soluzione). Attira l'attenzione istantaneamente. Elimina parole deboli o passive.
2. CERVELLO LIMBICO (Emozione & Chimica): Usa "Storytelling" ed empatia. Scegli vocaboli sensoriali che stimolino il rilascio di dopamina (curiosità/ricompensa) e ossitocina (fiducia/connessione). Fai percepire al lettore che comprendi esattamente il suo stato d'animo.
3. NEOCORTECCIA (Logica & Dati): Fornisci struttura, dati precisi, ragionamenti logici e prove che giustifichino razionalmente le emozioni suscitate dal sistema limbico.
"""
    else:
        modulo_stilistico = """
=== APPROCCIO ANALITICO E OGGETTIVO ===
Il genere e lo stile scelti richiedono un approccio neutrale e rigoroso. 
NON utilizzare manipolazioni emotive o neuromarketing. Mantieni un tono accademico, logico e fattuale. 
Fornisci dati, structures deduttive e un linguaggio pulito, tipico delle pubblicazioni di alto rigore tecnico-scientifico.
"""

    modulo_fonti = ""
    if st.session_state.get("conoscenza_extra"):
        modulo_fonti = """
=== FONTI ESTERNE: METODO DI ORIGINALITÀ OBBLIGATORIO ===
Hai a disposizione soltanto una mappa concettuale interna già separata dai documenti caricati. I documenti non sono un modello di stile, struttura o formulazione.
1. Usa le fonti per verificare concetti e terminologia, non per riscriverle frase per frase.
2. Progetta sempre una spiegazione indipendente: cambia ordine, angolazione, esempi, connessioni e sviluppo logico in funzione del brief di questo libro.
3. Non riprendere formulazioni distintive, titoli, elenchi, dialoghi, casi, metafore o sequenze di sei o più parole provenienti dalle fonti.
4. Produci contenuto nuovo e contestualizzato, non una parafrasi. Se un fatto deve restare invariato, spiega il suo significato con parole e un esempio propri.
5. Non inserire citazioni o riferimenti nel manoscritto salvo istruzione esplicita dell'utente e disponibilità dei relativi diritti.
"""

    # --- INIZIO NUOVE RIGHE PER ADATTAMENTO PROMPT IN BASE AL GENERE ---
    modulo_approfondimento_genere = ""
    # Aggiunto "Storico" a questo blocco logico
    if "Manuale" in val_genere or "Saggio" in val_genere or "Test Prep" in val_genere or "Economia" in val_genere or "Storico" in val_genere:
        modulo_approfondimento_genere = """
=== DIRETTIVA DI APPROFONDIMENTO PROPORZIONATO (MANUALISTICA E SAGGISTICA) ===
Trattandosi di un testo tecnico, didattico o manualistico, il tuo compito primario è ISTRUIRE. 
Ogni sottocapitolo deve essere accurato e utile, entro la lunghezza assegnata. 
- NON dare MAI nulla per scontato o rimanere in superficie.
- Spiega dettagliatamente "COSA è", "PERCHÉ funziona così" e "COME si applica" nella pratica.
- Inserisci esempi concreti, casi d'uso pratici, schemi logici o spiegazioni passo-passo.
- Il lettore deve acquisire una competenza reale, dettagliata e spendibile alla fine di questa sezione. Evita categoricamente la superficialità.
"""
    elif "Romanzo" in val_genere or "Narrativo" in val_genere or "Thriller" in val_genere or "Fantasy" in val_genere or "Fantascienza" in val_genere:
        modulo_approfondimento_genere = """
=== DIRETTIVA DI IMMERSIONE NARRATIVA (NARRATIVA E ROMANZO) ===
Trattandosi di un'opera narrativa, il tuo focus esclusivo è lo STORYTELLING e l'immersione.
- Mostra, non raccontare (Show, Don't Tell). Descrivi minuziosamente gli ambienti, le espressioni e le atmosfere usando i 5 sensi.
- Dai spessore profondo ai personaggi, ai dialoghi e gestisci il ritmo dell'azione o dell'introspezione.
- Evita totalmente lo stile accademico, saggistico o manualistico: il lettore deve "vivere" la scena in tempo reale, non studiarla.
"""
    # --- FINE NUOVE RIGHE ---

    # --- INIZIO NUOVE RIGHE PER APPROFONDIMENTO ED ESEMPI SPECIFICI PER GENERE ---
    modulo_esempi_specifici = """
=== DIRETTIVA DI PROFONDITÀ ED ESEMPI PRATICI (CALIBRATA SUL GENERE) ===
Il tuo compito non è solo descrivere, ma DIMOSTRARE e APPROFONDIRE. In base al genere selezionato, devi generare contenuti operativi o immersivi reali:
"""
    if "Test Prep" in val_genere or "Quiz" in val_genere:
        modulo_esempi_specifici += "- Crea VERE SIMULAZIONI D'ESAME o test complessi inerenti al capitolo. Inserisci domande a risposta multipla, scenari pratici e fornisci le soluzioni dettagliate con spiegazione logica per ogni opzione (perché è giusta o sbagliata).\n"
    elif "Manuale" in val_genere or "Business" in val_genere or "Economia" in val_genere or "Self-Help" in val_genere:
        modulo_esempi_specifici += "- Inserisci veri e propri CASI STUDIO (reali o verosimili), framework applicativi, checklist e scenari di 'Roleplay' o 'What-if' per mostrare come applicare la teoria nella realtà.\n"
    # Aggiunto blocco logico per "Storico"
    elif "Storico" in val_genere:
        modulo_esempi_specifici += "- Inserisci riferimenti storici accurati, date precise, eventi chiave, analisi del contesto socio-politico e cita documenti o figure di rilievo dell'epoca.\n"
    elif "Saggio" in val_genere or "Tecnico" in val_genere:
        modulo_esempi_specifici += "- Fornisci spiegazioni tecniche microscopiche, formule, dati statistici ed esempi concreti di applicazione nel mondo reale per supportare la tesi.\n"
    elif "Romanzo" in val_genere or "Narrativo" in val_genere or "Fantasy" in val_genere or "Thriller" in val_genere:
        modulo_esempi_specifici += "- Crea scene vissute. Mostra interazioni specifiche, dialoghi autentici tra personaggi e reazioni ambientali. Non riassumere gli eventi, falli accadere 'in camera'.\n"
    elif "Ricettario" in val_genere:
        modulo_esempi_specifici += "- Aggiungi varianti degli ingredienti, trucchi dello chef per rimediare agli errori comuni e cenni scientifici sul perché avvengono certe reazioni in cottura.\n"
    else:
        modulo_esempi_specifici += "- Non rimanere mai in superficie: ogni volta che introduci un concetto, fai subito un ESEMPIO PRATICO e dettagliato che lo esplichi al 100%.\n"
    # --- FINE NUOVE RIGHE ---

    # --- INIZIO NUOVE RIGHE PER ADERENZA TITOLO-GENERE (DETTAGLIO ESTREMO) ---
    modulo_aderenza_titolo_genere = f"""
=== DIRETTIVA DI ESECUZIONE REALE E SPECIFICA (INCROCIO GENERE E TITOLO SEZIONE) ===
Analizza attentamente il genere letterario ('{val_genere}') e il titolo esatto della sezione che stai per scrivere.
Non limitarti a fare un discorso generico: devi FORNIRE MATERIALMENTE ciò che il titolo della sezione suggerisce, declinato per quel genere.
- Esempio 1: Se il genere è "Test Prep" o simile e il titolo della sezione che stai scrivendo parla di "Test", "Quiz", "Simulazione" o "Autovalutazione", DEVI generare un VERO test (es. domande a risposta multipla, scenari, soluzioni e spiegazioni dettagliate del perché un'opzione è corretta o sbagliata). Non descrivere come si fa un test, FALLO E REDIGILO REALMENTE.
- Esempio 2: Se il genere è "Manuale", "Business" o "Economia" e il titolo suggerisce un "Piano d'Azione", un "Caso Studio" o un "Esercizio", scrivi i passaggi operativi completi o il caso studio di esempio con nomi, dati, calcoli finanziari e soluzioni.
- Esempio 3: Se il genere è "Ricettario" e la sezione è un piatto, scrivi la ricetta vera e propria.
- Regola Universale: Cogli l'intento pratico implicito nel titolo della sezione corrente in base al genere. Se la sezione richiede un contenuto specifico (una tabella, un test di autovalutazione, un esercizio, un dialogo d'esempio), REDIGI QUEL CONTENUTO FISICAMENTE. Il lettore deve trovarsi di fronte allo strumento o alla scena reale, non a un riassunto teorico.
"""
    # --- FINE NUOVE RIGHE ---

    modulo_operativita_universale = """
=== DIRETTIVA OBBLIGATORIA DI OPERATIVITÀ E DETTAGLIO ===
Ogni sottocapitolo deve insegnare concretamente l'argomento, non limitarsi a descriverlo.
Scegli soltanto gli elementi realmente necessari al titolo: procedura, esempio, errore comune,
esercizio o criterio verificabile. Non inserire tutti gli elementi nello stesso testo solo per allungarlo.
Se la sezione tratta prompt, scrivi prompt completi e copiabili. Se tratta software,
indica menu, comandi, pulsanti e sequenza esatta. Se tratta una procedura, eseguila
materialmente nel testo con dati o esempi realistici. Distingui fatti, esempi ipotetici
e indicazioni da verificare. Non promettere risultati garantiti e non usare formule vaghe.
Nel capitolo principale mantieni la visione d'insieme; nel sottocapitolo sviluppa il
dettaglio assegnato senza anticipare o ripetere gli altri sottocapitoli.
"""

    profilo_lunghezza_corrente = PROFILI_LUNGHEZZA_STESURA.get(
        val_lunghezza, PROFILI_LUNGHEZZA_STESURA["Standard KDP"]
    )
    minimo_parole_tolleranza, massimo_parole_tolleranza = vincolo_parole_con_tolleranza(val_lunghezza)
    modulo_lunghezza = f"""
=== LUNGHEZZA E DENSITÀ OBBLIGATORIE ===
Profilo scelto: {val_lunghezza}.
- Per una sezione autonoma, l'obiettivo è {profilo_lunghezza_corrente['parole']}.
- Tolleranza massima consentita: 5%. Non produrre meno di {minimo_parole_tolleranza} parole né più di {massimo_parole_tolleranza} parole.
- L'obiettivo del manoscritto completo è {profilo_lunghezza_corrente['pagine_minime']} pagine nel formato 6×9, con tolleranza del 15%. Contribuisci con contenuto utile, senza gonfiare il testo; se una sezione è già completa e di qualità, non aggiungere riempitivi solo per aumentare le pagine.
- Il limite di output dell'AI è configurato in coerenza con questa tolleranza: non aggirarlo con frasi riempitive o elenchi superflui.
- Prima di iniziare, distribuisci mentalmente lo spazio: sviluppa prima i passaggi essenziali e riserva sempre circa il 12% finale della sezione alla conclusione. Non avviare un nuovo esempio, elenco o periodo se non puoi terminarlo entro il limite.
- Completa sempre l'ultima idea con una frase significativa: se lo spazio non basta, riduci prima dettagli secondari, esempi o elenchi, senza interrompere ragionamenti, procedure o scene.
- Questa direttiva prevale su ogni invito generico a essere estremamente dettagliato o a includere molte categorie di esempi.
- Una sezione è valida quando risponde bene al suo titolo con informazioni nuove, non quando ripete o aggiunge dettagli non necessari.
- Un capitolo con sottocapitoli è solo una cornice breve: i contenuti completi appartengono ai sottocapitoli.
"""

    profilo_genere_corrente = profilo_genere_stesura(val_genere)
    profilo_tipologia_corrente = profilo_tipologia_stesura(val_stile)
    profilo_indice_corrente = profilo_struttura_indice(val_genere, val_titolo, val_trama, val_goal)

    # PROMPT POTENZIATO CON COERENZA POV, PULIZIA SINTATTICA E CONFORMITA' DI GENERE
    S_PROMPT = f"""
Sei un esperto Madrelingua in {lingua_sel}, Editor e Luminare mondiale nel campo '{val_genere}'. 
Stai redigendo l'ebook '{val_titolo}'. 

{modulo_fonti}

{modulo_approfondimento_genere}
{modulo_esempi_specifici}
{modulo_aderenza_titolo_genere}
{modulo_operativita_universale}
{modulo_lunghezza}

=== PROFILO EDITORIALE SPECIFICO ===
- Regole del genere '{val_genere}': {profilo_genere_corrente}
- Regole della tipologia '{val_stile}': {profilo_tipologia_corrente}
- Architettura richiesta: {profilo_indice_corrente}
Queste regole prevalgono sulle istruzioni generiche incompatibili con il genere. Non applicare
procedure, checklist, test, sottocapitoli o scene narrative quando non sono pertinenti.

PARAMETRI DI BASE (DA APPLICARE TASSATIVAMENTE IN OGNI SEZIONE):
- Stile di Racconto: {val_narrativa}
- Obiettivo Emozionale/Pratico: {val_goal}
- Tipologia di Scrittura: {val_stile}
- Punto di Vista (Relazione con il lettore): {val_pov}. Adatta coerentemente questo pronome alla grammatica della lingua {lingua_sel}.
- Conformità di Genere: Il testo DEVE rispecchiare in pieno le regole, la formattazione e la terminologia del genere '{val_genere}' (es. se è un ricettario, usa formati strutturati con ingredienti e step; se è un romanzo usa narrazione fluida; se è 'Test Prep', usa schemi, riassunti puntati, concetti chiave da memorizzare e simulazioni d'esame).
- Lingua di Output Categorica: {lingua_sel}

{specifica_editoriale}

{modulo_stilistico}

=== REGOLA DI FORMATTAZIONE E SINTASSI PULITA (CRITICO) ===
- Usa ESCLUSIVAMENTE una punteggiatura standard, tipografica e impeccabile. 
- SONO SEVERAMENTE VIETATE punteggiature anomale, artefatti markdown inutili, asterischi eccessivi, o emoji nel corpo del testo.
- Il testo deve scorrere con l'eleganza formale e la pulizia di un vero libro stampato (sintassi corretta, paragrafi chiari).

=== COMPLETEZZA DEL MANOSCRITTO (OBBLIGATORIA PER OGNI CERVELLO AI) ===
- Prima di consegnare una sezione, rileggi internamente ogni periodo e soprattutto le ultime due frasi.
- Non lasciare mai frasi spezzate, periodi sospesi, elenchi iniziati ma non terminati, parentesi o virgolette aperte,
  parole di collegamento finali, procedure interrotte, scene senza esito o ragionamenti privi di conclusione.
- Riserva una parte dello spazio disponibile alla chiusura: non iniziare un nuovo esempio o un nuovo passaggio se non
  puoi terminarlo. Se sei vicino al limite di lunghezza, elimina prima un dettaglio secondario e chiudi con una frase
  breve, grammaticale e significativa che completi l'idea della sezione.
- Questa regola vale indistintamente per GPT-5.4 e DeepSeek V4 Pro, per stesura singola, capitolo, sottocapitoli,
  rigenerazioni e scrittura dell'intero libro.

=== REGOLA AUREA: GERARCHIA E NON-RIPETIZIONE (CAPITOLO VS SOTTOCAPITOLO) ===
Dovrai analizzare l'indice fornito per capire la tua esatta posizione:
- SE STAI SCRIVENDO UN CAPITOLO PRINCIPALE (es. 1, 2, 3): Focalizzati sulla visione d'insieme, introduci l'argomento in modo macroscopico. NON rubare i dettagli tecnici, gli esempi specifici o i casi studio che appartengono ai tuoi sottocapitoli.
- SE STAI SCRIVENDO UN SOTTOCAPITOLO (es. 1.1, 1.2, 3.4): Entra inmediatamente nel dettaglio estremo, nell'azione pratica o nell'analisi profonda. NON ripetere mai le premesse o le introduzioni generali già spiegate nel capitolo padre. 
- MEMORIA GLOBALE: Leggi il contesto fornito. Non ripetere mai concetti, parole chiave o aneddoti già utilizzati in altre sezioni.

=== DIRETTIVA ANTI-RIPETIZIONE E BLACKLIST DEGLI ARGOMENTI ===
Il sistema anti-ripetizione è il parametro più critico di questa operazione:
1. DISTINZIONE PADRE/FIGLIO: Se stai scrivendo un Capitolo Principale (es. "Capitolo 1"), devi limitarti a una visione "dall'alto", introducendo i temi SENZA svelarne le meccaniche o gli esempi. Se stai scrivendo un Sottocapitolo (es. "1.1" o "1.2"), devi entrare nel micro-dettaglio e ti è SEVERAMENTE VIETATO riassumere o ripetere l'introduzione già fatta nel Capitolo Padre.
2. BLACKLIST DEI CONTENUTI PRECEDENTI: I contenuti presenti nella "MEMORIA CONTENUTI PRECEDENTI" sono da considerarsi in una BLACKLIST. Non usare MAI le stesse introduzioni, non riciclare esempi e non riproporre gli stessi concetti o checklist. Ogni sezione deve essere 100% inedita rispetto alle precedenti.

=== SILENZIO STAMPA ASSOLUTO SUI SOTTOCAPITOLI (MUTUAMENTE ESCLUSIVI) ===
Questa è la regola d'oro per evitare sovrapposizioni e non farti trattare lo stesso argomento due volte:
1. IL CAPITOLO PARLA DEL "PERCHÉ": Se l'indice ti posiziona nella stesura di un Capitolo Padre (es. "Capitolo 2"), il tuo UNICO compito è creare la cornice concettuale. Ti è IMPOSTO IL SILENZIO STAMPA su qualsiasi argomento, tecnica o dettaglio che abbia un Sottocapitolo dedicato (es. 2.1, 2.2). NON SPIEGARE NIENTE DI SPECIFICO NEL CAPITOLO PADRE.
2. IL SOTTOCAPITOLO PARLA DEL "COME" e del "COSA": Se la sezione è un Sottocapitolo (es. "2.1 L'argomento X"), l'intera spiegazione dell'Argomento X DEVE avvenire ESCLUSIVAMENTE lì. Nel Capitolo Padre, X non doveva essere spiegato, ma al massimo accennato come un titolo nel futuro.
3. CONTROLLO FINALE PRIMA DI GENERARE: Guarda la lista completa dei tuoi sottocapitoli e chiediti: "Sto spiegando in questo testo qualcosa che l'indice dice di spiegare nel prossimo paragrafo numerato?". Se la risposta è SÌ, CANCELLA e astieniti. Lascia vuoto informativo per permettere al Sottocapitolo di esistere senza ripetizioni.

=== DIVIETO DI ANTICIPAZIONE (SPOILER SUI SOTTOCAPITOLI) ===
ASCOLTA ATTENTAMENTE: Se l'indice prevede che un argomento specifico venga trattato in un Sottocapitolo (es. 1.1, 1.2, 1.3), è ASSOLUTAMENTE VIETATO parlarne, menzionarlo o spiegarlo nel Capitolo Padre (es. Capitolo 1).
Il Capitolo Padre deve fungere SOLO da cornice introduttiva generale. Non deve MAI svuotare di significato i sottocapitoli anticipandone i contenuti. Mantieni il vuoto informativo sulle questioni specifiche finché non arrivi a scrivere il sottocapitolo dedicato.

=== APPLICAZIONE DIRETTIVE (STESURA PULITA) ===
Devi interiorizzare e applicare alla lettera le seguenti istruzioni prima di generare il testo:
1. Il genere '{val_genere}'
2. La tipologia di scrittura '{val_stile}' e lo stile di racconto '{val_narrativa}'
3. Il POV '{val_pov}'
4. L'obiettivo '{val_goal}'
CRITICO: NON inserire alcun "ragionamento editoriale", commento, introduzione o meta-testo. L'output DEVE contenere ESCLUSIVAMENTE il contenuto finale del capitolo/sottocapitolo, pronto per la pubblicazione.

=== DIRETTIVA DI CONFORMITÀ ASSOLUTA (PUNTO DI VISTA E STILE) ===
È TASSATIVO e NON NEGOZIABILE che l'intero testo sia redatto utilizzando ESATTAMENTE il Punto di Vista (POV) impostato nella sidebar: "{val_pov}". 
- Se è impostato su "Tu", rivolgiti direttamente e informalmente al singolo lettore (es. "scoprirai che...").
- Se è impostato su "Voi", rivolgiti in modo plurale e autorevole (es. "scoprirete che...").
- Se è impostato su "Noi", usa un approccio inclusivo (es. "scopriremo che...").
- Se è "Impersonale", usa forme impersonali o passive, distaccate e oggettive (es. "si scoprirà che...").
L'intelligenza artificiale DEVE effettuare un controllo lessicale e grammaticale ad ogni fine paragrafo per assicurarsi che non ci siano "scivoloni" o cambi di pronome accidentali. Lo stile di scrittura "{val_stile}" deve permeare ogni singola scelta di vocabolario.

=== REGOLA DELLA DENSITÀ E APPROCCIO DIRETTO (NO FLUFF) ===
- VAI AL SODO: Elimina qualsiasi preambolo inutile, frasi fatte o giri di parole. Inizia immediatamente a trattare il cuore dell'argomento della sezione.
- ZERO VAGHEZZA: Sii estremamente descrittivo, specifico e dettagliato. Non limitarti a enunciare i concetti, ma sviscerali e dimostrali.
- PROFONDITÀ ARGOMENTATIVA: Tratta gli argomenti in maniera fortemente argomentativa. Se stai spiegando una teoria, una tecnica o un concetto pratico, fornisci il "come" e il "perché" con autorevolezza, supportando le tue affermazioni con logica ferrea, dati e dettagli concreti, mantenendo un focus laser sull'argomento.

=== OUTPUT OBBLIGATORI PER EVITARE GENERICITÀ ===
- Ogni sezione deve contenere almeno un elemento applicabile: procedura numerata, esempio concreto,
  checklist, tabella, esercizio, caso studio o criterio di verifica, in base al titolo della sezione.
- Non dichiarare soltanto che una tecnica è utile: mostra quando usarla, come applicarla e come controllare il risultato.
- Non introdurre strumenti, dati o risultati non supportati dalle fonti o dal contesto; segnala ciò che deve essere verificato.
- Rispetta i confini della sezione e non riempire spazio con ripetizioni o frasi motivazionali generiche.

=== APPROCCIO IPER-PRATICO E MICRO-DETTAGLIO ===
- OPERATIVITÀ IMMEDIATA: Spiega esattamente "COME" fare le cose. Inserisci step operativi, checklist, esempi concreti, casi studio reali o template applicativi.
- IPER-DETTAGLIO: Scendi in profondità nel micro-dettaglio. Se menzioni una tecnica, smontala nei suoi componenti base. Il lettore non deve mai chiedersi "Ok, ma in pratica come si fa?". La risposta deve essere già lì, sviscerata in ogni suo singolo passaggio logico e pratico.

=== DIVIETO ASSOLUTO DI RITRASCRIZIONE TITOLI (CRITICO) ===
- NON RITRASCRIVERE o ripetere MAI il nome del capitolo, del sottocapitolo o della sezione all'interno del testo generato o come intestazione (es. non scrivere mai "Capitolo 1" o "1.1 Introduzione" all'inizio).
- Inizia a scrivere DIRETTAMENTE il corpo del testo. L'applicazione impagina i titoli automaticamente; se tu li scrivi, verrà creato un brutto e fastidioso doppione visivo. Non usare `#` o `##` all'inizio per ripetere il titolo che ti è stato assegnato.

=== MAESTRIA LINGUISTICA E PROFONDITÀ DA LUMINARE (CRITICO) ===
- LIVELLO MADRELINGUA ASSOLUTO: Scrivi in {lingua_sel} con la naturalezza, il ritmo e la ricchezza di vocabolario di un autore locale di altissimo livello. Evita categoricamente frasi robotiche, traduzioni letterali o costrutti tipici dell'IA. Usa le sfumature linguistiche, le metafore e le espressioni idiomatiche proprie della lingua {lingua_sel}.
- COMPETENZA VERTICALE (ESPERTO DEL SETTORE): Comportati come un professionista con 30 anni di esperienza reale in questo esatto argomento. Sii chirurgico nei termini tecnici e fornisci dettagli, aneddoti o concetti avanzati che solo un vero "addetto ai lavori" conoscerebbe.
- NO SUPERFICIALITÀ: Non dare risposte generiche o banali. Ogni paragrafo deve trasudare competenza profonda, spiegando i meccanismi interni, le ragioni nascoste e i dettagli tecnici dell'argomento.
"""

    def esegui_collaudo_automatico():
        """Collauda il flusso reale: Prefazione, salvataggio, pausa e visibilità."""
        barra = st.progress(0, text="Collaudo automatico: preparazione...")
        try:
            st.session_state["admin_test_pausa_ripresa_verificata"] = False
            st.session_state["admin_test_prefazione_pausa_verificata"] = False
            st.session_state["admin_test_personalizzazione_verificata"] = False
            brief_test_personale = brief_personalizzazione_progetto()
            if "Voce o prospettiva dell'autore" not in brief_test_personale or "Priorità per il lettore" not in brief_test_personale:
                raise RuntimeError("La personalizzazione del progetto non è stata trasferita nel brief di collaudo.")
            st.session_state["admin_test_personalizzazione_verificata"] = True
            barra.progress(10, text="Collaudo automatico: ricerca e mappa delle fonti...")
            ricerca_preliminare_per_indice(
                val_titolo, val_genere, val_trama, val_goal, lingua_sel,
                f"{val_approfondimenti}\n\n{brief_personalizzazione_progetto()}", forza=True
            )
            barra.progress(30, text="Collaudo automatico: generazione e controllo dell'indice...")
            prompt_indice_test = f"""Crea esclusivamente l'indice gerarchico in {lingua_sel} per questo progetto di collaudo.

Titolo: {val_titolo}
Genere: {val_genere}
Tipologia: {val_stile}
Stile: {val_narrativa}
Punto di vista: {val_pov}
Obiettivo: {val_goal}
Risultato: {val_risultato}
Argomento: {val_trama}
Approfondimenti: {val_approfondimenti}
{brief_personalizzazione_progetto()}

VINCOLI INDEROGABILI: applica l'architettura editoriale effettiva del profilo {val_lunghezza}: {budget_struttura_indice}. Non superare {limite_voci_indice} voci d'indice. Ogni voce deve essere specifica, nuova e utile. Usa solo questo formato: Parte I:, Capitolo 1:, 1.1 Titolo. Non aggiungere spiegazioni, link, fonti, saluti o testo esterno all'indice.

Per Test Prep includi quiz o domande, simulazione e soluzioni separati. Per narrativa crea una progressione con situazione, ostacolo, scelta e risoluzione. Per manuali includi una sequenza concreta, esempio e controllo finale."""
            indice_test = genera_indice_controllato(
                prompt_indice_test,
                "Senior Book Architect: crea indici brevi, coerenti e verificabili.",
                val_genere, val_titolo, val_trama, val_goal, lingua_sel, val_stile, val_narrativa, val_pov,
                massimo_sezioni=limite_voci_indice,
                minimo_parti=minimi_struttura_indice[0],
                minimo_capitoli=minimi_struttura_indice[1],
                budget_strutturale=budget_struttura_indice,
                profilo_lunghezza=val_lunghezza,
                massimo_parti_editoriali=massimo_parti_editoriali,
                massimo_capitoli_editoriali=massimo_capitoli_editoriali,
                massimo_sottocapitoli=massimo_sottocapitoli_per_capitolo,
            )
            if not indice_test:
                raise RuntimeError(st.session_state.get("ultimo_controllo_indice", "L'indice di collaudo non ha superato il controllo."))
            imposta_indice_progetto(indice_test)
            sezioni_indice_test = list(st.session_state.get("lista_capitoli", []))
            prefazione_test = next((s for s in sezioni_indice_test if sezione_prefazione(s)), "")
            altre_sezioni = [
                sezione for sezione in sezioni_indice_test
                if not sezione_prefazione(sezione) and tipo_sezione_editoriale(sezione) != "parte"
            ]
            if val_genere == "Test Prep (Preparazione Esami)":
                altre_sezioni.sort(key=lambda sezione: 0 if re.search(r"quiz|domand|simulaz|soluz", sezione, re.I) else 1)
            if not prefazione_test or not altre_sezioni:
                raise RuntimeError("Il collaudo non ha ottenuto Prefazione e una sezione ordinaria dall'indice.")
            # Le prime due voci sono volutamente Prefazione + sezione normale:
            # testano il punto più fragile della coda senza alterare i parametri
            # reali dei due cervelli.
            sezioni_test = [prefazione_test, altre_sezioni[0]]
            for posizione, sezione in enumerate(sezioni_test, start=1):
                barra.progress(30 + int(posizione / len(sezioni_test) * 55), text=f"Collaudo automatico: stesura {posizione}/{len(sezioni_test)} — {sezione}...")
                prompt_sezione = crea_prompt_stesura_sezione(
                    sezione, indice_test, val_trama, val_genere, val_stile, val_narrativa,
                    val_pov, val_goal, lingua_sel, val_approfondimenti, val_lunghezza,
                )
                contenuto = genera_contenuto_editoriale(
                    prompt_sezione, S_PROMPT, sezione, indice_test, val_trama, val_genere,
                    val_goal, lingua_sel, val_lunghezza,
                )
                if not str(contenuto or "").strip() or str(contenuto).lstrip().upper().startswith("ERRORE:"):
                    raise RuntimeError(f"Risposta non valida per la sezione “{sezione}”.")
                # Stesso salvataggio atomico di SCRIVI TUTTO IL LIBRO.
                scrivi_sezione_stesura_completa(sezione, contenuto)
                st.session_state[CHIAVE_SELETTORE_EDITOR] = sezione
                st.session_state["job_scrittura_ultima_completata"] = sezione
                sezioni_da_proteggere = elenco_sezioni_progetto(st.session_state.get("lista_capitoli", []))
                salva_stesura_generata_in_cloud(sezioni_da_proteggere, "sezione del collaudo generata")
                if posizione == 1:
                    reidrata_sezioni_memorizzate(sezioni_da_proteggere)
                    if not str(leggi_sezione_memorizzata(sezione) or "").strip():
                        raise RuntimeError("La Prefazione non è rimasta disponibile dopo la simulazione della pausa.")
                    st.session_state["admin_test_pausa_ripresa_verificata"] = True
                    st.session_state["admin_test_prefazione_pausa_verificata"] = True
            barra.progress(100, text="Collaudo automatico completato: controlla ora le funzioni del browser.")
            st.session_state["admin_test_run_report"] = (
                f"Collaudo automatico completato con {st.session_state.get('admin_test_provider', provider_ia)}: "
                "fonti, indice, Prefazione e una sezione ordinaria sono stati prodotti e riletti con il motore reale."
            )
        except Exception as errore:
            st.session_state["admin_test_run_report"] = f"ERRORE: collaudo automatico interrotto — {errore}"
        finally:
            st.session_state.pop("admin_test_run_requested", None)

    guide_localizzate = {
        "Italiano": ("Come usare Scrittore Site", """1. Scegli prima il Cervello AI nella barra laterale. GPT-5.4 include tutte le funzioni, comprese ricerca web, verifica copyright web e immagini. DeepSeek V4 Pro usa invece un cervello distinto per ricerca delle fonti con registro visibile, indice, fonti caricate, stesura e controlli editoriali; non usa GPT. La verifica copyright web e le immagini restano disponibili solo con GPT. Poi compila titolo, autore, lingua, genere, stile, obiettivo, argomento e risultato finale. Usa Approfondimenti per priorità, vincoli ed esempi obbligatori. In Personalizza il tuo libro puoi aggiungere voce, casi, priorità e confini personali: sono facoltativi, non consumano crediti, vengono salvati nel progetto/CSV e guidano ricerca, indice e testo in modo originale.

2. Scegli Lunghezza delle sezioni: Compatto produce circa 480-560 parole per sezione, con massimo 65 sezioni totali e obiettivo di 75 pagine; Standard KDP (consigliato) circa 620-700 parole, con massimo 110 sezioni e obiettivo di 150 pagine; Approfondito circa 700-800 parole, con massimo 160 sezioni e obiettivo di 220 pagine. La Prefazione è inclusa in questi massimi. I riferimenti alle pagine si basano sul manoscritto Word 6×9 e hanno una tolleranza del 15%, perché immagini, tabelle e impaginazione possono modificarli. Il modello seleziona soltanto le fonti online più pertinenti al brief e costruisce l'indice con il minor numero di voci necessario: non aggiunge riempitivi e una struttura di qualità può essere pubblicata anche sotto la soglia. La scelta regola sia la dimensione del testo sia il tetto dell'indice. Un capitolo con sottocapitoli viene usato come breve cornice; il contenuto completo è sviluppato nei sottocapitoli, così il libro non ripete gli stessi argomenti.

3. Apri Indice e premi Genera Indice Professionale. Prima dell'indice il software cerca e studia fonti online pertinenti al brief, crea un dossier interno e lo usa per progettare la struttura; la ricerca costa 2 crediti ed è riutilizzata finché non cambi i dati della sidebar. Se carichi PDF o DOCX, vengono studiati insieme alla ricerca. Se modifichi l'indice a mano, usa Salva e Sincronizza Capitoli. Voto Indice lo valuta; Rigenera indice seguendo il voto propone una nuova versione da applicare soltanto se ti convince.

4. In Scrittura e Quiz scegli una sezione. Scrivi contenuto genera una sezione, Scrivi tutti i sottocapitoli del capitolo genera il blocco scelto e Scrivi tutto il libro completa tutte le sezioni ancora vuote dell'indice, comprese aperture, capitoli, sottocapitoli e chiusure. Pausa interrompe il lavoro prima della sezione successiva e Riprendi generazione lo continua. Se nella personalizzazione scegli una pausa guidata, il software si ferma prima di una Parte o della conclusione e ti permette di aggiungere un dettaglio contestuale; puoi anche proseguire senza scrivere nulla.

5. Rigenera con AI modifica solo la sezione scelta seguendo la tua istruzione. Quiz aggiunge domande, 10 Esempi aggiunge esempi, 10 Ricette è per i ricettari; Controlla i fatti e Report sintattico verificano qualità e leggibilità. Carica un'immagine inserisce la tua immagine in anteprima, Word e PDF.

6. In Anteprima leggi il libro e usa Controllo coerenza completo. La barra mostra l'avanzamento; il report indica capitolo, sottocapitolo, priorità e prompt da copiare in Rigenera con AI.

7. In Importa / Esporta / Copyright premi prima Controlla completezza del manoscritto: è gratuito, non usa API e segnala soltanto sezioni mancanti, troppo brevi o tecnicamente interrotte. Poi il controllo finale distingue sezioni mancanti, deboli e complete. Se rileva difetti, ricevi i prompt pronti per Rigenera con AI; il software non modifica nulla automaticamente. Dalla stessa area esporti Word, PDF o il CSV completo del progetto e puoi reimportare il CSV in seguito. In Formattazione carichi un manoscritto, crei metadati KDP e formatti un DOCX 6×9.

Notifiche sonore: sentirai il segnale quando la sidebar è pronta, quando parte o termina Scrivi tutto il libro, in caso di errore, al termine di Voto Indice, Controllo coerenza, formattazione ed esportazione. Controlla sempre testo e file finale prima di pubblicare."""),
        "English": ("How to use Scrittore Site", """1. Complete the sidebar: title, author, language, genre, style, goal, topic and desired final result. Use Further details for priorities, constraints and required examples.

2. Open Index and press Generate Professional Index. If you edit it manually, use Save and Sync Chapters. Index Score evaluates it; Regenerate index following the score creates a proposal that you apply only if you approve it.

3. In Write & Quiz choose a section. Write content creates one section, Write all subsections of the chapter creates the selected block, and Write the whole book completes empty sections. Pause stops before the next section; Resume continues.

4. Regenerate with AI changes only the selected section using your instruction. Quiz adds questions, 10 Examples adds examples, 10 Recipes is for cookbooks; Check facts and Syntax Report check quality and readability. Upload image places your image in Preview, Word and PDF.

5. In Preview, read the book and use Full consistency check. The progress bar shows completion; the report identifies chapter, subsection, priority and a prompt to copy into Regenerate with AI.

6. In Export, download Word or PDF, also as a draft. In Formatting, upload a manuscript, create KDP metadata and format a 6×9 DOCX.

Sound notifications alert you when the sidebar is ready, full-book writing starts or ends, an error occurs, and Index Score, consistency check, formatting or export finish. Always review the text and final file before publishing."""),
        "Español": ("Cómo usar Scrittore Site", """1. Completa la barra lateral: título, autor, idioma, género, estilo, objetivo, tema y resultado final deseado. Usa Profundizaciones para prioridades, límites y ejemplos obligatorios.

2. Abre Índice y pulsa Generar índice profesional. Si lo modificas manualmente, usa Guardar y sincronizar capítulos. Voto del índice lo evalúa; Regenerar índice según el voto crea una propuesta que aplicas solo si la apruebas.

3. En Escritura y cuestionarios elige una sección. Escribir contenido crea una sección; Escribir todos los subcapítulos crea el bloque elegido; Escribir todo el libro completa las secciones vacías. Pausa detiene el proceso y Reanudar lo continúa.

4. Regenerar con IA modifica solo la sección elegida. Cuestionario añade preguntas, 10 ejemplos añade ejemplos y 10 recetas es para recetarios. Verificar hechos y Reporte sintáctico revisan calidad y legibilidad. Cargar imagen la añade a Vista previa, Word y PDF.

5. En Vista previa lee el libro y usa Control completo de coherencia. La barra muestra el avance; el informe indica capítulo, subcapítulo, prioridad y un prompt para Regenerar con IA.

6. En Exportar descargas Word o PDF, también como borrador. En Formato cargas un manuscrito, creas metadatos KDP y formateas un DOCX de 6×9.

Las notificaciones sonoras avisan cuando la barra lateral está lista, al iniciar o terminar el libro completo, ante un error y al finalizar las verificaciones, el formato o la exportación. Revisa siempre el texto final antes de publicar."""),
        "Français": ("Comment utiliser Scrittore Site", """1. Remplissez la barre latérale : titre, auteur, langue, genre, style, objectif, sujet et résultat final souhaité. Utilisez Approfondissements pour les priorités, contraintes et exemples obligatoires.

2. Ouvrez Index puis cliquez sur Générer un index professionnel. Après une modification manuelle, utilisez Enregistrer et synchroniser les chapitres. Note de l’index l’évalue ; Régénérer selon la note crée une proposition à appliquer seulement si elle vous convient.

3. Dans Écriture et quiz, choisissez une section. Écrire le contenu crée une section ; Écrire tous les sous-chapitres crée le bloc choisi ; Écrire tout le livre complète les sections vides. Pause interrompt et Reprendre continue.

4. Régénérer avec l’IA modifie uniquement la section choisie. Quiz ajoute des questions, 10 exemples ajoute des exemples, 10 recettes est destiné aux livres de recettes. Vérifier les faits et Rapport syntaxique contrôlent qualité et lisibilité. Importer une image l’ajoute à l’aperçu, au Word et au PDF.

5. Dans Aperçu, lisez le livre puis lancez le contrôle complet de cohérence. La barre indique la progression ; le rapport donne chapitre, sous-chapitre, priorité et prompt à copier dans Régénérer avec l’IA.

6. Dans Exporter, téléchargez Word ou PDF, même comme brouillon. Dans Mise en forme, importez un manuscrit, créez les métadonnées KDP et formatez un DOCX 6×9.

Les notifications sonores signalent que la barre latérale est prête, le début ou la fin de l’écriture complète, une erreur et la fin des contrôles, de la mise en forme ou de l’export. Vérifiez toujours le texte final avant publication."""),
        "Deutsch": ("Scrittore Site verwenden", """1. Füllen Sie die Seitenleiste aus: Titel, Autor, Sprache, Genre, Stil, Ziel, Thema und gewünschtes Endergebnis. Nutzen Sie Vertiefungen für Prioritäten, Vorgaben und Pflichtbeispiele.

2. Öffnen Sie Index und klicken Sie auf Professionellen Index erstellen. Nach manuellen Änderungen wählen Sie Speichern und Kapitel synchronisieren. Indexbewertung prüft die Struktur; Index nach Bewertung neu erstellen erzeugt einen Vorschlag, den Sie nur bei Zustimmung übernehmen.

3. Wählen Sie unter Schreiben & Quiz einen Abschnitt. Inhalt schreiben erstellt einen Abschnitt, Alle Unterkapitel schreiben den gewählten Block und Ganzes Buch schreiben füllt leere Abschnitte. Pause hält vor dem nächsten Abschnitt an; Fortsetzen setzt fort.

4. Mit KI neu erstellen ändert nur den ausgewählten Abschnitt. Quiz fügt Fragen hinzu, 10 Beispiele fügt Beispiele hinzu und 10 Rezepte ist für Kochbücher. Fakten prüfen und Syntaxbericht prüfen Qualität und Lesbarkeit. Bild hochladen fügt Ihr Bild in Vorschau, Word und PDF ein.

5. Lesen Sie das Buch in Vorschau und starten Sie die vollständige Kohärenzprüfung. Der Balken zeigt den Fortschritt; der Bericht nennt Kapitel, Unterkapitel, Priorität und einen Prompt für Mit KI neu erstellen.

6. Unter Export laden Sie Word oder PDF auch als Entwurf herunter. Unter Formatierung laden Sie ein Manuskript hoch, erstellen KDP-Metadaten und formatieren eine DOCX-Datei im Format 6×9.

Tonbenachrichtigungen informieren über eine fertige Seitenleiste, Start oder Ende des ganzen Buchs, Fehler sowie Ende von Bewertung, Prüfung, Formatierung oder Export. Prüfen Sie Text und Enddatei immer vor der Veröffentlichung."""),
        "Română": ("Cum se folosește Scrittore Site", """1. Completează bara laterală: titlu, autor, limbă, gen, stil, obiectiv, subiect și rezultat final dorit. Folosește Aprofundări pentru priorități, condiții și exemple obligatorii.

2. Deschide Index și apasă Generare index profesional. Dacă îl modifici manual, folosește Salvează și sincronizează capitolele. Nota indexului îl evaluează; Regenerare după notă creează o propunere pe care o aplici doar dacă o aprobi.

3. În Scriere și chestionare alege o secțiune. Scrie conținut creează o secțiune, Scrie toate subcapitolele creează blocul ales, iar Scrie toată cartea completează secțiunile goale. Pauză oprește procesul; Reia îl continuă.

4. Regenerează cu IA modifică doar secțiunea aleasă. Chestionar adaugă întrebări, 10 exemple adaugă exemple, iar 10 rețete este pentru cărți de bucate. Verifică faptele și Raport sintactic verifică calitatea și lizibilitatea. Încarcă imagine o introduce în previzualizare, Word și PDF.

5. În Previzualizare citește cartea și folosește Control complet de coerență. Bara arată progresul; raportul indică capitolul, subcapitolul, prioritatea și un prompt pentru Regenerare cu IA.

6. În Export descarci Word sau PDF, inclusiv ca schiță. În Formatare încarci manuscrisul, creezi metadate KDP și formatezi un DOCX 6×9.

Notificările sonore anunță când bara laterală este gata, la începutul sau sfârșitul cărții, la erori și la finalul verificărilor, formatării sau exportului. Verifică mereu textul și fișierul final înainte de publicare."""),
        "Русский": ("Как пользоваться Scrittore Site", """1. Заполните боковую панель: название, автора, язык, жанр, стиль, цель, тему и желаемый итог. Используйте раздел «Углубления» для приоритетов, ограничений и обязательных примеров.

2. Откройте «Индекс» и нажмите создание профессионального индекса. После ручных изменений используйте сохранение и синхронизацию глав. Оценка индекса проверяет структуру; повторная генерация по оценке создаёт предложение, которое применяется только с вашего согласия.

3. В разделе «Написание и тесты» выберите секцию. Написать содержание создаёт одну секцию, написать все подразделы создаёт выбранный блок, а написать всю книгу заполняет пустые секции. Пауза останавливает работу, продолжить — возобновляет.

4. Перегенерировать с ИИ меняет только выбранную секцию. Тест добавляет вопросы, 10 примеров — примеры, 10 рецептов предназначен для кулинарных книг. Проверка фактов и синтаксический отчёт проверяют качество и читаемость. Загрузка изображения добавляет его в предпросмотр, Word и PDF.

5. В предпросмотре прочитайте книгу и запустите полную проверку согласованности. Полоса показывает прогресс; отчёт указывает главу, подраздел, приоритет и запрос для перегенерации с ИИ.

6. В экспорте скачивайте Word или PDF, в том числе черновик. В форматировании загружайте рукопись, создавайте метаданные KDP и форматируйте DOCX 6×9.

Звуковые уведомления сообщают о готовности боковой панели, начале или завершении всей книги, ошибках и окончании оценки, проверки, форматирования или экспорта. Всегда проверяйте текст и итоговый файл перед публикацией."""),
        "العربية": ("كيفية استخدام Scrittore Site", """1. املأ الشريط الجانبي: العنوان، المؤلف، اللغة، النوع، الأسلوب، الهدف، الموضوع والنتيجة النهائية المطلوبة. استخدم قسم التفاصيل الإضافية للأولويات والقيود والأمثلة المطلوبة.

2. افتح الفهرس واضغط إنشاء فهرس احترافي. بعد التعديل اليدوي استخدم حفظ ومزامنة الفصول. تقييم الفهرس يراجعه؛ إعادة إنشاء الفهرس وفق التقييم تنشئ اقتراحاً تطبقه فقط عند الموافقة.

3. في الكتابة والاختبارات اختر قسماً. كتابة المحتوى تنشئ قسماً واحداً، وكتابة كل الأقسام الفرعية تنشئ الكتلة المختارة، وكتابة الكتاب كاملاً تكمل الأقسام الفارغة. الإيقاف المؤقت يوقف العمل والاستئناف يتابعه.

4. إعادة الإنشاء بالذكاء الاصطناعي تعدل القسم المختار فقط. الاختبار يضيف أسئلة، و10 أمثلة تضيف أمثلة، و10 وصفات مخصصة لكتب الطبخ. فحص الحقائق والتقرير النحوي يفحصان الجودة وسهولة القراءة. تحميل صورة يضيفها إلى المعاينة وWord وPDF.

5. في المعاينة اقرأ الكتاب واستخدم فحص الاتساق الكامل. شريط التقدم يوضح الحالة؛ ويحدد التقرير الفصل والقسم الفرعي والأولوية وموجهاً جاهزاً لإعادة الإنشاء بالذكاء الاصطناعي.

6. في التصدير نزّل Word أو PDF حتى كمسودة. في التنسيق حمّل مخطوطة، وأنشئ بيانات KDP الوصفية ونسّق DOCX بمقاس 6×9.

تنبّهك الإشعارات الصوتية عند جاهزية الشريط الجانبي، وبدء أو انتهاء كتابة الكتاب، وعند الخطأ، وانتهاء التقييم أو الفحص أو التنسيق أو التصدير. راجع النص والملف النهائي دائماً قبل النشر."""),
        "中文": ("如何使用 Scrittore Site", """1. 填写侧边栏：书名、作者、语言、类型、写作风格、目标、主题和期望最终成果。使用“补充说明”填写重点、限制和必须包含的示例。

2. 打开“目录”，点击“生成专业目录”。手动修改后，使用“保存并同步章节”。“目录评分”会进行评估；“按评分重新生成目录”会创建一个方案，只有在你认可后才应用。

3. 在“写作与测验”中选择一个部分。“写内容”生成一个部分；“写本章所有小节”生成所选章节；“写整本书”补全空白部分。暂停会在下一部分前停止，继续会恢复生成。

4. “用 AI 重新生成”只修改所选部分。测验会添加题目，10 个示例会添加示例，10 个食谱用于食谱书。事实检查和句法报告检查质量与可读性。上传图片会将你的图片加入预览、Word 和 PDF。

5. 在“预览”中阅读图书，并使用“完整一致性检查”。进度条显示完成情况；报告会列出章节、小节、优先级，以及可复制到“用 AI 重新生成”的提示词。

6. 在“导出”中下载 Word 或 PDF，也可以导出草稿。在“格式化”中上传手稿、生成 KDP 元数据并格式化 6×9 DOCX。

声音通知会在侧边栏准备好、整本书开始或结束、发生错误以及目录评分、一致性检查、格式化或导出完成时提醒你。发布前务必检查文本和最终文件。""")
    }
    aggiornamenti_guida_localizzati = {
        "Italiano": """### Funzioni aggiornate: salvataggio, fonti e copyright

**Salvare e recuperare.** Il progetto resta pulito quando accedi. Quando vuoi conservarlo, premi **💾 SALVA SESSIONE** nella sidebar: vengono salvati campi della sidebar, indice, testi, fonti, immagini e impostazioni del progetto nel tuo account. Per recuperare deliberatamente il lavoro salvato premi **🔄 RIAGGIORNA ALL'ULTIMA STESURA**. **RESET PROGETTO** cancella invece il progetto aperto e svuota anche la sidebar.

**Importa / Esporta / Copyright.** Esporta progetto completo crea un CSV che contiene sidebar, indice, sezioni, fonti e immagini associate. Importa progetto CSV ripristina questi dati nella pagina; premi poi **SALVA SESSIONE** se vuoi conservarli anche nel tuo account. Word e PDF restano disponibili nella stessa area dopo il controllo finale.

**Fonti e testo originale.** Prima dell'indice, la ricerca preliminare raccoglie fonti pertinenti e le trasforma in una mappa concettuale interna: il libro viene scritto con struttura, esempi e formulazioni nuovi, non copiando le fonti. Le fonti caricate e quelle web sono consultabili nell'area Importa / Esporta / Copyright. Per argomenti aggiornabili verifica sempre l'accuratezza prima di pubblicare.

**Controllo originalità e copyright.** Il controllo locale confronta gratuitamente il testo con PDF/DOCX caricati; il controllo web rapido analizza campioni; quello completo controlla tutto il manoscritto a lotti. L'esito mostra subito solo le sezioni da rielaborare. Apri **Dettaglio tecnico** soltanto se vuoi leggere l'intero report. **RIELABORA LE SEZIONI SEGNALATE** riscrive solo quelle sezioni, senza toccare le altre. È uno screening editoriale di rischio, non una certificazione legale o antiplagio universale.

**Editor, anteprima e ascolto.** L'Editor professionale modifica una sezione alla volta; l'anteprima legge la memoria completa del progetto. Nel lettore vocale puoi scegliere il punto di partenza, avviare, mettere in pausa e riprendere; nella vista lettura il testo letto viene evidenziato quando il browser lo consente.

**Controllo completezza e finale.** Prima del controllo finale puoi usare **Controlla completezza del manoscritto**: è gratuito, non usa API e rileva testi mancanti, troppo brevi o tecnicamente interrotti senza modificare il libro. Prima di Word/PDF, il controllo finale segnala poi contenuti mancanti, deboli o completi e propone correzioni da applicare soltanto con la tua approvazione. Rileggi sempre il manoscritto e il file esportato prima della pubblicazione.""",
        "English": """### Updated features: saving, sources and copyright

**Save and recover.** The project opens clean after sign-in. Use **💾 SAVE SESSION** to store sidebar fields, index, texts, sources, images and project settings in your account. Use **🔄 REFRESH TO LATEST DRAFT** only when you want to restore that saved work. **RESET PROJECT** clears the open project and the sidebar.

**Import / Export / Copyright.** Complete project export creates a CSV containing the sidebar, index, sections, sources and linked images. Importing it restores those data in the page; then use **SAVE SESSION** to keep them in your account. Word and PDF remain available there after the final check.

**Sources and original writing.** Before the index, preliminary research turns relevant sources into an internal concept map. The book is written with new structure, examples and wording, not copied from sources. The local, quick web and complete web checks show only sections needing revision; technical details are optional. Rewriting affects only flagged sections. These are editorial risk checks, not legal certification.

**Editor, preview and listening.** The Professional Editor edits one section at a time; Preview reads the full project memory. In the voice reader choose a starting point, play, pause and resume. The reading view highlights spoken text when the browser supports it. Always review the manuscript and exported file before publishing.""",
        "Español": """### Funciones actualizadas: guardado, fuentes y copyright

**Guardar y recuperar.** El proyecto se abre limpio al iniciar sesión. Usa **💾 GUARDAR SESIÓN** para guardar en tu cuenta los campos, índice, textos, fuentes, imágenes y ajustes. Usa **🔄 ACTUALIZAR A LA ÚLTIMA VERSIÓN** solo para restaurar ese trabajo. **RESETEAR PROYECTO** borra el proyecto abierto y la barra lateral.

**Importar / Exportar / Copyright.** La exportación completa crea un CSV con barra lateral, índice, secciones, fuentes e imágenes. Al importarlo se restaura en la página; después pulsa **GUARDAR SESIÓN** para conservarlo en tu cuenta. Los controles de originalidad muestran solo las secciones que requieren cambios; el detalle técnico es opcional. No sustituye una certificación legal.

**Fuentes, edición y escucha.** La investigación previa crea un mapa conceptual interno; el libro se redacta con estructura y formulación nuevas, no copiando fuentes. El editor modifica una sección, la vista previa lee todo el proyecto y el lector de voz permite elegir inicio, reproducir, pausar y reanudar. Revisa siempre el archivo antes de publicar.""",
        "Français": """### Fonctions mises à jour : sauvegarde, sources et copyright

**Sauvegarder et récupérer.** Le projet s’ouvre vide après connexion. Utilisez **💾 SAUVEGARDER LA SESSION** pour enregistrer dans votre compte les champs, l’index, les textes, les sources, les images et les réglages. Utilisez **🔄 ACTUALISER LA DERNIÈRE VERSION** uniquement pour restaurer ce travail. **RÉINITIALISER LE PROJET** efface le projet ouvert et la barre latérale.

**Importer / Exporter / Copyright.** L’export complet crée un CSV avec barre latérale, index, sections, sources et images. L’import le restaure dans la page ; sauvegardez ensuite la session pour le conserver. Les contrôles d’originalité affichent seulement les sections à corriger ; le détail technique est facultatif. Ce n’est pas une certification juridique.

**Sources, édition et écoute.** La recherche prépare une carte conceptuelle interne ; le livre est écrit avec une structure et des formulations nouvelles. L’éditeur modifie une section, l’aperçu lit tout le projet et le lecteur vocal permet de choisir le départ, lire, mettre en pause et reprendre. Vérifiez toujours le fichier avant publication.""",
        "Deutsch": """### Aktualisierte Funktionen: Speichern, Quellen und Urheberrecht

**Speichern und Wiederherstellen.** Nach der Anmeldung öffnet sich ein leeres Projekt. Mit **💾 SITZUNG SPEICHERN** sichern Sie Seitenleiste, Index, Texte, Quellen, Bilder und Einstellungen im Konto. **🔄 LETZTEN ENTWURF AKTUALISIEREN** stellt diesen Stand bewusst wieder her. **PROJEKT ZURÜCKSETZEN** leert Projekt und Seitenleiste.

**Import / Export / Copyright.** Der vollständige Export erzeugt eine CSV mit allen Projektdaten. Der Import stellt sie auf der Seite wieder her; speichern Sie danach die Sitzung. Originalitätsprüfungen zeigen direkt nur zu überarbeitende Abschnitte, technische Details sind optional. Sie ersetzen keine rechtliche Zertifizierung.

**Quellen, Editor und Vorlesen.** Die Recherche erstellt eine interne Konzeptkarte; der Text wird neu strukturiert und formuliert. Der Editor bearbeitet einen Abschnitt, die Vorschau liest das gesamte Projekt, und der Sprachleser kann Startpunkt, Wiedergabe, Pause und Fortsetzen steuern. Prüfen Sie das Enddokument immer vor der Veröffentlichung.""",
        "Română": """### Funcții actualizate: salvare, surse și copyright

**Salvare și recuperare.** Proiectul se deschide gol după autentificare. Folosește **💾 SALVEAZĂ SESIUNEA** pentru câmpuri, cuprins, texte, surse, imagini și setări. **🔄 ACTUALIZEAZĂ LA ULTIMA VERSIUNE** restaurează intenționat versiunea salvată. **RESETARE PROIECT** golește proiectul și bara laterală.

**Import / Export / Copyright.** Exportul complet creează un CSV cu toate datele proiectului; importul le restaurează în pagină, apoi salvează sesiunea. Controalele de originalitate arată doar secțiunile de corectat, iar detaliul tehnic este opțional; nu reprezintă o certificare juridică.

**Surse, editor și citire vocală.** Cercetarea creează o hartă conceptuală internă, iar textul este formulat original. Editorul modifică o secțiune, previzualizarea citește proiectul integral, iar cititorul vocal permite alegerea începutului, redare, pauză și reluare. Verifică mereu fișierul final înainte de publicare.""",
        "Русский": """### Обновлённые функции: сохранение, источники и авторское право

**Сохранение и восстановление.** После входа проект открывается чистым. Кнопка **💾 СОХРАНИТЬ СЕССИЮ** сохраняет поля, оглавление, тексты, источники, изображения и настройки в аккаунте. **🔄 ОБНОВИТЬ ПОСЛЕДНЮЮ ВЕРСИЮ** намеренно восстанавливает сохранённую работу. Сброс очищает проект и боковую панель.

**Импорт / экспорт / copyright.** Полный экспорт создаёт CSV со всеми данными проекта; импорт восстанавливает их на странице, после чего сохраните сессию. Проверки оригинальности сразу показывают только разделы для исправления, а технические подробности открываются по желанию. Это не юридическая сертификация.

**Источники, редактор и озвучивание.** Исследование создаёт внутреннюю карту понятий, а текст формулируется заново. Редактор меняет один раздел, предпросмотр читает весь проект, голосовой читатель позволяет выбрать начало, воспроизвести, поставить на паузу и продолжить. Всегда проверяйте итоговый файл перед публикацией.""",
        "العربية": """### ميزات محدّثة: الحفظ والمصادر وحقوق النشر

**الحفظ والاستعادة.** يفتح المشروع فارغاً بعد تسجيل الدخول. استخدم **💾 حفظ الجلسة** لحفظ حقول الشريط الجانبي والفهرس والنصوص والمصادر والصور والإعدادات في حسابك. استخدم **🔄 تحديث إلى آخر نسخة** فقط لاستعادة العمل المحفوظ. إعادة ضبط المشروع تمسح المشروع والشريط الجانبي.

**استيراد / تصدير / حقوق النشر.** ينشئ التصدير الكامل ملف CSV يحوي كل بيانات المشروع، ويعيد الاستيراد هذه البيانات إلى الصفحة ثم احفظ الجلسة. تعرض فحوص الأصالة الأقسام التي تحتاج تعديلاً فقط، بينما التفاصيل التقنية اختيارية. لا تمثل هذه الفحوص شهادة قانونية.

**المصادر والتحرير والقراءة.** ينشئ البحث خريطة مفاهيم داخلية ويُكتب النص بصياغة جديدة. يعدّل المحرر قسماً واحداً، وتقرأ المعاينة المشروع كاملاً، ويتيح القارئ الصوتي اختيار نقطة البداية والتشغيل والإيقاف والاستئناف. راجع الملف النهائي دائماً قبل النشر.""",
        "中文": """### 更新功能：保存、来源与版权

**保存和恢复。** 登录后项目会以干净状态打开。使用 **💾 保存会话** 将侧边栏、目录、文本、来源、图片和设置保存到账号。仅在需要恢复时使用 **🔄 更新到最新草稿**。重置项目会清空当前项目和侧边栏。

**导入 / 导出 / 版权。** 完整导出会创建包含所有项目数据的 CSV；导入会在页面中恢复数据，随后请保存会话。原创性检查只显示需要修改的部分，技术详情可按需展开；它不是法律认证。

**来源、编辑与朗读。** 前期研究会建立内部概念图，正文会以新的结构和表述创作。专业编辑器一次编辑一个部分，预览读取整个项目，语音阅读器可选择起点、播放、暂停和继续。发布前请始终检查最终文件。""",
    }
    avviso_salvataggio_ia = {
        "Italiano": "🛡️ **Protezione del lavoro pagato:** ogni contenuto creato o rielaborato con l’IA viene salvato automaticamente nel tuo account dopo una generazione riuscita. **SALVA SESSIONE** resta utile per conservare anche modifiche manuali, immagini o altre modifiche fatte senza IA.",
        "English": "🛡️ **Protection for paid work:** every AI-created or AI-reworked item is automatically saved to your account after a successful generation. **SAVE SESSION** remains useful for manual edits, images and changes made without AI.",
        "Español": "🛡️ **Protección del trabajo pagado:** cada contenido creado o reelaborado con IA se guarda automáticamente en tu cuenta tras una generación correcta. **GUARDAR SESIÓN** sigue siendo útil para cambios manuales, imágenes y modificaciones sin IA.",
        "Français": "🛡️ **Protection du travail payé :** chaque contenu créé ou réécrit avec l’IA est enregistré automatiquement dans votre compte après une génération réussie. **SAUVEGARDER LA SESSION** reste utile pour les modifications manuelles, images et changements sans IA.",
        "Deutsch": "🛡️ **Schutz bezahlter Arbeit:** Jeder mit KI erstellte oder überarbeitete Inhalt wird nach erfolgreicher Generierung automatisch im Konto gespeichert. **SITZUNG SPEICHERN** bleibt für manuelle Änderungen, Bilder und Änderungen ohne KI wichtig.",
        "Română": "🛡️ **Protecția muncii plătite:** fiecare conținut creat sau rescris cu IA se salvează automat în cont după o generare reușită. **SALVEAZĂ SESIUNEA** rămâne util pentru modificări manuale, imagini și schimbări fără IA.",
        "Русский": "🛡️ **Защита оплаченной работы:** каждый созданный или переработанный ИИ материал автоматически сохраняется в аккаунте после успешной генерации. **СОХРАНИТЬ СЕССИЮ** по-прежнему нужен для ручных правок, изображений и изменений без ИИ.",
        "العربية": "🛡️ **حماية العمل المدفوع:** يُحفظ كل محتوى تم إنشاؤه أو إعادة صياغته بالذكاء الاصطناعي تلقائياً في حسابك بعد نجاح العملية. يظل **حفظ الجلسة** مفيداً للتعديلات اليدوية والصور والتغييرات دون الذكاء الاصطناعي.",
        "中文": "🛡️ **已付费内容保护：** 每次 AI 成功生成或改写的内容都会自动保存到你的账户。**保存会话** 仍适用于手动修改、图片和未使用 AI 的其他变更。",
    }
    avviso_cervelli_ia = {
        "Italiano": "🧠 **Due cervelli separati.** **GPT-5.4** include tutte le funzioni e il tariffario GPT. **DeepSeek V4 Pro** svolge ricerca delle fonti con registro visibile, indice, analisi delle fonti caricate, scrittura e controlli editoriali senza usare GPT: scrittura, quiz ed esempi costano 1 credito ogni 3 operazioni; indice circa 2 crediti; coerenza e dieci ricette circa 4. Il controllo copyright web e le immagini richiedono GPT.",
        "English": "🧠 **Two separate AI engines.** **GPT-5.4** includes every feature and the GPT pricing. **DeepSeek V4 Pro** performs source research with a visible register, outlines, uploaded-source analysis, writing and editorial checks without using GPT: writing, quizzes and examples cost 1 credit every 3 operations; an outline is about 2 credits; consistency and ten recipes are about 4. Web copyright checks and images require GPT.",
        "Español": "🧠 **Dos motores de IA separados.** **GPT-5.4** incluye todas las funciones y tarifas GPT. **DeepSeek V4 Pro** realiza investigación de fuentes con registro visible, índice, análisis de fuentes cargadas, escritura y controles editoriales sin usar GPT: escritura, cuestionarios y ejemplos cuestan 1 crédito cada 3 operaciones; el índice cuesta unos 2 créditos; coherencia y diez recetas, unos 4. El control de copyright web y las imágenes requieren GPT.",
        "Français": "🧠 **Deux moteurs IA séparés.** **GPT-5.4** inclut toutes les fonctions et le tarif GPT. **DeepSeek V4 Pro** effectue la recherche de sources avec registre visible, le plan, l’analyse des sources importées, la rédaction et les contrôles éditoriaux sans GPT : rédaction, quiz et exemples coûtent 1 crédit toutes les 3 opérations ; le plan coûte environ 2 crédits ; cohérence et dix recettes, environ 4. Le contrôle de copyright web et les images nécessitent GPT.",
        "Deutsch": "🧠 **Zwei getrennte KI-Engines.** **GPT-5.4** umfasst alle Funktionen und GPT-Tarife. **DeepSeek V4 Pro** recherchiert Quellen mit sichtbarem Register, erstellt die Gliederung, analysiert hochgeladene Quellen, schreibt und prüft redaktionell ohne GPT: Schreiben, Quiz und Beispiele kosten 1 Kredit je 3 Vorgänge; die Gliederung etwa 2 Kredite; Kohärenz und zehn Rezepte etwa 4. Web-Copyrightprüfung und Bilder erfordern GPT.",
        "Română": "🧠 **Două motoare IA separate.** **GPT-5.4** include toate funcțiile și tarifele GPT. **DeepSeek V4 Pro** face cercetare de surse cu registru vizibil, cuprins, analiză a surselor încărcate, scriere și controale editoriale fără GPT: scrierea, testele și exemplele costă 1 credit la 3 operațiuni; cuprinsul aproximativ 2 credite; coerența și zece rețete aproximativ 4. Controlul copyright web și imaginile necesită GPT.",
        "Русский": "🧠 **Два отдельных ИИ-движка.** **GPT-5.4** включает все функции и тариф GPT. **DeepSeek V4 Pro** ищет источники с видимым реестром, создаёт оглавление, анализирует загруженные источники, пишет и выполняет редакторские проверки без GPT: текст, тесты и примеры стоят 1 кредит за 3 операции; оглавление — около 2 кредитов; согласованность и десять рецептов — около 4. Проверка copyright в интернете и изображения требуют GPT.",
        "العربية": "🧠 **محركان منفصلان للذكاء الاصطناعي.** يشمل **GPT-5.4** كل الوظائف وتسعيرة GPT. ينفذ **DeepSeek V4 Pro** بحث المصادر مع سجل ظاهر والفهرس وتحليل المصادر المرفوعة والكتابة والفحوص التحريرية دون GPT: الكتابة والاختبارات والأمثلة تكلف رصيداً واحداً لكل 3 عمليات؛ الفهرس نحو رصيدين؛ الاتساق وعشر وصفات نحو 4. فحص حقوق النشر على الويب والصور يتطلبان GPT.",
        "中文": "🧠 **两个独立的 AI 引擎。** **GPT-5.4** 包含所有功能及 GPT 计费。**DeepSeek V4 Pro** 不使用 GPT 即可进行带可见记录的来源研究、目录、已上传来源分析、写作和编辑检查：写作/测验/示例每 3 次操作消耗 1 积分；目录约 2 积分；一致性检查和 10 道食谱约 4 积分。网页版权检查和图片需要 GPT。",
    }
    percorso_rapido = {
        "Italiano": """### ⚡ Percorso rapido in 4 passaggi
**1. Configura** — compila la sidebar e scegli il Cervello AI.  
**2. Progetta** — cerca le fonti se servono e genera l'indice; controllalo prima di scrivere.  
**3. Scrivi e salva** — genera una sezione o l'intero libro; le sezioni AI sono protette, mentre le modifiche manuali richiedono **SALVA SESSIONE**.  
**4. Controlla ed esporta** — usa Anteprima, completezza, controllo finale e poi Word, PDF o CSV.""",
        "English": """### ⚡ Quick 4-step path
**1. Configure** — complete the sidebar and choose the AI engine.  
**2. Plan** — search sources when needed and generate the index; review it before writing.  
**3. Write and save** — generate one section or the full book; AI sections are protected, while manual edits require **SAVE SESSION**.  
**4. Review and export** — use Preview, completeness check, final review, then Word, PDF or CSV.""",
        "Español": """### ⚡ Ruta rápida en 4 pasos
**1. Configura** — completa la barra lateral y elige el motor de IA.  
**2. Planifica** — busca fuentes si hace falta y genera el índice; revísalo antes de escribir.  
**3. Escribe y guarda** — genera una sección o todo el libro; las secciones IA están protegidas y los cambios manuales requieren **GUARDAR SESIÓN**.  
**4. Revisa y exporta** — usa Vista previa, control de completitud, control final y luego Word, PDF o CSV.""",
        "Français": """### ⚡ Parcours rapide en 4 étapes
**1. Configurez** — remplissez la barre latérale et choisissez le moteur IA.  
**2. Planifiez** — recherchez des sources si nécessaire et générez l’index; vérifiez-le avant la rédaction.  
**3. Rédigez et enregistrez** — générez une section ou le livre entier; les sections IA sont protégées, les modifications manuelles demandent **ENREGISTRER LA SESSION**.  
**4. Vérifiez et exportez** — utilisez l’aperçu, le contrôle de complétude, le contrôle final puis Word, PDF ou CSV.""",
        "Deutsch": """### ⚡ Schneller Ablauf in 4 Schritten
**1. Konfigurieren** — Seitenleiste ausfüllen und KI-Engine wählen.  
**2. Planen** — bei Bedarf Quellen suchen und Inhaltsverzeichnis erzeugen; vor dem Schreiben prüfen.  
**3. Schreiben und speichern** — einen Abschnitt oder das ganze Buch erzeugen; KI-Abschnitte sind geschützt, manuelle Änderungen brauchen **SITZUNG SPEICHERN**.  
**4. Prüfen und exportieren** — Vorschau, Vollständigkeitsprüfung und Endkontrolle nutzen, dann Word, PDF oder CSV.""",
        "Română": """### ⚡ Parcurs rapid în 4 pași
**1. Configurează** — completează bara laterală și alege motorul IA.  
**2. Planifică** — caută surse când este nevoie și generează cuprinsul; verifică-l înainte de scriere.  
**3. Scrie și salvează** — generează o secțiune sau cartea completă; secțiunile IA sunt protejate, iar modificările manuale cer **SALVEAZĂ SESIUNEA**.  
**4. Verifică și exportă** — folosește previzualizarea, controlul de completitudine și controlul final, apoi Word, PDF sau CSV.""",
        "Русский": """### ⚡ Быстрый путь из 4 шагов
**1. Настройте** — заполните боковую панель и выберите ИИ-движок.  
**2. Спланируйте** — при необходимости найдите источники и создайте оглавление; проверьте его перед написанием.  
**3. Пишите и сохраняйте** — создайте раздел или всю книгу; ИИ-разделы защищены, ручные правки требуют **СОХРАНИТЬ СЕССИЮ**.  
**4. Проверьте и экспортируйте** — используйте предпросмотр, проверку полноты, финальный контроль, затем Word, PDF или CSV.""",
        "العربية": """### ⚡ مسار سريع من 4 خطوات
**1. الإعداد** — أكمل الشريط الجانبي واختر محرك الذكاء الاصطناعي.  
**2. التخطيط** — ابحث عن المصادر عند الحاجة وأنشئ الفهرس ثم راجعه قبل الكتابة.  
**3. الكتابة والحفظ** — أنشئ قسماً أو الكتاب كاملاً؛ أقسام الذكاء الاصطناعي محمية بينما تتطلب التعديلات اليدوية **حفظ الجلسة**.  
**4. المراجعة والتصدير** — استخدم المعاينة وفحص الاكتمال والتحكم النهائي ثم Word أو PDF أو CSV.""",
        "中文": """### ⚡ 四步快速流程
**1. 配置** — 填写侧边栏并选择 AI 引擎。  
**2. 规划** — 需要时搜索来源并生成目录；写作前先检查目录。  
**3. 写作并保存** — 生成单个部分或整本书；AI 生成内容受保护，手动修改需要点击 **保存会话**。  
**4. 检查并导出** — 使用预览、完整性检查和最终检查，然后导出 Word、PDF 或 CSV。""",
    }
    titolo_guida, testo_guida = guide_localizzate.get(lingua_sel, guide_localizzate["Italiano"])
    if st.session_state.get("admin_test_run_requested") and st.session_state.get("admin_test_mode"):
        esegui_collaudo_automatico()
        st.rerun()

    tabs = st.tabs([f"📘 0. {titolo_guida}"] + L["tabs"] + [testo_ui("formattazione", lingua_sel)])

    # Una correzione preparata dal controllo finale o una stesura in corso
    # deve arrivare davvero all'editor. Streamlit non espone un'API Python per
    # attivare una tab: il piccolo script seleziona la terza scheda, che è
    # sempre Scrittura & Quiz in tutte le lingue. Durante una pausa il click
    # viene eseguito una volta e l'utente può poi consultare liberamente le
    # altre sezioni dell'app.
    apri_tab_scrittura = (
        st.session_state.pop("apri_tab_scrittura_da_correzione", False)
        or st.session_state.pop("apri_tab_scrittura_da_stesura", False)
        or bool(st.session_state.get("job_scrittura_attivo"))
        or bool(st.session_state.pop("apri_tab_scrittura_da_pausa", False))
    )
    if apri_tab_scrittura:
        components.html(
            """
            <script>
            setTimeout(function () {
              try {
                const schede = Array.from(window.parent.document.querySelectorAll('[role="tab"]'));
                const scrittura = schede[2] || schede.find(function (scheda) {
                  const testo = (scheda.innerText || '').toLowerCase();
                  return testo.includes('scrittura') || testo.includes('write');
                });
                if (scrittura) scrittura.click();
              } catch (errore) { /* il messaggio di conferma resta il fallback */ }
            }, 250);
            </script>
            """,
            height=0,
        )

    with tabs[0]:
        st.subheader(titolo_guida)
        st.markdown(percorso_rapido.get(lingua_sel, percorso_rapido["Italiano"]))
        chat_subito = {
            "Italiano": ("🤖 Preferisci essere guidato?", "La chat guidata prepara con te i campi della sidebar. È facoltativa: puoi sempre compilare tutto anche da solo.", "💬 APRI LA CHAT GUIDATA"),
            "English": ("🤖 Would you like guidance?", "Guided chat helps you prepare the sidebar fields. It is optional: you can always complete everything yourself.", "💬 OPEN GUIDED CHAT"),
            "Español": ("🤖 ¿Prefieres recibir ayuda?", "El chat guiado prepara contigo los campos de la barra lateral. Es opcional: siempre puedes completarlos por tu cuenta.", "💬 ABRIR CHAT GUIADO"),
            "Français": ("🤖 Vous préférez être guidé ?", "Le chat guidé prépare avec vous les champs de la barre latérale. Il est facultatif : vous pouvez toujours tout remplir vous-même.", "💬 OUVRIR LE CHAT GUIDÉ"),
            "Deutsch": ("🤖 Möchten Sie Unterstützung?", "Der geführte Chat bereitet mit Ihnen die Felder der Seitenleiste vor. Er ist optional: Sie können alles selbst ausfüllen.", "💬 GEFÜHRTEN CHAT ÖFFNEN"),
            "Română": ("🤖 Preferi să fii ghidat?", "Chatul ghidat pregătește cu tine câmpurile barei laterale. Este opțional: le poți completa oricând și singur.", "💬 DESCHIDE CHATUL GHIDAT"),
            "Русский": ("🤖 Хотите помощь?", "Управляемый чат поможет подготовить поля боковой панели. Это необязательно: всё можно заполнить самостоятельно.", "💬 ОТКРЫТЬ УПРАВЛЯЕМЫЙ ЧАТ"),
            "العربية": ("🤖 هل تفضّل المساعدة؟", "تساعدك الدردشة الموجّهة في إعداد حقول الشريط الجانبي. هي اختيارية: يمكنك دائماً إكمالها بنفسك.", "💬 افتح الدردشة الموجّهة"),
            "中文": ("🤖 想要引导帮助吗？", "引导式聊天会协助你准备侧边栏字段。这是可选功能：你随时可以自行填写。", "💬 打开引导式聊天"),
        }
        titolo_chat_subito, testo_chat_subito, pulsante_chat_subito = chat_subito.get(
            lingua_sel, chat_subito["Italiano"]
        )
        st.success(f"**{titolo_chat_subito}**\n\n{testo_chat_subito}")
        st.button(
            pulsante_chat_subito,
            key="avvia_chat_sidebar_subito",
            on_click=avvia_chat_sidebar_da_guida,
            use_container_width=True,
        )
        guida_personalizzazione = {
            "Italiano": "✍️ **Personalizza il tuo libro** è facoltativo: conserva voce, casi, priorità e confini nel progetto. Puoi anche scegliere una pausa guidata prima delle Parti o della conclusione; la risposta viene aggiunta al brief senza consumare crediti.",
            "English": "✍️ **Personalize your book** is optional: it keeps your voice, cases, priorities and boundaries with the project. You may also choose a guided pause before Parts or the conclusion; your answer is added to the brief without using credits.",
            "Español": "✍️ **Personaliza tu libro** es opcional: conserva voz, casos, prioridades y límites con el proyecto. También puedes elegir una pausa guiada antes de las Partes o de la conclusión; la respuesta se añade al brief sin consumir créditos.",
            "Français": "✍️ **Personnalisez votre livre** est facultatif : voix, cas, priorités et limites sont conservés dans le projet. Vous pouvez aussi choisir une pause guidée avant les parties ou la conclusion, sans crédit consommé.",
            "Deutsch": "✍️ **Personalisiere dein Buch** ist optional: Stimme, Fälle, Prioritäten und Grenzen werden mit dem Projekt gespeichert. Eine geführte Pause vor Teilen oder Schluss fügt Hinweise ohne Credits zum Brief hinzu.",
            "Română": "✍️ **Personalizează cartea** este opțional: vocea, cazurile, prioritățile și limitele se păstrează în proiect. Poți alege și o pauză ghidată înainte de părți sau concluzie, fără credite.",
            "Русский": "✍️ **Персонализируйте книгу** — необязательная функция: голос, примеры, приоритеты и ограничения сохраняются в проекте. Можно выбрать паузу перед частями или заключением без расхода кредитов.",
            "العربية": "✍️ **خصص كتابك** ميزة اختيارية تحفظ الصوت والأمثلة والأولويات والحدود داخل المشروع. يمكنك أيضاً اختيار توقف موجّه قبل الأجزاء أو الخاتمة من دون استهلاك أرصدة.",
            "中文": "✍️ **个性化你的图书**是可选功能：写作声音、案例、重点和边界会随项目保存。也可在各部分或结论前选择引导暂停，不消耗积分。",
        }
        st.info(guida_personalizzazione.get(lingua_sel, guida_personalizzazione["Italiano"]))
        avvisi_guida_base = {
            "Italiano": ("Completa la barra laterale, poi crea l'indice e inizia a scrivere. Le istruzioni complete restano disponibili qui sotto.", "📚 Guida completa e regole di utilizzo", "✨ Funzioni, salvataggio e copyright"),
            "English": ("Complete the sidebar, then create the index and start writing. The full instructions remain available below.", "📚 Complete guide and rules", "✨ Features, saving and copyright"),
            "Español": ("Completa la barra lateral, crea el índice y empieza a escribir. Las instrucciones completas están disponibles abajo.", "📚 Guía completa y reglas de uso", "✨ Funciones, guardado y copyright"),
            "Français": ("Remplissez la barre latérale, puis créez l'index et commencez à écrire. Les instructions complètes restent disponibles ci-dessous.", "📚 Guide complet et règles d'utilisation", "✨ Fonctions, enregistrement et copyright"),
            "Deutsch": ("Füllen Sie die Seitenleiste aus, erstellen Sie dann den Index und beginnen Sie zu schreiben. Die vollständige Anleitung steht unten bereit.", "📚 Vollständige Anleitung und Regeln", "✨ Funktionen, Speichern und Copyright"),
            "Română": ("Completează bara laterală, apoi generează cuprinsul și începe să scrii. Instrucțiunile complete rămân disponibile mai jos.", "📚 Ghid complet și reguli de utilizare", "✨ Funcții, salvare și copyright"),
            "Русский": ("Заполните боковую панель, затем создайте оглавление и начните писать. Полные инструкции доступны ниже.", "📚 Полное руководство и правила", "✨ Функции, сохранение и copyright"),
            "العربية": ("أكمل الشريط الجانبي ثم أنشئ الفهرس وابدأ الكتابة. التعليمات الكاملة متاحة أدناه.", "📚 الدليل الكامل وقواعد الاستخدام", "✨ الميزات والحفظ وحقوق النشر"),
            "中文": ("请先完成侧边栏，然后生成目录并开始写作。完整说明在下方提供。", "📚 完整指南和使用规则", "✨ 功能、保存和版权"),
        }
        avviso_guida, etichetta_guida_completa, etichetta_funzioni = avvisi_guida_base.get(lingua_sel, avvisi_guida_base["Italiano"])
        if not sidebar_pronta:
            st.info(avviso_guida)
        with st.expander(etichetta_guida_completa, expanded=False):
            st.markdown(testo_guida)
        with st.expander(etichetta_funzioni, expanded=False):
            st.markdown(aggiornamenti_guida_localizzati.get(lingua_sel, aggiornamenti_guida_localizzati["Italiano"]))
        st.info(avviso_cervelli_ia.get(lingua_sel, avviso_cervelli_ia["Italiano"]))
        st.info(avviso_salvataggio_ia.get(lingua_sel, avviso_salvataggio_ia["Italiano"]))
        etichette_prova_notifiche = {
            "Italiano": "🔔 PROVA NOTIFICHE", "English": "🔔 TEST NOTIFICATIONS", "Español": "🔔 PROBAR NOTIFICACIONES",
            "Français": "🔔 TESTER LES NOTIFICATIONS", "Deutsch": "🔔 BENACHRICHTIGUNGEN TESTEN", "Română": "🔔 TESTEAZĂ NOTIFICĂRILE",
            "Русский": "🔔 ПРОВЕРИТЬ УВЕДОМЛЕНИЯ", "العربية": "🔔 اختبار الإشعارات", "中文": "🔔 测试通知"
        }
        istruzioni_audio = {
            "Italiano": "Se non senti il suono, controlla che la scheda del browser non sia silenziata e consenti la riproduzione audio.",
            "English": "If you do not hear the sound, check that the browser tab is not muted and allow audio playback.",
            "Español": "Si no oyes el sonido, comprueba que la pestaña no esté silenciada y permite la reproducción de audio.",
            "Français": "Si vous n’entendez pas le son, vérifiez que l’onglet n’est pas muet et autorisez la lecture audio.",
            "Deutsch": "Wenn Sie keinen Ton hören, prüfen Sie, ob der Browser-Tab stummgeschaltet ist, und erlauben Sie die Audiowiedergabe.",
            "Română": "Dacă nu auzi sunetul, verifică dacă fila browserului nu este pe mut și permite redarea audio.",
            "Русский": "Если звука нет, проверьте, не отключён ли звук у вкладки браузера, и разрешите воспроизведение аудио.",
            "العربية": "إذا لم تسمع الصوت، تحقق من أن علامة تبويب المتصفح غير مكتومة واسمح بتشغيل الصوت.",
            "中文": "如果听不到声音，请检查浏览器标签页是否被静音，并允许音频播放。"
        }
        st.caption(istruzioni_audio.get(lingua_sel, istruzioni_audio["Italiano"]))
        if st.button(etichette_prova_notifiche.get(lingua_sel, etichette_prova_notifiche["Italiano"]), key="prova_notifiche"):
            notifica_sonora("test_notifiche", lingua_sel, ripeti=True)

        st.divider()
        guida_chat_sidebar = {
            "Italiano": {
                "titolo": "🤖 Crea la tua chat per compilare la sidebar",
                "nota": "Questa guida non è il prompt. Leggila prima; subito sotto trovi il prompt da copiare nella tua chat personale di GPT.",
                "passi": """1. Apri una nuova chat nel tuo GPT personale.

2. Copia tutto il testo nel riquadro scuro qui sotto, senza cambiarlo.

3. Incollalo come primo messaggio e invialo.

4. Dopo la tua idea, GPT chiederà sempre se vuoi aggiungere una personalizzazione del libro. Solo dopo la tua risposta preparerà la scheda; se manca un dato davvero decisivo, potrà fare un'ultima domanda breve.

5. Alla fine riceverai una scheda pronta. Copia ogni risposta nel campo con lo stesso nome nella sidebar di Scrittore Site.

6. Non chiedere a quella chat di scrivere il libro completo: serve solo a preparare bene la sidebar e poi Scrittore Site creerà l'indice.""",
                "etichetta": "PROMPT DA COPIARE NELLA TUA CHAT PERSONALE DI GPT",
            },
            "English": {
                "titolo": "🤖 Create your chat to complete the sidebar",
                "nota": "This guide is not the prompt. Read it first; immediately below is the prompt to copy into your personal GPT chat.",
                "passi": """1. Open a new chat in your personal GPT account.

2. Copy all the text in the dark box below, without changing it.

3. Paste it as the first message and send it.

4. GPT will begin with one simple question about the book you want to create. If the information is already clear, it will prepare the form immediately; otherwise, it will ask one short question about the most important choice.

5. At the end you will receive a ready-to-use form. Copy each answer into the field with the same name in the Scrittore Site sidebar.

6. Do not ask that chat to write the complete book: it only prepares the sidebar, then Scrittore Site creates the index.""",
                "etichetta": "PROMPT TO COPY INTO YOUR PERSONAL GPT CHAT",
            },
            "Español": {
                "titolo": "🤖 Crea tu chat para completar la barra lateral",
                "nota": "Esta guía no es el prompt. Léela primero; justo debajo está el prompt que debes copiar en tu chat personal de GPT.",
                "passi": """1. Abre un chat nuevo en tu cuenta personal de GPT.

2. Copia todo el texto del recuadro oscuro de abajo, sin cambiarlo.

3. Pégalo como primer mensaje y envíalo.

4. GPT empezará con una pregunta sencilla sobre el libro que quieres crear. Si la información ya está clara, preparará la ficha enseguida; si no, hará una sola pregunta breve sobre la decisión más importante.

5. Al final recibirás una ficha lista. Copia cada respuesta en el campo con el mismo nombre de la barra lateral de Scrittore Site.

6. No pidas a ese chat que escriba el libro completo: solo prepara la barra lateral y después Scrittore Site creará el índice.""",
                "etichetta": "PROMPT PARA COPIAR EN TU CHAT PERSONAL DE GPT",
            },
            "Français": {
                "titolo": "🤖 Créez votre chat pour remplir la barre latérale",
                "nota": "Ce guide n’est pas le prompt. Lisez-le d’abord ; le prompt à copier dans votre chat GPT personnel se trouve juste dessous.",
                "passi": """1. Ouvrez une nouvelle conversation dans votre compte GPT personnel.

2. Copiez tout le texte du cadre sombre ci-dessous, sans le modifier.

3. Collez-le comme premier message et envoyez-le.

4. GPT commencera par une question simple sur le livre à créer. Si les informations sont déjà claires, il préparera directement la fiche ; sinon, il posera une seule question courte sur le choix le plus important.

5. À la fin, vous recevrez une fiche prête à l’emploi. Copiez chaque réponse dans le champ du même nom de la barre latérale de Scrittore Site.

6. Ne demandez pas à ce chat d’écrire le livre complet : il prépare seulement la barre latérale, puis Scrittore Site créera l’index.""",
                "etichetta": "PROMPT À COPIER DANS VOTRE CHAT GPT PERSONNEL",
            },
            "Deutsch": {
                "titolo": "🤖 Erstellen Sie Ihren Chat zum Ausfüllen der Seitenleiste",
                "nota": "Diese Anleitung ist nicht der Prompt. Lesen Sie sie zuerst; darunter finden Sie den Prompt zum Kopieren in Ihren persönlichen GPT-Chat.",
                "passi": """1. Öffnen Sie einen neuen Chat in Ihrem persönlichen GPT-Konto.

2. Kopieren Sie den gesamten Text im dunklen Feld unten, ohne ihn zu ändern.

3. Fügen Sie ihn als erste Nachricht ein und senden Sie ihn ab.

4. GPT beginnt mit einer einfachen Frage zu dem Buch, das Sie erstellen möchten. Wenn die Angaben bereits klar sind, erstellt es sofort die Vorlage; andernfalls stellt es eine kurze Frage zur wichtigsten Entscheidung.

5. Am Ende erhalten Sie eine fertige Vorlage. Kopieren Sie jede Antwort in das gleichnamige Feld der Scrittore-Site-Seitenleiste.

6. Bitten Sie diesen Chat nicht, das ganze Buch zu schreiben: Er bereitet nur die Seitenleiste vor, danach erstellt Scrittore Site den Index.""",
                "etichetta": "PROMPT ZUM KOPIEREN IN IHREN PERSÖNLICHEN GPT-CHAT",
            },
            "Română": {
                "titolo": "🤖 Creează chatul tău pentru completarea barei laterale",
                "nota": "Acest ghid nu este promptul. Citește-l mai întâi; imediat mai jos găsești promptul de copiat în chatul tău personal GPT.",
                "passi": """1. Deschide un chat nou în contul tău personal GPT.

2. Copiază tot textul din caseta întunecată de mai jos, fără să îl modifici.

3. Lipește-l ca primul mesaj și trimite-l.

4. GPT va începe cu o întrebare simplă despre cartea pe care vrei să o creezi. Dacă informațiile sunt deja clare, va pregăti imediat fișa; altfel, va pune o singură întrebare scurtă despre alegerea cea mai importantă.

5. La final vei primi o fișă pregătită. Copiază fiecare răspuns în câmpul cu același nume din bara laterală Scrittore Site.

6. Nu cere acelui chat să scrie cartea completă: el pregătește doar bara laterală, apoi Scrittore Site va crea indexul.""",
                "etichetta": "PROMPT DE COPIAT ÎN CHATUL TĂU PERSONAL GPT",
            },
            "Русский": {
                "titolo": "🤖 Создайте чат для заполнения боковой панели",
                "nota": "Это руководство — не промпт. Сначала прочитайте его; ниже находится промпт для копирования в ваш личный чат GPT.",
                "passi": """1. Откройте новый чат в своём личном аккаунте GPT.

2. Скопируйте весь текст из тёмного блока ниже, ничего не меняя.

3. Вставьте его первым сообщением и отправьте.

4. GPT начнёт с простого вопроса о книге, которую вы хотите создать. Если информации достаточно, он сразу подготовит карточку; иначе задаст один короткий вопрос о самом важном выборе.

5. В конце вы получите готовую карточку. Скопируйте каждый ответ в поле с таким же названием на боковой панели Scrittore Site.

6. Не просите этот чат написать всю книгу: он только готовит боковую панель, после чего Scrittore Site создаст оглавление.""",
                "etichetta": "ПРОМПТ ДЛЯ КОПИРОВАНИЯ В ВАШ ЛИЧНЫЙ ЧАТ GPT",
            },
            "العربية": {
                "titolo": "🤖 أنشئ محادثتك لملء الشريط الجانبي",
                "nota": "هذا الدليل ليس هو الموجّه. اقرأه أولاً؛ ستجد أسفله مباشرة الموجّه لنسخه في محادثتك الشخصية مع GPT.",
                "passi": """1. افتح محادثة جديدة في حسابك الشخصي على GPT.

2. انسخ كل النص الموجود في المربع الداكن أدناه دون تغييره.

3. الصقه كأول رسالة ثم أرسله.

4. سيبدأ GPT بسؤال بسيط عن الكتاب الذي تريد إنشاءه. إذا كانت المعلومات واضحة، فسيُعدّ البطاقة مباشرة؛ وإلا فسيسأل سؤالاً قصيراً واحداً عن أهم اختيار.

5. في النهاية ستحصل على بطاقة جاهزة. انسخ كل إجابة إلى الحقل الذي يحمل الاسم نفسه في الشريط الجانبي لـ Scrittore Site.

6. لا تطلب من تلك المحادثة كتابة الكتاب كاملاً: دورها فقط إعداد الشريط الجانبي، وبعد ذلك ينشئ Scrittore Site الفهرس.""",
                "etichetta": "موجّه لنسخه في محادثتك الشخصية مع GPT",
            },
            "中文": {
                "titolo": "🤖 创建聊天以填写侧边栏",
                "nota": "这不是提示词本身。请先阅读本说明；下面就是可复制到你个人 GPT 聊天中的提示词。",
                "passi": """1. 在你的个人 GPT 账户中打开一个新聊天。

2. 完整复制下方深色框中的所有文字，不要修改。

3. 将它作为第一条消息粘贴并发送。

4. GPT 会先询问一个有关你想创作的图书的简单问题。如果信息已足够清晰，它会立即准备好资料卡；否则只会就最重要的选择再问一个简短问题。

5. 最后你会得到一张可直接使用的资料卡。将每项答案复制到 Scrittore Site 侧边栏中同名的字段。

6. 不要让该聊天直接写完整本书：它只负责准备侧边栏，然后 Scrittore Site 会生成目录。""",
                "etichetta": "复制到你的个人 GPT 聊天中的提示词",
            },
        }
        guida_chat = guida_chat_sidebar.get(lingua_sel, guida_chat_sidebar["Italiano"])
        etichette_tendine_chat = {
            "Italiano": {"manuale": "📋 Crea la tua chat per compilare la sidebar", "guidata": "Chat guidata per compilare la sidebar (opzionale)"},
            "English": {"manuale": "📋 Create your chat to complete the sidebar", "guidata": "Guided chat to complete the sidebar (optional)"},
            "Español": {"manuale": "📋 Crea tu chat para completar la barra lateral", "guidata": "Chat guiado para completar la barra lateral (opcional)"},
            "Français": {"manuale": "📋 Créez votre chat pour remplir la barre latérale", "guidata": "Chat guidé pour remplir la barre latérale (facultatif)"},
            "Deutsch": {"manuale": "📋 Erstellen Sie Ihren Chat zum Ausfüllen der Seitenleiste", "guidata": "Geführter Chat zum Ausfüllen der Seitenleiste (optional)"},
            "Română": {"manuale": "📋 Creează chatul tău pentru completarea barei laterale", "guidata": "Chat ghidat pentru completarea barei laterale (opțional)"},
            "Русский": {"manuale": "📋 Создайте чат для заполнения боковой панели", "guidata": "Управляемый чат для заполнения боковой панели (необязательно)"},
            "العربية": {"manuale": "📋 أنشئ محادثتك لملء الشريط الجانبي", "guidata": "دردشة موجّهة لملء الشريط الجانبي (اختياري)"},
            "中文": {"manuale": "📋 创建聊天以填写侧边栏", "guidata": "引导式聊天填写侧边栏（可选）"},
        }
        etichette_chat = etichette_tendine_chat.get(lingua_sel, etichette_tendine_chat["Italiano"])
        st.subheader(guida_chat["titolo"])
        st.info(guida_chat["nota"])
        st.markdown(guida_chat["passi"])
        istruzioni_chat_guidata = {
            "Italiano": "**Alternativa opzionale:** apri la chat guidata qui sotto per parlare direttamente con Scrittore Site. Usa lingua e cervello selezionati nella sidebar. L'avvio è gratuito; ogni risposta costa 1 credito con GPT oppure 1/3 di credito con DeepSeek (1 credito ogni 3 risposte). Alla fine puoi compilare automaticamente solo i campi vuoti oppure sostituire quelli già presenti.",
            "English": "**Optional alternative:** open the guided chat below to speak directly with Scrittore Site. It uses the language and brain selected in the sidebar. Starting is free; each reply costs 1 credit with GPT or 1/3 credit with DeepSeek (1 credit every 3 replies). At the end you may fill only empty fields automatically or replace existing ones.",
            "Español": "**Alternativa opcional:** abre el chat guiado de abajo para hablar directamente con Scrittore Site. Usa el idioma y el motor elegidos en la barra lateral. Iniciar es gratuito; cada respuesta cuesta 1 crédito con GPT o 1/3 de crédito con DeepSeek (1 crédito cada 3 respuestas). Al final puedes completar solo campos vacíos o sustituir los existentes.",
            "Français": "**Alternative facultative :** ouvrez le chat guidé ci-dessous pour dialoguer directement avec Scrittore Site. Il utilise la langue et le moteur choisis dans la barre latérale. Le démarrage est gratuit ; chaque réponse coûte 1 crédit avec GPT ou 1/3 de crédit avec DeepSeek (1 crédit toutes les 3 réponses). À la fin, vous pouvez remplir seulement les champs vides ou remplacer ceux qui existent.",
            "Deutsch": "**Optionale Alternative:** Öffnen Sie unten den geführten Chat und sprechen Sie direkt mit Scrittore Site. Er verwendet die in der Seitenleiste gewählte Sprache und KI. Der Start ist kostenlos; jede Antwort kostet 1 Guthabenpunkt mit GPT oder 1/3 Punkt mit DeepSeek (1 Punkt je 3 Antworten). Am Ende können Sie nur leere Felder füllen oder vorhandene ersetzen.",
            "Română": "**Alternativă opțională:** deschide chatul ghidat de mai jos pentru a vorbi direct cu Scrittore Site. Folosește limba și motorul selectate în bara laterală. Pornirea este gratuită; fiecare răspuns costă 1 credit cu GPT sau 1/3 credit cu DeepSeek (1 credit la 3 răspunsuri). La final poți completa doar câmpurile goale sau le poți înlocui pe cele existente.",
            "Русский": "**Необязательная альтернатива:** откройте управляемый чат ниже, чтобы общаться напрямую со Scrittore Site. Он использует язык и модель, выбранные на боковой панели. Запуск бесплатен; каждый ответ стоит 1 кредит с GPT или 1/3 кредита с DeepSeek (1 кредит за 3 ответа). В конце можно заполнить только пустые поля или заменить существующие.",
            "العربية": "**بديل اختياري:** افتح الدردشة الموجّهة أدناه للتحدث مباشرة مع Scrittore Site. تستخدم اللغة والمحرك المختارين في الشريط الجانبي. البدء مجاني؛ كل رد يكلف رصيدًا واحدًا مع GPT أو ثلث رصيد مع DeepSeek (رصيد واحد كل 3 ردود). في النهاية يمكنك ملء الحقول الفارغة فقط أو استبدال الحقول الموجودة.",
            "中文": "**可选方式：** 打开下方的引导式聊天，直接与 Scrittore Site 对话。它使用侧边栏中选择的语言和模型。启动免费；GPT 每次回复消耗 1 积分，DeepSeek 每次消耗 1/3 积分（每 3 次回复消耗 1 积分）。最后可自动仅填写空字段，或替换已有字段。",
        }
        st.info(istruzioni_chat_guidata.get(lingua_sel, istruzioni_chat_guidata["Italiano"]))
        istruzione_lingua_prompt = {
            "Italiano": f"LINGUA OPERATIVA OBBLIGATORIA: comunica sempre con l'utente e restituisci la scheda finale esclusivamente in {lingua_sel}. Non tradurre i nomi esatti delle opzioni di Scrittore Site indicate sotto.",
            "English": f"MANDATORY WORKING LANGUAGE: communicate with the user and return the final form only in {lingua_sel}. Do not translate the exact Scrittore Site option names shown below.",
            "Español": f"IDIOMA DE TRABAJO OBLIGATORIO: comunícate con el usuario y entrega la ficha final únicamente en {lingua_sel}. No traduzcas los nombres exactos de las opciones de Scrittore Site indicadas abajo.",
            "Français": f"LANGUE DE TRAVAIL OBLIGATOIRE : communiquez avec l’utilisateur et fournissez la fiche finale uniquement en {lingua_sel}. Ne traduisez pas les noms exacts des options de Scrittore Site ci-dessous.",
            "Deutsch": f"VERBINDLICHE ARBEITSSPRACHE: Kommunizieren Sie mit dem Nutzer und geben Sie die fertige Vorlage ausschließlich auf {lingua_sel} zurück. Übersetzen Sie die unten stehenden exakten Namen der Scrittore-Site-Optionen nicht.",
            "Română": f"LIMBĂ DE LUCRU OBLIGATORIE: comunică cu utilizatorul și oferă fișa finală exclusiv în {lingua_sel}. Nu traduce denumirile exacte ale opțiunilor Scrittore Site de mai jos.",
            "Русский": f"ОБЯЗАТЕЛЬНЫЙ РАБОЧИЙ ЯЗЫК: общайтесь с пользователем и выдавайте итоговую карточку только на {lingua_sel}. Не переводите точные названия параметров Scrittore Site ниже.",
            "العربية": f"لغة العمل الإلزامية: تواصل مع المستخدم وقدّم البطاقة النهائية باللغة {lingua_sel} فقط. لا تترجم أسماء خيارات Scrittore Site الدقيقة الواردة أدناه.",
            "中文": f"强制工作语言：仅使用 {lingua_sel} 与用户交流并返回最终资料卡。不要翻译下方 Scrittore Site 选项的准确名称。",
        }
        prompt_chat_sidebar = f"""{istruzione_lingua_prompt.get(lingua_sel, istruzione_lingua_prompt['Italiano'])}

Agisci esclusivamente come Assistente per la Sidebar di Scrittore Site.

Il tuo unico compito è trasformare l'idea dell'utente in una scheda completa, pronta da copiare nei campi della sidebar di Scrittore Site.

Non scrivere il libro.
Non creare l'indice.
Non aggiungere campi non presenti nella sidebar.
Non proporre sottotitoli, prezzo, formato, numero di pagine, capitoli, piano marketing o altri elementi esterni alla sidebar.

FASE 1 — INIZIO SEMPLICE

All'inizio scrivi soltanto:

“Che libro vuoi creare?

Puoi scrivere anche una sola frase, per esempio: ‘Vorrei una guida per principianti sulle tartarughe di terra’.

Se sai già qualcosa in più, puoi indicare lettore, lingua, risultato desiderato o un titolo. Se non li sai, non preoccuparti: li scelgo io in modo coerente per Scrittore Site.”

FASE 2 — POCHE DOMANDE

Dopo la risposta dell'utente:

1. Prima di redigere la scheda, fai SEMPRE questa domanda esplicita e soltanto questa:

“Vuoi personalizzare questo libro?

A. No, procedi con un progetto editoriale generale
B. Sì, voglio aggiungere voce personale, episodi, esempi, materiali o priorità
C. Sì, voglio indicare soprattutto confini da rispettare o punti in cui fermare la scrittura

Quale scegli?”

2. Se l'utente sceglie A, non aggiungere campi personali e passa alla scheda.
3. Se sceglie B o C, chiedi in un solo messaggio i dettagli indispensabili: cosa includere, cosa evitare e, solo se utile, quando fermare la stesura per aggiungere una nota. Non inventare mai dati personali.
4. Solo dopo la risposta sulla personalizzazione analizza l'idea e ricava autonomamente tutto ciò che è ragionevole dedurre.
5. Se pubblico, obiettivo e argomento sono abbastanza chiari, prepara subito la scheda finale. Fai un'ulteriore domanda soltanto se manca una scelta che cambia davvero il progetto; deve essere breve e offrire al massimo tre alternative concrete.
6. Non chiedere titolo, autore, tipologia di scrittura, stile di racconto, punto di vista, lunghezza o Cervello AI: sceglili tu.
7. Non fare più di due domande successive alla risposta iniziale dell'utente, salvo che l'idea sia troppo vaga per creare una scheda affidabile.

Esempio di unica domanda aggiuntiva utile:

“Questo libro può essere rivolto soprattutto a:
A. Principianti
B. Appassionati già esperti
C. Professionisti o utenti tecnici

Quale scegli?”

REGOLE DI CONFORMITÀ CON SCRITTORE SITE

Usa esclusivamente i campi e le opzioni disponibili nella sidebar di Scrittore Site.

Per il GENERE LETTERARIO scegli un solo valore tra:

- Saggio Scientifico
- Quiz Scientifico
- Manuale Tecnico
- Religioso / Teologico
- Spirituale / Esoterico
- Meditazione / Mindfulness
- Business & Marketing
- Economia e Finanza
- Romanzo Rosa
- Thriller / Noir
- Fantasy
- Fantascienza
- Manuale Psicologico
- Biografia
- Ricettario
- Test Prep (Preparazione Esami)
- Narrativo
- Romanzo Classico
- Contemporaneo
- Self-Help
- Manuale Pratico
- Storico

Non proporre mai un genere secondario, alternativo o inventato.

Per la TIPOLOGIA SCRITTURA scegli un solo valore tra:

- Standard
- Professionale Accademico
- Persuasivo (Neuromarketing Applicato)
- Conversazionale ed Empatico
- Scientifico Divulgativo
- Storytelling Immersivo
- Giornalistico d'Inchiesta
- Socratico (Dialogico / Riflessivo)
- Epico ed Evocativo
- Minimalista ed Essenziale

Per lo STILE DI RACCONTO scegli un solo valore tra:

- Coinvolgente e Narrativo
- Tecnico e Analitico
- Ispirazionale e Motivante
- Socratico (Domanda/Risposta)
- Storytelling Emozionale
- Diretto e Pratico (Action-oriented)
- Storico e Documentale

Per il PUNTO DI VISTA scegli un solo valore tra:

- Tu (Diretto, confidenziale e personale)
- Voi (Plurale, autorevole e rispettoso)
- Noi (Inclusivo, partecipativo e didattico)
- Impersonale / Terza Persona (Distaccato, analitico, oggettivo)

Per LUNGHEZZA DELLE SEZIONI scegli un solo valore tra:

- Compatto — circa 480-560 parole per sezione, massimo 65 sezioni totali inclusa la Prefazione, obiettivo 75 pagine
- Standard KDP — circa 620-700 parole per sezione, massimo 110 sezioni totali inclusa la Prefazione, obiettivo 150 pagine
- Approfondito — circa 700-800 parole per sezione, massimo 160 sezioni totali inclusa la Prefazione, obiettivo 220 pagine

Scegli Standard KDP come impostazione predefinita. I limiti si riferiscono a tutte le sezioni dell'indice, Prefazione inclusa. Gli obiettivi di pagine 6×9 hanno una tolleranza del 15%: se il libro è completo e di qualità, non aggiungere mai riempitivi solo per raggiungerli. Usa Compatto per guide rapide o libri brevi. Usa Approfondito solo per argomenti tecnici, esami, procedure o materie che richiedono più spiegazione.

Usa esattamente i nomi delle opzioni qui riportate. Non modificarli e non crearne di nuovi.

Per CERVELLO AI scegli un solo valore tra:

- GPT-5.4 (OpenAI)
- DeepSeek V4 Pro

Scegli GPT-5.4 (OpenAI) come valore predefinito. Scegli DeepSeek V4 Pro soltanto se il lettore vuole ridurre il consumo di crediti e non ha bisogno di verifica copyright web o generazione immagini. Con DeepSeek è disponibile la ricerca delle fonti con registro visibile. Il Cervello AI non modifica lingua, genere, stile o contenuto del libro: indica solo il motore che Scrittore Site utilizzerà.

PERSONALIZZAZIONE DEL LIBRO (OPZIONALE, DA CHIEDERE PRIMA DELLA SCHEDA)

La sidebar dispone anche di questi campi facoltativi: voce o prospettiva dell'autore; episodi, casi, esempi o materiali personali; priorità personali per il lettore; confini da rispettare; eventuali note aggiunte durante una pausa guidata.

- Poni sempre la domanda A/B/C prevista nella FASE 2 prima di redigere la scheda. La personalizzazione resta facoltativa: la scelta A non richiede altri dati.
- Se l'utente sceglie B o C, aiutalo a formularli in modo concreto e coerente con il progetto.
- Non inventare mai esperienze, testimonianze, risultati o dettagli personali non dichiarati dall'utente.
- Se i dati non sono disponibili, non aggiungere il blocco facoltativo nella scheda finale e non sostituirli con formule generiche.
- Se sono disponibili, restituisci il blocco facoltativo dopo APPROFONDIMENTI: deve essere pronto da copiare nei campi di Personalizza il tuo libro. La modalità di pausa può essere soltanto: Continua automaticamente; Fermati prima di ogni Parte; Fermati prima della conclusione; Fermati prima delle Parti e della conclusione.

REGOLE DI QUALITÀ

- Proponi un titolo unico, senza sottotitolo, se l'utente non ne ha già fornito uno.
- Se l'autore non è noto, inserisci: [Inserisci il tuo nome].
- Scrivi obiettivo, risultato e argomento in modo dettagliato, concreto e attinente al progetto.
- Evita formule vaghe come “diventare esperto” o “migliorare la vita”.
- Spiega cosa il lettore saprà fare, applicare, produrre o comprendere alla fine.
- Mantieni confini chiari per evitare un libro troppo ampio o ripetitivo.
- In TRAMA O ARGOMENTO descrivi pubblico, problema affrontato, contenuti principali, progressione logica e ciò che deve restare fuori dal libro. Non creare un indice né elencare capitoli.
- In APPROFONDIMENTI indica gli aspetti che l'indice dovrà sviluppare con maggiore precisione: procedure, esempi, esercizi, casi pratici, errori comuni, verifiche o vincoli.
- Ricorda che un capitolo con sottocapitoli sarà una breve cornice: il contenuto dettagliato verrà sviluppato nei sottocapitoli, evitando ripetizioni.
- Adatta il progetto al genere: procedure per manuali, ricette per ricettari, quiz e simulazioni per test prep, trama e conflitto per narrativa.
- Non promettere risultati garantiti.
- Per salute, animali, norme, leggi, software, esami, prezzi, dati o altri temi aggiornabili, specifica negli Approfondimenti che le informazioni devono essere verificate e aggiornate prima della pubblicazione.

RISPOSTA FINALE

Quando hai dati sufficienti, non fare altre domande.

Restituisci soltanto questa scheda, nella lingua scelta per il libro, senza commenti prima o dopo:

TITOLO DEL LIBRO:

NOME AUTORE:

LINGUA:

GENERE LETTERARIO:

TIPOLOGIA SCRITTURA:

STILE DI RACCONTO:

PUNTO DI VISTA:

LUNGHEZZA DELLE SEZIONI:

CERVELLO AI:

OBIETTIVO DEL LIBRO:

RISULTATO FINALE DESIDERATO:

TRAMA O ARGOMENTO:

APPROFONDIMENTI (FACOLTATIVO):

Se l'utente ha fornito dati di personalizzazione, aggiungi soltanto allora anche:

VOCE O PROSPETTIVA PERSONALE (FACOLTATIVO):

MATERIALI PERSONALI (FACOLTATIVO):

PRIORITÀ PERSONALI PER IL LETTORE (FACOLTATIVO):

CONFINI PERSONALI DA RISPETTARE (FACOLTATIVO):

PAUSA GUIDATA DURANTE SCRIVI TUTTO IL LIBRO (FACOLTATIVO):

Ora copia ogni voce nel campo con lo stesso nome nella sidebar di Scrittore Site e genera l'indice."""
        # Per le lingue diverse dall'italiano sostituiamo l'intero testo operativo,
        # non solo l'intestazione della guida. Le opzioni restano in italiano perché
        # sono valori tecnici da selezionare esattamente nella sidebar.
        istruzioni_multilingue = {
            "English": "Act only as the Scrittore Site Sidebar Assistant. Turn the user's idea into a complete form ready to copy into the sidebar. Do not write the book or outline, add external fields, or suggest subtitles, prices, formats, page counts, chapters, or marketing plans.\n\nPHASE 1 — Start with only: ‘What book would you like to create? You may write just one sentence. If you know more, mention reader, language, desired result, or title. If not, I will choose coherently.’\n\nPHASE 2 — After the answer, infer everything reasonable. If audience, goal, and topic are clear, prepare the form immediately. Ask only one short follow-up question with at most three choices when a decision would materially change the project. Do not ask for title, author, writing type, narrative style, point of view, or length: choose them yourself.\n\nQUALITY — Propose one title without subtitle; use [Enter your name] if the author is unknown. Make goal, result, and topic concrete and detailed. In TRAMA O ARGOMENTO include audience, problem, contents, progression, and boundaries. In APPROFONDIMENTI include procedures, examples, exercises, cases, mistakes, checks, or constraints. For changing topics, state that information must be verified before publication. Do not promise guaranteed results.",
            "Español": "Actúa solo como asistente de la barra lateral de Scrittore Site. Convierte la idea del usuario en una ficha completa lista para copiar. No escribas el libro ni el índice, no añadas campos externos ni propongas subtítulos, precios, formatos, páginas, capítulos o marketing.\n\nFASE 1 — Empieza solo con: ‘¿Qué libro quieres crear? Puedes escribir una sola frase. Si sabes más, indica lector, idioma, resultado o título; si no, lo elegiré de forma coherente.’\n\nFASE 2 — Después de la respuesta, deduce lo razonable. Si público, objetivo y tema están claros, prepara la ficha. Haz una sola pregunta breve, con tres alternativas como máximo, solo si falta una decisión importante. No preguntes título, autor, tipo de escritura, estilo, punto de vista o longitud: elígelos tú.\n\nCALIDAD — Propón un título único sin subtítulo; usa [Introduce tu nombre] si no se conoce el autor. Redacta objetivo, resultado y tema de forma concreta y detallada. En TRAMA O ARGOMENTO incluye público, problema, contenidos, progresión y límites. En APPROFONDIMENTI incluye procedimientos, ejemplos, ejercicios, casos, errores, verificaciones o restricciones. Para temas actualizables indica que deben verificarse antes de publicar. No prometas resultados garantizados.",
            "Français": "Agissez uniquement comme assistant de la barre latérale de Scrittore Site. Transformez l'idée de l'utilisateur en fiche complète prête à copier. N'écrivez ni le livre ni le sommaire, n'ajoutez pas de champs externes et ne proposez ni sous-titre, ni prix, ni format, ni pages, ni chapitres, ni marketing.\n\nPHASE 1 — Commencez seulement par : « Quel livre voulez-vous créer ? Vous pouvez écrire une seule phrase. Si vous en savez plus, indiquez lecteur, langue, résultat ou titre ; sinon, je les choisirai de façon cohérente. »\n\nPHASE 2 — Après la réponse, déduisez ce qui est raisonnable. Si public, objectif et sujet sont clairs, préparez la fiche. Posez une seule question courte avec trois choix au maximum si une décision importante manque. Ne demandez ni titre, ni auteur, ni type d'écriture, ni style, ni point de vue, ni longueur : choisissez-les.\n\nQUALITÉ — Proposez un titre unique sans sous-titre ; utilisez [Saisissez votre nom] si l'auteur est inconnu. Rendez objectif, résultat et sujet concrets et détaillés. Dans TRAMA O ARGOMENTO, indiquez public, problème, contenu, progression et limites. Dans APPROFONDIMENTI, indiquez procédures, exemples, exercices, cas, erreurs, vérifications ou contraintes. Pour les sujets évolutifs, précisez la vérification avant publication. Ne promettez aucun résultat garanti.",
            "Deutsch": "Handeln Sie nur als Assistent für die Scrittore-Site-Seitenleiste. Wandeln Sie die Idee des Nutzers in ein vollständiges, kopierfertiges Formular um. Schreiben Sie weder Buch noch Inhaltsverzeichnis, fügen Sie keine externen Felder hinzu und schlagen Sie keine Untertitel, Preise, Formate, Seitenzahlen, Kapitel oder Marketing vor.\n\nPHASE 1 — Beginnen Sie nur mit: „Welches Buch möchten Sie erstellen? Sie können einen einzigen Satz schreiben. Wenn Sie mehr wissen, nennen Sie Leser, Sprache, Ergebnis oder Titel; andernfalls wähle ich es passend aus.“\n\nPHASE 2 — Leiten Sie nach der Antwort alles Sinnvolle ab. Sind Zielgruppe, Ziel und Thema klar, erstellen Sie das Formular. Stellen Sie nur eine kurze Zusatzfrage mit höchstens drei Optionen, wenn eine wichtige Entscheidung fehlt. Fragen Sie nicht nach Titel, Autor, Schreibtyp, Stil, Perspektive oder Länge: Wählen Sie diese selbst.\n\nQUALITÄT — Schlagen Sie einen Titel ohne Untertitel vor; verwenden Sie [Ihren Namen eingeben], wenn der Autor unbekannt ist. Formulieren Sie Ziel, Ergebnis und Thema konkret und detailliert. In TRAMA O ARGOMENTO gehören Zielgruppe, Problem, Inhalte, Ablauf und Grenzen. In APPROFONDIMENTI gehören Verfahren, Beispiele, Übungen, Fälle, Fehler, Prüfungen oder Vorgaben. Bei veränderlichen Themen muss eine Prüfung vor Veröffentlichung genannt werden. Versprechen Sie keine garantierten Ergebnisse.",
            "Română": "Acționează doar ca asistent pentru bara laterală Scrittore Site. Transformă ideea utilizatorului într-o fișă completă gata de copiat. Nu scrie cartea sau cuprinsul, nu adăuga câmpuri externe și nu propune subtitluri, prețuri, formate, pagini, capitole sau marketing.\n\nFAZA 1 — Începe doar cu: «Ce carte vrei să creezi? Poți scrie o singură propoziție. Dacă știi mai multe, indică cititorul, limba, rezultatul sau titlul; dacă nu, le aleg eu coerent.»\n\nFAZA 2 — După răspuns, dedu tot ce este rezonabil. Dacă publicul, obiectivul și subiectul sunt clare, pregătește fișa. Pune o singură întrebare scurtă cu cel mult trei opțiuni numai când lipsește o alegere importantă. Nu întreba despre titlu, autor, tip, stil, punct de vedere sau lungime: alege-le tu.\n\nCALITATE — Propune un titlu unic fără subtitlu; folosește [Introduceți numele] dacă autorul nu este cunoscut. Scrie concret și detaliat obiectivul, rezultatul și subiectul. În TRAMA O ARGOMENTO include publicul, problema, conținutul, progresia și limitele. În APPROFONDIMENTI include proceduri, exemple, exerciții, cazuri, erori, verificări sau constrângeri. Pentru subiecte actualizabile menționează verificarea înainte de publicare. Nu promite rezultate garantate.",
            "Русский": "Действуйте только как помощник боковой панели Scrittore Site. Превратите идею пользователя в полную карточку для копирования. Не пишите книгу или оглавление, не добавляйте внешние поля и не предлагайте подзаголовки, цены, форматы, страницы, главы или маркетинг.\n\nЭТАП 1 — Начните только с: «Какую книгу вы хотите создать? Можно написать одно предложение. Если вы знаете больше, укажите читателя, язык, результат или название; иначе я выберу их согласованно.»\n\nЭТАП 2 — После ответа выведите всё разумно возможное. Если аудитория, цель и тема ясны, подготовьте карточку. Задайте только один короткий вопрос максимум с тремя вариантами, когда не хватает важного решения. Не спрашивайте название, автора, тип письма, стиль, точку зрения или длину: выберите сами.\n\nКАЧЕСТВО — Предложите одно название без подзаголовка; если автор неизвестен, используйте [Введите ваше имя]. Пишите цель, результат и тему конкретно и подробно. В TRAMA O ARGOMENTO укажите аудиторию, проблему, содержание, последовательность и границы. В APPROFONDIMENTI укажите процедуры, примеры, упражнения, случаи, ошибки, проверки или ограничения. Для изменяемых тем укажите проверку до публикации. Не обещайте гарантированных результатов.",
            "العربية": "اعمل فقط كمساعد للشريط الجانبي في Scrittore Site. حوّل فكرة المستخدم إلى بطاقة كاملة جاهزة للنسخ. لا تكتب الكتاب أو الفهرس، ولا تضف حقولاً خارجية، ولا تقترح عناوين فرعية أو أسعاراً أو صيغاً أو صفحات أو فصولاً أو تسويقاً.\n\nالمرحلة 1 — ابدأ فقط بعبارة: «ما الكتاب الذي تريد إنشاءه؟ يمكنك كتابة جملة واحدة. إذا كنت تعرف المزيد فاذكر القارئ أو اللغة أو النتيجة أو العنوان؛ وإلا سأختارها بشكل متسق.»\n\nالمرحلة 2 — بعد الإجابة استنتج كل ما هو معقول. إذا كان الجمهور والهدف والموضوع واضحين فجهز البطاقة. اطرح سؤالاً إضافياً واحداً قصيراً بثلاثة بدائل كحد أقصى فقط عند غياب قرار مهم. لا تسأل عن العنوان أو المؤلف أو نوع الكتابة أو الأسلوب أو وجهة النظر أو الطول: اخترها بنفسك.\n\nالجودة — اقترح عنواناً واحداً بلا عنوان فرعي؛ استخدم [أدخل اسمك] إذا كان المؤلف غير معروف. اكتب الهدف والنتيجة والموضوع بوضوح وتفصيل. في TRAMA O ARGOMENTO اذكر الجمهور والمشكلة والمحتوى والتدرج والحدود. في APPROFONDIMENTI اذكر الإجراءات والأمثلة والتمارين والحالات والأخطاء والتحققات والقيود. للموضوعات المتغيرة اذكر التحقق قبل النشر. لا تعد بنتائج مضمونة.",
            "中文": "仅作为 Scrittore Site 侧边栏助手工作。将用户想法转换为可直接复制的完整资料卡。不要撰写书籍或目录，不要添加外部字段，也不要建议副标题、价格、格式、页数、章节或营销计划。\n\n阶段一——仅以此开始：“你想创作什么书？你可以只写一句话。如果你知道更多信息，可以说明读者、语言、结果或标题；如果不知道，我会合理选择。”\n\n阶段二——收到回答后，合理推断所有信息。如果读者、目标和主题明确，立即准备资料卡。只有在缺少会实质改变项目的重要选择时，才提出一个简短的附加问题，最多三个具体选项。不要询问标题、作者、写作类型、风格、视角或篇幅：请自行选择。\n\n质量——提出一个无副标题的唯一标题；作者未知时使用[填写你的姓名]。具体、详细地写出目标、结果和主题。在 TRAMA O ARGOMENTO 中包括读者、问题、内容、推进和边界。在 APPROFONDIMENTI 中包括流程、示例、练习、案例、错误、核查或限制。对于会更新的主题，注明出版前必须核实。不要承诺保证性结果。",
        }
        # Funzione mantenuta inattiva: il prompt mostrato resta quello precedente finché non verrà richiesta una revisione multilingue completa.
        if os.getenv("ENABLE_MULTILINGUAL_SIDEBAR_PROMPT", "0") == "1" and lingua_sel in istruzioni_multilingue:
            opzioni_esatte = """Use exclusively these exact sidebar option values; do not translate or invent them.\nGENERE LETTERARIO: Saggio Scientifico; Quiz Scientifico; Manuale Tecnico; Religioso / Teologico; Spirituale / Esoterico; Meditazione / Mindfulness; Business & Marketing; Economia e Finanza; Romanzo Rosa; Thriller / Noir; Fantasy; Fantascienza; Manuale Psicologico; Biografia; Ricettario; Test Prep (Preparazione Esami); Narrativo; Romanzo Classico; Contemporaneo; Self-Help; Manuale Pratico; Storico.\nTIPOLOGIA SCRITTURA: Standard; Professionale Accademico; Persuasivo (Neuromarketing Applicato); Conversazionale ed Empatico; Scientifico Divulgativo; Storytelling Immersivo; Giornalistico d'Inchiesta; Socratico (Dialogico / Riflessivo); Epico ed Evocativo; Minimalista ed Essenziale.\nSTILE DI RACCONTO: Coinvolgente e Narrativo; Tecnico e Analitico; Ispirazionale e Motivante; Socratico (Domanda/Risposta); Storytelling Emozionale; Diretto e Pratico (Action-oriented); Storico e Documentale.\nPUNTO DI VISTA: Tu (Diretto, confidenziale e personale); Voi (Plurale, autorevole e rispettoso); Noi (Inclusivo, partecipativo e didattico); Impersonale / Terza Persona (Distaccato, analitico, oggettivo).\nLUNGHEZZA DELLE SEZIONI: Compatto (480-560 words, max 50 sections); Standard KDP (620-700 words, max 80 sections); Approfondito (700-800 words, max 110 sections). Choose Standard KDP by default; Compatto for short guides and Approfondito for technical subjects, exams, or procedures.\nCERVELLO AI: GPT-5.4 (OpenAI); DeepSeek V4 Pro. Choose GPT-5.4 (OpenAI) by default. Choose DeepSeek V4 Pro only when lower credit consumption is preferred and web copyright checks or image generation are not required. DeepSeek performs source research with a visible register."""
            personalizzazione_prompt_multilingue = """OPTIONAL PERSONALIZATION — Before drafting the final form, always ask explicitly whether the user wants to personalize the book: A. no, proceed with a general editorial project; B. yes, add personal voice, episodes, examples, materials or reader priorities; C. yes, mainly define boundaries or guided pauses. If A is selected, omit all optional personalization fields. If B or C is selected, ask in one message only for the indispensable details, then formulate the relevant fields concretely. Never invent personal experiences, testimonials or facts. Append personalization fields after APPROFONDIMENTI only when the user has chosen personalization and supplied relevant details. Guided-pause mode can only be: Continue automatically; Pause before every Part; Pause before the conclusion; Pause before Parts and conclusion."""
            prompt_chat_sidebar = f"""{istruzione_lingua_prompt[lingua_sel]}

{istruzioni_multilingue[lingua_sel]}

{opzioni_esatte}

{personalizzazione_prompt_multilingue}

When sufficient information is available, ask no more questions. Return only the following form. Keep its labels and the selected option values exactly unchanged; write all descriptive values in {lingua_sel}.

TITOLO DEL LIBRO:

NOME AUTORE:

LINGUA:

GENERE LETTERARIO:

TIPOLOGIA SCRITTURA:

STILE DI RACCONTO:

PUNTO DI VISTA:

LUNGHEZZA DELLE SEZIONI:

CERVELLO AI:

OBIETTIVO DEL LIBRO:

RISULTATO FINALE DESIDERATO:

TRAMA O ARGOMENTO:

APPROFONDIMENTI (FACOLTATIVO):

If the user has provided relevant personalization details, append only then:

VOCE O PROSPETTIVA PERSONALE (FACOLTATIVO):

MATERIALI PERSONALI (FACOLTATIVO):

PRIORITÀ PERSONALI PER IL LETTORE (FACOLTATIVO):

CONFINI PERSONALI DA RISPETTARE (FACOLTATIVO):

PAUSA GUIDATA DURANTE SCRIVI TUTTO IL LIBRO (FACOLTATIVO):"""
        # Il prompt manuale resta disponibile per chi vuole usarlo nella
        # propria ChatGPT, ma non occupa più tutta la pagina.
        with st.expander(etichette_chat["manuale"], expanded=False):
            st.caption(guida_chat["etichetta"])
            st.code(prompt_chat_sidebar, language=None)

        # Alternativa facoltativa al copia/incolla: una conversazione interna
        # che usa il cervello già selezionato nella sidebar. L'avvio è gratuito;
        # si scala credito soltanto quando il cervello restituisce una risposta.
        with st.expander(
            f"🟢 {etichette_chat['guidata']}",
            expanded=bool(st.session_state.get("chat_sidebar_attiva")),
        ):
            if st.session_state.pop("chat_sidebar_scroll_request", False):
                # Il comando in alto avvia subito la chat e porta l'utente al
                # pannello aperto, senza aggiungere una seconda chat o azione.
                components.html(
                    """
                    <script>
                    setTimeout(function () {
                      try {
                        const pannelli = Array.from(
                          window.parent.document.querySelectorAll('[data-testid="stExpander"]')
                        ).filter(function (elemento) {
                          return (elemento.innerText || '').toLowerCase().includes('chat');
                        });
                        const pannello = pannelli[pannelli.length - 1];
                        if (pannello) {
                          pannello.scrollIntoView({ behavior: 'smooth', block: 'start' });
                        }
                      } catch (errore) {}
                    }, 250);
                    </script>
                    """,
                    height=0,
                )
            usa_deepseek_chat = usa_deepseek_pro()
            costo_chat_visibile = (
                "1/3 di credito (1 credito ogni 3 risposte)"
                if usa_deepseek_chat else
                f"{CREDIT_COSTS['chat_sidebar_guidata']} credito"
            )
            cervello_chat = st.session_state.get("provider_ia", "GPT-5.4 (OpenAI)")
            st.info(
                f"La chat usa il cervello selezionato: **{cervello_chat}**. "
                f"L'avvio e i tuoi messaggi sono gratuiti; ogni risposta della chat costa "
                f"**{costo_chat_visibile}**."
            )
            st.caption(
                "Puoi cambiare cervello nella sidebar anche durante la conversazione: "
                "la risposta successiva userà la nuova scelta e il relativo costo."
            )

            if not st.session_state.get("chat_sidebar_attiva"):
                st.button(
                    "💬 AVVIA CHAT GUIDATA",
                    key="avvia_chat_sidebar_guidata",
                    on_click=avvia_chat_sidebar_guidata,
                    use_container_width=True,
                )
            else:
                messaggi_chat = list(st.session_state.get("chat_sidebar_messaggi", []) or [])
                for messaggio in messaggi_chat:
                    with st.chat_message(messaggio.get("role", "assistant")):
                        st.markdown(messaggio.get("content", ""))

                nonce_chat = int(st.session_state.get("chat_sidebar_input_nonce", 0))
                chiave_input_chat = f"chat_sidebar_input_{nonce_chat}"
                testo_utente_chat = st.text_area(
                    campo_ui("messaggio", lingua_sel),
                    key=chiave_input_chat,
                    placeholder="Descrivi l'idea del libro oppure rispondi alla domanda della chat.",
                    height=90,
                )
                if st.button(
                    f"➤ INVIA — risposta IA: {costo_chat_visibile}",
                    key=f"invia_chat_sidebar_{nonce_chat}",
                    use_container_width=True,
                ):
                    testo_utente_chat = str(testo_utente_chat or "").strip()
                    if not testo_utente_chat:
                        st.warning("Scrivi prima un messaggio per la chat.")
                    else:
                        messaggi_chat.append({"role": "user", "content": testo_utente_chat})
                        cronologia = messaggi_chat[-12:]
                        conversazione = "\n\n".join(
                            f"{'UTENTE' if voce.get('role') == 'user' else 'ASSISTENTE'}:\n{voce.get('content', '')}"
                            for voce in cronologia
                        )
                        risposte_utente_chat = sum(
                            1 for voce in messaggi_chat
                            if voce.get("role") == "user"
                        )
                        ultimo_messaggio_chat = testo_utente_chat.casefold()
                        conferma_chat = bool(re.search(
                            r"\b(confermo|conferma|procedi|prosegui|vai|ok|okay|yes|si|sì|continue|go ahead|continúa|continuez|weiter|продолж|افعل|继续)\b",
                            ultimo_messaggio_chat,
                        ))
                        autonomia_creativa_chat = bool(re.search(
                            r"(fai tu|scegli tu|decidi tu|usa la tua immaginazione|fa tu|make it up|you decide|hazlo tú|elige tú|fais[- ]le|choisis|mach du|entscheide du|tu decizi|alege tu|делай сам|реши сам|افعلها أنت|قرر أنت|你来决定|你决定)",
                            ultimo_messaggio_chat,
                        ))
                        istruzione_chiusura_chat = (
                            "L'ultimo messaggio contiene una conferma o richiesta di procedere: produci ADESSO la scheda finale completa. Non porre domande, non chiedere conferme e non offrire alternative."
                            if conferma_chat else
                            "L'utente non ha ancora dato una conferma conclusiva: segui il protocollo sotto e chiedi soltanto il prossimo dato davvero necessario."
                        )
                        istruzione_autonomia_chat = (
                            "L'utente ti ha chiesto di scegliere autonomamente: usa immaginazione editoriale coerente con tutta la conversazione e compila tutti i campi descrittivi. Puoi inventare soltanto impostazioni narrative, esempi generici, tono, obiettivo, pubblico plausibile e priorità editoriali; non attribuire mai all'autore esperienze, testimonianze, risultati, biografia, fonti o fatti reali non dichiarati. Produci subito la scheda finale."
                            if autonomia_creativa_chat else
                            "Se l'utente non ti chiede di scegliere autonomamente, non inventare informazioni personali o fatti non dichiarati."
                        )
                        istruzioni_chat_interattiva = f"""{prompt_chat_sidebar}

MODALITÀ CHAT INTERATTIVA — PRIORITÀ ASSOLUTA
Il tuo scopo non è intrattenere una conversazione generica: devi accompagnare con naturalezza l'utente fino alla compilazione affidabile della sidebar. In questa conversazione l'utente ha già inviato {risposte_utente_chat} messaggi.

SEGNALI DELL'ULTIMO MESSAGGIO
{istruzione_chiusura_chat}
{istruzione_autonomia_chat}

PROTOCOLLO DI GUIDA OBBLIGATORIO
1. Leggi tutta la cronologia e memorizza ogni dato già espresso. Non chiedere mai di nuovo titolo, argomento, pubblico, obiettivo o preferenze che sono già chiari.
2. Dopo la prima descrizione concreta del libro, poni una sola volta la domanda A/B/C sulla personalizzazione prevista dal prompt principale. Non discutere altri argomenti prima di aver ricevuto questa scelta, salvo una risposta di una frase che aiuti davvero il progetto.
3. Se l'utente sceglie A o comunica chiaramente di non volere personalizzazioni, prepara subito la scheda completa. Se sceglie B o C, chiedi in un unico messaggio soltanto i dettagli personali indispensabili. Non proporre liste generiche, non chiedere di confermare una tua proposta e non aprire un secondo giro: alla risposta successiva prepara la scheda. Spiega che può anche rispondere “fai tu”.
4. Se l'utente risponde “confermo”, “ok”, “procedi”, “vai”, o esprime una richiesta equivalente dopo B/C, interpreta la risposta come autorizzazione a chiudere subito: completa la scheda con ciò che è noto e lascia vuoti solo i campi personali che non possono essere costruiti senza attribuire falsamente esperienze all'autore.
5. Puoi porre al massimo un chiarimento editoriale aggiuntivo, e solo se senza quel dato il progetto cambierebbe materialmente. Offri fino a tre scelte concrete. Altrimenti deduci in modo ragionevole e completa tu i campi.
6. Se l'utente apre un discorso laterale, rispondi con intelligenza in non più di due frasi solo quando è utile al suo libro; poi riporta immediatamente al prossimo passo necessario per la scheda. Non aprire ricerche, dibattiti, consigli generici o nuove conversazioni indipendenti.
7. Non creare mai indice, capitoli, testo del libro, piani marketing, prezzi, pagine o campi esterni. Non inventare esperienze personali, fatti o preferenze.

CRITERIO DI CHIUSURA
Appena esistono argomento, pubblico o lettore plausibile, obiettivo e scelta sulla personalizzazione, i dati sono sufficienti: non fare altre domande. Produci la scheda finale completa, anche se alcuni dettagli devono essere scelti da te in modo coerente.

Il cervello AI è già scelto dall'utente nella sidebar: **{cervello_chat}**. Non suggerire di cambiarlo e, nella scheda finale, riporta esattamente questo valore per CERVELLO AI.
Comunica sempre nella lingua operativa selezionata dall'utente. Quando produci la scheda, restituisci tutte e sole le etichette previste, una per riga, senza blocchi di codice né commenti prima o dopo. Non produrre mai una scheda parziale."""
                        with st.spinner("La chat sta preparando la risposta..."):
                            risposta_chat = chiedi_gpt(
                                conversazione,
                                istruzioni_chat_interattiva,
                                addebita=True,
                                amount=CREDIT_COSTS["chat_sidebar_guidata"],
                                max_completion_tokens=1700,
                                model=MODELLO_EDITORIALE,
                                reason="chat_sidebar_guidata",
                                timeout_seconds=90,
                            )
                        if str(risposta_chat or "").startswith("ERRORE:"):
                            st.error(
                                "La chat non ha prodotto una risposta. Nessun credito è stato trattenuto; "
                                "puoi inviare di nuovo il messaggio."
                            )
                        else:
                            messaggi_chat.append({"role": "assistant", "content": risposta_chat})
                            st.session_state["chat_sidebar_messaggi"] = messaggi_chat
                            scheda_estratta = estrai_scheda_chat_guidata(risposta_chat)
                            if CAMPI_CHAT_GUIDATA_OBBLIGATORI.issubset(scheda_estratta):
                                # Il provider viene sempre preso dalla scelta corrente, non dalla risposta IA.
                                scheda_estratta["provider_ia"] = cervello_chat
                                st.session_state["chat_sidebar_scheda_pronta"] = scheda_estratta
                            st.session_state["chat_sidebar_input_nonce"] = nonce_chat + 1
                            st.rerun()

                scheda_chat_pronta = dict(st.session_state.get("chat_sidebar_scheda_pronta", {}) or {})
                if scheda_chat_pronta:
                    st.success("La scheda è pronta. Puoi applicarla alla sidebar senza copiare i singoli campi.")
                    with st.expander("Vedi i dati che verranno applicati", expanded=False):
                        for campo, valore in scheda_chat_pronta.items():
                            if campo != "provider_ia":
                                st.write(f"**{campo.replace('_', ' ').capitalize()}:** {valore}")
                        st.caption(f"Cervello AI mantenuto: {cervello_chat}")
                    st.button(
                        "✍️ COMPILA SOLO I CAMPI VUOTI",
                        key="applica_chat_sidebar_solo_vuoti",
                        on_click=applica_scheda_chat_guidata,
                        args=(False,),
                        use_container_width=True,
                    )
                    st.caption("Questa scelta conserva tutti i campi della sidebar già compilati.")
                    st.button(
                        "⚠️ SOSTITUISCI I CAMPI GIÀ COMPILATI",
                        key="applica_chat_sidebar_sovrascrivi",
                        on_click=applica_scheda_chat_guidata,
                        args=(True,),
                        use_container_width=True,
                    )
                    st.caption("Sostituisce i dati editoriali già presenti, ma non cambia il cervello scelto nella sidebar.")

                esito_applicazione_chat = st.session_state.get("chat_sidebar_esito_applicazione", "")
                if esito_applicazione_chat:
                    st.success(esito_applicazione_chat)
                st.button(
                    "↺ CHIUDI E REIMPOSTA QUESTA CHAT",
                    key="reset_chat_sidebar_guidata",
                    on_click=reimposta_chat_sidebar_guidata,
                    use_container_width=True,
                )

    # TAB 1: INDICE (CHIRURGIA: FIX SENSO LOGICO E PULIZIA ASSOLUTA DELL'INDICE E CONNESSIONE SARTORIALE)
    with tabs[1]:
        if not sidebar_pronta:
            st.info(
                "Completa tutti i campi obbligatori della barra laterale prima di generare l'indice. "
                "Mancano: " + ", ".join(campi_sidebar_mancanti) + "."
            )
        descrizione_indice = (
            "DeepSeek Pro studia il brief, svolge la ricerca delle fonti con registro visibile e analizza le fonti caricate; poi genera, valuta e corregge l'indice senza usare GPT."
            if usa_deepseek_pro() else
            "Include ricerca preliminare online, generazione, valutazione editoriale e possibili correzioni automatiche dell'indice."
        )
        if pulsante_con_preventivo("genera_indice", L["btn_idx"], CREDIT_COSTS["indice_ricerca_web"] + CREDIT_COSTS["indice_generazione_editoriale"],
                                   descrizione_indice,
                                   disabled=not sidebar_pronta):
            with st.spinner("Ricerca preliminare delle fonti e progettazione dell'indice in corso..."):
                barra_indice = st.progress(
                    5,
                    text="Preparazione del brief editoriale e delle fonti...",
                )

                def aggiorna_avanzamento_indice(percentuale, messaggio):
                    barra_indice.progress(
                        max(0, min(100, int(percentuale))), text=str(messaggio)
                    )

                dossier_ricerca_web = ricerca_preliminare_per_indice(
                val_titolo, val_genere, val_trama, val_goal, lingua_sel,
                f"{val_approfondimenti}\n\n{brief_personalizzazione_progetto()}"
                )
                aggiorna_avanzamento_indice(
                    30,
                    "Fonti studiate: costruzione dell'architettura dell'indice...",
                )
                if dossier_ricerca_web:
                    st.session_state["ultimo_esito_ricerca_preliminare"] = (
                        "Ricerca preliminare completata: il dossier interno guiderà indice e stesura."
                    )
                
                # --- INIZIO NUOVE RIGHE PER TRADUZIONE TERMINI INDICE ---
                trad_termini = {
                    "Italiano": {"parte": "Parte", "cap": "Capitolo"},
                    "English": {"parte": "Part", "cap": "Chapter"},
                    "Español": {"parte": "Parte", "cap": "Capítulo"},
                    "Français": {"parte": "Partie", "cap": "Chapitre"},
                    "Deutsch": {"parte": "Teil", "cap": "Kapitel"},
                    "Română": {"parte": "Partea", "cap": "Capitolul"},
                    "Русский": {"parte": "Часть", "cap": "Глава"},
                    "العربية": {"parte": "الجزء", "cap": "الفصل"},
                    "中文": {"parte": "部分", "cap": "章节"}
                }
                t_parte = trad_termini.get(lingua_sel, trad_termini["Italiano"])["parte"]
                t_cap = trad_termini.get(lingua_sel, trad_termini["Italiano"])["cap"]
                limite_sezioni_totali = PROFILI_LUNGHEZZA_STESURA[val_lunghezza]["max_sezioni"]
                # La Prefazione è aggiunta dal software: riserviamo sempre
                # una voce del limite totale senza sovrascrivere il calcolo
                # già effettuato nella sidebar.
                limite_voci_indice = max(1, limite_sezioni_totali - 1)
                # --- FINE NUOVE RIGHE ---

                # PROMPT BLINDATO PER L'INDICE: Ora prende in carico TUTTI i parametri della sidebar per coerenza assoluta.
                # E include i termini tradotti (f-string per iniezione variabili)
                prompt_idx = f"""Crea l'indice per il libro '{val_titolo}' rigorosamente in lingua {lingua_sel}. 

PARAMETRI EDITORIALI (L'indice deve essere costruito su misura e strettamente attinente a queste caratteristiche):
- Trama/Argomento Centrale: {val_trama}
- Genere Letterario: {val_genere}
- Tipologia di Scrittura: {val_stile}
- Stile di Racconto: {val_narrativa}
- Punto di Vista: {val_pov}
- Obiettivo Emozionale/Pratico: {val_goal}
- Risultato finale desiderato: {val_risultato or "Non dichiarato"}
- Approfondimenti prioritari: {val_approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}

{brief_personalizzazione_progetto()}

{specifica_editoriale}

Gli approfondimenti prioritari devono essere considerati prima di distribuire gli altri argomenti nell'indice. Trasformali in capitoli o sottocapitoli soltanto quando sono pertinenti al libro e assegna loro una collocazione logica, senza creare duplicazioni o voci generiche.

=== DIRETTIVA SPECIFICA DELLA TIPOLOGIA DI SCRITTURA ===
Tipologia selezionata: {val_stile}
{direttiva_indice_selezionata}

=== ARCHITETTURA ADATTIVA AL GENERE ===
{profilo_struttura_indice(val_genere, val_titolo, val_trama, val_goal)}

=== REGOLA DI STESURA DEL GENERE ===
{profilo_genere_stesura(val_genere)}

=== SPECIFICA OPERATIVA PER LA PROGETTAZIONE DELL'INDICE ===
Costruisci l'indice come un progetto editoriale eseguibile, non come un elenco generico.
Ricava dal brief il risultato finale promesso, il pubblico e il livello di partenza, i problemi
concreti, il metodo didattico, i deliverable e i limiti del libro.
Definisci una sequenza dal livello iniziale al risultato finale. Ogni Parte deve avere una
funzione distinta e ogni Capitolo un obiettivo autonomo. Crea un sottocapitolo soltanto quando
apre un argomento davvero diverso: non creare voci separate per esempio, variante, errore comune,
strumento, checklist, verifica o caso studio se possono essere sviluppati bene nella medesima sezione.
Ogni voce deve quindi contenere un argomento completo e concreto, non un singolo passaggio minimo.
Per questo profilo non superare {massimo_parti_editoriali} Parti, {massimo_capitoli_editoriali} Capitoli e
{massimo_sottocapitoli_per_capitolo} sottocapitoli per Capitolo.
Mantieni separati solo passaggi che cambiano realmente competenza, scena, decisione, problema o risultato.
Distribuisci gli argomenti dell'obiettivo e della trama senza anticipare tutto nell'introduzione.
Per strumenti o software soggetti ad aggiornamento, separa principi stabili, funzioni da verificare
e applicazioni. Mantieni coerenza con genere, tipologia, stile, POV, obiettivo e argomento.
L'indice deve permettere di scrivere sezioni dettagliate senza riempitivi.

=== LIMITE ASSOLUTO DI ESTENSIONE ===
Profilo scelto: {val_lunghezza}. L'intero libro può contenere al massimo {limite_sezioni_totali} sezioni, Prefazione inclusa.
La Prefazione viene aggiunta dal software: genera quindi AL MASSIMO {limite_voci_indice} voci qui sotto.
Non superare mai {limite_voci_indice} voci. Non esiste un numero minimo di voci da raggiungere: se il brief non giustifica
questa estensione con argomenti davvero distinti, pubblica un indice più corto ma completo e di qualità. È vietato aggiungere
voci simili, generiche o riempitive soltanto per aumentare la lunghezza del libro.
Conta internamente tutte le Parti, i Capitoli e i sottocapitoli prima di rispondere.

=== BUDGET STRUTTURALE OBBLIGATORIO ===
Per questo profilo usa: {budget_struttura_indice}. Questa fascia è un obiettivo editoriale, non un pretesto per riempitivi.
Se un argomento è collegato a un altro, trattalo nello stesso sottocapitolo invece di creare una nuova voce. Prima dell'output
verifica che ogni voce abbia un contenuto autonomo e che la somma di Parti + Capitoli + sottocapitoli resti nel limite.
"""
                if st.session_state.get("conoscenza_extra"):
                    dossier_fonti = st.session_state.get("brief_fonti_originale") or st.session_state.get("dossier_fonti_ai", "")
                    prompt_idx += (
                        "\n\nMAPPA CONCETTUALE INTERNA DELLE FONTI (GIA' RIELABORATA):\n"
                        "Usa solo principi, verifiche e competenze presenti nella mappa; non riprodurre struttura, "
                        "formulazioni, esempi distintivi o ordine delle fonti. Progetta un percorso nuovo e autonomo, "
                        "senza citare fonti e senza aggiungere argomenti non supportati.\n"
                        f"{dossier_fonti[:7000]}\n"
                    )
                if dossier_ricerca_web:
                    prompt_idx += (
                        "\n\nDOSSIER DELLA RICERCA PRELIMINARE ONLINE (USO INTERNO):\n"
                        "Usa questo dossier per controllare definizioni, priorità, limiti e progressione dell'indice. "
                        "Non citare fonti, URL, siti o riferimenti bibliografici nell'indice né nel libro.\n"
                        f"{dossier_ricerca_web[:7000]}\n"
                    )

                prompt_idx += f"""
REGOLE FONDAMENTALI ED ESCLUSIVE:
0. ATTINENZA ASSOLUTA: Inserisci esclusivamente capitoli e sottocapitoli direttamente pertinenti al titolo, alla trama, al pubblico e all'obiettivo del libro. Non aggiungere sezioni generiche o accessorie come glossario dei termini, elenco di risorse, checklist generiche, bibliografia, link, ringraziamenti, conclusioni vaghe o suggerimenti finali. Ogni voce dell'indice deve sviluppare un argomento reale del libro e poter essere trasformata in contenuto sostanziale.
1. SOLO L'INDICE: Non inserire convenevoli, saluti, introduzioni o conclusioni. L'output deve contenere ESCLUSIVAMENTE la lista dell'indice. Nient'altro.
2. COERENZA ASSOLUTA: I titoli dei capitoli e sottocapitoli devono riflettere perfectly lo stile, il genere e la trama richiesta. Se è un ricettario, l'indice deve sembrare un menu; se è un thriller, i capitoli devono creare suspense.
3. ESTENSIONE PROPORZIONATA: Applica l'architettura adattiva indicata sopra. Non imporre 15-18 capitoli, 3-5 sottocapitoli o 100 pagine a generi che richiedono una struttura diversa. Per un ricettario, il numero di Capitoli deve coincidere con il numero di ricette richiesto e ogni Capitolo deve essere una ricetta, senza sottocapitoli artificiosi. Per la narrativa, non frammentare scene in sottocapitoli riempitivi.
4. STRUTTURA GERARCHICA RIGIDA E PULITA: Usa unicamente ed esattamente questo formato di elencazione, SENZA ASTERISCHI O SIMBOLI STRANI:
   {t_parte} I: [Nome Parte]
   {t_cap} 1: [Nome Capitolo]
   1.1 [Sottocapitolo]
   1.2 [Sottocapitolo]
5. SENSO LOGICO SEQUENZIALE: Il flusso narrativo/didattico deve essere ineccepibile. Parti dalle basi/introduzione, sviluppa il cuore del problema, e concludi con soluzioni o risoluzioni finali.
6. PULIZIA VISIVA: Nessuna descrizione sotto i capitoli. Nessuna punteggiatura anomala. Solo l'elenco nudo e crudo.

7. APPLICAZIONE SILENZIOSA DEI PARAMETRI: Applica rigorosamente le istruzioni della sidebar garantendo una perfetta coerenza editoriale. CRITICO: NON inserire alcun "ragionamento strutturale", commento preliminare o spiegazione. Stampa SOLO ed ESCLUSIVAMENTE la lista dell'indice nuda e cruda.

8. PRATICITÀ ESTREMA E IPER-DETTAGLIO: I titoli devono essere estremamente pratici e orientati all'azione. Niente macro-concetti vaghi. Ogni capitolo e sottocapitolo deve puntare a risolvere un problema specifico, mostrando il "come fare" passo dopo passo, con un taglio estremamente operativo e profondo.

9. COMPLETEZZA SENZA RIEMPITIVI: Rispetta il numero e il formato stabiliti dall'architettura adattiva. Ogni Capitolo deve avere una funzione autonoma. Crea sottocapitoli soltanto quando sviluppano aspetti distinti e non quando ripetono ingredienti, procedimenti, esempi o scene già assegnati. Prima di concludere, conta internamente le voci richieste e verifica che nessuna sia vuota o solo un titolo.

10. ADATTAMENTO AL TIPO DI LIBRO E OUTPUT FINALE: Per manuali tecnici separa fondamenti, strumenti, procedure, verifiche e progetto applicativo. Per manuali pratici inserisci esercizi, checklist e risultati misurabili. Per business, marketing, economia e self-help inserisci framework, casi studio, piani d'azione e criteri di valutazione. Per saggi scientifici o storici separa contesto, tesi, prove, fonti e conclusioni. Per ricettari con un numero dichiarato di ricette, ogni Capitolo deve essere una ricetta e non sono ammessi Capitoli introduttivi su tecniche, ingredienti o sicurezza. Per test prep inserisci teoria, esercizi, simulazioni e soluzioni. Per narrativa costruisci sviluppo di trama, personaggi, conflitto e risoluzione, senza imporre procedure tecniche e con titoli di capitolo specifici del brief. In ogni caso prevedi un output finale coerente con il genere: progetto, piano, esercizio completato, ricetta, simulazione, decisione applicativa, sintesi o conclusione narrativa. Gli esempi devono essere concreti e verificabili secondo il tipo di libro.
"""
                aggiorna_avanzamento_indice(
                    40,
                    "Brief pronto: generazione dell'indice professionale...",
                )
                indice_generato = genera_indice_controllato(
                    prompt_idx, "Senior Book Architect esperto in flow logico-narrativo e design editoriale pulito.",
                    val_genere, val_titolo, val_trama, val_goal, lingua_sel, val_stile, val_narrativa, val_pov,
                    massimo_sezioni=limite_voci_indice,
                    minimo_parti=minimi_struttura_indice[0],
                    minimo_capitoli=minimi_struttura_indice[1],
                    budget_strutturale=budget_struttura_indice,
                    profilo_lunghezza=val_lunghezza,
                    massimo_parti_editoriali=massimo_parti_editoriali,
                    massimo_capitoli_editoriali=massimo_capitoli_editoriali,
                    massimo_sottocapitoli=massimo_sottocapitoli_per_capitolo,
                    aggiorna_stato=aggiorna_avanzamento_indice,
                )
                st.session_state.pop("analisi_voto_indice", None)
                if indice_generato:
                    imposta_indice_progetto(indice_generato)
                    aggiorna_avanzamento_indice(100, "Indice salvato e capitoli sincronizzati.")
                    salva_stesura_generata_in_cloud(opzioni_editor, "indice generato"); st.rerun()
                else:
                    aggiorna_avanzamento_indice(100, "Generazione conclusa: controlla l'esito editoriale sotto.")
                    # Non cancellare mai un indice già presente se la nuova proposta non supera i controlli.
                    mostra_avviso_operativo(
                        "errore", "Indice non pubblicato",
                        "La proposta non ha superato il controllo editoriale oppure non ha restituito una struttura utilizzabile.",
                        "Rivedi il brief o riprova la generazione: l'indice precedente, se presente, e le sezioni già scritte restano invariati.",
                    )
                    with st.expander("Dettaglio dell'esito dell'indice", expanded=False):
                        st.caption(st.session_state.get("ultimo_controllo_indice", "Indice non approvato: riprova con un brief più specifico."))
                
        # Ogni indice esistente viene normalizzato una sola volta con la
        # Prefazione iniziale, anche quando è stato importato o scritto a mano.
        testo_corrente = str(st.session_state.get("indice_raw", "") or "")
        righe_indice_corrente = [riga.strip() for riga in testo_corrente.splitlines() if riga.strip()]
        if righe_indice_corrente and not sezione_prefazione(righe_indice_corrente[0]):
            testo_corrente = imposta_indice_progetto(testo_corrente)
        if st.session_state.get("ultimo_controllo_indice"):
            esito_indice = st.session_state["ultimo_controllo_indice"]
            if re.search(r"\b(?:8|9|10)/10\b", esito_indice):
                st.success(esito_indice)
            elif "richiede" in esito_indice or "non ha raggiunto" in esito_indice:
                st.warning(esito_indice)
        # Una chiave con versione impedisce a Streamlit di riutilizzare la
        # textarea vuota rimasta nel browser prima di un import CSV o di un
        # ripristino cloud. Il contenuto mostrato coincide così sempre con
        # indice_raw, che è la fonte persistente del progetto.
        versione_indice = int(st.session_state.get("indice_widget_version", 0))
        chiave_widget_indice = f"indice_editoriale_{versione_indice}"
        testo_input = st.text_area(
            campo_ui("indice", lingua_sel),
            value=testo_corrente,
            height=400,
            key=chiave_widget_indice,
        )
        
        if testo_input != testo_corrente:
            # Se la UI ricarica e invia stringa vuota per errore, ignoriamo l'aggiornamento, preservando i dati
            if testo_input.strip() == "" and testo_corrente != "":
                pass
            else:
                st.session_state["indice_raw"] = testo_input
                st.session_state.pop("analisi_voto_indice", None)
                
        if st.button(L["btn_sync"]):
            imposta_indice_progetto(st.session_state.get("indice_raw", ""))
            salva_stesura_immediata(opzioni_editor)
            st.rerun()

        indice_da_valutare = st.session_state.get("indice_raw", "").strip()
        if indice_da_valutare:
            if pulsante_con_preventivo("voto_indice", "⭐ VOTO INDICE", CREDIT_COSTS["voto_indice"],
                                       "Analizza l'indice rispetto al brief editoriale.", use_container_width=True):
                with st.spinner("Analisi editoriale dell'indice in corso..."):
                    st.session_state["analisi_voto_indice"] = valuta_indice_editoriale(
                        indice_da_valutare, val_titolo, val_trama, val_genere, val_stile,
                        val_narrativa, val_pov, val_goal, lingua_sel, val_approfondimenti
                    )
                    notifica_sonora("voto_indice_completato", lingua_sel, ripeti=True)
            if st.session_state.get("analisi_voto_indice"):
                st.text_area(
                    campo_controllo_ui(1, lingua_sel),
                    value=st.session_state["analisi_voto_indice"],
                    height=320,
                    key="output_voto_indice"
                )
                if pulsante_con_preventivo("rigenera_indice_voto", "🔄 RIGENERA INDICE SEGUENDO IL VOTO", CREDIT_COSTS["rigenera_indice"],
                                           "Include rigenerazione, valutazione editoriale e possibili correzioni automatiche.", use_container_width=True):
                    with st.spinner("Creazione della proposta migliorata in corso..."):
                        prompt_rigenerazione = f"""Riscrivi esclusivamente l'indice del libro sotto indicato, rigorosamente in lingua {lingua_sel}.

BRIEF
Titolo: {val_titolo}
Argomento: {val_trama}
Genere: {val_genere}
Tipologia: {val_stile}
Stile: {val_narrativa}
POV: {val_pov}
Obiettivo: {val_goal}
Risultato finale desiderato: {val_risultato}
Approfondimenti: {val_approfondimenti or "Nessuno"}
{brief_personalizzazione_progetto()}

LIMITE OBBLIGATORIO: mantieni al massimo {limite_voci_indice} voci nell'output; la Prefazione sarà aggiunta dal software.
Non esiste un numero minimo di voci: pubblica una struttura più breve se è la scelta più completa, leggibile e non ripetitiva.
BUDGET STRUTTURALE: {budget_struttura_indice}. Non introdurre voci inutili per raggiungere un numero.

INDICE ATTUALE
{indice_da_valutare}

VOTO E SUGGERIMENTI DELL'EDITOR
{st.session_state["analisi_voto_indice"]}

Applica tutti i miglioramenti utili, senza introdurre capitoli generici, glossari, bibliografie, risorse o checklist vuote. Mantieni solo argomenti pertinenti. Restituisci soltanto l'indice gerarchico pulito, senza spiegazioni, Markdown o URL."""
                        proposta = genera_indice_controllato(
                            prompt_rigenerazione,
                            "Sei un editor senior. Correggi l'indice con precisione e conserva la coerenza con il brief.",
                            val_genere, val_titolo, val_trama, val_goal, lingua_sel, val_stile, val_narrativa, val_pov,
                            indice_da_superare=indice_da_valutare,
                            massimo_sezioni=limite_voci_indice,
                            minimo_parti=minimi_struttura_indice[0],
                            minimo_capitoli=minimi_struttura_indice[1],
                            budget_strutturale=budget_struttura_indice,
                            profilo_lunghezza=val_lunghezza,
                            massimo_parti_editoriali=massimo_parti_editoriali,
                            massimo_capitoli_editoriali=massimo_capitoli_editoriali,
                            massimo_sottocapitoli=massimo_sottocapitoli_per_capitolo,
                        )
                        if proposta:
                            st.session_state["indice_proposto_dal_voto"] = proposta
                        else:
                            st.error(
                                "La proposta di indice non ha superato il controllo automatico e non è stata applicata. "
                                "Modifica il brief o l'indice attuale e riprova. " +
                                st.session_state.get("ultimo_controllo_indice", "")
                            )
                if st.session_state.get("indice_proposto_dal_voto"):
                    st.text_area(
                        campo_controllo_ui(3, lingua_sel),
                        value=st.session_state["indice_proposto_dal_voto"],
                        height=400,
                        key="output_indice_proposto_dal_voto"
                    )
                    if st.button("✅ APPLICA E SINCRONIZZA LA PROPOSTA", use_container_width=True, key="applica_indice_voto"):
                        st.session_state["indice_backup_prima_della_proposta"] = st.session_state.get("indice_raw", "")
                        imposta_indice_progetto(st.session_state["indice_proposto_dal_voto"])
                        st.session_state.pop("indice_proposto_dal_voto", None)
                        st.session_state.pop("analisi_voto_indice", None)
                        salva_stesura_immediata(opzioni_editor)
                        st.rerun()

    # TAB 2: SCRITTURA E QUIZ (E ORA ANCHE RICETTE)
    with tabs[2]:
        if not lista_cap_base: st.warning(L["msg_err_idx"])
        else:
            # La stesura completa è una coda controllata: una sezione viene
            # generata, salvata e mostrata all'utente prima di autorizzare la
            # successiva. Così pausa e stop operano davvero tra due richieste
            # AI, e nessun testo concluso resta invisibile nel browser.
            sezioni_intero_libro = elenco_sezioni_progetto(lista_cap_base)
            # Un indice valido contiene sempre la Prefazione come prima voce.
            # Se una vecchia sessione/importazione ha una lista non aggiornata,
            # la reinseriamo nella coda senza toccare i capitoli già presenti.
            prefazione_indice = next(
                (sezione for sezione in lista_cap_base if sezione_prefazione(sezione)),
                titolo_prefazione(lingua_sel),
            )
            if not any(sezione_prefazione(sezione) for sezione in sezioni_intero_libro):
                sezioni_intero_libro = [prefazione_indice, *sezioni_intero_libro]
            else:
                sezioni_intero_libro = [
                    prefazione_indice,
                    *[
                        sezione for sezione in sezioni_intero_libro
                        if not sezione_prefazione(sezione)
                    ],
                ]

            def sezione_pronta_per_la_coda(sezione):
                """Usa il medesimo esito del pulsante di scrittura dettagliata.

                Nessuna sezione già salvata come valida può tornare nella coda
                per una soglia più alta applicata soltanto alla stesura intera.
                """
                stesura_in_corso = bool(
                    st.session_state.get("job_scrittura_attivo")
                    or st.session_state.get("job_scrittura_pausa")
                    or st.session_state.get("job_scrittura_in_attesa")
                )
                prefazione_confermata = pulisci_testo_editoriale(
                    st.session_state.get("job_scrittura_prefazione_confermata", "")
                ).strip()
                if sezione_prefazione(sezione) and stesura_in_corso and prefazione_confermata:
                    # Questa è la barriera definitiva contro il bug: prima di
                    # decidere la coda ricostruiamo tutte le copie visibili
                    # dalla Prefazione già approvata. Non si interroga l'AI.
                    testo_corrente = pulisci_testo_editoriale(
                        contenuto_memorizzato_puro(sezione)
                    ).strip()
                    if testo_corrente != prefazione_confermata:
                        scrivi_sezione_stesura_completa(sezione, prefazione_confermata)
                    testo = prefazione_confermata
                else:
                    testo = pulisci_testo_editoriale(
                        leggi_sezione_memorizzata(sezione)
                    ).strip()
                if not testo:
                    return False
                if sezione_prefazione(sezione):
                    return len(testo.split()) >= 30 and not motivo_chiusura_tecnica(testo)
                return True

            manoscritto = memoria_progetto_unica().get("contenuti", {})
            da_generare_libro = [
                sezione for sezione in sezioni_intero_libro
                if not sezione_pronta_per_la_coda(sezione)
            ]
            st.caption(
                f"Stesura completa disponibile: {len(sezioni_intero_libro)} sezioni rilevate. "
                "Il libro viene elaborato con un passaggio visibile e salvato per ogni sezione."
            )
            stima_libro = sum(
                stima_massima_crediti_stesura(sezione, st.session_state["indice_raw"], val_trama, val_goal, val_genere)
                for sezione in da_generare_libro
            )
            if pulsante_con_preventivo(
                "scrivi_tutto_libro", "📚 SCRIVI TUTTO IL LIBRO", f"fino a {stima_libro}",
                f"Saranno scritte {len(da_generare_libro)} sezioni ancora vuote; i contenuti già presenti non verranno modificati.",
                use_container_width=True,
            ):
                if not da_generare_libro:
                    st.info("Il libro risulta già scritto: nessun contenuto è stato sovrascritto.")
                else:
                    st.session_state["job_scrittura_sezioni"] = list(sezioni_intero_libro)
                    st.session_state["job_scrittura_coda"] = list(da_generare_libro)
                    st.session_state["job_scrittura_totale"] = len(da_generare_libro)
                    # Il pulsante "Scrivi tutto il libro" integra la
                    # Prefazione come primo passaggio effettivo. Se era già
                    # presente, ne conserva il testo; se è vuota, la riceverà
                    # una sola volta dal generatore iniziale della coda.
                    testo_prefazione_esistente = pulisci_testo_editoriale(
                        leggi_sezione_memorizzata(prefazione_indice)
                    ).strip()
                    st.session_state["job_scrittura_prefazione_confermata"] = (
                        testo_prefazione_esistente
                        if sezione_pronta_per_la_coda(prefazione_indice)
                        else ""
                    )
                    # La Prefazione non dipende dalla lista ricostruita nei
                    # rerun: viene registrata come passaggio obbligatorio
                    # dell'intero libro prima di ogni capitolo.
                    st.session_state["job_scrittura_prefazione_obbligatoria"] = (
                        prefazione_indice
                        if not sezione_pronta_per_la_coda(prefazione_indice)
                        else ""
                    )
                    st.session_state["job_scrittura_attivo"] = True
                    st.session_state["job_scrittura_pausa"] = False
                    st.session_state["job_scrittura_in_attesa"] = True
                    st.session_state["job_scrittura_prossimo_avvio"] = time.time() + 3.0
                    st.session_state["job_scrittura_fermato"] = False
                    st.session_state["job_scrittura_checkpoint_superati"] = []
                    # Ogni nuova stesura parte con il proprio conteggio di
                    # recuperi: un errore di un progetto precedente non puo'
                    # mai condizionare la prima sezione del nuovo libro.
                    st.session_state["job_scrittura_tentativi"] = {}
                    # Ricevute della stesura corrente: restano valide anche
                    # se Streamlit ridisegna editor o tabs durante la coda.
                    st.session_state["job_scrittura_sezioni_confermate"] = {}
                    # La Prefazione viene sempre avviata e verificata nel
                    # flusso principale prima del timer. Usa però lo stesso
                    # generatore del pulsante "Scrivi contenuto dettagliato".
                    st.session_state["job_scrittura_avvio_protetto_rimanenti"] = 1
                    st.session_state.pop("job_scrittura_checkpoint_richiesto", None)
                    st.session_state.pop("job_scrittura_errore", None)
                    st.session_state.pop("job_scrittura_ultima_completata", None)
                    st.session_state["apri_tab_scrittura_da_stesura"] = True
                    notifica_sonora("avvio_scrittura_completa", lingua_sel, ripeti=True)
                    st.rerun()

            coda_scrittura = [
                sezione for sezione in (st.session_state.get("job_scrittura_coda", []) or [])
                if sezione in sezioni_intero_libro
                and not sezione_pronta_per_la_coda(sezione)
            ]
            prefazione_obbligatoria = str(
                st.session_state.get("job_scrittura_prefazione_obbligatoria", "") or ""
            ).strip()
            if (
                not prefazione_obbligatoria
                and st.session_state.get("job_scrittura_attivo")
                and not sezione_pronta_per_la_coda(prefazione_indice)
            ):
                # Compatibilità con una generazione avviata prima di questa
                # correzione: la prima voce mancante viene recuperata senza
                # obbligare l'utente a cancellare o riavviare il progetto.
                prefazione_obbligatoria = prefazione_indice
                st.session_state["job_scrittura_prefazione_obbligatoria"] = prefazione_indice
            if prefazione_obbligatoria and not sezione_pronta_per_la_coda(prefazione_obbligatoria):
                # Anche se una sessione vecchia, un CSV o una cache avesse
                # costruito la coda senza Prefazione, questa riga la riporta
                # davanti a tutto il resto. Non può essere saltata.
                coda_scrittura = [
                    prefazione_obbligatoria,
                    *[
                        sezione for sezione in coda_scrittura
                        if not sezione_prefazione(sezione)
                    ],
                ]
            elif prefazione_obbligatoria:
                st.session_state.pop("job_scrittura_prefazione_obbligatoria", None)
            st.session_state["job_scrittura_coda"] = list(coda_scrittura)
            totale = max(1, int(st.session_state.get("job_scrittura_totale", len(coda_scrittura) or 1)))
            completati = max(0, totale - len(coda_scrittura))

            # Comandi sempre disponibili nel corpo normale della pagina, non
            # nel fragment che si aggiorna ogni secondo. In questo modo il
            # clic non può scomparire fra due rerun automatici. Una richiesta
            # arrivata mentre l'AI sta rispondendo viene elaborata al termine
            # sicuro della risposta, prima che parta qualsiasi altra sezione.
            if st.session_state.get("job_scrittura_attivo") and coda_scrittura:
                st.caption(
                    "Controlli della stesura: Pausa ferma la coda dopo la sezione eventualmente già in corso; "
                    "Stop definitivo conserva il creato e non avvia altre richieste AI."
                )
                comando_pausa, comando_stop = st.columns(2)
                with comando_pausa:
                    if st.button(
                        "⏸ METTI IN PAUSA", use_container_width=True,
                        key="pausa_scrittura_sempre_visibile",
                    ):
                        st.session_state["job_scrittura_attivo"] = False
                        st.session_state["job_scrittura_pausa"] = True
                        st.session_state["job_scrittura_in_attesa"] = True
                        st.session_state["apri_tab_scrittura_da_pausa"] = True
                        salva_stesura_generata_in_cloud(
                            sezioni_intero_libro, "stesura messa in pausa"
                        )
                        st.rerun()
                with comando_stop:
                    if st.button(
                        "⏹ STOP DEFINITIVO", use_container_width=True,
                        key="stop_scrittura_sempre_visibile",
                    ):
                        st.session_state["job_scrittura_attivo"] = False
                        st.session_state["job_scrittura_pausa"] = False
                        st.session_state["job_scrittura_in_attesa"] = False
                        st.session_state["job_scrittura_fermato"] = True
                        st.session_state["job_scrittura_interrotte"] = list(coda_scrittura)
                        st.session_state["job_scrittura_coda"] = []
                        st.session_state.pop("job_scrittura_checkpoint_richiesto", None)
                        salva_stesura_generata_in_cloud(
                            sezioni_intero_libro, "stesura interrotta definitivamente"
                        )
                        st.rerun()

            # Avvio protetto: le prime voci sono salvate dal normale ciclo
            # dell'app, non dal fragment. Questo elimina il caso in cui il
            # timer ricaricava la prima sezione prima che editor, anteprima e
            # memoria unica potessero registrarla.
            avvio_protetto = int(
                st.session_state.get("job_scrittura_avvio_protetto_rimanenti", 0) or 0
            )
            if st.session_state.get("job_scrittura_attivo") and coda_scrittura and avvio_protetto > 0:
                sezione_iniziale = coda_scrittura[0]
                # L'avvio protetto salva le prime sezioni fuori dal timer, ma
                # non puo' ignorare una scelta esplicita dell'utente. Se e'
                # stata selezionata una modalita checkpoint, la pausa avviene
                # anche davanti alla prima Parte o alla conclusione iniziale.
                if richiede_checkpoint_personale(sezione_iniziale):
                    st.session_state["job_scrittura_attivo"] = False
                    st.session_state["job_scrittura_pausa"] = True
                    st.session_state["job_scrittura_in_attesa"] = True
                    st.session_state["job_scrittura_checkpoint_richiesto"] = sezione_iniziale
                    salva_stesura_generata_in_cloud(
                        sezioni_intero_libro, "pausa guidata per personalizzazione"
                    )
                    st.rerun()
                st.info(
                    f"Scrittura iniziale: {sezione_iniziale}. "
                    "Usa lo stesso controllo e salvataggio di Scrivi contenuto dettagliato."
                )
                try:
                    contenuto_iniziale = genera_e_conferma_sezione_del_libro(
                        sezione_iniziale, st.session_state["indice_raw"], val_trama, val_genere,
                        val_stile, val_narrativa, val_pov, val_goal, lingua_sel,
                        val_approfondimenti, val_lunghezza,
                    )
                    if not str(contenuto_iniziale or "").strip():
                        raise RuntimeError("il cervello selezionato non ha restituito testo")
                    st.session_state[CHIAVE_SELETTORE_EDITOR] = sezione_iniziale
                    st.session_state[CHIAVE_SEZIONE_EDITOR_ATTIVA] = None
                    st.session_state["job_scrittura_ultima_completata"] = sezione_iniziale
                    if sezione_prefazione(sezione_iniziale):
                        st.session_state.pop("job_scrittura_prefazione_obbligatoria", None)
                    st.session_state["job_scrittura_coda"] = coda_scrittura[1:]
                    st.session_state["job_scrittura_avvio_protetto_rimanenti"] = avvio_protetto - 1
                    st.session_state["job_scrittura_prossimo_avvio"] = time.time() + 1.0
                    st.session_state.setdefault("job_scrittura_tentativi", {}).pop(sezione_iniziale, None)
                    salva_stesura_generata_in_cloud(
                        sezioni_intero_libro, "sezione iniziale del libro generata"
                    )
                    st.rerun()
                except Exception as exc:
                    st.session_state["job_scrittura_attivo"] = False
                    st.session_state["job_scrittura_pausa"] = True
                    st.session_state["job_scrittura_in_attesa"] = True
                    st.session_state["job_scrittura_errore"] = (
                        f"{sezione_iniziale}: {exc}. La sezione non e' stata saltata; "
                        "puoi riprendere la generazione senza perdere le altre."
                    )
                    salva_stesura_generata_in_cloud(
                        sezioni_intero_libro, "errore nella sezione iniziale"
                    )
                    st.rerun()

            if st.session_state.get("job_scrittura_attivo") and coda_scrittura:
                # ``st.fragment`` è il timer nativo di Streamlit: aggiorna il
                # server senza clic simulati dal browser. La compatibilità con
                # versioni meno recenti usa il vecchio nome ufficiale.
                decoratore_fragment = getattr(st, "fragment", None) or getattr(st, "experimental_fragment", None)
                if not decoratore_fragment:
                    st.error("Questa versione di Streamlit non supporta l'avanzamento automatico. Aggiorna Streamlit e riprova.")
                else:
                    @decoratore_fragment(run_every=1.0)
                    def gestisci_stesura_automatica():
                        """Esegue una sola sezione e ricostruisce la coda ad ogni ciclo.

                        Un fragment Streamlit conserva le variabili del suo
                        primo disegno. Per questo non usa mai ``coda_scrittura``
                        catturata fuori dalla funzione: dopo ogni salvataggio
                        legge di nuovo la memoria stabile e passa davvero alla
                        voce successiva, incluse Prefazione e Parte iniziale.
                        """
                        coda_attuale = [
                            sezione
                            for sezione in (st.session_state.get("job_scrittura_coda", []) or [])
                            if sezione in sezioni_intero_libro
                            and not sezione_pronta_per_la_coda(sezione)
                        ]
                        st.session_state["job_scrittura_coda"] = list(coda_attuale)
                        if not coda_attuale:
                            st.session_state["job_scrittura_attivo"] = False
                            st.session_state["job_scrittura_pausa"] = False
                            st.session_state["job_scrittura_in_attesa"] = False
                            st.rerun(scope="app")

                        totale_attuale = max(
                            1,
                            int(st.session_state.get("job_scrittura_totale", len(coda_attuale) or 1)),
                        )
                        completati_attuali = max(0, totale_attuale - len(coda_attuale))
                        sezione_corrente = coda_attuale[0]
                        st.progress(
                            int(completati_attuali / totale_attuale * 100),
                            text=(
                                f"Stesura automatica: completate {completati_attuali} di "
                                f"{totale_attuale} sezioni."
                            ),
                        )
                        ultima_completata = st.session_state.get("job_scrittura_ultima_completata", "")
                        prossima_esecuzione = float(
                            st.session_state.get("job_scrittura_prossimo_avvio", 0) or 0
                        )
                        if time.time() < prossima_esecuzione:
                            if ultima_completata:
                                testo_ultima = pulisci_testo_editoriale(
                                    contenuto_memorizzato_puro(ultima_completata)
                                )
                                st.success(
                                    f"Sezione salvata e pronta: {ultima_completata}. "
                                    "È già selezionata nell'Editor di Testo Professionale ed è presente nell'Anteprima."
                                )
                                with st.expander("📄 Visualizza l'ultima sezione generata", expanded=True):
                                    st.markdown(f"#### {ultima_completata}")
                                    st.write(testo_ultima)
                            else:
                                st.info(f"Pronto a iniziare dalla sezione: {sezione_corrente}.")
                            st.caption(
                                f"La prossima sezione partirà automaticamente: {sezione_corrente}. "
                                "Usa i comandi sempre visibili sopra per metterla in pausa o fermarla."
                            )
                            return

                        if richiede_checkpoint_personale(sezione_corrente):
                            st.session_state["job_scrittura_attivo"] = False
                            st.session_state["job_scrittura_pausa"] = True
                            st.session_state["job_scrittura_in_attesa"] = True
                            st.session_state["job_scrittura_checkpoint_richiesto"] = sezione_corrente
                            salva_stesura_generata_in_cloud(
                                sezioni_intero_libro, "pausa guidata per personalizzazione"
                            )
                            st.rerun(scope="app")

                        st.info(
                            f"Elaborazione in corso: {sezione_corrente}. "
                            "Attendi il salvataggio della sezione corrente."
                        )
                        try:
                            contenuto_generato = genera_e_conferma_sezione_del_libro(
                                sezione_corrente, st.session_state["indice_raw"], val_trama, val_genere,
                                val_stile, val_narrativa, val_pov, val_goal, lingua_sel,
                                val_approfondimenti, val_lunghezza,
                            )
                            if not str(contenuto_generato or "").strip():
                                raise RuntimeError("il cervello selezionato non ha restituito testo")
                            st.session_state[CHIAVE_SELETTORE_EDITOR] = sezione_corrente
                            st.session_state[CHIAVE_SEZIONE_EDITOR_ATTIVA] = None
                            st.session_state["job_scrittura_ultima_completata"] = sezione_corrente
                            st.session_state["job_scrittura_coda"] = coda_attuale[1:]
                            st.session_state["job_scrittura_in_attesa"] = True
                            st.session_state["job_scrittura_prossimo_avvio"] = time.time() + 3.0
                            st.session_state.setdefault("job_scrittura_tentativi", {}).pop(sezione_corrente, None)
                            salva_stesura_generata_in_cloud(
                                sezioni_intero_libro, "sezione del libro generata"
                            )
                            st.rerun(scope="app")
                        except Exception as exc:
                            tentativi = st.session_state.setdefault("job_scrittura_tentativi", {})
                            numero_tentativi = int(tentativi.get(sezione_corrente, 0))
                            if numero_tentativi < 1:
                                tentativi[sezione_corrente] = numero_tentativi + 1
                                st.session_state["job_scrittura_attivo"] = True
                                st.session_state["job_scrittura_pausa"] = False
                                st.session_state["job_scrittura_in_attesa"] = True
                                st.session_state["job_scrittura_errore"] = (
                                    f"{sezione_corrente}: primo tentativo non riuscito ({exc}). "
                                    "Il software riprova automaticamente senza saltare la sezione."
                                )
                                st.session_state["job_scrittura_prossimo_avvio"] = time.time() + 2.0
                                st.rerun(scope="app")

                            st.session_state["job_scrittura_attivo"] = False
                            st.session_state["job_scrittura_pausa"] = True
                            st.session_state["job_scrittura_in_attesa"] = True
                            st.session_state["job_scrittura_errore"] = f"{sezione_corrente}: {exc}"
                            notifica_sonora("errore_scrittura", lingua_sel, ripeti=True)
                            st.rerun(scope="app")

                    gestisci_stesura_automatica()

            if st.session_state.get("job_scrittura_pausa") and st.session_state.get("job_scrittura_coda"):
                rimanenti = len(st.session_state["job_scrittura_coda"])
                st.warning(f"Generazione in pausa: restano {rimanenti} sezioni. Le sezioni concluse sono già leggibili e salvate.")
                checkpoint_richiesto = str(st.session_state.get("job_scrittura_checkpoint_richiesto", "") or "").strip()
                chiave_nota_checkpoint = ""
                if checkpoint_richiesto:
                    testi_personalizzazione = etichette_personalizzazione(lingua_sel)
                    st.info(
                        f"Pausa guidata prima di: {checkpoint_richiesto}. "
                        "Puoi aggiungere un'indicazione specifica oppure continuare senza modifiche."
                    )
                    chiave_nota_checkpoint = (
                        "nota_checkpoint_" + hashlib.sha256(checkpoint_richiesto.encode("utf-8")).hexdigest()[:12]
                    )
                    st.text_area(
                        testi_personalizzazione["nota"], height=100, key=chiave_nota_checkpoint,
                        placeholder="Esempio: usa un caso realistico, chiarisci questo dubbio, evita questo esempio.",
                    )
                if st.session_state.get("job_scrittura_errore"):
                    sezione_con_errore = checkpoint_richiesto or (
                        coda_scrittura[0] if coda_scrittura else "Sezione in attesa"
                    )
                    mostra_avviso_operativo(
                        "errore", f"Generazione sospesa — {sezione_con_errore}",
                        "La sezione corrente non ha completato la richiesta. "
                        "Il sistema non l'ha saltata e non ha cancellato le sezioni precedenti.",
                        "Premi RIPRENDI GENERAZIONE: verrà riprovata prima la sezione indicata.",
                    )
                    with st.expander("Dettaglio tecnico dell'ultimo errore", expanded=False):
                        st.caption(str(st.session_state["job_scrittura_errore"]))
                riprendi, stop_pausa = st.columns(2)
                with riprendi:
                    if st.button("▶ RIPRENDI GENERAZIONE", use_container_width=True, key="riprendi_scrittura_libro"):
                        sezione_da_riavviare = checkpoint_richiesto or (
                            coda_scrittura[0] if coda_scrittura else ""
                        )
                        if checkpoint_richiesto:
                            registra_nota_checkpoint(
                                checkpoint_richiesto, st.session_state.get(chiave_nota_checkpoint, "")
                            )
                            superati = list(st.session_state.get("job_scrittura_checkpoint_superati", []) or [])
                            if checkpoint_richiesto not in superati:
                                superati.append(checkpoint_richiesto)
                            st.session_state["job_scrittura_checkpoint_superati"] = superati
                            # Lasciapassare immediato per evitare che un
                            # vecchio stato del widget riattivi la stessa pausa
                            # nel rerun subito successivo al clic su Riprendi.
                            st.session_state["job_scrittura_checkpoint_da_superare"] = checkpoint_richiesto
                            st.session_state.pop("job_scrittura_checkpoint_richiesto", None)
                        st.session_state["job_scrittura_attivo"] = True
                        st.session_state["job_scrittura_pausa"] = False
                        st.session_state["job_scrittura_in_attesa"] = True
                        st.session_state["apri_tab_scrittura_da_stesura"] = True
                        st.session_state["job_scrittura_prossimo_avvio"] = time.time() + 0.2
                        # La voce sospesa riparte dal flusso principale, che
                        # la scrive e la salva prima di restituire il controllo
                        # al timer automatico. Non viene mai saltata.
                        st.session_state["job_scrittura_avvio_protetto_rimanenti"] = 1
                        if sezione_da_riavviare:
                            st.session_state.setdefault("job_scrittura_tentativi", {}).pop(
                                sezione_da_riavviare, None
                            )
                        st.session_state.pop("job_scrittura_errore", None)
                        st.rerun()
                with stop_pausa:
                    if st.button("⏹ STOP DEFINITIVO", use_container_width=True, key="stop_scrittura_in_pausa"):
                        st.session_state["job_scrittura_attivo"] = False
                        st.session_state["job_scrittura_pausa"] = False
                        st.session_state["job_scrittura_in_attesa"] = False
                        st.session_state["job_scrittura_fermato"] = True
                        st.session_state["job_scrittura_interrotte"] = list(st.session_state["job_scrittura_coda"])
                        st.session_state["job_scrittura_coda"] = []
                        salva_stesura_generata_in_cloud(sezioni_intero_libro, "stesura interrotta dall'utente")
                        st.rerun()
            elif st.session_state.get("job_scrittura_fermato"):
                fermate = len(st.session_state.get("job_scrittura_interrotte", []) or [])
                st.warning(
                    f"Stesura interrotta definitivamente. Le sezioni già generate restano intatte; "
                    f"ne risultano non avviate {fermate}. Puoi ripartire in futuro con SCRIVI TUTTO IL LIBRO."
                )
            elif not st.session_state.get("job_scrittura_coda") and st.session_state.get("job_scrittura_totale"):
                st.session_state["job_scrittura_attivo"] = False
                st.session_state["job_scrittura_pausa"] = False
                st.session_state["job_scrittura_in_attesa"] = False
                st.success("Libro completato: tutte le sezioni previste sono state generate e salvate.")

            with st.expander("🔎 Ricerca e sostituzione nel libro", expanded=False):
                st.caption("Cerca un termine nelle sezioni già scritte e sostituiscilo in tutto il manoscritto. L'operazione non avvia l'AI e non consuma crediti.")
                cerca_globale = st.text_input(campo_ui("cerca", lingua_sel), key="ricerca_globale_testo")
                sostituisci_globale = st.text_input(campo_ui("sostituisci", lingua_sel), key="ricerca_globale_sostituzione")
                rispetta_maiuscole = st.checkbox(campo_ui("maiuscole", lingua_sel), key="ricerca_globale_maiuscole")
                if cerca_globale:
                    flag_ricerca = 0 if rispetta_maiuscole else re.IGNORECASE
                    occorrenze, sezioni_trovate = 0, []
                    for sezione in opzioni_editor:
                        testo_sezione = leggi_sezione_memorizzata(sezione)
                        trovate = len(re.findall(re.escape(cerca_globale), testo_sezione, flags=flag_ricerca))
                        if trovate:
                            occorrenze += trovate
                            sezioni_trovate.append(sezione)
                    st.info(f"Trovate {occorrenze} occorrenze in {len(sezioni_trovate)} sezioni.")
                    conferma_sostituzione = st.checkbox(
                        campo_controllo_ui(4, lingua_sel),
                        key="ricerca_globale_conferma",
                    )
                    if st.button(
                        "🔁 SOSTITUISCI IN TUTTO IL LIBRO",
                        type="primary",
                        disabled=not conferma_sostituzione or occorrenze == 0,
                        key="ricerca_globale_applica",
                    ):
                        for sezione in sezioni_trovate:
                            testo_sezione = leggi_sezione_memorizzata(sezione)
                            scrivi_sezione_memorizzata(sezione, re.sub(
                                re.escape(cerca_globale),
                                lambda _match: sostituisci_globale,
                                testo_sezione,
                                flags=flag_ricerca,
                            ))
                        st.success(f"Sostituzione completata: {occorrenze} modifiche in {len(sezioni_trovate)} sezioni. Il salvataggio automatico verrà aggiornato.")
                        salva_stesura_immediata(opzioni_editor)
                        st.rerun()

            # Il selettore ha una chiave esplicita e una callback: così il
            # testo della sezione scelta viene caricato dal progetto prima che
            # l'Editor di Testo Professionale sia disegnato.
            correzione_preparata = st.session_state.get("correzione_finale_da_preparare")
            if correzione_preparata:
                sezione_da_correggere = str(correzione_preparata.get("sezione", "")).strip()
                if sezione_da_correggere in opzioni_editor:
                    st.session_state[CHIAVE_SELETTORE_EDITOR] = sezione_da_correggere
                    st.session_state[CHIAVE_SEZIONE_EDITOR_ATTIVA] = None
                    st.session_state[f"mod_{chiave_sezione(sezione_da_correggere)}"] = correzione_preparata.get("istruzione", "")
                    st.session_state["messaggio_correzione_finale"] = (
                        f"Sezione pronta nell'editor: {sezione_da_correggere}. Controlla l'istruzione e premi RIELABORA CON IA quando vuoi."
                    )
                    st.session_state.pop("correzione_finale_da_preparare", None)
                else:
                    st.session_state["messaggio_correzione_finale"] = (
                        "La sezione segnalata non è presente nell'indice attuale: sincronizza l'indice e riprova."
                    )
            if st.session_state.get(CHIAVE_SELETTORE_EDITOR) not in opzioni_editor:
                st.session_state[CHIAVE_SELETTORE_EDITOR] = opzioni_editor[0]
            sez_scelta = st.selectbox(
                L["lbl_sec"], opzioni_editor,
                key=CHIAVE_SELETTORE_EDITOR,
                on_change=prepara_sezione_editor_selezionata,
            )
            if st.session_state.get(CHIAVE_SEZIONE_EDITOR_ATTIVA) != sez_scelta:
                # Primo caricamento oppure indice modificato: il widget di
                # testo non esiste ancora in questo rerun, quindi può essere
                # riempito con certezza senza sovrascrivere modifiche manuali.
                st.session_state[chiave_sezione(sez_scelta)] = contenuto_memorizzato_puro(sez_scelta)
                st.session_state[CHIAVE_SEZIONE_EDITOR_ATTIVA] = sez_scelta
            k_sessione = chiave_sezione(sez_scelta)
            mostra_suggerimento_editoriale_contestuale(
                sez_scelta, opzioni_editor, val_goal, val_trama
            )
            sottocapitoli_capitolo = individua_sottocapitoli_del_capitolo(sez_scelta, lista_cap_base)
            if sottocapitoli_capitolo:
                # Il capitolo è una vera sezione editoriale introduttiva, non
                # un semplice contenitore: il comando collettivo deve redigerlo
                # insieme ai suoi sottocapitoli.
                sezioni_del_capitolo = [sez_scelta] + sottocapitoli_capitolo
                chiave_audit_capitolo = f"audit_fatti_{sez_scelta.replace(' ', '_').replace('.', '')}"
                st.caption(f"Capitolo selezionato: verranno elaborati il testo del capitolo e {len(sottocapitoli_capitolo)} sottocapitoli ancora vuoti.")
                esito_sottocapitoli = st.session_state.pop("messaggio_stesura_sottocapitoli", None)
                if esito_sottocapitoli:
                    if esito_sottocapitoli.get("errori"):
                        st.warning(esito_sottocapitoli["testo"])
                    else:
                        st.success(esito_sottocapitoli["testo"])
                stima_sottocapitoli = sum(
                    stima_massima_crediti_stesura(sezione_capitolo, st.session_state['indice_raw'], val_trama, val_goal, val_genere)
                    for sezione_capitolo in sezioni_del_capitolo
                    if not leggi_sezione_memorizzata(sezione_capitolo).strip()
                ) + 1
                if pulsante_con_preventivo(f"scrivi_sottocapitoli_{k_sessione}", "📝 SCRIVI CAPITOLO E TUTTI I SOTTOCAPITOLI", f"fino a {stima_sottocapitoli}",
                                           "Il totale esatto dipende dal capitolo e dai sottocapitoli ancora vuoti; include un eventuale controllo dei fatti.", use_container_width=True):
                    da_generare = []
                    gia_presenti = 0
                    for sezione_capitolo in sezioni_del_capitolo:
                        # La memoria stabile ha la precedenza sul singolo widget: in questo
                        # modo un cambio di sezione non rende di nuovo "vuoto" un testo esistente.
                        if leggi_sezione_memorizzata(sezione_capitolo).strip():
                            gia_presenti += 1
                        else:
                            da_generare.append(sezione_capitolo)
                    if not da_generare:
                        st.info("Il capitolo e tutti i suoi sottocapitoli sono già presenti: nessun contenuto è stato sovrascritto.")
                    else:
                        avanzamento = st.progress(0, text="Preparazione della stesura del capitolo...")
                        completati = []
                        errori = []
                        for posizione, sottocapitolo in enumerate(da_generare, start=1):
                            avanzamento.progress(
                                int((posizione - 1) / len(da_generare) * 100),
                                text=f"Scrittura di {sottocapitolo} ({posizione}/{len(da_generare)})..."
                            )
                            prompt = crea_prompt_stesura_sezione(
                                sottocapitolo, st.session_state['indice_raw'], val_trama, val_genere,
                                val_stile, val_narrativa, val_pov, val_goal, lingua_sel, val_approfondimenti, val_lunghezza
                            )
                            try:
                                contenuto = genera_contenuto_editoriale(
                                    prompt, S_PROMPT, sottocapitolo, st.session_state['indice_raw'], val_trama,
                                    val_genere, val_goal, lingua_sel, val_lunghezza
                                )
                                if not str(contenuto or "").strip() or str(contenuto).lstrip().upper().startswith("ERRORE:"):
                                    errori.append(f"{sottocapitolo}: {str(contenuto or 'risposta vuota')[:180]}")
                                    continue
                                # Non sostituiamo mai testi già presenti. Qui arrivano soltanto
                                # sottocapitoli dichiarati vuoti al momento del click.
                                scrivi_sezione_memorizzata(sottocapitolo, contenuto)
                                # Protezione immediata: non attendere che il
                                # capitolo intero termini. Se la pagina si
                                # aggiorna durante il ciclo, ogni sezione già
                                # completata è recuperabile dall'account.
                                salva_stesura_generata_in_cloud(
                                    opzioni_editor, "sezione generata"
                                )
                                completati.append(sottocapitolo)
                            except Exception as exc:
                                errori.append(f"{sottocapitolo}: {exc}")

                        avanzamento.progress(100, text="Elaborazione dei sottocapitoli conclusa.")
                        if completati:
                            contenuti_capitolo = [
                                (sezione_capitolo, leggi_sezione_memorizzata(sezione_capitolo))
                                for sezione_capitolo in sezioni_del_capitolo
                            ]
                            # Il controllo fatti è accessorio: se non è disponibile non può
                            # annullare né nascondere i sottocapitoli appena generati.
                            try:
                                st.session_state[chiave_audit_capitolo] = audit_fatti_capitolo(
                                    sez_scelta, contenuti_capitolo, lingua_sel
                                )
                            except Exception as exc:
                                st.session_state[chiave_audit_capitolo] = f"Controllo fatti non disponibile: {exc}"

                        capitolo_completato = sez_scelta in completati
                        sottocapitoli_completati = len([s for s in completati if s != sez_scelta])
                        messaggio = (
                            f"Capitolo {'completato' if capitolo_completato else 'già presente'}; "
                            f"completati {sottocapitoli_completati} sottocapitoli."
                        )
                        if gia_presenti:
                            messaggio += f" Conservati senza modifiche: {gia_presenti}."
                        if errori:
                            messaggio += " Alcune sezioni non sono state generate; riprova solo quelle indicate."
                        if completati:
                            salva_stesura_generata_in_cloud(opzioni_editor, "capitolo generato")
                            st.session_state["messaggio_stesura_sottocapitoli"] = {
                                "testo": messaggio,
                                "errori": bool(errori),
                            }
                            st.session_state["dettaglio_errori_sottocapitoli"] = errori
                            st.rerun()
                        st.error("Nessuna sezione del capitolo è stata scritta. " + (" ".join(errori) if errori else "Riprova tra poco."))
                dettaglio_errori_sottocapitoli = st.session_state.pop("dettaglio_errori_sottocapitoli", [])
                if dettaglio_errori_sottocapitoli:
                    with st.expander("Dettaglio dei sottocapitoli da riprovare", expanded=False):
                        for dettaglio in dettaglio_errori_sottocapitoli:
                            st.write("- " + dettaglio)
                if pulsante_con_preventivo(f"controlla_fatti_{k_sessione}", "🔎 CONTROLLA I FATTI DEL CAPITOLO", CREDIT_COSTS["audit_fatti_capitolo"],
                                           "Verifica online solo i dati aggiornabili del capitolo selezionato.", use_container_width=True):
                    contenuti_capitolo = [
                        (sottocapitolo, leggi_sezione_memorizzata(sottocapitolo))
                        for sottocapitolo in sottocapitoli_capitolo
                    ]
                    with st.spinner("Controllo online mirato dei soli dati aggiornabili del capitolo..."):
                        st.session_state[chiave_audit_capitolo] = audit_fatti_capitolo(
                            sez_scelta, contenuti_capitolo, lingua_sel
                        )
                if st.session_state.get(chiave_audit_capitolo):
                    with st.expander("🔎 Esito del controllo fatti del capitolo", expanded=False):
                        st.write(st.session_state[chiave_audit_capitolo])
            messaggio_stesura = st.session_state.pop("messaggio_stesura_sezione", "")
            if messaggio_stesura:
                if messaggio_stesura.startswith("La sezione non è stata salvata:"):
                    mostra_avviso_operativo(
                        "errore", "Generazione della sezione sospesa", messaggio_stesura,
                        "Riprova la stessa sezione oppure usa SALVA SESSIONE dopo avere verificato il testo.",
                    )
                else:
                    st.success(messaggio_stesura)
            messaggio_correzione_finale = st.session_state.pop("messaggio_correzione_finale", "")
            if messaggio_correzione_finale:
                st.info(messaggio_correzione_finale)
            c1, c2, c3 = st.columns([2, 2, 1])
            with c1:
                if pulsante_con_preventivo(f"scrivi_sezione_{k_sessione}", L["btn_write"], f"fino a {stima_massima_crediti_stesura(sez_scelta, st.session_state['indice_raw'], val_trama, val_goal, val_genere)}",
                                           f"Verrà generata o sostituita la sezione selezionata: {sez_scelta}."):
                    sezione_salvata = False
                    with st.spinner(L["msg_run"]):
                        try:
                            contenuto_generato = scrivi_contenuto_dettagliato(
                                sez_scelta, st.session_state['indice_raw'], val_trama, val_genere,
                                val_stile, val_narrativa, val_pov, val_goal, lingua_sel,
                                val_approfondimenti, val_lunghezza,
                            )
                            salvata_nel_cloud = salva_stesura_generata_in_cloud(opzioni_editor, "sezione generata")
                            st.session_state["messaggio_stesura_sezione"] = (
                                f"Sezione salvata nel tuo account: {sez_scelta}."
                                if salvata_nel_cloud else
                                f"Sezione creata: {sez_scelta}. Salvataggio nel tuo account da riprovare con SALVA SESSIONE."
                            )
                            sezione_salvata = True
                        except Exception as exc:
                            st.session_state["messaggio_stesura_sezione"] = (
                                f"La sezione non è stata salvata: {exc}. Le altre sezioni restano invariate."
                            )
                            mostra_avviso_operativo(
                                "errore", f"Sezione: {sez_scelta}",
                                "La richiesta non ha prodotto un contenuto che possa essere salvato.",
                                "Riprova questa sola sezione: le altre restano disponibili e non vengono riscritte.",
                            )
                    if sezione_salvata:
                        # Il rerun richiama l'editor prima che il widget venga
                        # disegnato e applica la copia protetta appena scritta.
                        st.rerun()
            with c2:
                istr = st.text_input(L["btn_edit"], key=f"mod_{k_sessione}", placeholder="Es: Potenzia l'esposizione...")
                if pulsante_con_preventivo(f"rigenera_sezione_{k_sessione}", L["btn_edit"] + " 🪄", f"fino a {stima_massima_crediti_stesura(sez_scelta, st.session_state['indice_raw'], val_trama, val_goal, val_genere)}",
                                           f"Verrà rielaborata solo la sezione selezionata: {sez_scelta}."):
                    if k_sessione in st.session_state:
                        prompt_rigenerazione = (
                            f"Rielabora con focus su: {istr}. Mantieni categoricamente la lingua {lingua_sel}, "
                            f"il POV ({val_pov}) e una lunghezza di {PROFILI_LUNGHEZZA_STESURA[val_lunghezza]['parole']}. "
                            "Non allungare il testo con ripetizioni; conserva solo dettagli pertinenti alla sezione. "
                            "Non usare punteggiatura anomala né riscrivere il titolo all'inizio. Non inserire URL, link, "
                            f"citazioni o sezioni bibliografiche. Testo da modificare:\n{st.session_state[k_sessione]}"
                        )
                        scrivi_sezione_memorizzata(sez_scelta, genera_contenuto_editoriale(
                            prompt_rigenerazione, S_PROMPT, sez_scelta, st.session_state['indice_raw'], val_trama,
                            val_genere, val_goal, lingua_sel, val_lunghezza,
                        ))
                        salva_stesura_generata_in_cloud(opzioni_editor, "sezione rielaborata")
                        st.rerun()
            with c3:
                if pulsante_con_preventivo(f"quiz_{k_sessione}", "🧠 QUIZ", CREDIT_COSTS["scrittura_sezione"],
                                           "Saranno aggiunte 10 domande alla sezione selezionata."):
                    if k_sessione in st.session_state:
                        with st.spinner("Generazione Quiz didattico..."):
                            res_q = chiedi_gpt(f"Crea quiz di 10 domande in lingua {lingua_sel} dando del {val_pov} al lettore su:\n{st.session_state[k_sessione]}", "Learning Expert.")
                            scrivi_sezione_memorizzata(sez_scelta, st.session_state[k_sessione] + f"\n\nTEST DI VALUTAZIONE\n\n" + pulisci_testo_editoriale(res_q))
                            salva_stesura_generata_in_cloud(opzioni_editor, "quiz generato"); st.rerun()

                # --- INIZIO NUOVE RIGHE PER TRADUZIONE ESEMPI ---
                trad_esempi = {
                    "Italiano": {"btn": "💡 10 ESEMPI", "titolo": "### 💡 10 ESEMPI PRATICI"},
                    "English": {"btn": "💡 10 EXAMPLES", "titolo": "### 💡 10 PRACTICAL EXAMPLES"},
                    "Español": {"btn": "💡 10 EJEMPLOS", "titolo": "### 💡 10 EJEMPLOS PRÁCTICOS"},
                    "Français": {"btn": "💡 10 EXEMPLES", "titolo": "### 💡 10 EXEMPLES PRATIQUES"},
                    "Deutsch": {"btn": "💡 10 BEISPIELE", "titolo": "### 💡 10 PRAKTISCHE BEISPIELE"},
                    "Română": {"btn": "💡 10 EXEMPLE", "titolo": "### 💡 10 EXEMPLE PRACTICE"},
                    "Русский": {"btn": "💡 10 ПРИМЕРОВ", "titolo": "### 💡 10 ПРАКТИЧЕСКИХ ПРИМЕРОВ"},
                    "العربية": {"btn": "💡 10 أمثلة", "titolo": "### 💡 10 أمثلة عملية"},
                    "中文": {"btn": "💡 10 个例子", "titolo": "### 💡 10 个实际例子"}
                }
                t_btn_ese = trad_esempi.get(lingua_sel, trad_esempi["Italiano"])["btn"]
                t_tit_ese = trad_esempi.get(lingua_sel, trad_esempi["Italiano"])["titolo"]
                # --- FINE NUOVE RIGHE ---

                # --- AGGIUNTA PULSANTE GENERATORE ESEMPI ---
                if pulsante_con_preventivo(f"esempi_{k_sessione}", t_btn_ese, CREDIT_COSTS["scrittura_sezione"],
                                           "Saranno aggiunti 10 esempi pratici alla sezione selezionata."):
                    if k_sessione in st.session_state:
                        with st.spinner(f"Creazione 10 esempi in {lingua_sel}..."):
                            mem_esempi = st.session_state.get(k_sessione, "")
                            p_esempi = f"""Genera ESATTAMENTE 10 ESEMPI PRATICI, unici e dettagliati rigorosamente in lingua {lingua_sel} per la sezione '{sez_scelta}', perfettamente coerenti con l'argomento: '{val_trama}' e il genere '{val_genere}'.
                            Usa il punto di vista '{val_pov}'.
                            
                            [ATTENZIONE ALLA LINGUA]: È TASSATIVO che l'intero contenuto, inclusi i titoli, sia scritto in {lingua_sel}.
                            
                            [REGOLA ANTI-RIPETIZIONE]: Leggi i contenuti già generati qui sotto e NON RIPETERLI MAI. Crea scenari, casi studio o applicazioni completamente nuovi:
                            
                            {mem_esempi[-4000:]}"""
                            
                            res_e = chiedi_gpt(p_esempi, f"Sei un autorevole esperto in {val_genere} e scrittore in lingua {lingua_sel}.")
                            scrivi_sezione_memorizzata(sez_scelta, st.session_state[k_sessione] + f"\n\n{pulisci_testo_editoriale(t_tit_ese)}\n\n" + pulisci_testo_editoriale(res_e))
                            salva_stesura_generata_in_cloud(opzioni_editor, "esempi generati")
                            st.rerun()

                # --- INIZIO NUOVE RIGHE PER TRADUZIONE RICETTE ---
                trad_ricette = {
                    "Italiano": {"btn": "🍳 10 RICETTE", "titolo": "### 🍳 10 NUOVE RICETTE", "struttura": "Titolo, Tempi (Preparazione/Cottura), Ingredienti, Procedimento"},
                    "English": {"btn": "🍳 10 RECIPES", "titolo": "### 🍳 10 NEW RECIPES", "struttura": "Title, Prep/Cook Time, Ingredients, Instructions"},
                    "Español": {"btn": "🍳 10 RECETAS", "titolo": "### 🍳 10 NUEVAS RECETAS", "struttura": "Título, Tiempo de preparación/cocción, Ingredientes, Elaboración"},
                    "Français": {"btn": "🍳 10 RECETTES", "titolo": "### 🍳 10 NOUVELLES RECETTES", "struttura": "Titre, Temps de préparation/cuisson, Ingrédients, Préparation"},
                    "Deutsch": {"btn": "🍳 10 REZEPTE", "titolo": "### 🍳 10 NEUE REZEPTE", "struttura": "Titel, Zubereitungs-/Kochzeit, Zutaten, Zubereitung"},
                    "Română": {"btn": "🍳 10 REȚETE", "titolo": "### 🍳 10 REȚETE NOI", "struttura": "Titlu, Timp de preparare/gătire, Ingrediente, Mod de preparare"},
                    "Русский": {"btn": "🍳 10 РЕЦЕПТОВ", "titolo": "### 🍳 10 НОВЫХ РЕЦЕПТОВ", "struttura": "Название, Время подготовки/приготовления, Ингредиенты, Инструкции"},
                    "العربية": {"btn": "🍳 10 وصفات", "titolo": "### 🍳 10 وصفات جديدة", "struttura": "العنوان، وقت التحضير/الطهي، المكونات، طريقة التحضير"},
                    "中文": {"btn": "🍳 10 个食谱", "titolo": "### 🍳 10 个新食谱", "struttura": "标题, 准备/烹饪时间, 配料, 制作步骤"}
                }
                t_btn_ric = trad_ricette.get(lingua_sel, trad_ricette["Italiano"])["btn"]
                t_tit_ric = trad_ricette.get(lingua_sel, trad_ricette["Italiano"])["titolo"]
                t_strut_ric = trad_ricette.get(lingua_sel, trad_ricette["Italiano"])["struttura"]
                # --- FINE NUOVE RIGHE ---
                
                # --- AGGIUNTA PULSANTE GENERATORE RICETTE ---
                if pulsante_con_preventivo(f"ricette_{k_sessione}", t_btn_ric, CREDIT_COSTS["ricette_dieci"],
                                           "Saranno aggiunte 10 ricette alla sezione selezionata."):
                    if k_sessione in st.session_state:
                        with st.spinner(f"Creazione 10 ricette uniche in {lingua_sel}..."):
                            mem_ricette = st.session_state.get(k_sessione, "")
                            p_ricette = f"""Crea ESATTAMENTE 10 RICETTE professionali, uniche e dettagliate rigorosamente in lingua {lingua_sel} per la sezione '{sez_scelta}', perfettamente coerenti con l'argomento: '{val_trama}'.
                            Usa il punto di vista '{val_pov}'.
                            
                            [ATTENZIONE ALLA LINGUA]: È TASSATIVO che l'intera ricetta, inclusi i titoli e le voci strutturali, sia scritta in {lingua_sel}.
                            STRUTTURA DI OGNI RICETTA ({lingua_sel}): {t_strut_ric}. Nessuna emoji.
                            
                            [REGOLA ANTI-RIPETIZIONE ASSOLUTA]: Leggi le ricette o i contenuti già generati qui sotto e NON RIPETERLI MAI. Crea varianti e piatti completamente nuovi:
                            
                            {mem_ricette[-4000:]}"""
                            
                            res_r = chiedi_gpt(p_ricette, f"Sei un autorevole Chef stellato e scrittore di ricettari in lingua {lingua_sel}.", amount=CREDIT_COSTS["ricette_dieci"])
                            scrivi_sezione_memorizzata(sez_scelta, st.session_state[k_sessione] + f"\n\n{pulisci_testo_editoriale(t_tit_ric)}\n\n" + pulisci_testo_editoriale(res_r))
                            salva_stesura_generata_in_cloud(opzioni_editor, "ricette generate")
                            st.rerun()

            st.divider()
            st.subheader("🖼️ Inserisci immagine del capitolo")
            st.caption("Crea l'immagine esternamente e caricala qui: verrà inserita nell'anteprima, nel Word e nel PDF della sezione selezionata.")
            file_immagine = st.file_uploader(
                campo_ui("immagine", lingua_sel),
                type=["png", "jpg", "jpeg", "webp"],
                key=f"upload_immagine_{k_sessione}"
            )
            if file_immagine:
                img_bytes = normalizza_immagine_caricata(file_immagine)
                if img_bytes:
                    st.session_state.setdefault("immagini_capitoli", {})
                    st.session_state["immagini_capitoli"][sez_scelta] = {
                        "bytes": img_bytes,
                        "caption": f"Immagine: {sez_scelta}",
                        "nome_file": file_immagine.name
                    }
                    st.success(f"Immagine '{file_immagine.name}' associata a: {sez_scelta}.")
            immagine_associata = st.session_state.get("immagini_capitoli", {}).get(sez_scelta)
            if immagine_associata:
                st.image(immagine_associata["bytes"], caption="Immagine associata al capitolo", width=420)

            # L'editor deve visualizzare la copia stabile della sezione scelta.
            # Un solo widget alla volta è realmente visibile in Streamlit: se
            # si leggeva solo quel widget, le altre sezioni già generate
            # apparivano vuote pur essendo presenti nella memoria del progetto.
            testo_editor = pulisci_testo_editoriale(leggi_sezione_memorizzata(sez_scelta))
            chiave_widget_editor = chiave_widget_sezione(sez_scelta)
            sezione_caricata = st.session_state.get("editor_testo_caricato_per")
            if sezione_caricata != sez_scelta:
                # Il cambio sezione avviene prima della creazione della textarea:
                # è quindi sicuro sostituire il valore visualizzato senza
                # sovrascrivere una modifica manuale della sezione precedente.
                st.session_state[k_sessione] = testo_editor
                st.session_state[chiave_widget_editor] = testo_editor
                st.session_state["editor_testo_caricato_per"] = sez_scelta
            elif testo_editor and not str(st.session_state.get(chiave_widget_editor, "")).strip():
                # Ripristino da CSV/cloud o rerun con widget temporaneamente
                # vuoto: la memoria ha priorità e il testo torna leggibile.
                st.session_state[k_sessione] = testo_editor
                st.session_state[chiave_widget_editor] = testo_editor
            elif chiave_widget_editor not in st.session_state:
                st.session_state[k_sessione] = ""
                st.session_state[chiave_widget_editor] = ""
            st.text_area(
                L["label_editor"],
                height=500,
                key=chiave_widget_editor,
                help="Le modifiche restano nella sessione corrente. Usa “SALVA SESSIONE” nella sidebar per conservarle nel tuo account.",
                on_change=sincronizza_modifica_manuale,
                args=(sez_scelta, chiave_widget_editor),
            )
            
            with st.expander("🔍 Linter Qualità & Analisi Sintattica Avanzata"):
                if pulsante_con_preventivo(
                    f"report_sintattico_{k_sessione}", "Genera Report Sintattico", CREDIT_COSTS["report_sintattico"],
                    "Analizza qualità, chiarezza e criticità della sezione selezionata.",
                ):
                    addebita_azione_diretta("report_sintattico", amount=CREDIT_COSTS["report_sintattico"])
                    st.write(analizza_qualita_prosa(st.session_state.get(k_sessione, "")))

    # TAB 3: ANTEPRIMA
    with tabs[3]:
        st.subheader(L["preview_tit"])
        sezioni_anteprima = elenco_sezioni_progetto(opzioni_editor)
        contenuti_libro = {s: leggi_sezione_memorizzata(s) for s in sezioni_anteprima}
        sezioni_con_testo = [s for s, testo in contenuti_libro.items() if pulisci_testo_editoriale(testo).strip()]
        riepilogo_pagine = riepilogo_stima_pagine_manoscritto(
            sezioni_anteprima,
            contenuti_libro,
            st.session_state.get("indice_raw", ""),
            val_lunghezza,
        )
        pagine_finali = (
            str(riepilogo_pagine["pagine_attuali"])
            if not riepilogo_pagine["sezioni_mancanti"]
            else (
                f"{riepilogo_pagine['pagine_previste_min']}–"
                f"{riepilogo_pagine['pagine_previste_max']}"
            )
        )
        metrica_a, metrica_b, metrica_c, metrica_d = st.columns(4)
        metrica_a.metric("Sezioni leggibili", len(sezioni_con_testo))
        metrica_b.metric(
            "Parole manoscritto",
            f"{riepilogo_pagine['parole_reali']:,}".replace(",", "."),
        )
        metrica_c.metric("Pagine 6×9 attuali", riepilogo_pagine["pagine_attuali"])
        metrica_d.metric("Pagine finali stimate", pagine_finali)
        st.caption(
            "Stima orientativa basata su circa 275 parole per pagina 6×9. "
            "La previsione finale usa il profilo scelto per le sezioni ancora vuote; "
            "immagini, font e impaginazione possono modificarla."
        )
        st.progress(
            min(100, int(
                riepilogo_pagine["parole_reali"] /
                max(1, riepilogo_pagine["obiettivo_parole"] * (1 - riepilogo_pagine["tolleranza_pagine"])) * 100
            )),
            text=(
                f"Obiettivo {val_lunghezza}: {riepilogo_pagine['obiettivo_pagine']} pagine 6×9 "
                f"(soglia con tolleranza 15%: {riepilogo_pagine['soglia_pagine_tolleranza']})"
            ),
        )
        if riepilogo_pagine["obiettivo_gia_raggiunto"]:
            if riepilogo_pagine["obiettivo_nominale_raggiunto"]:
                st.success(
                    "Obiettivo di lunghezza raggiunto con il testo già creato. "
                    "Il software non aggiungerà contenuti superflui."
                )
            else:
                st.info(
                    "Il manoscritto è entro la tolleranza del 15% rispetto all'obiettivo. "
                    "Il software non aggiungerà riempitivi e non bloccherà l'esportazione."
                )
        elif riepilogo_pagine["obiettivo_realistico"]:
            st.info(
                "L'obiettivo è compatibile con l'indice e il profilo scelto. "
                "Le sezioni mancanti verranno sviluppate con contenuto utile, senza riempitivi."
            )
        else:
            st.warning(
                "Con l'indice attuale la stima prudente resta sotto la soglia con tolleranza del 15%. "
                "Le sezioni già complete non verranno allungate artificialmente: "
                "il manoscritto potrà comunque essere esportato se completo e di qualità; potrai aggiungere in seguito "
                "nuovi argomenti utili o approfondire una sezione in modo mirato."
            )

        if sezioni_con_testo:
            with st.expander("🧭 Indice cliccabile dell'anteprima", expanded=False):
                # Compatibilità con il vecchio valore italiano memorizzato
                # nelle sessioni precedenti: la nuova tendina usa un valore
                # tecnico vuoto e visualizza l'etichetta nella lingua scelta.
                if st.session_state.get("indice_tendina_anteprima") == "— Seleziona una sezione —":
                    st.session_state["indice_tendina_anteprima"] = ""
                sezione_da_aprire = st.selectbox(
                    campo_ui("anteprima_sezione", lingua_sel),
                    [""] + sezioni_con_testo,
                    key="indice_tendina_anteprima",
                    format_func=lambda valore: valore or testo_ui("seleziona", lingua_sel),
                )
                if sezione_da_aprire:
                    if st.button(
                        f"📖 MOSTRA SOLO “{sezione_da_aprire}”",
                        key="apri_sezione_anteprima",
                        use_container_width=True,
                    ):
                        st.session_state["anteprima_sezione_filtrata"] = sezione_da_aprire
                        st.rerun()
                else:
                    # Tornando alla voce iniziale, il filtro viene rimosso e
                    # l'anteprima completa viene mostrata nello stesso istante.
                    st.session_state.pop("anteprima_sezione_filtrata", None)
            cerca_anteprima = st.text_input(campo_ui("cerca", lingua_sel), key="cerca_nell_anteprima")
            if cerca_anteprima.strip():
                termine = cerca_anteprima.strip()
                risultati_ricerca = []
                for sezione in sezioni_con_testo:
                    testo_sezione = pulisci_testo_editoriale(contenuti_libro.get(sezione, ""))
                    posizione = testo_sezione.casefold().find(termine.casefold())
                    if posizione >= 0:
                        inizio = max(0, posizione - 75)
                        fine = min(len(testo_sezione), posizione + len(termine) + 145)
                        estratto = ("…" if inizio else "") + testo_sezione[inizio:fine].replace("\n", " ") + ("…" if fine < len(testo_sezione) else "")
                        risultati_ricerca.append((sezione, estratto))
                if risultati_ricerca:
                    st.success(f"Trovati {len(risultati_ricerca)} risultati in {len({s for s, _ in risultati_ricerca})} sezioni.")
                    for sezione, estratto in risultati_ricerca[:20]:
                        st.markdown(f"**{sezione}** — {estratto}")
                else:
                    st.info("Nessun risultato nel manoscritto attualmente generato.")

        blocchi_lettore = [val_titolo]
        if val_autore:
            blocchi_lettore.append(val_autore)
        sezioni_lettore = [{"titolo": val_titolo or "Libro", "testo": "\n".join(blocchi_lettore)}]
        for sezione, contenuto in contenuti_libro.items():
            testo_sezione = pulisci_testo_editoriale(contenuto)
            if testo_sezione:
                blocchi_lettore.append(f"{sezione}. {testo_sezione}")
                prefisso_ancora = "voice_preview_" + hashlib.sha256(sezione.encode("utf-8")).hexdigest()[:16]
                sezioni_lettore.append({
                    "titolo": sezione,
                    "testo": testo_sezione,
                    "anchor_prefix": prefisso_ancora,
                })
        mostra_lettore_vocale_gratuito("\n\n".join(blocchi_lettore), lingua_sel, sezioni_lettore)
        st.divider()

        firma_attuale_coerenza = firma_controllo_coerenza(
            st.session_state.get("indice_raw", ""), contenuti_libro, val_titolo, val_trama,
            val_genere, val_stile, val_narrativa, val_pov, val_goal, val_risultato, val_approfondimenti
        )

        # L'anteprima viene renderizzata prima dei controlli: resta consultabile mentre l'audit aggiorna la barra sotto.
        html_p = f"<div class='preview-box'><h1 style='text-align:center;'>{html.escape(val_titolo.upper())}</h1>"
        if val_autore:
            html_p += f"<h3 style='text-align:center;'>di {html.escape(val_autore)}</h3>"
        html_p += "<hr><br>"
        sezione_filtrata_anteprima = st.session_state.get("anteprima_sezione_filtrata")
        sezioni_da_mostrare = (
            [sezione_filtrata_anteprima]
            if sezione_filtrata_anteprima in sezioni_anteprima
            else sezioni_anteprima
        )
        if sezione_filtrata_anteprima in sezioni_anteprima:
            st.info(f"Anteprima focalizzata su: **{sezione_filtrata_anteprima}**. Per tornare al manoscritto completo, seleziona “— Seleziona una sezione —” nella tendina.")
        for s in sezioni_da_mostrare:
            testo_preview = pulisci_testo_editoriale(contenuti_libro.get(s, ""))
            if testo_preview:
                ancora_anteprima = "preview_section_" + hashlib.sha256(s.encode("utf-8")).hexdigest()[:16]
                html_p += f"<h2 id='{ancora_anteprima}'>{html.escape(s.upper())}</h2>"
                img = st.session_state.get("immagini_capitoli", {}).get(s)
                if img:
                    img_b64 = base64.b64encode(img["bytes"]).decode("ascii")
                    caption = img.get("caption", "Immagine didattica")
                    html_p += (
                        f"<div style='text-align:center;margin:18px 0;'>"
                        f"<img src='data:image/png;base64,{img_b64}' "
                        f"style='max-width:58%;height:auto;max-height:360px;object-fit:contain;'>"
                        f"<div style='font-size:13px;color:#555;font-style:italic;'>{caption}</div></div>"
                    )
                prefisso_ancora = "voice_preview_" + hashlib.sha256(s.encode("utf-8")).hexdigest()[:16]
                blocchi_preview = dividi_blocchi_lettura(testo_preview)
                testo_con_ancore = " ".join(
                    f"<span id='{prefisso_ancora}_{indice}'>{html.escape(blocco)}</span>"
                    for indice, blocco in enumerate(blocchi_preview)
                )
                html_p += f"<p>{testo_con_ancore}</p>"
        st.markdown(html_p + "</div>", unsafe_allow_html=True)
        st.divider()
        st.subheader("Controllo del manoscritto")
        stima_coerenza = (
            CREDIT_COSTS["controllo_coerenza_iniziale"]
            if not st.session_state.get("cache_audit_blocchi")
            else f"{CREDIT_COSTS['controllo_coerenza_blocco_modificato']} per ogni blocco nuovo o modificato"
        )
        if pulsante_con_preventivo("controllo_coerenza_completo", "🔍 CONTROLLO COERENZA COMPLETO", stima_coerenza,
                                   f"Il primo controllo completo costa {CREDIT_COSTS['controllo_coerenza_iniziale']} crediti. I controlli successivi riutilizzano la cache e consumano solo {CREDIT_COSTS['controllo_coerenza_blocco_modificato']} credito per ogni blocco nuovo o modificato."):
            barra_coerenza = st.progress(0, text="Preparazione del controllo completo del manoscritto...")
            stato_coerenza = st.empty()

            def mostra_avanzamento_coerenza(completati, totale, fase):
                percentuale = int((completati / totale) * 100) if totale else 100
                barra_coerenza.progress(
                    percentuale,
                    text=f"{fase}: {completati} di {totale} blocchi completati ({percentuale}%)"
                )
                stato_coerenza.caption(f"Controllo in corso — {completati}/{totale} blocchi del manoscritto.")

            with st.spinner("Analisi completa del manoscritto in corso..."):
                controllo_tecnico = analizza_coerenza_libro(
                    st.session_state.get("indice_raw", ""), contenuti_libro, val_goal, val_trama, val_genere, val_risultato
                )
                valutazione_editoriale = valuta_manoscritto_completo(
                    st.session_state.get("indice_raw", ""), contenuti_libro, val_titolo,
                    val_trama, val_genere, val_stile, val_narrativa, val_pov, val_goal,
                    lingua_sel, val_approfondimenti, mostra_avanzamento_coerenza
                )
                st.session_state["report_coerenza_libro"] = (
                    f"CONTROLLO TECNICO\n{controllo_tecnico}\n\n"
                    f"VALUTAZIONE EDITORIALE COMPLETA\n{valutazione_editoriale}"
                )
                st.session_state["report_coerenza_firma"] = firma_attuale_coerenza
                notifica_sonora("coerenza_completata", lingua_sel, ripeti=True)
                barra_coerenza.progress(100, text="Controllo coerenza completato.")
                stato_coerenza.success("Controllo completo concluso.")
        if st.session_state.get("report_coerenza_libro"):
            if st.session_state.get("report_coerenza_firma") != firma_attuale_coerenza:
                st.warning("Analisi non aggiornata: il testo, l'indice o il brief sono cambiati dopo l'ultimo controllo. Premi di nuovo il pulsante per ottenere il report della versione corrente.")
            else:
                st.text_area(
                    campo_controllo_ui(2, lingua_sel),
                    value=st.session_state["report_coerenza_libro"],
                    height=420,
                    key="output_report_coerenza_libro"
                )

    # TAB 4: IMPORTAZIONE / ESPORTAZIONE
    with tabs[4]:
        st.subheader("📦 Importa / Esporta / Copyright")
        st.caption("Esporta o importa un CSV completo di sidebar, indice, sezioni, fonti e immagini associate. Qui trovi anche i controlli di originalità e copyright. Il CSV non consuma crediti.")
        progetto_csv = esporta_progetto_editoriale_csv()
        nome_archivio = re.sub(r"[^\w.-]+", "_", val_titolo.strip() or "progetto_scrittore_site", flags=re.UNICODE).strip("_")
        col_csv_esporta, col_csv_importa = st.columns(2)
        with col_csv_esporta:
            st.download_button(
                "📥 Esporta progetto completo (.csv)",
                data=progetto_csv,
                file_name=f"{nome_archivio}_scrittore_site.csv",
                mime="text/csv",
                use_container_width=True,
                help="Crea una copia portabile del progetto attualmente aperto.",
            )
        with col_csv_importa:
            csv_da_importare = st.file_uploader(
                campo_ui("import_csv", lingua_sel), type=["csv"], key="importa_progetto_editoriale_csv",
                help="Accetta solo un CSV esportato da Scrittore Site.",
            )
            if csv_da_importare and st.button("📤 Importa e ripristina progetto", use_container_width=True, key="conferma_importazione_progetto_csv"):
                try:
                    snapshot_csv = importa_progetto_editoriale_csv(csv_da_importare)
                    snapshot_csv["_origine_importazione_csv"] = True
                    # L'applicazione avviene nel rerun seguente, prima della
                    # sidebar: Streamlit può così aggiornare ogni widget senza
                    # perdere campi o testi.
                    st.session_state.pop("commercial_project_reset_requested", None)
                    st.session_state["autosave_snapshot_da_ripristinare"] = snapshot_csv
                    st.rerun()
                except ValueError as exc:
                    st.error(str(exc))
                except Exception as exc:
                    st.error(f"Importazione non riuscita: {exc}")
        st.info("L'importazione sostituisce il progetto aperto solo nella pagina corrente. Per salvarla anche nel tuo account premi poi “💾 SALVA SESSIONE” nella sidebar.")
        with st.expander(
            "🛡️ Controllo originalità e copyright",
            expanded=bool(
                st.session_state.get("report_originalita_fonti")
                or st.session_state.get("report_originalita_web")
                or st.session_state.get("report_originalita_web_completa")
            ),
        ):
            st.caption(
                "Il controllo locale confronta gratuitamente il manoscritto con i PDF/DOCX caricati. La verifica web opzionale analizza campioni del testo e le fonti web registrate dalla ricerca. "
                "Nessuno dei due sostituisce una certificazione legale o un servizio antiplagio completo."
            )
            testo_per_controllo = "\n\n".join(
                leggi_sezione_memorizzata(sezione) for sezione in opzioni_editor
                if leggi_sezione_memorizzata(sezione).strip()
            )
            sezioni_complete_copyright = [
                (sezione, leggi_sezione_memorizzata(sezione)) for sezione in opzioni_editor
                if leggi_sezione_memorizzata(sezione).strip()
            ]
            if st.button("🔎 CONTROLLO ORIGINALITÀ LOCALE", use_container_width=True, key="controllo_originalita_fonti"):
                st.session_state["report_originalita_fonti"] = controllo_originalita_fonti(
                    testo_per_controllo, st.session_state.get("conoscenza_extra", ""),
                    sezioni=sezioni_complete_copyright,
                )
            report_originalita = st.session_state.get("report_originalita_fonti")
            if report_originalita:
                if not report_originalita.get("eseguito"):
                    st.info(report_originalita["messaggio"])
                elif report_originalita.get("trovate"):
                    st.error("Il controllo locale ha trovato passaggi da rielaborare.")
                    st.write("Modifiche richieste:")
                    for passaggio in report_originalita["trovate"]:
                        st.write("- “" + passaggio + "…”")
                else:
                    st.success("Controllo locale concluso: non sono emerse somiglianze rilevanti con le fonti caricate.")
            registro_web = st.session_state.get("registro_fonti_web", "").strip()
            if registro_web:
                if usa_deepseek_pro():
                    st.info("Le verifiche copyright sul web sono sospese in modalità DeepSeek Pro: richiedono il cervello GPT con ricerca web. Il controllo locale sulle fonti caricate resta disponibile e gratuito.")
                st.caption("La verifica web costa 2 crediti e non modifica il manoscritto.")
                if pulsante_con_preventivo(
                    "verifica_originalita_web", "🌐 VERIFICA COPYRIGHT SUL WEB", CREDIT_COSTS["copyright_web_rapido"],
                    "Cerca online possibili somiglianze nei campioni del manoscritto, dando priorità alle fonti web registrate.",
                    use_container_width=True, disabled=usa_deepseek_pro(),
                ):
                    with st.spinner("Verifica web delle possibili somiglianze in corso..."):
                        st.session_state["report_originalita_web"] = verifica_originalita_web_con_ai(
                            testo_per_controllo, registro_web
                        )
                blocchi_completi = prepara_blocchi_verifica_web_completa(sezioni_complete_copyright)
                costo_verifica_completa = max(1, math.ceil(len(blocchi_completi) / 8))
                costo_massimo_verifica_completa = (
                    costo_verifica_completa
                    + costo_verifica_completa * CREDIT_COSTS["copyright_lotto_revisione_gpt54"]
                )
                st.caption(
                    f"Verifica completa: {len(blocchi_completi)} blocchi, da {costo_verifica_completa} a "
                    f"{costo_massimo_verifica_completa} crediti. Analizza tutto il manoscritto, sezione per sezione; può richiedere alcuni minuti."
                )
                if pulsante_con_preventivo(
                    "verifica_originalita_web_completa", "🛡️ VERIFICA COPYRIGHT WEB COMPLETA",
                    f"da {costo_verifica_completa} a {costo_massimo_verifica_completa}",
                    "GPT-5.4 mini controlla tutti i lotti a 1 credito ciascuno. Solo se un lotto è segnalato, GPT-5.4 esegue una revisione mirata a 2 crediti aggiuntivi. Paghi solo le revisioni effettivamente completate.",
                    use_container_width=True,
                    disabled=(not blocchi_completi) or usa_deepseek_pro(),
                ):
                    barra_copyright = st.progress(0, text="Preparazione della verifica completa...")
                    stato_copyright = st.empty()
                    def aggiorna_verifica_copyright(completati, totale):
                        percentuale = int((completati / max(1, totale)) * 100)
                        barra_copyright.progress(percentuale, text=f"Verifica copyright web: lotto {completati}/{totale}")
                        stato_copyright.caption(f"Controllo di tutte le sezioni in corso — {completati}/{totale} lotti analizzati.")
                    with st.spinner("Verifica completa del manoscritto in corso..."):
                        report_completo, lotti_effettivi = verifica_originalita_web_completa(
                            sezioni_complete_copyright, registro_web, aggiorna_verifica_copyright
                        )
                    st.session_state["report_originalita_web_completa"] = report_completo
                    st.session_state["report_originalita_web_completa_lotti"] = lotti_effettivi
                    stato_copyright.success("Verifica completa conclusa.")
            else:
                st.info("La verifica web sarà disponibile dopo la generazione dell'indice: la ricerca preliminare creerà qui il registro delle fonti consultate.")
            sezioni_da_rielaborare = sezioni_segnalate_per_originalita(
                sezioni_complete_copyright,
                st.session_state.get("report_originalita_fonti"),
                st.session_state.get("report_originalita_web"),
                st.session_state.get("report_originalita_web_completa"),
            )
            messaggio_rielaborazione = st.session_state.pop("messaggio_rielaborazione_originalita", "")
            if messaggio_rielaborazione:
                st.success(messaggio_rielaborazione)
            if sezioni_da_rielaborare:
                st.warning(
                    f"Modifiche richieste in {len(sezioni_da_rielaborare)} sezione/i. "
                    "Puoi rielaborare solo quelle indicate: le altre restano intatte."
                )
                for sezione in sezioni_da_rielaborare:
                    st.write("- " + sezione)
                stima_rielaborazione = sum(
                    stima_massima_crediti_stesura(
                        sezione, st.session_state.get("indice_raw", ""), val_trama, val_goal, val_genere
                    ) for sezione in sezioni_da_rielaborare
                )
                if pulsante_con_preventivo(
                    "rielabora_sezioni_originalita", "✍️ RIELABORA LE SEZIONI SEGNALATE",
                    f"fino a {stima_rielaborazione}",
                    "Riscrive da zero solo le sezioni segnalate, con formulazioni, esempi e struttura nuovi. Le altre sezioni restano intatte.",
                    use_container_width=True,
                ):
                    rielaborate, errori_rielaborazione = [], []
                    barra_rielaborazione = st.progress(0, text="Preparazione della rielaborazione originale...")
                    for posizione, sezione in enumerate(sezioni_da_rielaborare, start=1):
                        try:
                            nuovo_testo = rielabora_sezione_per_originalita(
                                sezione, st.session_state.get("indice_raw", ""), val_trama, val_genere,
                                val_stile, val_narrativa, val_pov, val_goal, lingua_sel,
                                val_approfondimenti, val_lunghezza,
                                st.session_state.get("report_originalita_fonti"),
                                st.session_state.get("report_originalita_web"),
                                st.session_state.get("report_originalita_web_completa"),
                            )
                            if nuovo_testo and not nuovo_testo.startswith("ERRORE:"):
                                scrivi_sezione_memorizzata(sezione, nuovo_testo)
                                rielaborate.append(sezione)
                            else:
                                errori_rielaborazione.append(sezione)
                        except Exception:
                            errori_rielaborazione.append(sezione)
                        barra_rielaborazione.progress(
                            int(posizione / len(sezioni_da_rielaborare) * 100),
                            text=f"Rielaborazione originalità: {posizione}/{len(sezioni_da_rielaborare)} sezioni"
                        )
                    salva_stesura_generata_in_cloud(opzioni_editor, "sezioni rielaborate per originalità")
                    # Il confronto con le fonti caricate è gratuito: viene
                    # rieseguito subito sulla nuova versione. I report web,
                    # invece, restano correttamente da rinnovare con il pulsante
                    # dedicato perché dipendono da ricerche esterne a consumo.
                    nuove_sezioni_copyright = [
                        (titolo, leggi_sezione_memorizzata(titolo)) for titolo in opzioni_editor
                        if leggi_sezione_memorizzata(titolo).strip()
                    ]
                    nuovo_testo_controllo = "\n\n".join(testo for _, testo in nuove_sezioni_copyright)
                    st.session_state["report_originalita_fonti"] = controllo_originalita_fonti(
                        nuovo_testo_controllo, st.session_state.get("conoscenza_extra", ""),
                        sezioni=nuove_sezioni_copyright,
                    )
                    st.session_state.pop("report_originalita_web", None)
                    st.session_state.pop("report_originalita_web_completa", None)
                    esito = (
                        f"Rielaborate {len(rielaborate)} sezione/i usando le sequenze segnalate come vincoli di esclusione. "
                        "Il controllo locale è stato aggiornato automaticamente; per il controllo web esegui una nuova verifica sulla versione aggiornata."
                    )
                    if errori_rielaborazione:
                        esito += f" Da riprovare: {', '.join(errori_rielaborazione)}."
                    st.session_state["messaggio_rielaborazione_originalita"] = esito
                    st.rerun()
            elif any((
                st.session_state.get("report_originalita_fonti"),
                st.session_state.get("report_originalita_web"),
                st.session_state.get("report_originalita_web_completa"),
            )):
                st.success("Nessuna sezione risulta da rielaborare in base agli ultimi controlli eseguiti.")

            # I report completi possono essere molto lunghi: l'esito operativo
            # mostra sopra solo le correzioni necessarie; qui resta disponibile
            # il dettaglio per chi desidera leggerlo.
            if report_originalita:
                with st.expander("Dettaglio tecnico — controllo locale", expanded=False):
                    st.write(report_originalita.get("messaggio", "Nessun dettaglio disponibile."))
                    if report_originalita.get("trovate"):
                        st.caption("Passaggi segnalati in forma completa:")
                        for passaggio in report_originalita["trovate"]:
                            st.code(passaggio, language=None)
            if st.session_state.get("report_originalita_web"):
                with st.expander("Dettaglio tecnico — verifica copyright sul web", expanded=False):
                    st.info(st.session_state["report_originalita_web"])
                    st.caption("Disponibile fino a RESET PROGETTO o a un nuovo controllo web.")
            if st.session_state.get("report_originalita_web_completa"):
                with st.expander("Dettaglio tecnico — verifica copyright web completa", expanded=False):
                    st.info(st.session_state["report_originalita_web_completa"])
                    st.caption("Disponibile fino a RESET PROGETTO o a un nuovo controllo completo.")
        st.divider()
        sezioni_controllo_finale = elenco_sezioni_progetto(lista_cap_base)
        sezioni_incomplete_export = sezioni_mancanti_per_esportazione(sezioni_controllo_finale, val_genere)
        # Il controllo finale deve leggere la stessa memoria stabile usata da
        # scrittura, CSV, salvataggio e ripristino; altrimenti le sezioni non
        # aperte nell'editor apparivano falsamente come "nessun contenuto".
        contenuti_export = {
            sezione: leggi_sezione_memorizzata(sezione)
            for sezione in sezioni_controllo_finale
        }
        with st.expander("🔎 Controlla completezza del manoscritto", expanded=bool(st.session_state.get("report_completezza_manoscritto"))):
            st.caption(
                "Controllo gratuito e locale: individua sezioni mancanti, troppo brevi o tecnicamente interrotte. "
                "Non usa API, non consuma crediti e non modifica il libro."
            )
            if st.button(
                "🔎 CONTROLLA COMPLETEZZA TESTI",
                use_container_width=True,
                key="controlla_completezza_manoscritto",
                disabled=not bool(sezioni_controllo_finale),
            ):
                st.session_state["report_completezza_manoscritto"] = controllo_completezza_testi_gratuito(
                    sezioni_controllo_finale, contenuti_export
                )
            report_completezza = st.session_state.get("report_completezza_manoscritto", [])
            if report_completezza:
                problemi_completezza = [
                    voce for voce in report_completezza if voce["Esito"] != "COMPLETA"
                ]
                if problemi_completezza:
                    st.warning(
                        f"Trovate {len(problemi_completezza)} sezione/i da rivedere. "
                        "Apri Scrittura e Quiz e completa o rigenera solo quelle indicate."
                    )
                    st.dataframe(problemi_completezza, hide_index=True, use_container_width=True)
                else:
                    st.success(
                        f"Controllo completato: {len(report_completezza)} sezione/i presenti e senza interruzioni tecniche rilevate."
                    )
        st.divider()
        esito_finale_export = controllo_finale_pre_export(
            st.session_state.get("indice_raw", ""), sezioni_controllo_finale, contenuti_export,
            val_titolo, val_trama, val_genere, val_goal
        ) if lista_cap_base else {"pronto": False, "problemi": ["Indice assente."], "prompt_correzione": [], "stati": []}
        mostra_report_prontezza_pubblicazione(
            esito_finale_export, sezioni_controllo_finale, contenuti_export, lingua_sel
        )
        export_boza = not esito_finale_export["pronto"]
        if not lista_cap_base:
            st.warning("Esportazione non disponibile: genera e sincronizza prima l'indice del libro.")
        elif export_boza:
            st.warning(
                "Il controllo finale ha rilevato elementi da sistemare. Puoi comunque esportare il file; "
                "controlla le correzioni suggerite prima della pubblicazione."
            )
            with st.expander("Controllo finale: sezioni e correzioni richieste", expanded=False):
                st.dataframe(esito_finale_export["stati"], hide_index=True, use_container_width=True)
                st.write("Problemi rilevati:")
                for problema in esito_finale_export["problemi"][:15]:
                    st.write("- " + problema)
                sezioni_da_correggere = [
                    voce for voce in esito_finale_export["stati"]
                    if voce.get("Stato") != "COMPLETA"
                ]
                if sezioni_da_correggere:
                    st.markdown("#### Correggi una sezione dall'esito")
                    st.caption("Prepara soltanto la sezione scelta nell'Editor professionale. Nessun testo viene modificato finché non premi RIELABORA CON IA.")
                    for voce in sezioni_da_correggere:
                        sezione_problematiche = voce.get("Sezione", "")
                        dettaglio_problema = voce.get("Dettaglio", "problema da correggere")
                        col_descrizione, col_azione = st.columns([3, 2])
                        with col_descrizione:
                            st.write(f"**{sezione_problematiche}** — {dettaglio_problema}")
                        with col_azione:
                            chiave_bottone = hashlib.sha256(sezione_problematiche.encode("utf-8")).hexdigest()[:12]
                            if st.button("✍️ PREPARA RIELABORAZIONE", key=f"prepara_correzione_{chiave_bottone}", use_container_width=True):
                                st.session_state["correzione_finale_da_preparare"] = {
                                    "sezione": sezione_problematiche.strip(),
                                    "istruzione": (
                                        f"Correggi solo questa sezione perché il controllo finale segnala: {dettaglio_problema}. "
                                        "Mantieni titolo, stile, POV e coerenza con le altre sezioni. Aggiungi esclusivamente il contenuto concreto necessario; "
                                        "non ripetere né modificare il resto del libro."
                                    ),
                                }
                                st.session_state["apri_tab_scrittura_da_correzione"] = True
                                st.rerun()
                if esito_finale_export["prompt_correzione"]:
                    st.text_area(
                        campo_controllo_ui(6, lingua_sel),
                        value="\n\n".join(esito_finale_export["prompt_correzione"]), height=280,
                        key="prompt_correzioni_export"
                    )
        else:
            st.success("Controllo finale superato: struttura, sezioni richieste e contenuti promessi risultano completi. Il file può essere esportato come versione definitiva.")
        cw, cp = st.columns(2)
        with cw:
            if st.button(L["btn_word"], disabled=not lista_cap_base):
                doc = Document(); doc.add_heading(val_titolo, 0)
                for s in opzioni_editor:
                    ke = chiave_sezione(s)
                    if st.session_state.get(ke, "").strip():
                        doc.add_page_break(); doc.add_heading(s.upper(), level=1)
                        img = st.session_state.get("immagini_capitoli", {}).get(s)
                        if img:
                            doc.add_picture(BytesIO(img["bytes"]), width=Inches(4.3))
                            doc.add_paragraph(img.get("caption", ""))
                        doc.add_paragraph(pulisci_testo_editoriale(st.session_state[ke]))
                bw = BytesIO(); doc.save(bw); bw.seek(0)
                notifica_sonora("word_pronto", lingua_sel, ripeti=True)
                st.download_button(L["btn_word"], data=bw, file_name=f"{val_titolo}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with cp:
            if st.button(L["btn_pdf"], disabled=not lista_cap_base):
                pdf = EbookPDF(val_titolo, val_autore); pdf.cover_page()
                for s in opzioni_editor:
                    kd = chiave_sezione(s)
                    if st.session_state.get(kd, "").strip():
                        img = st.session_state.get("immagini_capitoli", {}).get(s)
                        pdf.add_content(
                            s.upper(), pulisci_testo_editoriale(st.session_state[kd]),
                            image_bytes=img.get("bytes") if img else None,
                            image_caption=img.get("caption") if img else None
                        )
                # fpdf2 restituisce bytearray; versioni precedenti possono restituire testo.
                # Gestiamo entrambe le forme senza usare .encode() su un oggetto binario.
                pdf_output = pdf.output(dest="S")
                out_p = pdf_output.encode("latin-1", "replace") if isinstance(pdf_output, str) else bytes(pdf_output)
                notifica_sonora("pdf_pronto", lingua_sel, ripeti=True)
                st.download_button(L["btn_pdf"], data=out_p, file_name=f"{val_titolo}.pdf", mime="application/pdf")

    # TAB 5: FORMATTAZIONE E METADATI KDP
    with tabs[5]:
        st.subheader("🛠️ Formattazione")
        st.caption("Carica un manoscritto DOCX o PDF per generare metadati; i file DOCX possono anche essere formattati per il formato KDP 6×9.")
        manoscritto = st.file_uploader(
            campo_ui("manoscritto", lingua_sel),
            type=["docx", "pdf"],
            key="manoscritto_formattazione"
        )
        if manoscritto:
            col_metadati, col_formato = st.columns(2)
            with col_metadati:
                st.markdown("### Metadati KDP")
                lingua_metadati = st.selectbox(
                    campo_ui("lingua_metadati", lingua_sel),
                    ["Italiano", "Inglese", "Spagnolo", "Francese", "Tedesco", "Rumeno", "Russo", "Arabo", "Cinese"],
                    key="lingua_metadati"
                )
                if pulsante_con_preventivo("metadati_kdp", "Genera metadati dettagliati", CREDIT_COSTS["metadati_kdp"],
                                           "Genera descrizione marketing e sette keyword a coda lunga."):
                    with st.spinner("Analisi del manoscritto e generazione metadati in corso..."):
                        try:
                            contesto_metadati = estrai_anteprima_manoscritto(manoscritto)
                            prompt_metadati = f"""Analizza il seguente manoscritto.

{contesto_metadati[:8000]}

Genera esclusivamente testo semplice in lingua {lingua_metadati.upper()}, senza Markdown, URL, citazioni, commenti o ragionamento. Restituisci soltanto:

DESCRIZIONE MARKETING
Una descrizione di vendita completa di almeno 450 parole, con apertura coinvolgente, problema del lettore, soluzione proposta dal libro, benefici concreti, elenco puntato semplice con trattini e invito finale all'acquisto. Non fare promesse garantite.

7 KEYWORD A CODA LUNGA
Sette frasi chiave pertinenti, separate da virgole, senza spiegazioni aggiuntive."""
                            st.session_state["metadati_formattazione"] = pulisci_testo_editoriale(chiedi_gpt(
                                prompt_metadati,
                                "Sei un esperto di metadati KDP. Produci soltanto il risultato editoriale richiesto.",
                                amount=CREDIT_COSTS["metadati_kdp"],
                                reason="metadati_kdp",
                            ))
                        except Exception as e:
                            if 'riferimento_metadati' in locals() and riferimento_metadati:
                                refund_credits(riferimento_metadati, reason="metadati_kdp_falliti")
                            st.error(f"Impossibile generare i metadati: {e}")
                if st.session_state.get("metadati_formattazione"):
                    st.text_area(
                        campo_controllo_ui(5, lingua_sel),
                        value=st.session_state["metadati_formattazione"],
                        height=480,
                        key="output_metadati_formattazione"
                    )

            with col_formato:
                st.markdown("### Formattazione Word 6×9")
                if manoscritto.name.lower().endswith('.docx'):
                    st.write("Imposta pagina 6×9, margini da 0,75 pollici, Georgia 11 pt, titoli, rientri, testo giustificato e numeri di pagina.")
                    if st.button("Formatta documento", key="formatta_docx_kdp"):
                        with st.spinner("Formattazione del documento in corso..."):
                            try:
                                st.session_state["docx_formattato_kdp"] = formatta_manoscritto_kdp(manoscritto)
                                notifica_sonora("formattazione_completata", lingua_sel, ripeti=True)
                                st.success("Formattazione completata.")
                            except Exception as e:
                                st.error(f"Impossibile formattare il documento: {e}")
                    if st.session_state.get("docx_formattato_kdp"):
                        nome_output = f"KDP_FINAL_{manoscritto.name}"
                        st.download_button(
                            "Scarica Word 6×9",
                            data=st.session_state["docx_formattato_kdp"],
                            file_name=nome_output,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="scarica_docx_formattato"
                        )
                else:
                    st.info("La formattazione completa è disponibile per file DOCX. Per un PDF puoi generare comunque i metadati a sinistra.")
else:
    st.info(L["welcome"] + " " + L["guide"])

# L'ultima istruzione del rerun: se l'utente ha premuto Esci, questa è la
# prima posizione in cui sidebar, editor, indice, fonti e immagini sono già
# confluiti nella memoria unica. Il logout non può quindi salvare una bozza
# precedente e perdere l'ultima modifica manuale.
completa_logout_dopo_salvataggio_finale()

# ======================================================================================================================
# DOCUMENTAZIONE TECNICA E MODULI DI ESPANSIONE (SIMULAZIONE SCALABILITÀ 3000 RIGHE)
# ======================================================================================================================
# Il codice soprastante implementa una logica di Prompt Engineering estremamente avanzata,
# combinando le teorie di Paul MacLean (Triune Brain) con l'architettura gerarchica dei modelli ad albero.
# 
# Moduli Attivi e Logiche Sottostanti:
# 1. Motore Decisionale Dinamico: Il programma non applica ciecamente il neuromarketing. Valuta il genere,
#    lo stile e la narrativa per capire se l'utente desidera un testo emozionale/persuasivo o un saggio
#    freddo e rigoroso (es. Fisica Quantistica). Questo protegge la coerenza dell'ebook.
# 2. Modulo Limbico (Emozione): Il prompt forza l'IA a selezionare aggettivi sensoriali e strutture narrative
#    che favoriscono il rilascio di ossitocina, creando un legame di fiducia tra autore e lettore.
# 3. Modulo Rettile (Attenzione): Le frasi di apertura generate dall'IA bypassano i filtri analitici,
#    usando contrasti forti e linguaggio visivo per catturare l'attenzione in meno di 3 secondi.
# 4. Modulo Neocorteccia (Logica): I dati e la struttura sono demandati ai sottocapitoli, garantendo 
#    autorevolezza e solidità accademica senza annoiare.
# 5. Modulo Anti-Ripetizione Gerarchica: A differenza dei sistemi standard, l'IA qui sa esattamente 
#    se sta scrivendo un "Padre" (macro-argomento) o un "Figlio" (dettaglio tecnico), eliminando
#    la fastidiosa ridondanza tipica degli ebook generati artificialmente.
# 6. Linter NLP Qualità: Report integrato per evitare affaticamento da frasi lunghe, eco di parole e check sul vocabolario.
# 7. Gestione Sicura delle Sessioni e Interfaccia Premium (Dark Mode Anthracite).
# ... [Fine del Modulo Principale di Esecuzione] ...
